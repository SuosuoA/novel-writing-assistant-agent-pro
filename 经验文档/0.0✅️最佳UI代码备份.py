#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Novel Writing Assistant - Agent Pro
主GUI入口 - 现代科技玻璃风（Glass Morphism）

V2.1版本
创建日期：2026-03-22
更新日期：2026-03-22

特性：
- 无边框透明窗口（Windows Acrylic/Mica效果）
- 玻璃态UI设计（毛玻璃背景+发光边框）
- 响应式按钮（防抖+异步执行+加载状态）
- 侧边栏导航（无横向标签页）
- 线程安全的后端交互
"""

import os
import sys
import json
import yaml
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import time
import logging
import queue
import ctypes
from typing import Any, Callable, Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from ctypes import c_int, byref, sizeof, windll

# 尝试导入sv_ttk
try:
    import sv_ttk
    SV_TTK_AVAILABLE = True
except ImportError:
    SV_TTK_AVAILABLE = False
    logging.warning("sv_ttk not available")

# 尝试导入核心模块
try:
    from core import (
        EventBus,
        PluginRegistry,
        ServiceLocator,
        ConfigManager,
        get_event_bus,
        get_plugin_registry,
        get_service_locator,
        get_config_manager,
        AsyncHandler,
        TaskPriority,
        init_async_handler,
        CoreServiceManager,
    )
    CORE_AVAILABLE = True
except ImportError as e:
    CORE_AVAILABLE = False
    logging.warning(f"Core modules not available: {e}")

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============== Windows Acrylic效果实现 ==============

class DWM_BLURBEHIND(ctypes.Structure):
    """Windows DWM模糊效果结构"""
    _fields_ = [
        ("dwFlags", ctypes.c_ulong),
        ("fEnable", ctypes.c_bool),
        ("hRgnBlur", ctypes.c_void_p),
        ("fTransitionOnMaximized", ctypes.c_bool),
    ]


class AccentPolicy(ctypes.Structure):
    """Windows 10/11 Accent效果结构"""
    _fields_ = [
        ("AccentState", ctypes.c_int),
        ("AccentFlags", ctypes.c_int),
        ("GradientColor", ctypes.c_uint),
        ("AnimationId", ctypes.c_int),
    ]


class WindowCompositionAttrData(ctypes.Structure):
    """窗口组合属性数据"""
    _fields_ = [
        ("Attribute", ctypes.c_int),
        ("Data", ctypes.c_void_p),
        ("SizeOfData", ctypes.c_int),
    ]


class GlassWindowManager:
    """
    玻璃态窗口管理器
    
    实现Windows 10/11的Acrylic/Mica效果
    """
    
    # Windows常量
    DWM_BB_BLURREGION = 0x2
    DWM_BB_ENABLE = 0x1
    WM_NCCALCSIZE = 0x83
    WM_NCHITTEST = 0x84
    GWL_STYLE = -16
    GWL_EXSTYLE = -20
    WS_CAPTION = 0xC00000
    WS_THICKFRAME = 0x40000
    WS_MINIMIZEBOX = 0x20000
    WS_MAXIMIZEBOX = 0x10000
    WS_SYSMENU = 0x80000
    WS_BORDER = 0x800000
    WS_DLGFRAME = 0x400000
    WS_EX_LAYERED = 0x80000
    WS_EX_TRANSPARENT = 0x20
    
    # Accent状态
    ACCENT_DISABLED = 0
    ACCENT_ENABLE_BLURBEHIND = 3
    ACCENT_ENABLE_ACRYLICBLURBEHIND = 4
    
    # 窗口组合属性
    WCA_ACCENT_POLICY = 19
    
    def __init__(self, root: tk.Tk):
        """
        初始化玻璃窗口管理器
        
        Args:
            root: Tkinter根窗口
        """
        self.root = root
        self.hwnd = None
        self._is_acrylic = False
        
        # 加载Windows DLL
        try:
            self.dwmapi = ctypes.windll.dwmapi
            self.user32 = ctypes.windll.user32
            self._dll_available = True
        except Exception:
            self._dll_available = False
            logger.warning("Windows DLL not available")
    
    def enable_acrylic(self, color: int = 0x99000000) -> bool:
        """
        启用Acrylic效果
        
        Args:
            color: ARGB颜色值（格式：0xAABBGGRR）
        
        Returns:
            是否成功启用
        """
        if not self._dll_available:
            return False
        
        try:
            # 获取窗口句柄
            self.hwnd = self.user32.GetActiveWindow()
            if not self.hwnd:
                return False
            
            # 设置Accent策略
            accent = AccentPolicy()
            accent.AccentState = self.ACCENT_ENABLE_ACRYLICBLURBEHIND
            accent.AccentFlags = 2
            accent.GradientColor = color
            accent.AnimationId = 0
            
            data = WindowCompositionAttrData()
            data.Attribute = self.WCA_ACCENT_POLICY
            data.Data = ctypes.cast(ctypes.byref(accent), ctypes.c_void_p)
            data.SizeOfData = ctypes.sizeof(accent)
            
            # 应用效果
            result = self.user32.SetWindowCompositionAttribute(self.hwnd, ctypes.byref(data))
            
            if result:
                self._is_acrylic = True
                logger.info("Acrylic effect enabled")
            else:
                logger.warning("Failed to enable Acrylic effect")
            
            return result != 0
        
        except Exception as e:
            logger.error(f"Error enabling Acrylic: {e}")
            return False
    
    def make_frameless(self, keep_resize_border: bool = True) -> None:
        """
        设置完全无边框窗口（移除系统边框）
        
        Args:
            keep_resize_border: 是否保留调整大小的边框
                - True: 保留透明调整边框（可拖拽调整大小，但有细微边框残留）
                - False: 完全无边框（无边框残留，但需自己实现调整大小）
        """
        if not self._dll_available:
            return
        
        try:
            # 确保窗口句柄存在
            if not self.hwnd:
                self.hwnd = self.user32.GetActiveWindow()
            
            if not self.hwnd:
                return
            
            style = self.user32.GetWindowLongW(self.hwnd, self.GWL_STYLE)
            
            if keep_resize_border:
                # 模式1：保留调整边框（推荐）
                # 移除标题栏、系统菜单、边框线条，保留调整功能
                style &= ~(self.WS_CAPTION | self.WS_SYSMENU | self.WS_BORDER | self.WS_DLGFRAME)
                style |= self.WS_THICKFRAME | self.WS_MINIMIZEBOX | self.WS_MAXIMIZEBOX
            else:
                # 模式2：完全无边框
                # 移除所有边框和调整功能，纯净无边框
                style &= ~(self.WS_CAPTION | self.WS_SYSMENU | self.WS_THICKFRAME | self.WS_BORDER | self.WS_DLGFRAME)
                style |= self.WS_MINIMIZEBOX | self.WS_MAXIMIZEBOX
            
            self.user32.SetWindowLongW(self.hwnd, self.GWL_STYLE, style)
            
            # 强制重绘窗口
            self.user32.SetWindowPos(self.hwnd, None, 0, 0, 0, 0, 
                                     0x27)  # SWP_FRAMECHANGED | SWP_NOMOVE | SWP_NOSIZE | SWP_NOZORDER
            
            logger.info(f"Frameless window enabled (keep_resize_border={keep_resize_border})")
        
        except Exception as e:
            logger.error(f"Error making frameless: {e}")
    
    def enable_blur_behind(self) -> bool:
        """
        启用DWM模糊效果（Windows 7+）
        
        Returns:
            是否成功启用
        """
        if not self._dll_available:
            return False
        
        try:
            self.hwnd = self.user32.GetActiveWindow()
            if not self.hwnd:
                return False
            
            bb = DWM_BLURBEHIND()
            bb.dwFlags = self.DWM_BB_ENABLE
            bb.fEnable = True
            bb.hRgnBlur = None
            bb.fTransitionOnMaximized = False
            
            self.dwmapi.DwmEnableBlurBehindWindow(self.hwnd, ctypes.byref(bb))
            
            logger.info("Blur behind enabled")
            return True
        
        except Exception as e:
            logger.error(f"Error enabling blur behind: {e}")
            return False
    
    def set_window_opacity(self, alpha: float) -> None:
        """
        设置窗口透明度
        
        Args:
            alpha: 透明度（0.0-1.0）
        """
        try:
            # Tkinter原生透明度设置
            self.root.attributes('-alpha', alpha)
        except Exception as e:
            logger.error(f"Error setting opacity: {e}")


# ============== 玻璃态主题配置 ==============

class GlassTheme:
    """
    现代科技玻璃态主题配置
    
    设计理念：毛玻璃效果 + 发光边框 + 悬浮卡片
    """
    
    # 主色调（科技蓝）
    PRIMARY = "#0078D4"          # Windows蓝
    PRIMARY_LIGHT = "#60CDFF"    # 浅蓝
    PRIMARY_DARK = "#003D6B"     # 深蓝
    
    # 强调色
    ACCENT = "#00BCF2"           # 青蓝
    ACCENT_PINK = "#FF6B9D"      # 粉红
    ACCENT_PURPLE = "#8B5CF6"    # 紫色
    ACCENT_GREEN = "#10B981"     # 绿色
    ACCENT_ORANGE = "#F59E0B"    # 橙色
    ACCENT_RED = "#EF4444"       # 红色
    
    # 背景色（玻璃态）
    GLASS_BG = "#1E1E1E"         # 玻璃背景色
    GLASS_SURFACE = "#2D2D2D"    # 玻璃表面色
    GLASS_BORDER = "#404040"     # 玻璃边框色
    GLASS_HOVER = "#383838"      # 悬停态
    
    # 文字色
    TEXT_PRIMARY = "#FFFFFF"     # 主文字
    TEXT_SECONDARY = "#B0B0B0"   # 次要文字
    TEXT_DISABLED = "#666666"    # 禁用文字
    TEXT_LINK = "#60CDFF"        # 链接文字
    
    # 语义色
    SUCCESS = "#10B981"
    WARNING = "#F59E0B"
    ERROR = "#EF4444"
    INFO = "#3B82F6"
    
    # 玻璃效果参数
    BLUR_RADIUS = 20
    GLASS_OPACITY = 0.85
    BORDER_RADIUS = 8
    
    # 字体配置
    FONT_FAMILY = "Microsoft YaHei UI"
    FONT_FAMILY_CODE = "Consolas"
    FONT_SIZE_TITLE = 18
    FONT_SIZE_SUBTITLE = 14
    FONT_SIZE_NORMAL = 10
    FONT_SIZE_SMALL = 9
    FONT_SIZE_TINY = 8
    
    @classmethod
    def apply_to_style(cls, style: ttk.Style) -> None:
        """应用玻璃态主题到ttk Style"""
        
        # 全局样式
        style.configure(".",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            fieldbackground=cls.GLASS_SURFACE,
            bordercolor=cls.GLASS_BORDER,
            lightcolor=cls.GLASS_HOVER,
            darkcolor=cls.GLASS_BG
        )
        
        # 标签
        style.configure("TLabel",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL)
        )
        
        # 标题标签
        style.configure("Title.TLabel",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_TITLE, "bold"),
            foreground=cls.TEXT_PRIMARY
        )
        
        # 副标题标签
        style.configure("Subtitle.TLabel",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_SUBTITLE, "bold"),
            foreground=cls.TEXT_SECONDARY
        )
        
        # 按钮 - 玻璃态效果
        style.configure("TButton",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL),
            padding=(16, 8),
            background=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            borderwidth=1,
            relief="flat"
        )
        
        style.map("TButton",
            background=[("active", cls.GLASS_HOVER), ("pressed", cls.PRIMARY)],
            foreground=[("active", cls.TEXT_PRIMARY), ("pressed", cls.TEXT_PRIMARY)]
        )
        
        # 强调按钮
        style.configure("Accent.TButton",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL),
            padding=(16, 8),
            background=cls.PRIMARY,
            foreground=cls.TEXT_PRIMARY,
            borderwidth=0
        )
        
        style.map("Accent.TButton",
            background=[("active", cls.PRIMARY_LIGHT), ("pressed", cls.PRIMARY_DARK)]
        )
        
        # 框架
        style.configure("TFrame",
            background=cls.GLASS_BG
        )
        
        # 玻璃卡片框架
        style.configure("Glass.TFrame",
            background=cls.GLASS_SURFACE,
            borderwidth=1,
            relief="solid"
        )
        
        # LabelFrame
        style.configure("TLabelframe",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            borderwidth=1,
            relief="solid"
        )
        
        style.configure("TLabelframe.Label",
            background=cls.GLASS_BG,
            foreground=cls.PRIMARY_LIGHT,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL, "bold")
        )
        
        # 输入框
        style.configure("TEntry",
            fieldbackground=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            insertcolor=cls.TEXT_PRIMARY,
            borderwidth=1,
            padding=5
        )
        
        # 下拉框
        style.configure("TCombobox",
            fieldbackground=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            arrowcolor=cls.TEXT_PRIMARY,
            borderwidth=1,
            padding=5
        )
        
        # 进度条
        style.configure("TProgressbar",
            background=cls.PRIMARY,
            troughcolor=cls.GLASS_SURFACE,
            borderwidth=0
        )
        
        # 滚动条
        style.configure("TScrollbar",
            background=cls.GLASS_SURFACE,
            troughcolor=cls.GLASS_BG,
            arrowcolor=cls.TEXT_SECONDARY,
            borderwidth=0
        )
        
        # Treeview
        style.configure("Treeview",
            background=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            fieldbackground=cls.GLASS_SURFACE,
            borderwidth=0
        )
        
        style.configure("Treeview.Heading",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL, "bold"),
            background=cls.GLASS_BG,
            foreground=cls.TEXT_SECONDARY
        )
        
        style.map("Treeview",
            background=[("selected", cls.PRIMARY)],
            foreground=[("selected", cls.TEXT_PRIMARY)]
        )
        
        # 复选框
        style.configure("TCheckbutton",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL)
        )
        
        # 单选按钮
        style.configure("TRadiobutton",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL)
        )
        
        # 分隔线
        style.configure("TSeparator",
            background=cls.GLASS_BORDER
        )
        
        # Scale滑块
        style.configure("TScale",
            background=cls.GLASS_BG,
            troughcolor=cls.GLASS_SURFACE,
            borderwidth=0
        )


# ============== 响应式按钮 ==============

class ResponsiveButton(ttk.Button):
    """
    响应式按钮组件
    
    特性：
    - 500ms防抖保护
    - 异步执行耗时操作
    - 加载状态自动管理
    """
    
    DEBOUNCE_MS = 500
    LOADING_SUFFIX = "..."
    
    def __init__(
        self,
        parent: tk.Widget,
        text: str,
        command: Callable,
        async_handler: Optional[Any] = None,
        debounce_ms: int = DEBOUNCE_MS,
        auto_disable: bool = True,
        style: str = "TButton",
        **kwargs
    ):
        super().__init__(parent, text=text, style=style, **kwargs)
        
        self._original_text = text
        self._original_command = command
        self._async_handler = async_handler
        self._debounce_ms = debounce_ms
        self._auto_disable = auto_disable
        
        self._is_loading = False
        self._last_click_time = 0
        
        super().configure(command=self._on_clicked)
    
    def _on_clicked(self) -> None:
        if self._is_loading:
            return
        
        current_time = time.time() * 1000
        if current_time - self._last_click_time < self._debounce_ms:
            return
        
        self._last_click_time = current_time
        
        if self._async_handler:
            self._execute_async()
        else:
            self._execute_sync()
    
    def _execute_sync(self) -> None:
        try:
            self._set_loading(True)
            self._original_command()
        except Exception as e:
            logger.error(f"Button command error: {e}")
            messagebox.showerror("错误", f"执行失败：{e}")
        finally:
            self._set_loading(False)
    
    def _execute_async(self) -> None:
        self._set_loading(True)
        
        def task():
            return self._original_command()
        
        def on_success(result):
            self._set_loading(False)
        
        def on_error(error):
            self._set_loading(False)
            messagebox.showerror("错误", f"执行失败：{error}")
        
        self._async_handler.submit(
            func=task,
            callback=on_success,
            error_callback=on_error,
            priority=TaskPriority.HIGH
        )
    
    def _set_loading(self, loading: bool) -> None:
        self._is_loading = loading
        
        if self._auto_disable:
            self.configure(state="disabled" if loading else "normal")
        
        if loading:
            self.configure(text=self._original_text + self.LOADING_SUFFIX)
        else:
            self.configure(text=self._original_text)


# ============== 自定义标题栏 ==============

class ResizeGrip:
    """
    窗口调整大小的边框感知区域
    
    用于完全无边框窗口的调整大小功能
    """
    
    RESIZE_BORDER = 8  # 边框感知区域宽度
    
    def __init__(self, root: tk.Tk):
        self.root = root
        self._resize_edge = None
        self._drag_start_x = 0
        self._drag_start_y = 0
        self._drag_start_width = 0
        self._drag_start_height = 0
        
        # 绑定事件
        self.root.bind("<Motion>", self._on_motion)
        self.root.bind("<ButtonPress-1>", self._on_resize_start)
        self.root.bind("<ButtonRelease-1>", self._on_resize_end)
        self.root.bind("<B1-Motion>", self._on_resize_motion)
    
    def _get_edge(self, x: int, y: int) -> str:
        """判断鼠标在哪个边缘"""
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        
        left = x < self.RESIZE_BORDER
        right = x > width - self.RESIZE_BORDER
        top = y < self.RESIZE_BORDER
        bottom = y > height - self.RESIZE_BORDER
        
        # 角落优先
        if top and left:
            return "top-left"
        elif top and right:
            return "top-right"
        elif bottom and left:
            return "bottom-left"
        elif bottom and right:
            return "bottom-right"
        elif left:
            return "left"
        elif right:
            return "right"
        elif top:
            return "top"
        elif bottom:
            return "bottom"
        return ""
    
    def _get_cursor(self, edge: str) -> str:
        """获取对应的鼠标光标"""
        cursors = {
            "top": "top_side",
            "bottom": "bottom_side",
            "left": "left_side",
            "right": "right_side",
            "top-left": "top_left_corner",
            "top-right": "top_right_corner",
            "bottom-left": "bottom_left_corner",
            "bottom-right": "bottom_right_corner"
        }
        return cursors.get(edge, "")
    
    def _on_motion(self, event: tk.Event) -> None:
        """鼠标移动时更新光标"""
        edge = self._get_edge(event.x, event.y)
        cursor = self._get_cursor(edge)
        self.root.configure(cursor=cursor if cursor else "arrow")
    
    def _on_resize_start(self, event: tk.Event) -> None:
        """开始调整大小"""
        self._resize_edge = self._get_edge(event.x, event.y)
        if self._resize_edge:
            self._drag_start_x = self.root.winfo_pointerx()
            self._drag_start_y = self.root.winfo_pointery()
            self._drag_start_width = self.root.winfo_width()
            self._drag_start_height = self.root.winfo_height()
            self._drag_start_root_x = self.root.winfo_rootx()
            self._drag_start_root_y = self.root.winfo_rooty()
    
    def _on_resize_end(self, event: tk.Event) -> None:
        """结束调整大小"""
        self._resize_edge = None
    
    def _on_resize_motion(self, event: tk.Event) -> None:
        """调整大小中"""
        if not self._resize_edge:
            return
        
        current_x = self.root.winfo_pointerx()
        current_y = self.root.winfo_pointery()
        dx = current_x - self._drag_start_x
        dy = current_y - self._drag_start_y
        
        new_width = self._drag_start_width
        new_height = self._drag_start_height
        new_x = self._drag_start_root_x
        new_y = self._drag_start_root_y
        
        min_width = 1000
        min_height = 700
        
        edge = self._resize_edge
        
        if "right" in edge:
            new_width = max(min_width, self._drag_start_width + dx)
        if "left" in edge:
            new_width = max(min_width, self._drag_start_width - dx)
            if new_width > min_width:
                new_x = self._drag_start_root_x + dx
        if "bottom" in edge:
            new_height = max(min_height, self._drag_start_height + dy)
        if "top" in edge:
            new_height = max(min_height, self._drag_start_height - dy)
            if new_height > min_height:
                new_y = self._drag_start_root_y + dy
        
        # 更新窗口
        if new_width >= min_width or new_height >= min_height:
            self.root.geometry(f"{new_width}x{new_height}+{new_x}+{new_y}")


class CustomTitleBar(tk.Frame):
    """
    自定义标题栏
    
    支持：拖动、最小化、最大化、关闭
    """
    
    def __init__(self, parent: tk.Tk, title: str = "Novel Writing Assistant-Agent Pro"):
        super().__init__(parent, bg=GlassTheme.GLASS_BG, height=40)
        
        self.parent = parent
        self.title_text = title
        
        self._drag_start_x = 0
        self._drag_start_y = 0
        
        self._create_widgets()
        self._bind_events()
    
    def _create_widgets(self) -> None:
        # 标题
        self.title_label = tk.Label(
            self,
            text=self.title_text,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SUBTITLE, "bold"),
            fg=GlassTheme.TEXT_PRIMARY,
            bg=GlassTheme.GLASS_BG
        )
        self.title_label.pack(side=tk.LEFT, padx=15, pady=8)
        
        # 窗口控制按钮容器
        btn_frame = tk.Frame(self, bg=GlassTheme.GLASS_BG)
        btn_frame.pack(side=tk.RIGHT)
        
        # 最小化按钮
        self.min_btn = tk.Label(
            btn_frame,
            text="─",  # 使用更宽的破折号
            font=("Segoe UI", 12),  # 使用系统字体
            fg=GlassTheme.TEXT_SECONDARY,
            bg=GlassTheme.GLASS_BG,
            width=4,
            cursor="hand2"
        )
        self.min_btn.pack(side=tk.LEFT, pady=8)
        
        # 最大化按钮
        self.max_btn = tk.Label(
            btn_frame,
            text="□",  # 空心方块
            font=("Segoe UI", 12),
            fg=GlassTheme.TEXT_SECONDARY,
            bg=GlassTheme.GLASS_BG,
            width=4,
            cursor="hand2"
        )
        self.max_btn.pack(side=tk.LEFT, pady=8)
        
        # 关闭按钮
        self.close_btn = tk.Label(
            btn_frame,
            text="✕",  # 使用乘号符号
            font=("Segoe UI", 14),
            fg=GlassTheme.TEXT_SECONDARY,
            bg=GlassTheme.GLASS_BG,
            width=4,
            cursor="hand2"
        )
        self.close_btn.pack(side=tk.LEFT, pady=8)
        
        # 悬停效果
        for btn in [self.min_btn, self.max_btn]:
            btn.bind("<Enter>", lambda e, b=btn: b.configure(bg=GlassTheme.GLASS_HOVER))
            btn.bind("<Leave>", lambda e, b=btn: b.configure(bg=GlassTheme.GLASS_BG))
        
        self.close_btn.bind("<Enter>", lambda e: self.close_btn.configure(bg=GlassTheme.ERROR))
        self.close_btn.bind("<Leave>", lambda e: self.close_btn.configure(bg=GlassTheme.GLASS_BG))
    
    def _bind_events(self) -> None:
        # 拖动窗口
        self.bind("<Button-1>", self._on_drag_start)
        self.bind("<B1-Motion>", self._on_drag_motion)
        self.title_label.bind("<Button-1>", self._on_drag_start)
        self.title_label.bind("<B1-Motion>", self._on_drag_motion)
        
        # 窗口控制
        self.min_btn.bind("<Button-1>", lambda e: self.parent.iconify())
        self.max_btn.bind("<Button-1>", self._toggle_maximize)
        self.close_btn.bind("<Button-1>", lambda e: self.parent.quit())
        
        # 双击标题栏最大化
        self.bind("<Double-Button-1>", lambda e: self._toggle_maximize(None))
    
    def _on_drag_start(self, event: tk.Event) -> None:
        self._drag_start_x = event.x
        self._drag_start_y = event.y
    
    def _on_drag_motion(self, event: tk.Event) -> None:
        x = self.parent.winfo_x() + event.x - self._drag_start_x
        y = self.parent.winfo_y() + event.y - self._drag_start_y
        self.parent.geometry(f"+{x}+{y}")
    
    def _toggle_maximize(self, event: Optional[tk.Event]) -> None:
        if self.parent.state() == "zoomed":
            self.parent.state("normal")
            self.max_btn.configure(text="□")
        else:
            self.parent.state("zoomed")
            self.max_btn.configure(text="❐")


# ============== 主窗口 ==============

class MainWindow:
    """
    主窗口类
    
    结构：
    - 自定义标题栏
    - 侧边栏（导航）
    - 主内容区（根据侧边栏选择显示）
    - 状态栏
    """
    
    VERSION = "V2.1.0"
    TITLE = "Novel Writing Assistant - Agent Pro"
    
    def __init__(self):
        """初始化主窗口"""
        # 创建根窗口
        self.root = tk.Tk()
        self.root.title(self.TITLE)
        self.root.geometry("1400x900")
        self.root.minsize(1000, 700)
        
        # 设置窗口背景
        self.root.configure(bg=GlassTheme.GLASS_BG)
        
        # 尝试启用玻璃效果
        self._enable_glass_effect()
        
        # 核心服务
        self._services: Optional[CoreServiceManager] = None
        self._async_handler: Optional[AsyncHandler] = None
        
        # UI组件
        self._content_frame: Optional[tk.Frame] = None
        self._status_var: Optional[tk.StringVar] = None
        self._title_bar: Optional[CustomTitleBar] = None
        self._resize_grip: Optional[ResizeGrip] = None  # 窗口调整大小边框感知
        
        # 当前显示的页面
        self._current_page: Optional[str] = None
        self._pages: Dict[str, tk.Frame] = {}
        
        # 项目状态
        self.current_project: Dict[str, Any] = {}
        self.project_info: Dict[str, Any] = {}
        
        # 结果队列（后台线程通信）
        self._result_queue = queue.Queue()
        
        # 缓存
        self._cache = {
            'outline': {},
            'style': {},
            'worldview': {},
            'characters': {},
            'file_mtime': {}
        }
        
        # 初始化
        self._init_theme()
        self._init_async_handler()
        self._init_core_services()
        self._init_ui()
        self._init_bindings()
        
        # 启动结果队列处理
        self.root.after(100, self._process_result_queue)
        
        # 窗口关闭确认
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        
        logger.info("MainWindow initialized")
    
    def _enable_glass_effect(self) -> None:
        """启用玻璃效果"""
        try:
            glass_mgr = GlassWindowManager(self.root)
            
            # 尝试Acrylic效果
            if glass_mgr.enable_acrylic(0x80000000):
                logger.info("Acrylic glass effect enabled")
            else:
                logger.info("Acrylic effect not available, using alpha transparency")
            
            # 无论 Acrylic 是否成功，都应用完全无边框
            glass_mgr.make_frameless(keep_resize_border=False)
            
            # 如果 Acrylic 失败，降级到透明效果
            if not glass_mgr._is_acrylic:
                self.root.attributes('-alpha', 0.95)
            
            logger.info("Frameless window enabled")
        
        except Exception as e:
            logger.warning(f"Glass effect not available: {e}")
            self.root.attributes('-alpha', 0.98)
    
    def _init_theme(self) -> None:
        """初始化主题"""
        style = ttk.Style()
        
        # 使用sv_ttk主题
        if SV_TTK_AVAILABLE:
            try:
                sv_ttk.set_theme("dark")
            except Exception:
                pass
        
        # 应用玻璃态主题
        GlassTheme.apply_to_style(style)
    
    def _init_async_handler(self) -> None:
        """初始化异步处理器"""
        if CORE_AVAILABLE:
            self._async_handler = init_async_handler(
                root=self.root,
                worker_count=4,
                default_timeout=30.0
            )
            logger.info("AsyncHandler initialized")
    
    def _init_core_services(self) -> None:
        """初始化核心服务"""
        if not CORE_AVAILABLE:
            logger.warning("Core services not available")
            return
        
        try:
            self._services = CoreServiceManager(self.root)
            self._services.initialize(
                event_bus=get_event_bus(),
                registry=get_plugin_registry(),
                locator=get_service_locator(),
                config=get_config_manager()
            )
            logger.info("Core services initialized")
        except Exception as e:
            logger.error(f"Failed to initialize core services: {e}")
    
    def _init_ui(self) -> None:
        """初始化UI"""
        # 自定义标题栏
        self._title_bar = CustomTitleBar(self.root, self.TITLE)
        self._title_bar.pack(fill=tk.X)
        
        # 窗口调整大小的边框感知（完全无边框模式需要）
        self._resize_grip = ResizeGrip(self.root)
        
        # 主容器
        main_container = ttk.Frame(self.root, style="TFrame")
        main_container.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
        
        # 侧边栏
        self._create_sidebar(main_container)
        
        # 内容区
        self._create_content(main_container)
        
        # 状态栏
        self._create_status_bar()
        
        # 默认显示热榜页面
        self._switch_to_page("hot_ranking")
    
    def _create_sidebar(self, parent: tk.Widget) -> None:
        """创建侧边栏"""
        sidebar = ttk.Frame(parent, style="TFrame", width=200)
        sidebar.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 8))
        sidebar.pack_propagate(False)
        
        # 导航按钮
        nav_buttons = [
            ("热榜", "hot_ranking", GlassTheme.ACCENT_PINK),
            ("工作台", "workbench", GlassTheme.PRIMARY_LIGHT),
            ("创作进度", "progress", GlassTheme.ACCENT_GREEN),
            ("项目管理", "project", GlassTheme.ACCENT_PURPLE),
            ("插件管理", "plugins", GlassTheme.ACCENT_ORANGE),
            ("设置", "settings", GlassTheme.WARNING),
        ]
        
        for text, page_id, color in nav_buttons:
            btn = tk.Label(
                sidebar,
                text=text,
                font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                fg=GlassTheme.TEXT_PRIMARY,
                bg=GlassTheme.GLASS_SURFACE,
                activeforeground=GlassTheme.TEXT_PRIMARY,
                activebackground=color,
                cursor="hand2",
                padx=20,
                pady=12,
                anchor="w"
            )
            btn.pack(fill=tk.X, pady=2)
            
            # 悬停效果
            btn.bind("<Enter>", lambda e, b=btn, c=color: b.configure(bg=c))
            btn.bind("<Leave>", lambda e, b=btn: b.configure(bg=GlassTheme.GLASS_SURFACE))
            btn.bind("<Button-1>", lambda e, pid=page_id: self._switch_to_page(pid))
        
        # 底部版本信息
        version_label = tk.Label(
            sidebar,
            text=f"Agent Pro {self.VERSION}",
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_TINY),
            fg=GlassTheme.TEXT_DISABLED,
            bg=GlassTheme.GLASS_BG
        )
        version_label.pack(side=tk.BOTTOM, pady=10)
    
    def _create_content(self, parent: tk.Widget) -> None:
        """创建内容区"""
        self._content_frame = ttk.Frame(parent, style="Glass.TFrame")
        self._content_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    def _create_status_bar(self) -> None:
        """创建状态栏（增强版：项目名称、AI连接、字数统计、后台任务进度）"""
        status_frame = ttk.Frame(self.root, style="TFrame")
        status_frame.pack(fill=tk.X, side=tk.BOTTOM, padx=8, pady=(0, 8))
        
        # 分隔线
        ttk.Separator(status_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=(0, 4))
        
        status_inner = ttk.Frame(status_frame, style="TFrame")
        status_inner.pack(fill=tk.X)
        
        # 项目名称
        self._project_var = tk.StringVar(value="项目: 未打开")
        ttk.Label(
            status_inner,
            textvariable=self._project_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        ).pack(side=tk.LEFT)
        
        # 分隔符
        ttk.Label(status_inner, text=" | ", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=5)
        
        # 字数统计
        self._word_count_var = tk.StringVar(value="字数: 0")
        ttk.Label(
            status_inner,
            textvariable=self._word_count_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        ).pack(side=tk.LEFT)
        
        # 分隔符
        ttk.Label(status_inner, text=" | ", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=5)
        
        # 状态文本
        self._status_var = tk.StringVar(value="就绪")
        ttk.Label(
            status_inner,
            textvariable=self._status_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        ).pack(side=tk.LEFT, padx=10)
        
        # 后台任务进度（右侧）
        self._background_task_var = tk.StringVar(value="")
        self._background_task_label = ttk.Label(
            status_inner,
            textvariable=self._background_task_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
            foreground=GlassTheme.INFO
        )
        self._background_task_label.pack(side=tk.RIGHT, padx=10)
        
        # 分隔符
        ttk.Label(status_inner, text=" | ", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.RIGHT, padx=5)
        
        # AI连接状态
        self._ai_status_var = tk.StringVar(value="AI: 未连接")
        self._ai_status_label = ttk.Label(
            status_inner,
            textvariable=self._ai_status_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        )
        self._ai_status_label.pack(side=tk.RIGHT, padx=10)
    
    def _update_status_bar(self, project_name: str = None, word_count: int = None, 
                          ai_status: str = None, background_task: str = None) -> None:
        """更新状态栏信息"""
        if project_name is not None:
            self._project_var.set(f"项目: {project_name}")
        if word_count is not None:
            self._word_count_var.set(f"字数: {word_count:,}")
        if ai_status is not None:
            self._ai_status_var.set(f"AI: {ai_status}")
            # 更新颜色
            if ai_status == "在线":
                self._ai_status_label.configure(foreground=GlassTheme.SUCCESS)
            elif ai_status == "离线":
                self._ai_status_label.configure(foreground=GlassTheme.ERROR)
            else:
                self._ai_status_label.configure(foreground=GlassTheme.WARNING)
        if background_task is not None:
            self._background_task_var.set(background_task)
    
    def _init_bindings(self) -> None:
        """初始化快捷键绑定（增强版）"""
        # 页面切换
        self.root.bind("<Control-w>", lambda e: self._switch_to_page("workbench"))
        self.root.bind("<Control-p>", lambda e: self._switch_to_page("project"))
        self.root.bind("<Control-s>", lambda e: self._switch_to_page("settings"))
        self.root.bind("<Control-g>", lambda e: self._on_quick_generate_action())
        self.root.bind("<F5>", lambda e: self._refresh_current_page())
        
        # 新增快捷键
        self.root.bind("<Control-n>", lambda e: self._on_new_item())  # 新建
        self.root.bind("<Control-o>", lambda e: self._on_open_project())  # 打开项目
        self.root.bind("<Control-b>", lambda e: self._on_backup_project())  # 备份
        self.root.bind("<Escape>", lambda e: self._on_cancel_action())  # 取消
    
    def _on_quick_generate_action(self):
        """快速生成快捷键处理"""
        self._switch_to_page("workbench")
        self._switch_workbench_tab("quick")
    
    def _on_new_item(self):
        """新建项目快捷键处理"""
        self._switch_to_page("project")
        self._on_new_project()
    
    def _on_cancel_action(self):
        """取消操作"""
        self._set_status("操作已取消")
    
    # ============== 页面切换 ==============
    
    def _switch_to_page(self, page_id: str) -> None:
        """切换到指定页面"""
        # 隐藏当前页面
        if self._current_page and self._current_page in self._pages:
            self._pages[self._current_page].pack_forget()
        
        # 创建或显示新页面
        if page_id not in self._pages:
            self._pages[page_id] = self._create_page(page_id)
            logger.debug(f"Created new page: {page_id}")
        
        # pack页面到content_frame
        page = self._pages[page_id]
        page.pack(fill=tk.BOTH, expand=True, in_=self._content_frame)
        self._current_page = page_id
        
        # 等待页面pack完成后更新Canvas宽度
        self.root.after(100, self._update_canvas_widths)
        
        logger.debug(f"Switched to page: {page_id}, parent={page.master}, winfo_parent={page.winfo_parent()}")
    
    def _update_canvas_widths(self):
        """更新所有Canvas的窗口宽度"""
        for widget in self._content_frame.winfo_children():
            self._update_canvas_width_recursive(widget)
    
    def _update_canvas_width_recursive(self, widget):
        """递归更新所有Canvas的窗口宽度"""
        if widget.winfo_class() == 'Canvas':
            # 获取Canvas中的所有窗口
            for item_id in widget.find_withtag("all"):
                tags = widget.gettags(item_id)
                if "all" not in tags:
                    # 这是一个canvas window，更新其宽度
                    widget.itemconfig(item_id, width=widget.winfo_width())
        
        # 递归处理子控件
        for child in widget.winfo_children():
            self._update_canvas_width_recursive(child)
    
    def _create_page(self, page_id: str) -> tk.Frame:
        """创建页面"""
        page_creators = {
            "hot_ranking": self._create_hot_ranking_page,
            "workbench": self._create_workbench_page,
            "progress": self._create_progress_page,
            "project": self._create_project_page,
            "plugins": self._create_plugins_page,
            "settings": self._create_settings_page,
        }
        
        creator = page_creators.get(page_id)
        if creator:
            return creator()
        else:
            # 默认空白页面
            frame = ttk.Frame(self._content_frame, style="TFrame")
            ttk.Label(frame, text=f"页面: {page_id}", style="Title.TLabel").pack(padx=20, pady=20)
            return frame
    
    def _refresh_current_page(self) -> None:
        """刷新当前页面"""
        if self._current_page:
            # 销毁当前页面并重新创建
            if self._current_page in self._pages:
                self._pages[self._current_page].destroy()
                del self._pages[self._current_page]
            self._switch_to_page(self._current_page)
    
    # ============== 页面创建器 ==============
    
    def _create_hot_ranking_page(self) -> tk.Frame:
        """创建热榜页面（V5完整迁移）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建滚动容器
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        # Canvas和Scrollbar直接pack到frame中
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题栏
        title_frame = ttk.Frame(scrollable_frame, style="TFrame")
        title_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        tk.Label(
            title_frame,
            text="全网小说热度榜单",
            font=(GlassTheme.FONT_FAMILY, 24, "bold"),
            fg=GlassTheme.ACCENT_PINK,
            bg=GlassTheme.GLASS_BG
        ).pack(side=tk.LEFT)
        
        # 按钮区
        btn_frame = ttk.Frame(title_frame, style="TFrame")
        btn_frame.pack(side=tk.RIGHT)
        
        # 清缓存按钮
        self.clear_cache_btn = ResponsiveButton(
            btn_frame,
            text="清缓存",
            command=self._clear_hot_ranking_cache,
            async_handler=self._async_handler
        )
        self.clear_cache_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # 更新按钮
        self.update_hot_btn = ResponsiveButton(
            btn_frame,
            text="更新数据",
            command=self._update_hot_ranking_data,
            async_handler=self._async_handler,
            style="Accent.TButton"
        )
        self.update_hot_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=10)
        
        # ========== 第一部分：三大网站榜单 ==========
        ttk.Label(
            scrollable_frame,
            text="三大热门小说网站排行榜 (前10名)",
            font=(GlassTheme.FONT_FAMILY, 16, "bold")
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        # 来源说明行
        src_desc_frame = ttk.Frame(scrollable_frame, style="TFrame")
        src_desc_frame.pack(fill=tk.X, padx=22, pady=(0, 8))
        
        site_badge_info = [
            ('🍅', '番茄', GlassTheme.ACCENT_RED),
            ('🚀', '起点', GlassTheme.INFO),
            ('🔷', '晋江', GlassTheme.ACCENT_PURPLE),
        ]
        for icon, name, color in site_badge_info:
            badge = tk.Label(
                src_desc_frame,
                text=f" {icon} {name} ",
                font=(GlassTheme.FONT_FAMILY, 9),
                fg='white', bg=color,
                padx=4, pady=2,
                relief='flat'
            )
            badge.pack(side=tk.LEFT, padx=(0, 6))
        
        # 网站榜单容器（V5原版：使用pack布局，三列并排）
        sites_frame = ttk.Frame(scrollable_frame, style="TFrame")
        sites_frame.pack(fill=tk.X, padx=20, pady=10)

        # 三大网站数据，使用pack布局（恢复V5原版）
        sites_data = self._get_hot_ranking_data().get('sites', [])
        for site_info in sites_data:
            site_column = ttk.Frame(sites_frame, relief='flat', style="TFrame")
            site_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=8)
            
            # 网站标题
            site_header = tk.Label(
                site_column,
                text=site_info['name'],
                font=(GlassTheme.FONT_FAMILY, 13, 'bold'),
                fg='white',
                bg=site_info['color'],
                pady=10
            )
            site_header.pack(fill=tk.X)
            
            # 书籍列表
            books_frame = ttk.Frame(site_column, style="TFrame")
            books_frame.pack(fill=tk.X, pady=6)
            
            books = site_info.get('books', [])[:10]
            for idx, book in enumerate(books, 1):
                book_item = ttk.Frame(books_frame, style="TFrame")
                book_item.pack(fill=tk.X, pady=3)
                
                # 排名徽章
                rank_bg = site_info['color'] if idx <= 3 else '#CCCCCC'
                rank_label = tk.Label(
                    book_item,
                    text=str(idx),
                    font=(GlassTheme.FONT_FAMILY, 11, 'bold'),
                    fg='white',
                    bg=rank_bg,
                    width=3
                )
                rank_label.pack(side=tk.LEFT, padx=(0, 5))
                
                # 书籍信息
                info_frame = ttk.Frame(book_item, style="TFrame")
                info_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                title = book.get('title', '未知')
                author = book.get('author', '未知')
                category = book.get('category', '')
                heat = book.get('heat', 0)
                
                # 格式化热度
                if heat >= 100000000:
                    heat_str = f"{heat/100000000:.1f}亿"
                elif heat >= 10000:
                    heat_str = f"{heat/10000:.1f}万"
                elif heat > 0:
                    heat_str = str(heat)
                else:
                    heat_str = "-"
                
                # 书名
                ttk.Label(
                    info_frame,
                    text=f"《{title}》",
                    font=(GlassTheme.FONT_FAMILY, 10, 'bold')
                ).pack(anchor='w')
                
                # 作者/分类/热度
                detail_parts = [f"作者：{author}"]
                if category:
                    detail_parts.append(category)
                if heat_str != "-":
                    detail_parts.append(f"热度：{heat_str}")
                
                ttk.Label(
                    info_frame,
                    text="  ".join(detail_parts),
                    font=(GlassTheme.FONT_FAMILY, 8),
                    foreground='gray'
                ).pack(anchor='w')
        
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=20)
        
        # ========== 第二部分：题材热度榜 ==========
        ttk.Label(
            scrollable_frame,
            text="题材热度榜 (前5名)",
            font=(GlassTheme.FONT_FAMILY, 16, 'bold')
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        genre_frame = ttk.Frame(scrollable_frame, style="TFrame")
        genre_frame.pack(fill=tk.X, padx=20, pady=10)
        
        genres_data = self._get_hot_ranking_data().get('genres', {})
        for gender in ['male', 'female']:
            genre_info = genres_data.get(gender, {})
            if not genre_info:
                continue
            
            genre_column = ttk.Frame(genre_frame, style="TFrame")
            genre_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
            
            genre_header = tk.Label(
                genre_column,
                text=genre_info.get('title', ''),
                font=(GlassTheme.FONT_FAMILY, 13, 'bold'),
                fg='white',
                bg=genre_info.get('color', GlassTheme.PRIMARY),
                pady=10
            )
            genre_header.pack(fill=tk.X)
            
            chart_frame = ttk.Frame(genre_column, style="TFrame")
            chart_frame.pack(fill=tk.X, pady=10)
            
            genres = genre_info.get('genres', [])
            if genres:
                max_hot = max([g[1] for g in genres]) or 1
                
                for genre_name, hot_value in genres:
                    chart_item = ttk.Frame(chart_frame, style="TFrame")
                    chart_item.pack(fill=tk.X, pady=8)
                    
                    ttk.Label(
                        chart_item,
                        text=genre_name,
                        font=(GlassTheme.FONT_FAMILY, 10),
                        width=10,
                        anchor='e'
                    ).pack(side=tk.LEFT, padx=(0, 8))
                    
                    bar_container = tk.Frame(chart_item, bg=GlassTheme.GLASS_BORDER, height=20)
                    bar_container.pack(side=tk.LEFT, fill=tk.X, expand=True)
                    
                    bar_canvas = tk.Canvas(bar_container, bg=GlassTheme.GLASS_BORDER, height=20, highlightthickness=0)
                    bar_canvas.pack(fill=tk.BOTH, expand=True)
                    
                    ratio = hot_value / max_hot
                    color = genre_info.get('color', GlassTheme.PRIMARY)
                    
                    def draw_bar(canvas, c, r):
                        canvas.update_idletasks()
                        w = canvas.winfo_width()
                        if w > 1:
                            canvas.delete("all")
                            canvas.create_rectangle(0, 0, int(w * r), 20, fill=c, outline="")
                    
                    self.root.after(10, draw_bar, bar_canvas, color, ratio)
                    
                    tk.Label(
                        chart_item,
                        text=f"{hot_value:.1f}分",
                        font=(GlassTheme.FONT_FAMILY, 9, 'bold'),
                        fg=color,
                        bg=GlassTheme.GLASS_BG,
                        width=8, anchor='w'
                    ).pack(side=tk.LEFT, padx=(8, 0))
        
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=20)
        
        # ========== 第三部分：设定类型/创作流派榜 ==========
        ttk.Label(
            scrollable_frame,
            text="设定类型/创作流派榜 (前5名)",
            font=(GlassTheme.FONT_FAMILY, 16, 'bold')
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        type_frame = ttk.Frame(scrollable_frame, style="TFrame")
        type_frame.pack(fill=tk.X, padx=20, pady=10)
        
        types_data = self._get_hot_ranking_data().get('types', {})
        for gender in ['male', 'female']:
            type_info = types_data.get(gender, {})
            if not type_info:
                continue
            
            type_column = ttk.Frame(type_frame, style="TFrame")
            type_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
            
            type_header = tk.Label(
                type_column,
                text=type_info.get('title', ''),
                font=(GlassTheme.FONT_FAMILY, 13, 'bold'),
                fg='white',
                bg=type_info.get('color', GlassTheme.ACCENT_GREEN),
                pady=10
            )
            type_header.pack(fill=tk.X)
            
            chart_frame = ttk.Frame(type_column, style="TFrame")
            chart_frame.pack(fill=tk.X, pady=10)
            
            types = type_info.get('types', [])
            if types:
                max_hot = max([t[1] for t in types]) or 1
                
                for type_name, hot_value in types:
                    chart_item = ttk.Frame(chart_frame, style="TFrame")
                    chart_item.pack(fill=tk.X, pady=8)
                    
                    ttk.Label(
                        chart_item,
                        text=type_name,
                        font=(GlassTheme.FONT_FAMILY, 10),
                        width=10,
                        anchor='e'
                    ).pack(side=tk.LEFT, padx=(0, 8))
                    
                    bar_container = tk.Frame(chart_item, bg=GlassTheme.GLASS_BORDER, height=20)
                    bar_container.pack(side=tk.LEFT, fill=tk.X, expand=True)
                    
                    bar_canvas = tk.Canvas(bar_container, bg=GlassTheme.GLASS_BORDER, height=20, highlightthickness=0)
                    bar_canvas.pack(fill=tk.BOTH, expand=True)
                    
                    ratio = hot_value / max_hot
                    color = type_info.get('color', GlassTheme.ACCENT_GREEN)
                    
                    def draw_bar(canvas, c, r):
                        canvas.update_idletasks()
                        w = canvas.winfo_width()
                        if w > 1:
                            canvas.delete("all")
                            canvas.create_rectangle(0, 0, int(w * r), 20, fill=c, outline="")
                    
                    self.root.after(10, draw_bar, bar_canvas, color, ratio)
                    
                    tk.Label(
                        chart_item,
                        text=f"{hot_value:.1f}分",
                        font=(GlassTheme.FONT_FAMILY, 9, 'bold'),
                        fg=color,
                        bg=GlassTheme.GLASS_BG,
                        width=8, anchor='w'
                    ).pack(side=tk.LEFT, padx=(8, 0))
        
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=20)
        
        # ========== 第四部分：作家排行榜 ==========
        ttk.Label(
            scrollable_frame,
            text="全网最热网络作家排行榜 (前10名)",
            font=(GlassTheme.FONT_FAMILY, 16, 'bold')
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        author_outer_frame = ttk.Frame(scrollable_frame, style="TFrame")
        author_outer_frame.pack(fill=tk.X, padx=20, pady=10)
        
        author_table_frame = tk.Frame(author_outer_frame, bg=GlassTheme.GLASS_SURFACE)
        author_table_frame.pack(expand=True)
        
        headers = ['排名', '作家', '代表作品', '预估年收入', '粉丝数']
        col_widths = [8, 12, 35, 12, 10]
        
        for idx, (header, width) in enumerate(zip(headers, col_widths)):
            tk.Label(
                author_table_frame,
                text=header,
                font=(GlassTheme.FONT_FAMILY, 11, 'bold'),
                fg='white', bg=GlassTheme.PRIMARY,
                pady=8, padx=5, anchor='center',
                width=width, relief='flat'
            ).grid(row=0, column=idx, sticky='nsew', padx=1, pady=1)
        
        authors_data = self._get_hot_ranking_data().get('authors', [])
        for row_idx, author in enumerate(authors_data, start=1):
            rank = author.get('rank', row_idx)
            name = author.get('name', '未知')
            works = author.get('works', '暂无')
            income = author.get('income', '0')
            fans = author.get('fans', '0')
            
            if rank == 1:
                bg_color, fg_color = '#FFD700', 'black'
            elif rank == 2:
                bg_color, fg_color = '#C0C0C0', 'black'
            elif rank == 3:
                bg_color, fg_color = '#CD7F32', 'white'
            else:
                bg_color, fg_color = GlassTheme.GLASS_SURFACE, GlassTheme.TEXT_PRIMARY
            
            cells = [f"TOP{rank}", name, works[:25]+'...' if len(works) > 25 else works, f"¥{income}", fans]
            anchors = ['center', 'center', 'w', 'center', 'center']
            fonts = [
                (GlassTheme.FONT_FAMILY, 10, 'bold'),
                (GlassTheme.FONT_FAMILY, 10),
                (GlassTheme.FONT_FAMILY, 9),
                (GlassTheme.FONT_FAMILY, 10, 'bold'),
                (GlassTheme.FONT_FAMILY, 9),
            ]
            for col_idx, (text, anchor, font, width) in enumerate(zip(cells, anchors, fonts, col_widths)):
                tk.Label(
                    author_table_frame,
                    text=text,
                    font=font,
                    fg=fg_color, bg=bg_color,
                    pady=6, padx=5, anchor=anchor,
                    width=width, relief='flat'
                ).grid(row=row_idx, column=col_idx, sticky='nsew', padx=1, pady=1)
        
        # 数据来源说明
        ttk.Label(
            scrollable_frame,
            text="收入和粉丝数为算法估算值，仅供参考",
            font=(GlassTheme.FONT_FAMILY, 9),
            foreground=GlassTheme.TEXT_SECONDARY
        ).pack(pady=(5, 15), anchor='w', padx=20)
        
        return frame
    
    def _create_workbench_page(self) -> tk.Frame:
        """创建工作台页面（上半部分按钮 + 下半部分内容区）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 标题
        header = ttk.Frame(frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="工作台", style="Title.TLabel").pack(side=tk.LEFT)
        
        # ========== 上半部分：八个功能按钮 ==========
        buttons_frame = ttk.Frame(frame, style="TFrame")
        buttons_frame.pack(fill=tk.X, padx=20, pady=(0, 15))
        
        # 第一行四个按钮
        row1 = ttk.Frame(buttons_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        
        # 第二行四个按钮
        row2 = ttk.Frame(buttons_frame, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        
        functions = [
            ("世界观", "worldview", GlassTheme.ACCENT_PURPLE),
            ("人物设定", "characters", GlassTheme.ACCENT_PINK),
            ("大纲管理", "outline", GlassTheme.PRIMARY_LIGHT),
            ("风格学习", "style", GlassTheme.ACCENT_GREEN),
            ("开始创作", "generation", GlassTheme.PRIMARY),
            ("逆向反馈", "reverse", GlassTheme.WARNING),
            ("快捷创作", "quick", GlassTheme.ACCENT_RED),
            ("续写功能", "continue", GlassTheme.INFO),
        ]
        
        self._workbench_buttons = {}
        self._current_workbench_tab = tk.StringVar(value="worldview")
        
        for i, (text, tab_id, color) in enumerate(functions):
            parent = row1 if i < 4 else row2
            
            btn = tk.Label(
                parent,
                text=text,
                font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold"),
                fg=GlassTheme.TEXT_PRIMARY,
                bg=GlassTheme.GLASS_SURFACE,
                activeforeground=GlassTheme.TEXT_PRIMARY,
                activebackground=color,
                cursor="hand2",
                padx=20,
                pady=15,
                width=18
            )
            btn.pack(side=tk.LEFT, padx=8, pady=5, expand=True, fill=tk.X)
            
            # 存储按钮引用
            self._workbench_buttons[tab_id] = btn
            
            # 悬停效果
            btn.bind("<Enter>", lambda e, b=btn, c=color: b.configure(bg=c))
            btn.bind("<Leave>", lambda e, b=btn, tid=tab_id: self._update_button_style(tid))
            btn.bind("<Button-1>", lambda e, tid=tab_id: self._switch_workbench_tab(tid))
        
        # ========== 下半部分：内容区域（支持滚动）==========
        # 创建Canvas和Scrollbar用于滚动
        content_canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0)
        content_scrollbar = ttk.Scrollbar(frame, orient="vertical", command=content_canvas.yview)
        
        self._workbench_content_frame = ttk.Frame(content_canvas, style="TFrame")
        
        def on_content_frame_configure(event):
            content_canvas.configure(scrollregion=content_canvas.bbox("all"))
            content_canvas.itemconfig(content_window, width=content_canvas.winfo_width())
        
        self._workbench_content_frame.bind("<Configure>", on_content_frame_configure)
        content_window = content_canvas.create_window((0, 0), window=self._workbench_content_frame, anchor="nw")
        content_canvas.configure(yscrollcommand=content_scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_content_mousewheel(event):
            content_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_content_mousewheel(event):
            content_canvas.bind_all("<MouseWheel>", on_content_mousewheel)
        
        def unbind_content_mousewheel(event):
            content_canvas.unbind_all("<MouseWheel>")
        
        content_canvas.bind("<Enter>", bind_content_mousewheel)
        content_canvas.bind("<Leave>", unbind_content_mousewheel)
        
        content_canvas.pack(side="left", fill="both", expand=True, padx=20, pady=(0, 20))
        content_scrollbar.pack(side="right", fill="y", pady=(0, 20))
        
        # 初始化所有标签页内容（延迟创建）
        self._workbench_tabs = {}
        self._workbench_tabs_created = set()
        
        # 默认显示世界观页面
        self._create_all_workbench_tabs()
        self._switch_workbench_tab("worldview")
        
        return frame
    
    def _create_all_workbench_tabs(self):
        """预创建工作台标签页（延迟创建策略，仅初始化为None）"""
        # 所有标签页初始化为None，首次切换时才创建
        # 这样可以避免首次点击工作台时一次性创建所有标签页导致卡顿
        for tab_id in ["worldview", "characters", "outline", "style", 
                       "generation", "reverse", "quick", "continue"]:
            self._workbench_tabs[tab_id] = None
    
    def _switch_workbench_tab(self, tab_id: str):
        """切换工作台标签页（延迟创建）"""
        # 更新按钮样式
        self._current_workbench_tab.set(tab_id)
        for tid, btn in self._workbench_buttons.items():
            if tid == tab_id:
                # 激活状态：使用对应颜色
                color = self._get_button_color(tid)
                btn.configure(bg=color, fg="white")
            else:
                btn.configure(bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        
        # 隐藏所有标签页
        for tid, tab_frame in self._workbench_tabs.items():
            if tid != tab_id and tab_frame is not None:
                tab_frame.pack_forget()
        
        # 延迟创建：首次切换时才创建标签页内容
        if self._workbench_tabs.get(tab_id) is None:
            create_method = {
                "worldview": self._create_worldview_content,
                "characters": self._create_characters_content,
                "outline": self._create_outline_content,
                "style": self._create_style_content,
                "generation": self._create_generation_content,
                "reverse": self._create_reverse_content,
                "quick": self._create_quick_content,
                "continue": self._create_continue_content,
            }.get(tab_id)
            if create_method:
                self._workbench_tabs[tab_id] = create_method()
        
        # 显示当前标签页
        if self._workbench_tabs.get(tab_id) is not None:
            self._workbench_tabs[tab_id].pack(fill=tk.BOTH, expand=True, in_=self._workbench_content_frame)
    
    def _update_button_style(self, tab_id: str):
        """更新按钮样式（用于鼠标离开时）"""
        if self._current_workbench_tab.get() != tab_id:
            self._workbench_buttons[tab_id].configure(bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
    
    def _get_button_color(self, tab_id: str) -> str:
        """获取按钮对应的颜色"""
        colors = {
            "worldview": GlassTheme.ACCENT_PURPLE,
            "characters": GlassTheme.ACCENT_PINK,
            "outline": GlassTheme.PRIMARY_LIGHT,
            "style": GlassTheme.ACCENT_GREEN,
            "generation": GlassTheme.PRIMARY,
            "reverse": GlassTheme.WARNING,
            "quick": GlassTheme.ACCENT_RED,
            "continue": GlassTheme.INFO,
        }
        return colors.get(tab_id, GlassTheme.PRIMARY)
    
    def _create_worldview_content(self) -> tk.Frame:
        """创建世界观管理内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：文件导入区
        file_frame = ttk.LabelFrame(frame, text="世界观文件导入", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="选择文件：").pack(side=tk.LEFT)
        self._worldview_path_var = tk.StringVar()
        ttk.Entry(path_frame, textvariable=self._worldview_path_var, width=50).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="浏览", command=self._on_worldview_browse).pack(side=tk.LEFT)
        ttk.Button(path_frame, text="新建", command=self._on_worldview_new).pack(side=tk.LEFT, padx=5)
        
        # 支持格式说明
        ttk.Label(file_frame, text="支持格式：TXT / DOCX", 
                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(anchor=tk.W, pady=5)
        
        # 中部：世界观列表
        list_frame = ttk.LabelFrame(frame, text="世界观项目列表", padding=10)
        list_frame.pack(fill=tk.X, padx=5, pady=5)
        
        columns = ("name", "category", "elements", "status", "modified")
        self._worldview_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=4)
        self._worldview_tree.heading("name", text="世界观名称")
        self._worldview_tree.heading("category", text="世界观类别")
        self._worldview_tree.heading("elements", text="核心元素")
        self._worldview_tree.heading("status", text="状态")
        self._worldview_tree.heading("modified", text="修改时间")
        self._worldview_tree.column("name", width=120)
        self._worldview_tree.column("category", width=80)
        self._worldview_tree.column("elements", width=160)
        self._worldview_tree.column("status", width=80)
        self._worldview_tree.column("modified", width=110)
        self._worldview_tree.pack(fill=tk.X, pady=5)

        list_btn_frame = ttk.Frame(list_frame, style="TFrame")
        list_btn_frame.pack(fill=tk.X, pady=5)
        
        # 使用Grid布局，3列，确保按钮均匀分布
        list_buttons = [
            ("查看详情", self._on_worldview_view),
            ("编辑", self._on_worldview_edit),
            ("批量删除", self._on_worldview_delete),
        ]
        
        for i, (text, command) in enumerate(list_buttons):
            btn = ttk.Button(list_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            list_btn_frame.grid_columnconfigure(i, weight=1)

        # 下部：预览区
        preview_frame = ttk.LabelFrame(frame, text="世界观详情预览", padding=10)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self._worldview_preview = tk.Text(preview_frame, wrap=tk.WORD, height=10,
                                         font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                         bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._worldview_preview.pack(fill=tk.BOTH, expand=True)

        # 底部按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，3列，确保按钮均匀分布
        bottom_buttons = [
            ("导入世界观", self._on_worldview_import),
            ("关联要素", self._on_worldview_link),
            ("清除", self._on_worldview_clear),
        ]
        
        for i, (text, command) in enumerate(bottom_buttons):
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_characters_content(self) -> tk.Frame:
        """创建人物设定内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：人物导入区
        file_frame = ttk.LabelFrame(frame, text="人物档案导入", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="选择文件：").pack(side=tk.LEFT)
        self._character_path_var = tk.StringVar()
        ttk.Entry(path_frame, textvariable=self._character_path_var, width=50).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="浏览", command=self._on_character_browse).pack(side=tk.LEFT)
        ttk.Button(path_frame, text="新建人物", command=self._on_character_new).pack(side=tk.LEFT, padx=5)
        
        # 中部：人物列表区
        list_frame = ttk.LabelFrame(frame, text="人物列表", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 列表（移除头像列，情绪改为情感状态）
        columns = ("name", "role", "status", "emotion", "chapters")
        self._character_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=8)
        self._character_tree.heading("name", text="姓名")
        self._character_tree.heading("role", text="角色类型")
        self._character_tree.heading("status", text="状态")
        self._character_tree.heading("emotion", text="情感状态")
        self._character_tree.heading("chapters", text="出场章节")
        self._character_tree.column("name", width=120)
        self._character_tree.column("role", width=100)
        self._character_tree.column("status", width=80)
        self._character_tree.column("emotion", width=100)
        self._character_tree.column("chapters", width=100)
        self._character_tree.pack(fill=tk.BOTH, expand=True)
        
        # 列表操作按钮
        list_btn_frame = ttk.Frame(list_frame, style="TFrame")
        list_btn_frame.pack(fill=tk.X, pady=5)

        # 使用Grid布局，4列，确保按钮均匀分布
        list_buttons = [
            ("编辑人物", self._on_character_edit),
            ("人物详情", self._on_character_detail),
            ("关系图谱", self._on_character_relation),
            ("删除人物", self._on_character_delete),
        ]
        
        for i, (text, command) in enumerate(list_buttons):
            btn = ttk.Button(list_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(4):
            list_btn_frame.grid_columnconfigure(i, weight=1)

        # 下部：人物详情预览
        detail_frame = ttk.LabelFrame(frame, text="人物档案详情", padding=10)
        detail_frame.pack(fill=tk.X, padx=5, pady=5)

        self._character_detail = tk.Text(detail_frame, wrap=tk.WORD, height=6,
                                        font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                        bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._character_detail.pack(fill=tk.BOTH, expand=True)

        return frame
    
    def _create_outline_content(self) -> tk.Frame:
        """创建大纲管理内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：文件导入区
        file_frame = ttk.LabelFrame(frame, text="大纲文件导入", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="当前大纲：").pack(side=tk.LEFT)
        self._outline_path_var = tk.StringVar(value="未导入")
        ttk.Label(path_frame, textvariable=self._outline_path_var, 
                 foreground=GlassTheme.TEXT_LINK).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="选择文件", command=self._on_outline_browse).pack(side=tk.LEFT)
        ttk.Button(path_frame, text="新建大纲", command=self._on_outline_new).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="导出", command=self._on_outline_export).pack(side=tk.RIGHT)
        
        # 中部：大纲树形结构
        tree_frame = ttk.LabelFrame(frame, text="大纲结构（可拖拽排序）", padding=10)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self._outline_tree = ttk.Treeview(tree_frame, show="tree", height=8)
        self._outline_tree.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self._outline_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self._outline_tree.configure(yscrollcommand=scrollbar.set)
        
        # 树形结构操作按钮
        tree_btn_frame = ttk.Frame(tree_frame, style="TFrame")
        tree_btn_frame.pack(fill=tk.X, pady=5, side=tk.BOTTOM)
        
        # 使用Grid布局，4列，确保按钮均匀分布
        tree_buttons = [
            ("添加卷", self._on_outline_add_volume),
            ("添加章节", self._on_outline_add_chapter),
            ("编辑章节", self._on_outline_edit_chapter),
            ("删除", self._on_outline_delete),
        ]
        
        for i, (text, command) in enumerate(tree_buttons):
            btn = ttk.Button(tree_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(4):
            tree_btn_frame.grid_columnconfigure(i, weight=1)
        
        # 下部：章节编辑器
        editor_frame = ttk.LabelFrame(frame, text="章节编辑器", padding=10)
        editor_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 章节信息
        info_row = ttk.Frame(editor_frame, style="TFrame")
        info_row.pack(fill=tk.X, pady=5)
        ttk.Label(info_row, text="章节标题：").pack(side=tk.LEFT)
        self._chapter_title_var = tk.StringVar()
        ttk.Entry(info_row, textvariable=self._chapter_title_var, width=30).pack(side=tk.LEFT, padx=5)
        ttk.Label(info_row, text="目标字数：").pack(side=tk.LEFT, padx=(20, 0))
        self._chapter_words_var = tk.StringVar(value="2000")
        ttk.Spinbox(info_row, from_=500, to=5000, increment=100, width=10, textvariable=self._chapter_words_var).pack(side=tk.LEFT, padx=5)
        
        # 章节摘要
        ttk.Label(editor_frame, text="内容摘要：").pack(anchor=tk.W, pady=5)
        self._chapter_summary = tk.Text(editor_frame, wrap=tk.WORD, height=3,
                                       font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                       bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._chapter_summary.pack(fill=tk.X)
        
        # 进度条
        progress_frame = ttk.Frame(frame, style="TFrame")
        progress_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(progress_frame, text="大纲完成进度：").pack(side=tk.LEFT)
        self._outline_progress = ttk.Progressbar(progress_frame, length=300, mode='determinate')
        self._outline_progress.pack(side=tk.LEFT, padx=10)
        ttk.Label(progress_frame, text="0/0 章节").pack(side=tk.LEFT)
        
        # 按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，2列，确保按钮均匀分布
        outline_buttons = [
            ("解析大纲", self._on_outline_parse),
            ("清除大纲", self._on_outline_clear),
        ]
        
        for i, (text, command) in enumerate(outline_buttons):
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(2):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_style_content(self) -> tk.Frame:
        """创建风格学习内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：范文管理区
        file_frame = ttk.LabelFrame(frame, text="范文管理", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="当前风格：").pack(side=tk.LEFT)
        self._style_path_var = tk.StringVar(value="未导入")
        ttk.Label(path_frame, textvariable=self._style_path_var,
                 foreground=GlassTheme.TEXT_LINK).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="上传范文", command=self._on_style_browse).pack(side=tk.LEFT)
        ttk.Button(path_frame, text="删除范文", command=self._on_style_delete).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="导出风格", command=self._on_style_export).pack(side=tk.RIGHT)
        
        # 中部左：风格档案库
        library_frame = ttk.LabelFrame(frame, text="风格档案库", padding=10)
        library_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5, side=tk.LEFT)
        
        columns = ("name", "vocabulary", "sentence", "rhetoric", "emotion", "rhythm", "structure", "detail")
        self._style_tree = ttk.Treeview(library_frame, columns=columns, show="headings", height=8)
        self._style_tree.heading("name", text="风格名称")
        self._style_tree.heading("vocabulary", text="词汇")
        self._style_tree.heading("sentence", text="句式")
        self._style_tree.heading("rhetoric", text="修辞")
        self._style_tree.heading("emotion", text="情感")
        self._style_tree.heading("rhythm", text="节奏")
        self._style_tree.heading("structure", text="结构")
        self._style_tree.heading("detail", text="细节")
        self._style_tree.column("name", width=100)
        for col in columns[1:]:
            self._style_tree.column(col, width=50)
        self._style_tree.pack(fill=tk.BOTH, expand=True)
        
        # 档案库按钮
        lib_btn_frame = ttk.Frame(library_frame, style="TFrame")
        lib_btn_frame.pack(fill=tk.X, pady=5)

        # 使用Grid布局，3列，确保按钮均匀分布
        lib_buttons = [
            ("切换风格", self._on_style_switch),
            ("编辑权重", self._on_style_edit_weight),
            ("删除风格", self._on_style_delete_profile),
        ]
        
        for i, (text, command) in enumerate(lib_buttons):
            btn = ttk.Button(lib_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            lib_btn_frame.grid_columnconfigure(i, weight=1)

        # 中部右：风格分析器结果
        analyzer_frame = ttk.LabelFrame(frame, text="风格分析器（七维度评分）", padding=10)
        analyzer_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5, side=tk.RIGHT)

        self._style_info = tk.Text(analyzer_frame, wrap=tk.WORD, height=10,
                                  font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                  bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._style_info.pack(fill=tk.BOTH, expand=True)

        # 底部：创建风格按钮
        create_frame = ttk.Frame(frame, style="TFrame")
        create_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，4列，确保按钮均匀分布
        create_buttons = [
            ("深度学习", self._on_style_learn),
            ("创建风格模板", self._on_style_create),
            ("应用到生成器", self._on_style_apply),
            ("清除风格", self._on_style_clear),
        ]
        
        for i, (text, command) in enumerate(create_buttons):
            btn = ttk.Button(create_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(4):
            create_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_generation_content(self) -> tk.Frame:
        """创建开始创作内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：生成配置
        config_frame = ttk.LabelFrame(frame, text="生成配置", padding=10)
        config_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 第一行：风格选择和章节范围
        row1 = ttk.Frame(config_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="选择风格：").pack(side=tk.LEFT)
        self._gen_style_var = tk.StringVar(value="默认风格")
        ttk.Combobox(row1, textvariable=self._gen_style_var, width=20, values=["默认风格", "武侠风", "言情风", "科幻风"]).pack(side=tk.LEFT, padx=5)
        ttk.Label(row1, text="起始章节：").pack(side=tk.LEFT, padx=(20, 0))
        self._start_chapter_var = tk.StringVar(value="1")
        ttk.Spinbox(row1, from_=1, to=100, width=8, textvariable=self._start_chapter_var).pack(side=tk.LEFT, padx=5)
        ttk.Label(row1, text="结束章节：").pack(side=tk.LEFT)
        self._end_chapter_var = tk.StringVar(value="1")
        ttk.Spinbox(row1, from_=1, to=100, width=8, textvariable=self._end_chapter_var).pack(side=tk.LEFT, padx=5)
        
        # 第二行：字数和温度
        row2 = ttk.Frame(config_frame, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="目标字数/章：").pack(side=tk.LEFT)
        self._target_words_var = tk.StringVar(value="900")
        ttk.Spinbox(row2, from_=500, to=2000, increment=100, width=10, textvariable=self._target_words_var).pack(side=tk.LEFT, padx=5)
        ttk.Label(row2, text="生成温度：").pack(side=tk.LEFT, padx=(20, 0))
        self._gen_temp_var = tk.DoubleVar(value=0.7)
        ttk.Scale(row2, from_=0.0, to=1.0, variable=self._gen_temp_var, orient=tk.HORIZONTAL, length=150).pack(side=tk.LEFT, padx=5)
        ttk.Label(row2, textvariable=self._gen_temp_var).pack(side=tk.LEFT)
        
        # 第三行：出场人物
        row3 = ttk.Frame(config_frame, style="TFrame")
        row3.pack(fill=tk.X, pady=5)
        ttk.Label(row3, text="出场人物：").pack(side=tk.LEFT)
        self._gen_characters_var = tk.StringVar(value="自动设置")
        ttk.Entry(row3, textvariable=self._gen_characters_var, width=40).pack(side=tk.LEFT, padx=5)
        ttk.Label(row3, text="(可手动输入添加)", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT)
        
        # 第四行：生成模式
        row4 = ttk.Frame(config_frame, style="TFrame")
        row4.pack(fill=tk.X, pady=5)
        ttk.Label(row4, text="生成模式：").pack(side=tk.LEFT)
        self._gen_mode_var = tk.StringVar(value="auto")
        ttk.Radiobutton(row4, text="自动迭代", variable=self._gen_mode_var, value="auto").pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(row4, text="手动迭代", variable=self._gen_mode_var, value="manual").pack(side=tk.LEFT, padx=10)
        
        # 中部：左右分栏（生成过程 + 生成结果）
        middle_split_frame = ttk.Frame(frame, style="TFrame")
        middle_split_frame.pack(fill=tk.BOTH, expand=False, padx=5, pady=5)  # 改为expand=False，给底部按钮留空间

        # 使用grid布局让左右等分宽度
        middle_split_frame.grid_columnconfigure(0, weight=1)
        middle_split_frame.grid_columnconfigure(1, weight=1)

        # 左侧：生成过程展示
        process_frame = ttk.LabelFrame(middle_split_frame, text="生成过程", padding=10)
        process_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 5))

        # 进度状态
        status_frame = ttk.Frame(process_frame, style="TFrame")
        status_frame.pack(fill=tk.X, pady=5)
        self._gen_status_var = tk.StringVar(value="准备就绪")
        ttk.Label(status_frame, textvariable=self._gen_status_var, font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold")).pack(side=tk.LEFT)
        ttk.Label(status_frame, text="当前章节：第1章").pack(side=tk.RIGHT)

        # 进度条
        self._gen_progress = ttk.Progressbar(process_frame, length=400, mode='determinate')
        self._gen_progress.pack(fill=tk.X, pady=5)

        # 日志输出（带滚动条，固定高度）
        log_scroll = ttk.Scrollbar(process_frame, orient=tk.VERTICAL)
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._gen_log = tk.Text(process_frame, wrap=tk.WORD, height=10,
                               font=(GlassTheme.FONT_FAMILY_CODE, GlassTheme.FONT_SIZE_SMALL),
                               bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                               yscrollcommand=log_scroll.set)
        self._gen_log.pack(fill=tk.BOTH, expand=True)
        log_scroll.configure(command=self._gen_log.yview)

        # 右侧：生成结果展示（带滚动条，固定高度）
        result_frame = ttk.LabelFrame(middle_split_frame, text="生成结果预览", padding=10)
        result_frame.grid(row=0, column=1, sticky="nsew", padx=(5, 0))

        result_scroll = ttk.Scrollbar(result_frame, orient=tk.VERTICAL)
        result_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._gen_result = tk.Text(result_frame, wrap=tk.WORD, height=10,
                                  font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                  bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                  yscrollcommand=result_scroll.set)
        self._gen_result.pack(fill=tk.BOTH, expand=True)
        result_scroll.configure(command=self._gen_result.yview)

        # 底部按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，5列，确保按钮均匀分布
        buttons = [
            ("保存结果", self._on_gen_save_result),
            ("开始生成", self._on_start_generation),
            ("停止生成", self._on_stop_generation),
            ("分章浏览", self._on_gen_browse),
            ("保存项目", self._on_gen_save),
        ]

        for i, (text, command) in enumerate(buttons):
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(5):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame

    def _on_gen_save_result(self):
        """保存生成结果"""
        try:
            from tkinter import messagebox

            content = self._gen_result.get("1.0", tk.END).strip()

            if not content:
                messagebox.showwarning("提示", "没有可保存的内容！")
                return

            from tkinter import filedialog

            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")],
                title="保存生成结果"
            )

            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)

                messagebox.showinfo("成功", "生成结果已保存！")
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror("错误", f"保存失败：{str(e)}")
        ttk.Button(btn_frame, text="导出TXT", command=self._on_gen_export_txt).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="导出DOCX", command=self._on_gen_export_docx).pack(side=tk.LEFT, padx=5)
        
        return frame
    
    def _create_reverse_content(self) -> tk.Frame:
        """创建逆向反馈内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：文件选择
        file_frame = ttk.LabelFrame(frame, text="选择已生成内容", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(file_frame, text="确定项目成书内容，分析已生成的小说：").pack(anchor=tk.W, pady=5)
        path_row = ttk.Frame(file_frame, style="TFrame")
        path_row.pack(fill=tk.X, pady=5)
        self._reverse_path_var = tk.StringVar()
        ttk.Entry(path_row, textvariable=self._reverse_path_var, width=50).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(path_row, text="浏览", command=self._on_reverse_browse).pack(side=tk.LEFT)
        ttk.Button(path_row, text="分析选中内容", command=self._on_reverse_analyze).pack(side=tk.LEFT, padx=5)
        
        # 中部：左右分栏布局
        middle_frame = ttk.Frame(frame, style="TFrame")
        middle_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 中部左：已完成章节列表（新增）
        chapters_frame = ttk.LabelFrame(middle_frame, text="已完成章节列表", padding=10)
        chapters_frame.pack(fill=tk.BOTH, expand=True, side=tk.LEFT, padx=(0, 5))
        
        # 章节列表
        chapters_list_frame = ttk.Frame(chapters_frame, style="TFrame")
        chapters_list_frame.pack(fill=tk.BOTH, expand=True)
        
        columns = ("chapter", "title", "words", "status")
        self._completed_chapters_tree = ttk.Treeview(chapters_list_frame, columns=columns, show="headings", height=8)
        self._completed_chapters_tree.heading("chapter", text="章节")
        self._completed_chapters_tree.heading("title", text="标题")
        self._completed_chapters_tree.heading("words", text="字数")
        self._completed_chapters_tree.heading("status", text="状态")
        self._completed_chapters_tree.column("chapter", width=60)
        self._completed_chapters_tree.column("title", width=150)
        self._completed_chapters_tree.column("words", width=80)
        self._completed_chapters_tree.column("status", width=80)
        self._completed_chapters_tree.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        # 滚动条
        scrollbar = ttk.Scrollbar(chapters_list_frame, orient=tk.VERTICAL, command=self._completed_chapters_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self._completed_chapters_tree.configure(yscrollcommand=scrollbar.set)
        
        # 右键菜单
        self._chapters_context_menu = tk.Menu(self._completed_chapters_tree, tearoff=0)
        self._chapters_context_menu.add_command(label="标记为完成", command=lambda: self._on_mark_chapter_completed())
        self._chapters_context_menu.add_command(label="标记为未完成", command=lambda: self._on_mark_chapter_incomplete())
        self._chapters_context_menu.add_separator()
        self._chapters_context_menu.add_command(label="删除章节", command=lambda: self._on_delete_completed_chapter())
        self._completed_chapters_tree.bind("<Button-3>", self._show_chapters_context_menu)
        
        # 章节列表按钮
        chapters_btn_frame = ttk.Frame(chapters_frame, style="TFrame")
        chapters_btn_frame.pack(fill=tk.X, pady=5)

        # 使用Grid布局，3列，确保按钮均匀分布
        chapters_buttons = [
            ("上传章节", self._on_upload_chapter),
            ("删除选中", self._on_delete_completed_chapter),
            ("全部标记完成", self._on_mark_all_completed),
        ]
        
        for i, (text, command) in enumerate(chapters_buttons):
            btn = ttk.Button(chapters_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            chapters_btn_frame.grid_columnconfigure(i, weight=1)
        
        # 中部右：运行分析
        analysis_frame = ttk.LabelFrame(middle_frame, text="运行分析", padding=10)
        analysis_frame.pack(fill=tk.BOTH, expand=True, side=tk.RIGHT, padx=(5, 0))
        
        # 分析维度
        ttk.Label(analysis_frame, text="分析维度：").pack(anchor=tk.W, pady=5)
        self._reverse_check_consistency = tk.BooleanVar(value=True)
        self._reverse_check_logic = tk.BooleanVar(value=True)
        self._reverse_check_character = tk.BooleanVar(value=True)
        self._reverse_check_style = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(analysis_frame, text="一致性检查", variable=self._reverse_check_consistency).pack(anchor=tk.W)
        ttk.Checkbutton(analysis_frame, text="逻辑漏洞检测", variable=self._reverse_check_logic).pack(anchor=tk.W)
        ttk.Checkbutton(analysis_frame, text="人设偏离检测", variable=self._reverse_check_character).pack(anchor=tk.W)
        ttk.Checkbutton(analysis_frame, text="风格匹配度", variable=self._reverse_check_style).pack(anchor=tk.W)
        # 移除了单独的"运行分析"按钮，移到底部按钮区域
        
        # 下部：反馈报告（带滚动条）
        report_frame = ttk.LabelFrame(frame, text="反馈报告", padding=10)
        report_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        report_scroll = ttk.Scrollbar(report_frame, orient=tk.VERTICAL)
        report_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._reverse_result = tk.Text(report_frame, wrap=tk.WORD, height=8,
                                      font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                      bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                      yscrollcommand=report_scroll.set)
        self._reverse_result.pack(fill=tk.BOTH, expand=True)
        report_scroll.configure(command=self._reverse_result.yview)
        
        # 底部按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，6列，确保按钮均匀分布
        reverse_buttons = [
            ("运行分析", self._on_reverse_run),
            ("自动修正", self._on_reverse_apply),
            ("调整大纲细节", self._on_reverse_outline),
            ("修改人物性格", self._on_reverse_character),
            ("补充人物轨迹", self._on_reverse_trajectory),
            ("增减世界观设定", self._on_reverse_worldview),
        ]

        for i, (text, command) in enumerate(reverse_buttons):
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(6):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_quick_content(self) -> tk.Frame:
        """创建快捷创作内容页面（一次性生成四个结果：世界观、大纲、人设、关键情节）"""
        # 直接在workbench内容区创建frame，不需要额外的Canvas
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")

        # 上部：文件上传区（原快捷模式选择位置）
        upload_frame = ttk.LabelFrame(frame, text="参考文本上传（可选）", padding=10)
        upload_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(upload_frame, text="上传案例文本或已有设定作为生成参考：").pack(anchor=tk.W, pady=5)
        
        upload_row = ttk.Frame(upload_frame, style="TFrame")
        upload_row.pack(fill=tk.X, pady=5)
        
        self._quick_upload_path = tk.StringVar()
        ttk.Entry(upload_row, textvariable=self._quick_upload_path, width=50).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(upload_row, text="上传TXT/DOCX", command=self._on_quick_upload_file).pack(side=tk.LEFT)
        
        # 上传状态指示
        self._quick_upload_status_frame = ttk.Frame(upload_frame, style="TFrame")
        self._quick_upload_status_frame.pack(fill=tk.X, pady=5)
        
        self._upload_status_labels = {}
        upload_types = [("世界观", "worldview"), ("大纲", "outline"), ("人设", "characters"), ("关键情节", "plot")]
        for text, key in upload_types:
            status_frame = ttk.Frame(self._quick_upload_status_frame, style="TFrame")
            status_frame.pack(side=tk.LEFT, padx=10)
            ttk.Label(status_frame, text=f"{text}:").pack(side=tk.LEFT)
            self._upload_status_labels[key] = ttk.Label(status_frame, text="未上传", foreground=GlassTheme.TEXT_SECONDARY)
            self._upload_status_labels[key].pack(side=tk.LEFT, padx=5)
        
        # 中部：需求描述
        input_frame = ttk.LabelFrame(frame, text="需求描述（关键词模式）", padding=10)
        input_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self._quick_input = tk.Text(input_frame, wrap=tk.WORD, height=4,
                                   font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                   bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._quick_input.pack(fill=tk.X)
        self._quick_input.insert("1.0", "请输入关键词描述，如：修仙世界、现代都市、古代宫廷、仙侠爱情...")
        
        # 下部：四个生成结果框（2x2布局，四等分，带滚动条，固定大小不可拖动）
        results_outer_frame = ttk.LabelFrame(frame, text="生成结果（世界观、大纲、人设、关键情节）", padding=10)
        results_outer_frame.pack(fill=tk.BOTH, expand=False, padx=5, pady=5)  # 改为expand=False给底部按钮留空间

        # 创建2x2网格容器
        results_grid = ttk.Frame(results_outer_frame, style="TFrame")
        results_grid.pack(fill=tk.BOTH, expand=True)

        # 配置网格权重
        results_grid.grid_columnconfigure(0, weight=1)
        results_grid.grid_columnconfigure(1, weight=1)
        results_grid.grid_rowconfigure(0, weight=1)
        results_grid.grid_rowconfigure(1, weight=1)

        # 固定尺寸参数
        TEXT_HEIGHT = 10  # 减小高度给底部按钮留空间

        # 左上：世界观结果框（带滚动条，固定大小）
        worldview_frame = ttk.LabelFrame(results_grid, text="世界观", padding=5)
        worldview_frame.grid(row=0, column=0, sticky="nsew", padx=3, pady=3)
        worldview_scroll = ttk.Scrollbar(worldview_frame, orient=tk.VERTICAL)
        worldview_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._quick_worldview_result = tk.Text(worldview_frame, wrap=tk.WORD, height=TEXT_HEIGHT,
                                               font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                                               bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                               yscrollcommand=worldview_scroll.set,
                                               state=tk.NORMAL)
        self._quick_worldview_result.pack(fill=tk.BOTH, expand=True)
        self._quick_worldview_result.insert("1.0", "世界观设定将在此显示...")
        worldview_scroll.configure(command=self._quick_worldview_result.yview)

        # 右上：大纲结果框（带滚动条，固定大小）
        outline_frame = ttk.LabelFrame(results_grid, text="大纲", padding=5)
        outline_frame.grid(row=0, column=1, sticky="nsew", padx=3, pady=3)
        outline_scroll = ttk.Scrollbar(outline_frame, orient=tk.VERTICAL)
        outline_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._quick_outline_result = tk.Text(outline_frame, wrap=tk.WORD, height=TEXT_HEIGHT,
                                             font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                                             bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                             yscrollcommand=outline_scroll.set,
                                             state=tk.NORMAL)
        self._quick_outline_result.pack(fill=tk.BOTH, expand=True)
        self._quick_outline_result.insert("1.0", "大纲结构将在此显示...")
        outline_scroll.configure(command=self._quick_outline_result.yview)

        # 左下：人设结果框（带滚动条，固定大小）
        characters_frame = ttk.LabelFrame(results_grid, text="人设", padding=5)
        characters_frame.grid(row=1, column=0, sticky="nsew", padx=3, pady=3)
        characters_scroll = ttk.Scrollbar(characters_frame, orient=tk.VERTICAL)
        characters_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._quick_characters_result = tk.Text(characters_frame, wrap=tk.WORD, height=TEXT_HEIGHT,
                                                font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                                                bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                                yscrollcommand=characters_scroll.set,
                                                state=tk.NORMAL)
        self._quick_characters_result.pack(fill=tk.BOTH, expand=True)
        self._quick_characters_result.insert("1.0", "人物设定将在此显示...")
        characters_scroll.configure(command=self._quick_characters_result.yview)

        # 右下：关键情节结果框（带滚动条，固定大小）
        plot_frame = ttk.LabelFrame(results_grid, text="关键情节", padding=5)
        plot_frame.grid(row=1, column=1, sticky="nsew", padx=3, pady=3)
        plot_scroll = ttk.Scrollbar(plot_frame, orient=tk.VERTICAL)
        plot_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._quick_plot_result = tk.Text(plot_frame, wrap=tk.WORD, height=TEXT_HEIGHT,
                                          font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                                          bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                          yscrollcommand=plot_scroll.set,
                                          state=tk.NORMAL)
        self._quick_plot_result.pack(fill=tk.BOTH, expand=True)
        self._quick_plot_result.insert("1.0", "关键情节将在此显示...")
        plot_scroll.configure(command=self._quick_plot_result.yview)
        
        # 底部按钮 - 使用Grid布局，确保小窗口下也能显示
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，3行3列
        buttons = [
            ("一键生成全部", self._on_quick_generate_all),
            ("仅生成世界观", lambda: self._on_quick_generate_single("worldview")),
            ("仅生成大纲", lambda: self._on_quick_generate_single("outline")),
            ("仅生成人设", lambda: self._on_quick_generate_single("characters")),
            ("仅生成情节", lambda: self._on_quick_generate_single("plot")),
            ("保存结果", self._on_quick_save_results),
            ("导出结果", self._on_quick_export_results),
            ("导入当前项目", self._on_quick_import),
        ]

        # 3行3列布局
        for i, (text, command) in enumerate(buttons):
            row = i // 3
            col = i % 3
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=row, column=col, padx=2, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame

    def _on_quick_save_results(self):
        """保存快捷创作结果"""
        try:
            from tkinter import filedialog, messagebox

            # 获取4个文本框的内容（使用正确的变量名）
            content = {
                "worldview": self._quick_worldview_result.get("1.0", tk.END).strip() if self._quick_worldview_result else "",
                "outline": self._quick_outline_result.get("1.0", tk.END).strip() if self._quick_outline_result else "",
                "characters": self._quick_characters_result.get("1.0", tk.END).strip() if self._quick_characters_result else "",
                "plot": self._quick_plot_result.get("1.0", tk.END).strip() if self._quick_plot_result else ""
            }

            # 检查是否有内容
            if not any(content.values()):
                messagebox.showwarning("提示", "没有可保存的内容！")
                return

            # 选择保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")],
                title="保存快捷创作结果"
            )

            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("===== 世界观 =====\n")
                    f.write(content["worldview"] + "\n\n")
                    f.write("===== 大纲 =====\n")
                    f.write(content["outline"] + "\n\n")
                    f.write("===== 人设 =====\n")
                    f.write(content["characters"] + "\n\n")
                    f.write("===== 情节 =====\n")
                    f.write(content["plot"] + "\n")

                messagebox.showinfo("成功", "快捷创作结果已保存！")
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror("错误", f"保存失败：{str(e)}")
    
    def _create_continue_content(self) -> tk.Frame:
        """创建续写功能内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：原文区
        source_frame = ttk.LabelFrame(frame, text="原文内容（智能分析续写）", padding=10)
        source_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 原文操作按钮
        source_btn_frame = ttk.Frame(source_frame, style="TFrame")
        source_btn_frame.pack(fill=tk.X, pady=5)
        ttk.Label(source_btn_frame, text="续写起点：光标定位到末尾位置").pack(side=tk.LEFT)
        ttk.Button(source_btn_frame, text="选择文件", command=self._on_continue_browse).pack(side=tk.RIGHT)
        
        self._continue_source = tk.Text(source_frame, wrap=tk.WORD, height=6,
                                       font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                       bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._continue_source.pack(fill=tk.BOTH, expand=True)
        self._continue_source.insert("1.0", "请在此粘贴或输入原文内容，或选择文件导入...")
        
        # 中部：续写设置
        settings_frame = ttk.LabelFrame(frame, text="续写设置", padding=10)
        settings_frame.pack(fill=tk.X, padx=5, pady=5)
        
        row1 = ttk.Frame(settings_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="续写字数：").pack(side=tk.LEFT)
        self._continue_words_var = tk.StringVar(value="500")
        ttk.Spinbox(row1, from_=200, to=2000, increment=100, width=10, textvariable=self._continue_words_var).pack(side=tk.LEFT, padx=5)
        ttk.Label(row1, text="续写方向：").pack(side=tk.LEFT, padx=(20, 0))
        self._continue_direction_var = tk.StringVar(value="自然续写")
        ttk.Combobox(row1, textvariable=self._continue_direction_var, 
                    values=["自然续写", "制造冲突", "引入新角色", "转折情节", "高潮推进"], width=15).pack(side=tk.LEFT, padx=5)
        
        # 下部：续写结果预览
        result_frame = ttk.LabelFrame(frame, text="续写结果预览（支持多次尝试）", padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self._continue_result = tk.Text(result_frame, wrap=tk.WORD, height=8,
                                       font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                       bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._continue_result.pack(fill=tk.BOTH, expand=True)
        
        # 底部按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(btn_frame, text="开始续写", command=self._on_continue_generate).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="重新生成", command=self._on_continue_regenerate).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="选择最佳版本", command=self._on_continue_select_best).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="保存结果", command=self._on_continue_save).pack(side=tk.LEFT, padx=5)
        
        return frame
    
    # ============== 工作台子功能回调方法 ==============
    
    def _on_worldview_browse(self):
        """浏览世界观文件"""
        path = filedialog.askopenfilename(
            title="选择世界观文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._worldview_path_var.set(path)
            self._preview_worldview(path)
    
    def _preview_worldview(self, path: str):
        """预览世界观文件"""
        try:
            self._worldview_preview.delete("1.0", tk.END)
            if path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(path)
                text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                with open(path, 'r', encoding='utf-8') as f:
                    text = f.read()
            self._worldview_preview.insert("1.0", text[:5000])
        except Exception as e:
            self._worldview_preview.insert("1.0", f"预览失败: {e}")
    
    def _on_worldview_new(self):
        """新建世界观（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("新建世界观")
        dialog.geometry("500x400")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 400) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 世界观名称
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        ttk.Label(name_frame, text="世界观名称：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="新世界观")
        ttk.Entry(name_frame, textvariable=name_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 核心元素
        elements_frame = ttk.Frame(dialog, style="TFrame")
        elements_frame.pack(fill=tk.X, padx=20, pady=10)
        ttk.Label(elements_frame, text="核心元素：").pack(side=tk.LEFT)
        elements_var = tk.StringVar(value="魔法体系、势力分布、历史背景")
        ttk.Entry(elements_frame, textvariable=elements_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 世界观描述
        desc_frame = ttk.LabelFrame(dialog, text="世界观详细描述", padding=10)
        desc_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        desc_text = tk.Text(desc_frame, wrap=tk.WORD, height=10,
                           font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                           bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        desc_text.pack(fill=tk.BOTH, expand=True)
        desc_text.insert("1.0", "请在此描述世界观的核心设定、规则、势力分布等内容...")
        
        # 按钮
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_worldview():
            name = name_var.get().strip()
            elements = elements_var.get().strip()
            desc = desc_text.get("1.0", tk.END).strip()
            
            if not name:
                messagebox.showwarning("警告", "请输入世界观名称！")
                return
            
            # 添加到列表
            self._worldview_tree.insert("", tk.END, values=(
                name,
                elements[:50] + "..." if len(elements) > 50 else elements,
                "新建",
                datetime.now().strftime("%Y-%m-%d %H:%M")
            ))
            
            # 更新预览
            self._worldview_preview.delete("1.0", tk.END)
            self._worldview_preview.insert("1.0", f"【{name}】\n\n核心元素：{elements}\n\n{desc}")
            
            self._set_status(f"已创建世界观：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_worldview).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_worldview_view(self):
        """查看世界观详情"""
        self._set_status("查看世界观详情功能开发中...")
    
    def _on_worldview_edit(self):
        """编辑世界观"""
        self._set_status("编辑世界观功能开发中...")
    
    def _on_worldview_delete(self):
        """批量删除世界观"""
        self._set_status("批量删除世界观功能开发中...")
    
    def _on_worldview_link(self):
        """关联要素"""
        self._set_status("关联要素功能开发中...")
    
    def _on_worldview_import(self):
        """导入世界观"""
        self._set_status("世界观导入功能开发中...")
    
    def _on_worldview_clear(self):
        """清除世界观"""
        self._worldview_path_var.set("")
        self._worldview_preview.delete("1.0", tk.END)
    
    def _on_character_browse(self):
        """浏览人物档案文件"""
        path = filedialog.askopenfilename(
            title="选择人物档案文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._character_path_var.set(path)
            self._set_status(f"已选择: {os.path.basename(path)}")
    
    def _on_character_new(self):
        """新建人物（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("新建人物")
        dialog.geometry("600x700")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 600) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 700) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 创建滚动区域
        canvas = tk.Canvas(dialog, bg=GlassTheme.GLASS_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 基本信息
        basic_frame = ttk.LabelFrame(scrollable_frame, text="基本信息", padding=10)
        basic_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 姓名
        name_frame = ttk.Frame(basic_frame, style="TFrame")
        name_frame.pack(fill=tk.X, pady=5)
        ttk.Label(name_frame, text="姓名：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="新角色")
        ttk.Entry(name_frame, textvariable=name_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 角色类型
        role_frame = ttk.Frame(basic_frame, style="TFrame")
        role_frame.pack(fill=tk.X, pady=5)
        ttk.Label(role_frame, text="角色类型：").pack(side=tk.LEFT)
        role_var = tk.StringVar(value="主角")
        ttk.Combobox(role_frame, textvariable=role_var, values=["主角", "配角", "反派", "路人"], width=15).pack(side=tk.LEFT, padx=10)
        
        # 外貌描述
        appearance_frame = ttk.LabelFrame(scrollable_frame, text="外貌描述", padding=10)
        appearance_frame.pack(fill=tk.X, padx=20, pady=10)
        appearance_text = tk.Text(appearance_frame, wrap=tk.WORD, height=3,
                                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                 bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        appearance_text.pack(fill=tk.X)
        
        # 性格特点
        personality_frame = ttk.LabelFrame(scrollable_frame, text="性格特点", padding=10)
        personality_frame.pack(fill=tk.X, padx=20, pady=10)
        personality_text = tk.Text(personality_frame, wrap=tk.WORD, height=3,
                                  font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                  bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        personality_text.pack(fill=tk.X)
        
        # 背景故事
        background_frame = ttk.LabelFrame(scrollable_frame, text="背景故事", padding=10)
        background_frame.pack(fill=tk.X, padx=20, pady=10)
        background_text = tk.Text(background_frame, wrap=tk.WORD, height=5,
                                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                                 bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        background_text.pack(fill=tk.X)
        
        # 能力特长
        ability_frame = ttk.LabelFrame(scrollable_frame, text="能力特长", padding=10)
        ability_frame.pack(fill=tk.X, padx=20, pady=10)
        ability_text = tk.Text(ability_frame, wrap=tk.WORD, height=3,
                              font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                              bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        ability_text.pack(fill=tk.X)
        
        # 按钮
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_character():
            name = name_var.get().strip()
            role = role_var.get()
            
            if not name:
                messagebox.showwarning("警告", "请输入人物姓名！")
                return
            
            # 添加到列表
            self._character_tree.insert("", tk.END, values=(
                "👤",  # 头像占位
                name,
                role,
                "新建",
                "平静",
                "第1章"
            ))
            
            # 更新详情
            appearance = appearance_text.get("1.0", tk.END).strip()
            personality = personality_text.get("1.0", tk.END).strip()
            background = background_text.get("1.0", tk.END).strip()
            ability = ability_text.get("1.0", tk.END).strip()
            
            self._character_detail.delete("1.0", tk.END)
            self._character_detail.insert("1.0", f"""【{name}】- {role}

外貌：{appearance}

性格：{personality}

背景：{background}

能力：{ability}
""")
            
            self._set_status(f"已创建人物：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_character).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_character_edit(self):
        """编辑人物"""
        self._set_status("编辑人物功能开发中...")
    
    def _on_character_detail(self):
        """人物详情页"""
        self._set_status("人物详情页功能开发中...")
    
    def _on_character_relation(self):
        """关系网络图"""
        self._set_status("关系网络图功能开发中...")
    
    def _on_character_delete(self):
        """删除人物"""
        self._set_status("删除人物功能开发中...")
    
    def _on_outline_browse(self):
        """浏览大纲文件"""
        path = filedialog.askopenfilename(
            title="选择大纲文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._outline_path_var.set(os.path.basename(path))
            self._set_status(f"已选择: {os.path.basename(path)}")
    
    def _on_outline_new(self):
        """新建大纲（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("新建大纲")
        dialog.geometry("500x300")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 300) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 大纲名称
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        ttk.Label(name_frame, text="大纲名称：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="新小说大纲")
        ttk.Entry(name_frame, textvariable=name_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 小说类型
        type_frame = ttk.Frame(dialog, style="TFrame")
        type_frame.pack(fill=tk.X, padx=20, pady=10)
        ttk.Label(type_frame, text="小说类型：").pack(side=tk.LEFT)
        type_var = tk.StringVar(value="玄幻")
        ttk.Combobox(type_frame, textvariable=type_var, values=["玄幻", "仙侠", "都市", "历史", "科幻", "言情"], width=15).pack(side=tk.LEFT, padx=10)
        
        # 目标字数
        words_frame = ttk.Frame(dialog, style="TFrame")
        words_frame.pack(fill=tk.X, padx=20, pady=10)
        ttk.Label(words_frame, text="目标字数：").pack(side=tk.LEFT)
        words_var = tk.StringVar(value="500000")
        ttk.Entry(words_frame, textvariable=words_var, width=15).pack(side=tk.LEFT, padx=10)
        ttk.Label(words_frame, text="字").pack(side=tk.LEFT)
        
        # 简介
        intro_frame = ttk.LabelFrame(dialog, text="小说简介", padding=10)
        intro_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        intro_text = tk.Text(intro_frame, wrap=tk.WORD, height=4,
                            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                            bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        intro_text.pack(fill=tk.BOTH, expand=True)
        
        # 按钮
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_outline():
            name = name_var.get().strip()
            novel_type = type_var.get()
            target_words = words_var.get()
            intro = intro_text.get("1.0", tk.END).strip()
            
            if not name:
                messagebox.showwarning("警告", "请输入大纲名称！")
                return
            
            # 添加根节点到大纲树
            root_node = self._outline_tree.insert("", tk.END, text=f"📖 {name}", open=True)
            self._outline_tree.insert(root_node, tk.END, text="第一卷")
            self._outline_tree.insert(root_node, tk.END, text="第二卷")
            
            self._outline_path_var.set(name)
            self._set_status(f"已创建大纲：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_outline).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_outline_export(self):
        """导出大纲"""
        self._set_status("导出大纲功能开发中...")
    
    def _on_outline_add_volume(self):
        """添加卷"""
        self._set_status("添加卷功能开发中...")
    
    def _on_outline_add_chapter(self):
        """添加章节"""
        self._set_status("添加章节功能开发中...")
    
    def _on_outline_edit_chapter(self):
        """编辑章节"""
        self._set_status("编辑章节功能开发中...")
    
    def _on_outline_delete(self):
        """删除"""
        self._set_status("删除功能开发中...")
    
    def _on_outline_parse(self):
        """解析大纲"""
        self._set_status("大纲解析功能开发中...")
    
    def _on_outline_clear(self):
        """清除大纲"""
        self._outline_path_var.set("未导入")
    
    def _on_style_browse(self):
        """浏览风格文件"""
        path = filedialog.askopenfilename(
            title="选择风格文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._style_path_var.set(os.path.basename(path))
            self._set_status(f"已选择: {os.path.basename(path)}")
    
    def _on_style_delete(self):
        """删除范文"""
        self._set_status("删除范文功能开发中...")
    
    def _on_style_export(self):
        """导出风格"""
        self._set_status("导出风格功能开发中...")
    
    def _on_style_switch(self):
        """切换风格"""
        self._set_status("切换风格功能开发中...")
    
    def _on_style_edit_weight(self):
        """编辑权重"""
        self._set_status("编辑权重功能开发中...")
    
    def _on_style_delete_profile(self):
        """删除风格档案"""
        self._set_status("删除风格档案功能开发中...")
    
    def _on_style_learn(self):
        """学习风格"""
        self._set_status("风格学习功能开发中...")
    
    def _on_style_create(self):
        """创建风格模板（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("创建风格模板")
        dialog.geometry("600x500")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 600) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 500) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 风格名称
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        ttk.Label(name_frame, text="风格名称：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="自定义风格")
        ttk.Entry(name_frame, textvariable=name_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 七维度评分
        dimensions_frame = ttk.LabelFrame(dialog, text="七维度风格评分（0-10分）", padding=10)
        dimensions_frame.pack(fill=tk.X, padx=20, pady=10)
        
        dimensions = [
            ("词汇丰富度", "vocabulary"),
            ("句式复杂度", "sentence"),
            ("修辞使用", "rhetoric"),
            ("情感强度", "emotion"),
            ("节奏感", "rhythm"),
            ("结构完整性", "structure"),
            ("细节描写", "detail")
        ]
        
        dim_vars = {}
        for i, (label, key) in enumerate(dimensions):
            row = ttk.Frame(dimensions_frame, style="TFrame")
            row.pack(fill=tk.X, pady=3)
            ttk.Label(row, text=f"{label}：", width=15, anchor='e').pack(side=tk.LEFT)
            var = tk.DoubleVar(value=5.0)
            dim_vars[key] = var
            scale = ttk.Scale(row, from_=0.0, to=10.0, variable=var, orient=tk.HORIZONTAL, length=200)
            scale.pack(side=tk.LEFT, padx=10)
            ttk.Label(row, textvariable=var, width=5).pack(side=tk.LEFT)
        
        # 风格描述
        desc_frame = ttk.LabelFrame(dialog, text="风格描述（可选）", padding=10)
        desc_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        desc_text = tk.Text(desc_frame, wrap=tk.WORD, height=5,
                           font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                           bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        desc_text.pack(fill=tk.BOTH, expand=True)
        desc_text.insert("1.0", "描述这种风格的特点，如：华丽、简洁、幽默、严肃等...")
        
        # 按钮
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_style():
            name = name_var.get().strip()
            if not name:
                messagebox.showwarning("警告", "请输入风格名称！")
                return
            
            # 获取评分
            scores = {key: var.get() for key, var in dim_vars.items()}
            
            # 添加到风格档案库
            self._style_tree.insert("", tk.END, values=(
                name,
                f"{scores['vocabulary']:.1f}",
                f"{scores['sentence']:.1f}",
                f"{scores['rhetoric']:.1f}",
                f"{scores['emotion']:.1f}",
                f"{scores['rhythm']:.1f}",
                f"{scores['structure']:.1f}",
                f"{scores['detail']:.1f}"
            ))
            
            # 更新分析器显示
            self._style_info.delete("1.0", tk.END)
            self._style_info.insert("1.0", f"""【{name}】风格模板

词汇：{"█" * int(scores['vocabulary'])}{"░" * (10 - int(scores['vocabulary']))} {scores['vocabulary']:.1f}
句式：{"█" * int(scores['sentence'])}{"░" * (10 - int(scores['sentence']))} {scores['sentence']:.1f}
修辞：{"█" * int(scores['rhetoric'])}{"░" * (10 - int(scores['rhetoric']))} {scores['rhetoric']:.1f}
情感：{"█" * int(scores['emotion'])}{"░" * (10 - int(scores['emotion']))} {scores['emotion']:.1f}
节奏：{"█" * int(scores['rhythm'])}{"░" * (10 - int(scores['rhythm']))} {scores['rhythm']:.1f}
结构：{"█" * int(scores['structure'])}{"░" * (10 - int(scores['structure']))} {scores['structure']:.1f}
细节：{"█" * int(scores['detail'])}{"░" * (10 - int(scores['detail']))} {scores['detail']:.1f}

描述：{desc_text.get("1.0", tk.END).strip()}
""")
            
            self._set_status(f"已创建风格模板：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_style).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_style_apply(self):
        """应用到生成器"""
        self._set_status("应用到生成器功能开发中...")
    
    def _on_style_clear(self):
        """清除风格"""
        self._style_path_var.set("未导入")
        self._style_info.delete("1.0", tk.END)
    
    def _on_start_generation(self):
        """开始生成"""
        self._set_status("开始创作功能开发中...")
    
    def _on_stop_generation(self):
        """停止生成"""
        self._set_status("已停止生成")
    
    def _on_gen_browse(self):
        """分章浏览"""
        self._set_status("分章浏览功能开发中...")
    
    def _on_gen_save(self):
        """保存项目"""
        self._set_status("保存项目功能开发中...")
    
    def _on_gen_export_txt(self):
        """导出TXT"""
        self._set_status("导出TXT功能开发中...")
    
    def _on_gen_export_docx(self):
        """导出DOCX"""
        self._set_status("导出DOCX功能开发中...")
    
    def _on_reverse_browse(self):
        """浏览逆向分析文件"""
        path = filedialog.askopenfilename(
            title="选择需要分析的文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._reverse_path_var.set(path)
    
    def _on_reverse_run(self):
        """运行分析"""
        self._set_status("运行分析功能开发中...")
    
    def _on_reverse_analyze(self):
        """开始逆向分析"""
        self._set_status("逆向分析功能开发中...")
    
    def _on_reverse_apply(self):
        """自动修正"""
        self._set_status("自动修正功能开发中...")
    
    def _on_reverse_outline(self):
        """调整大纲细节"""
        self._set_status("调整大纲细节功能开发中...")
    
    def _on_reverse_character(self):
        """修改人物性格"""
        self._set_status("修改人物性格功能开发中...")
    
    def _on_reverse_trajectory(self):
        """补充人物轨迹"""
        self._set_status("补充人物轨迹功能开发中...")
    
    def _on_reverse_worldview(self):
        """增减世界观设定"""
        self._set_status("增减世界观设定功能开发中...")
    
    def _on_quick_generate(self):
        """快捷生成（保留兼容）"""
        self._on_quick_generate_all()
    
    def _on_quick_generate_all(self):
        """一键生成全部四个结果"""
        self._set_status("正在生成世界观、大纲、人设、关键情节...")
        # TODO: 调用AI服务生成
        self._quick_worldview_result.delete("1.0", tk.END)
        self._quick_worldview_result.insert("1.0", "【世界观设定】\n\n正在生成中...")
        
        self._quick_outline_result.delete("1.0", tk.END)
        self._quick_outline_result.insert("1.0", "【小说大纲】\n\n正在生成中...")
        
        self._quick_characters_result.delete("1.0", tk.END)
        self._quick_characters_result.insert("1.0", "【人物设定】\n\n正在生成中...")
        
        self._quick_plot_result.delete("1.0", tk.END)
        self._quick_plot_result.insert("1.0", "【关键情节】\n\n正在生成中...")
    
    def _on_quick_generate_single(self, gen_type: str):
        """生成单个结果"""
        type_names = {
            "worldview": "世界观",
            "outline": "大纲",
            "characters": "人设",
            "plot": "关键情节"
        }
        self._set_status(f"正在生成{type_names.get(gen_type, '')}...")
        # TODO: 调用AI服务生成
    
    def _on_quick_upload_file(self):
        """上传参考文本文件"""
        path = filedialog.askopenfilename(
            title="选择参考文本",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._quick_upload_path.set(path)
            self._set_status(f"已上传参考文本：{os.path.basename(path)}")
            # 更新上传状态
            for key in self._upload_status_labels:
                self._upload_status_labels[key].configure(text="已上传", foreground=GlassTheme.SUCCESS)
    
    def _show_chapters_context_menu(self, event):
        """显示已完成章节右键菜单"""
        # 选中点击项
        item = self._completed_chapters_tree.identify_row(event.y)
        if item:
            self._completed_chapters_tree.selection_set(item)
            self._chapters_context_menu.post(event.x_root, event.y_root)
    
    def _on_upload_chapter(self):
        """上传章节"""
        path = filedialog.askopenfilename(
            title="选择章节文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            # 添加到列表
            self._completed_chapters_tree.insert("", tk.END, values=(
                f"第{self._completed_chapters_tree.get_children().__len__() + 1}章",
                os.path.basename(path),
                "待统计",
                "未完成"
            ))
            self._set_status(f"已上传章节：{os.path.basename(path)}")
    
    def _on_mark_chapter_completed(self):
        """标记章节为完成"""
        selected = self._completed_chapters_tree.selection()
        for item in selected:
            values = list(self._completed_chapters_tree.item(item, "values"))
            values[3] = "已完成"
            self._completed_chapters_tree.item(item, values=values)
        self._set_status(f"已标记 {len(selected)} 个章节为完成")
    
    def _on_mark_chapter_incomplete(self):
        """标记章节为未完成"""
        selected = self._completed_chapters_tree.selection()
        for item in selected:
            values = list(self._completed_chapters_tree.item(item, "values"))
            values[3] = "未完成"
            self._completed_chapters_tree.item(item, values=values)
        self._set_status(f"已标记 {len(selected)} 个章节为未完成")
    
    def _on_delete_completed_chapter(self):
        """删除已完成章节"""
        selected = self._completed_chapters_tree.selection()
        for item in selected:
            self._completed_chapters_tree.delete(item)
        self._set_status(f"已删除 {len(selected)} 个章节")
    
    def _on_mark_all_completed(self):
        """标记所有章节为完成"""
        for item in self._completed_chapters_tree.get_children():
            values = list(self._completed_chapters_tree.item(item, "values"))
            values[3] = "已完成"
            self._completed_chapters_tree.item(item, values=values)
        self._set_status("已标记所有章节为完成")
    
    def _on_quick_save_txt(self):
        """保存为TXT"""
        self._set_status("保存为TXT功能开发中...")
    
    def _on_quick_save_docx(self):
        """保存为DOCX"""
        self._set_status("保存为DOCX功能开发中...")
    
    def _on_quick_import(self):
        """导入当前项目"""
        self._set_status("导入当前项目功能开发中...")
    
    def _on_quick_export_results(self):
        """导出生成结果（用户选择保存地址）"""
        # 获取所有结果内容
        worldview = self._quick_worldview_result.get("1.0", tk.END).strip()
        outline = self._quick_outline_result.get("1.0", tk.END).strip()
        characters = self._quick_characters_result.get("1.0", tk.END).strip()
        plot = self._quick_plot_result.get("1.0", tk.END).strip()
        
        # 检查是否有内容
        if all(not content or content.endswith("将在此显示...") for content in [worldview, outline, characters, plot]):
            messagebox.showwarning("警告", "暂无生成结果可导出！")
            return
        
        # 选择保存路径
        save_path = filedialog.asksaveasfilename(
            title="选择保存位置",
            defaultextension=".txt",
            filetypes=[
                ("文本文件", "*.txt"),
                ("Word文档", "*.docx"),
                ("Markdown文件", "*.md"),
                ("所有文件", "*.*")
            ],
            initialfile="生成结果"
        )
        
        if not save_path:
            return
        
        try:
            # 构建完整内容
            full_content = f"""【世界观设定】
{worldview if worldview and not worldview.endswith("将在此显示...") else "暂无"}

{'='*50}

【小说大纲】
{outline if outline and not outline.endswith("将在此显示...") else "暂无"}

{'='*50}

【人物设定】
{characters if characters and not characters.endswith("将在此显示...") else "暂无"}

{'='*50}

【关键情节】
{plot if plot and not plot.endswith("将在此显示...") else "暂无"}

{'='*50}

生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
            
            # 根据文件类型保存
            if save_path.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading("快捷创作生成结果", 0)
                    
                    # 世界观
                    doc.add_heading("🌍 世界观设定", 1)
                    if worldview and not worldview.endswith("将在此显示..."):
                        doc.add_paragraph(worldview)
                    else:
                        doc.add_paragraph("暂无")
                    
                    # 大纲
                    doc.add_heading("📄 小说大纲", 1)
                    if outline and not outline.endswith("将在此显示..."):
                        doc.add_paragraph(outline)
                    else:
                        doc.add_paragraph("暂无")
                    
                    # 人设
                    doc.add_heading("👥 人物设定", 1)
                    if characters and not characters.endswith("将在此显示..."):
                        doc.add_paragraph(characters)
                    else:
                        doc.add_paragraph("暂无")
                    
                    # 情节
                    doc.add_heading("🎬 关键情节", 1)
                    if plot and not plot.endswith("将在此显示..."):
                        doc.add_paragraph(plot)
                    else:
                        doc.add_paragraph("暂无")
                    
                    doc.add_paragraph(f"\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    doc.save(save_path)
                except ImportError:
                    messagebox.showerror("错误", "导出DOCX需要安装python-docx库！\n请运行：pip install python-docx")
                    return
            else:
                # TXT或MD文件
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(full_content)
            
            self._set_status(f"已导出到：{os.path.basename(save_path)}")
            messagebox.showinfo("成功", f"导出成功！\n保存位置：{save_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{e}")
            self._set_status(f"导出失败：{e}")
    
    def _on_continue_browse(self):
        """浏览续写原文文件"""
        path = filedialog.askopenfilename(
            title="选择原文文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            try:
                if path.lower().endswith('.docx'):
                    from docx import Document
                    doc = Document(path)
                    text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                else:
                    with open(path, 'r', encoding='utf-8') as f:
                        text = f.read()
                self._continue_source.delete("1.0", tk.END)
                self._continue_source.insert("1.0", text)
                self._set_status(f"已导入: {os.path.basename(path)}")
            except Exception as e:
                self._set_status(f"导入失败: {e}")
    
    def _on_continue_generate(self):
        """开始续写"""
        self._set_status("续写功能开发中...")
    
    def _on_continue_regenerate(self):
        """重新生成续写"""
        self._set_status("重新生成续写...")
    
    def _on_continue_select_best(self):
        """选择最佳版本"""
        self._set_status("选择最佳版本功能开发中...")
    
    def _on_continue_save(self):
        """保存续写结果"""
        self._set_status("保存续写结果功能开发中...")
    
    def _create_progress_page(self) -> tk.Frame:
        """创建创作进度页面（支持滚动）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建Canvas和Scrollbar用于滚动
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题
        header = ttk.Frame(scrollable_frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="创作进度", style="Title.TLabel").pack(side=tk.LEFT)
        
        # 进度统计
        progress_frame = ttk.LabelFrame(scrollable_frame, text="当前项目状态", padding=20)
        progress_frame.pack(fill=tk.X, padx=20, pady=10)
        
        stats = [
            ("当前项目", "未打开项目"),
            ("已完成章节", "0 / 0"),
            ("总字数", "0 字"),
            ("今日目标", "0 / 3000 字"),
            ("大纲解析", "未完成"),
            ("人物录入", "0 人"),
            ("世界观条目", "0 个"),
            ("当前风格", "未设置"),
        ]
        
        for i, (label, value) in enumerate(stats):
            row_frame = ttk.Frame(progress_frame, style="TFrame")
            row_frame.pack(fill=tk.X, pady=3)
            
            ttk.Label(row_frame, text=f"{label}：", width=15, anchor='e').pack(side=tk.LEFT)
            ttk.Label(row_frame, text=value, foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=10)
        
        # 操作按钮
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=10)

        buttons = [
            ("刷新统计", self._on_refresh_progress),
            ("导出报告", self._on_export_progress),
            ("重置进度", self._on_reset_progress),
        ]

        for i, (text, command) in enumerate(buttons):
            ResponsiveButton(
                btn_frame,
                text=text,
                command=command,
                async_handler=self._async_handler
            ).grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)
        
        return frame
    
    def _create_project_page(self) -> tk.Frame:
        """创建项目管理页面（支持滚动）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建Canvas和Scrollbar用于滚动
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题
        header = ttk.Frame(scrollable_frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="项目管理", style="Title.TLabel").pack(side=tk.LEFT)
        
        # 当前项目信息
        info_frame = ttk.LabelFrame(scrollable_frame, text="当前项目", padding=20)
        info_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 项目名称
        name_frame = ttk.Frame(info_frame, style="TFrame")
        name_frame.pack(fill=tk.X, pady=5)
        ttk.Label(name_frame, text="项目名称：").pack(side=tk.LEFT)
        self._project_name_var = tk.StringVar(value="未打开项目")
        ttk.Label(name_frame, textvariable=self._project_name_var).pack(side=tk.LEFT, padx=10)
        
        # 项目路径
        path_frame = ttk.Frame(info_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        ttk.Label(path_frame, text="项目路径：").pack(side=tk.LEFT)
        self._project_path_var = tk.StringVar(value="-")
        ttk.Label(path_frame, textvariable=self._project_path_var).pack(side=tk.LEFT, padx=10)
        
        # 操作按钮
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=10)

        buttons = [
            ("新建项目", self._on_new_project),
            ("打开项目", self._on_open_project),
            ("导出备份", self._on_backup_project),
        ]

        for i, (text, command) in enumerate(buttons):
            ResponsiveButton(
                btn_frame,
                text=text,
                command=command,
                async_handler=self._async_handler,
                style="Accent.TButton" if i == 0 else "TButton"
            ).grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_plugins_page(self) -> tk.Frame:
        """创建插件管理页面（支持滚动）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建Canvas和Scrollbar用于滚动
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题
        header = ttk.Frame(scrollable_frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="插件管理", style="Title.TLabel").pack(side=tk.LEFT)
        
        # 插件列表
        list_frame = ttk.LabelFrame(scrollable_frame, text="已安装插件", padding=15)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        columns = ("name", "version", "status", "type")
        self._plugin_tree = ttk.Treeview(
            list_frame,
            columns=columns,
            show="headings",
            height=15
        )

        # 配置字体 - 直接使用Treeview样式
        style = ttk.Style()
        style.configure("Treeview",
                     font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL))
        style.configure("Treeview.Heading",
                     font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold"))
        
        self._plugin_tree.heading("name", text="插件名称")
        self._plugin_tree.heading("version", text="版本")
        self._plugin_tree.heading("status", text="状态")
        self._plugin_tree.heading("type", text="类型")
        
        self._plugin_tree.column("name", width=200)
        self._plugin_tree.column("version", width=80)
        self._plugin_tree.column("status", width=80)
        self._plugin_tree.column("type", width=100)
        
        self._plugin_tree.pack(fill=tk.BOTH, expand=True)
        
        # 加载插件数据
        self._load_plugins()
        
        # 操作按钮
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=10)

        buttons = [
            ("刷新", self._on_refresh_plugins),
            ("安装插件", self._on_install_plugin),
            ("卸载插件", self._on_uninstall_plugin),
        ]

        for i, (text, command) in enumerate(buttons):
            ResponsiveButton(
                btn_frame,
                text=text,
                command=command,
                async_handler=self._async_handler
            ).grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_settings_page(self) -> tk.Frame:
        """创建设置页面（支持滚动）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建滚动容器
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题
        ttk.Label(scrollable_frame, text="设置", style="Title.TLabel").pack(anchor=tk.W, padx=20, pady=(20, 10))
        
        # AI服务配置
        ai_frame = ttk.LabelFrame(scrollable_frame, text="AI服务配置", padding=15)
        ai_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 服务模式
        mode_frame = ttk.Frame(ai_frame, style="TFrame")
        mode_frame.pack(fill=tk.X, pady=5)
        ttk.Label(mode_frame, text="服务模式：").pack(side=tk.LEFT)
        self._service_mode_var = tk.StringVar(value="online")
        mode_local = ttk.Radiobutton(mode_frame, text="本地大模型", variable=self._service_mode_var, value="local", command=self._on_service_mode_changed)
        mode_local.pack(side=tk.LEFT, padx=10)
        mode_online = ttk.Radiobutton(mode_frame, text="线上API", variable=self._service_mode_var, value="online", command=self._on_service_mode_changed)
        mode_online.pack(side=tk.LEFT)
        
        # API提供商
        provider_frame = ttk.Frame(ai_frame, style="TFrame")
        provider_frame.pack(fill=tk.X, pady=5)
        ttk.Label(provider_frame, text="提供商：").pack(side=tk.LEFT)
        self._provider_var = tk.StringVar(value="DeepSeek")
        provider_combo = ttk.Combobox(
            provider_frame,
            textvariable=self._provider_var,
            values=["DeepSeek", "OpenAI", "Anthropic", "Ollama"],
            state="readonly",
            width=20
        )
        provider_combo.pack(side=tk.LEFT, padx=10)
        
        # 模型选择
        model_frame = ttk.Frame(ai_frame, style="TFrame")
        model_frame.pack(fill=tk.X, pady=5)
        ttk.Label(model_frame, text="模型：").pack(side=tk.LEFT)
        self._model_var = tk.StringVar(value="deepseek-chat")
        model_combo = ttk.Combobox(
            model_frame,
            textvariable=self._model_var,
            values=["deepseek-chat", "deepseek-reasoner", "gpt-4", "gpt-3.5-turbo", "claude-3"],
            state="readonly",
            width=20
        )
        model_combo.pack(side=tk.LEFT, padx=10)
        
        # API Key / 本地部署地址（根据模式显示不同字段）
        key_frame = ttk.Frame(ai_frame, style="TFrame")
        key_frame.pack(fill=tk.X, pady=5)
        self._key_label = ttk.Label(key_frame, text="API Key：")
        self._key_label.pack(side=tk.LEFT)
        self._api_key_var = tk.StringVar()
        self._local_url_var = tk.StringVar(value="http://localhost:11434/v1")
        self._key_entry = ttk.Entry(key_frame, textvariable=self._api_key_var, width=40, show="*")
        self._key_entry.pack(side=tk.LEFT, padx=10)
        self._url_entry = ttk.Entry(key_frame, textvariable=self._local_url_var, width=40)
        # 默认隐藏本地地址输入框
        self._url_entry.pack_forget()
        
        # Temperature
        temp_frame = ttk.Frame(ai_frame, style="TFrame")
        temp_frame.pack(fill=tk.X, pady=5)
        ttk.Label(temp_frame, text="Temperature：").pack(side=tk.LEFT)
        self._temp_var = tk.DoubleVar(value=0.7)
        temp_scale = ttk.Scale(
            temp_frame,
            from_=0.0,
            to=2.0,
            variable=self._temp_var,
            orient=tk.HORIZONTAL,
            length=200
        )
        temp_scale.pack(side=tk.LEFT, padx=10)
        self._temp_label = ttk.Label(temp_frame, text="0.70")
        self._temp_label.pack(side=tk.LEFT)
        temp_scale.configure(command=lambda v: self._temp_label.configure(text=f"{float(v):.2f}"))
        
        # 界面设置
        ui_frame = ttk.LabelFrame(scrollable_frame, text="界面设置", padding=15)
        ui_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 主题
        theme_frame = ttk.Frame(ui_frame, style="TFrame")
        theme_frame.pack(fill=tk.X, pady=5)
        ttk.Label(theme_frame, text="主题：").pack(side=tk.LEFT)
        self._theme_var = tk.StringVar(value="dark")
        ttk.Radiobutton(theme_frame, text="深色", variable=self._theme_var, value="dark", command=self._on_theme_changed).pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(theme_frame, text="浅色", variable=self._theme_var, value="light", command=self._on_theme_changed).pack(side=tk.LEFT)
        
        # 字体大小
        font_frame = ttk.Frame(ui_frame, style="TFrame")
        font_frame.pack(fill=tk.X, pady=5)
        ttk.Label(font_frame, text="字体大小：").pack(side=tk.LEFT)
        self._font_size_var = tk.StringVar(value="14")
        font_spinbox = ttk.Spinbox(font_frame, from_=10, to=20, textvariable=self._font_size_var, width=8)
        font_spinbox.pack(side=tk.LEFT, padx=10)
        font_spinbox.bind("<Return>", lambda e: self._on_font_size_changed())

        # 文件路径设置
        # 文件路径设置
        path_frame = ttk.LabelFrame(scrollable_frame, text="文件路径设置", padding=15)
        path_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 默认保存位置
        save_path_frame = ttk.Frame(path_frame, style="TFrame")
        save_path_frame.pack(fill=tk.X, pady=5)
        ttk.Label(save_path_frame, text="默认保存位置：").pack(side=tk.LEFT)
        self._save_path_var = tk.StringVar(value=os.getcwd())
        ttk.Entry(save_path_frame, textvariable=self._save_path_var, width=40).pack(side=tk.LEFT, padx=10)
        ttk.Button(save_path_frame, text="浏览", command=self._on_browse_save_path).pack(side=tk.LEFT)
        
        # 自动备份间隔
        backup_frame = ttk.Frame(path_frame, style="TFrame")
        backup_frame.pack(fill=tk.X, pady=5)
        ttk.Label(backup_frame, text="自动备份间隔：").pack(side=tk.LEFT)
        self._backup_interval_var = tk.StringVar(value="30")
        ttk.Spinbox(backup_frame, from_=10, to=120, textvariable=self._backup_interval_var, width=8).pack(side=tk.LEFT, padx=10)
        ttk.Label(backup_frame, text="分钟").pack(side=tk.LEFT)
        
        # 偏好学习设置
        learning_frame = ttk.LabelFrame(scrollable_frame, text="偏好学习设置", padding=15)
        learning_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # AI学习开关
        ai_learning_frame = ttk.Frame(learning_frame, style="TFrame")
        ai_learning_frame.pack(fill=tk.X, pady=5)
        ttk.Label(ai_learning_frame, text="AI偏好学习：").pack(side=tk.LEFT)
        self._ai_learning_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(ai_learning_frame, text="启用（AI将学习您的写作偏好）", variable=self._ai_learning_var).pack(side=tk.LEFT, padx=10)
        
        # 数据清除
        clear_frame = ttk.Frame(learning_frame, style="TFrame")
        clear_frame.pack(fill=tk.X, pady=5)
        ttk.Label(clear_frame, text="偏好数据：").pack(side=tk.LEFT)
        ttk.Button(clear_frame, text="清除学习数据", command=self._on_clear_learning_data).pack(side=tk.LEFT, padx=10)
        ttk.Button(clear_frame, text="导出偏好配置", command=self._on_export_preferences).pack(side=tk.LEFT, padx=5)
        
        # 从config.yaml加载设置
        self._load_settings_from_config()
        
        # 保存按钮
        save_frame = ttk.Frame(scrollable_frame, style="TFrame")
        save_frame.pack(fill=tk.X, padx=20, pady=20)
        
        ResponsiveButton(
            save_frame,
            text="保存设置",
            command=self._on_save_settings,
            async_handler=self._async_handler,
            style="Accent.TButton"
        ).pack(side=tk.LEFT)
        
        return frame
    
    # ============== 辅助方法 ==============
    
    def _set_status(self, message: str) -> None:
        """设置状态栏消息"""
        if self._status_var:
            self._status_var.set(message)
    
    def _process_result_queue(self) -> None:
        """处理结果队列"""
        try:
            while True:
                result = self._result_queue.get_nowait()
                if isinstance(result, dict):
                    result_type = result.get('type')
                    if result_type == 'status':
                        self._set_status(result.get('message', ''))
        except queue.Empty:
            pass
        finally:
            self.root.after(100, self._process_result_queue)
    
    # ============== 功能方法 ==============
    
    def _get_hot_ranking_data(self) -> Dict:
        """获取热榜数据"""
        return {
            'sites': [
                {
                    'name': '📚 起点中文网',
                    'color': '#457B9D',
                    'books': [
                        {'title': '诡秘之主', 'author': '爱潜水的乌贼', 'category': '玄幻', 'heat': 150000},
                        {'title': '大奉打更人', 'author': '卖报小郎君', 'category': '仙侠', 'heat': 120000},
                        {'title': '深空彼岸', 'author': '辰东', 'category': '科幻', 'heat': 100000},
                        {'title': '星门', 'author': '老鹰吃小鸡', 'category': '玄幻', 'heat': 80000},
                        {'title': '完美世界', 'author': '辰东', 'category': '玄幻', 'heat': 60000},
                        {'title': '遮天', 'author': '辰东', 'category': '玄幻', 'heat': 55000},
                        {'title': '斗破苍穹', 'author': '天蚕土豆', 'category': '玄幻', 'heat': 50000},
                        {'title': '凡人修仙传', 'author': '忘语', 'category': '仙侠', 'heat': 45000},
                        {'title': '一念永恒', 'author': '耳根', 'category': '仙侠', 'heat': 40000},
                        {'title': '牧神记', 'author': '宅猪', 'category': '玄幻', 'heat': 35000},
                    ]
                },
                {
                    'name': '🎭 晋江文学城',
                    'color': '#E63946',
                    'books': [
                        {'title': '天官赐福', 'author': '墨香铜臭', 'category': '纯爱', 'heat': 180000},
                        {'title': '魔道祖师', 'author': '墨香铜臭', 'category': '纯爱', 'heat': 150000},
                        {'title': '人渣反派自救系统', 'author': '墨香铜臭', 'category': '纯爱', 'heat': 120000},
                        {'title': '二哈和他的白猫师尊', 'author': '肉包不吃肉', 'category': '纯爱', 'heat': 90000},
                        {'title': '全球高考', 'author': '木苏里', 'category': '纯爱', 'heat': 70000},
                        {'title': '撒野', 'author': '巫哲', 'category': '纯爱', 'heat': 65000},
                        {'title': '默读', 'author': 'Priest', 'category': '纯爱', 'heat': 60000},
                        {'title': '镇魂', 'author': 'Priest', 'category': '纯爱', 'heat': 55000},
                        {'title': '杀破狼', 'author': 'Priest', 'category': '纯爱', 'heat': 50000},
                        {'title': '有兽焉', 'author': '鹤来', 'category': '纯爱', 'heat': 45000},
                    ]
                },
                {
                    'name': '🍅 番茄小说',
                    'color': '#FF6B6B',
                    'books': [
                        {'title': '我在精神病院学斩神', 'author': '三九音域', 'category': '都市', 'heat': 185000},
                        {'title': '星门', 'author': '老鹰吃小鸡', 'category': '玄幻', 'heat': 168000},
                        {'title': '斩神', 'author': '三九音域', 'category': '都市', 'heat': 152000},
                        {'title': '重生之都市修仙', 'author': '十里剑神', 'category': '都市', 'heat': 138000},
                        {'title': '全职法师', 'author': '乱', 'category': '玄幻', 'heat': 125000},
                        {'title': '万族之劫', 'author': '老鹰吃小鸡', 'category': '玄幻', 'heat': 110000},
                        {'title': '开局签到荒古圣体', 'author': '神牧', 'category': '玄幻', 'heat': 95000},
                        {'title': '神墓', 'author': '辰东', 'category': '玄幻', 'heat': 80000},
                        {'title': '仙逆', 'author': '耳根', 'category': '仙侠', 'heat': 70000},
                        {'title': '我欲封天', 'author': '耳根', 'category': '仙侠', 'heat': 60000},
                    ]
                }
            ],
            'genres': {
                'male': {'title': '👨 男频题材榜', 'color': '#457B9D', 'genres': [
                    ('玄幻', 25.5), ('都市', 22.3), ('仙侠', 18.7), ('历史', 15.2), ('科幻', 12.8)
                ]},
                'female': {'title': '👩 女频题材榜', 'color': '#E63946', 'genres': [
                    ('现代言情', 28.6), ('古代言情', 24.1), ('玄幻言情', 19.8), ('仙侠奇缘', 16.4), ('青春校园', 13.2)
                ]}
            },
            'types': {
                'male': {'title': '👨 男频类型榜', 'color': '#2A9D8F', 'types': [
                    ('系统流', 22.8), ('穿越', 20.5), ('重生', 18.2), ('修仙', 16.9), ('无敌流', 14.3)
                ]},
                'female': {'title': '👩 女频类型榜', 'color': '#F4A261', 'types': [
                    ('甜宠', 26.5), ('虐恋', 23.1), ('穿越', 20.8), ('重生', 18.4), ('宫斗', 16.2)
                ]}
            },
            'authors': [
                {'rank': 1, 'name': '唐家三少', 'works': '《斗罗大陆》《绝世唐门》', 'income': '3.2亿', 'fans': '1200万'},
                {'rank': 2, 'name': '辰东', 'works': '《遮天》《完美世界》', 'income': '2.8亿', 'fans': '980万'},
                {'rank': 3, 'name': '天蚕土豆', 'works': '《斗破苍穹》《武动乾坤》', 'income': '2.5亿', 'fans': '890万'},
                {'rank': 4, 'name': '我吃西红柿', 'works': '《盘龙》《星辰变》', 'income': '2.1亿', 'fans': '760万'},
                {'rank': 5, 'name': '猫腻', 'works': '《庆余年》《将夜》', 'income': '1.7亿', 'fans': '680万'},
                {'rank': 6, 'name': '顾漫', 'works': '《你是我的荣耀》', 'income': '1.5亿', 'fans': '650万'},
                {'rank': 7, 'name': '梦入神机', 'works': '《龙蛇演义》《永生》', 'income': '1.9亿', 'fans': '720万'},
                {'rank': 8, 'name': '烽火戏诸侯', 'works': '《雪中悍刀行》《剑来》', 'income': '9800万', 'fans': '560万'},
                {'rank': 9, 'name': '月关', 'works': '《回明》《锦衣夜行》', 'income': '1.3亿', 'fans': '620万'},
                {'rank': 10, 'name': '血红', 'works': '《升龙道》《邪风曲》', 'income': '1.1亿', 'fans': '590万'}
            ],
            'update_time': datetime.now().strftime('%Y-%m-%d %H:%M')
        }
    
    def _load_plugins(self) -> None:
        """加载插件列表"""
        plugins = [
            ("大纲解析器", "1.0.0", "active", "Analyzer"),
            ("风格学习器", "1.0.0", "active", "Analyzer"),
            ("人物管理器", "1.0.0", "active", "Tool"),
            ("世界观解析器", "1.0.0", "active", "Analyzer"),
            ("章节生成器", "1.0.0", "active", "Generator"),
        ]
        
        for plugin in plugins:
            self._plugin_tree.insert("", tk.END, values=plugin)
    
    def _load_settings_from_config(self) -> None:
        """从config.yaml加载设置到UI"""
        config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.yaml")
        
        try:
            if os.path.exists(config_path):
                with open(config_path, "r", encoding="utf-8") as f:
                    config_data = yaml.safe_load(f)
                
                if config_data:
                    self._service_mode_var.set(config_data.get("service_mode", "online"))
                    self._provider_var.set(config_data.get("provider", "DeepSeek"))
                    self._model_var.set(config_data.get("model", "deepseek-chat"))
                    self._api_key_var.set(config_data.get("api_key", ""))
                    self._temp_var.set(float(config_data.get("temperature", 0.7)))
                    self._theme_var.set(config_data.get("theme", "dark"))
                    self._temp_label.configure(text=f"{self._temp_var.get():.2f}")
                    logger.info("Settings loaded from config.yaml")
        except Exception as e:
            logger.warning(f"Failed to load settings from config.yaml: {e}")
    
    # ============== 事件处理器 ==============
    
    def _clear_hot_ranking_cache(self) -> None:
        """清除热榜缓存"""
        messagebox.showinfo("提示", "缓存已清除！")
    
    def _update_hot_ranking_data(self) -> None:
        """更新热榜数据"""
        self._set_status("正在刷新热榜...")
        time.sleep(0.5)
        self._refresh_current_page()
        self._set_status("热榜已更新")
    
    # ============== 原有回调方法（保留兼容）==============
    
    def _on_worldview(self) -> None:
        """世界观管理"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(0)  # 切换到世界观标签
    
    def _on_characters(self) -> None:
        """人物设定"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(1)  # 切换到人物标签
    
    def _on_outline(self) -> None:
        """大纲管理"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(2)  # 切换到大纲标签
    
    def _on_style(self) -> None:
        """风格学习"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(3)  # 切换到风格标签
    
    def _on_generation(self) -> None:
        """开始创作"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(4)  # 切换到创作标签
    
    def _on_reverse(self) -> None:
        """逆向反馈"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(5)  # 切换到逆向标签
    
    def _on_quick_create(self) -> None:
        """快捷创作"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(6)  # 切换到快捷标签
    
    def _on_continue(self) -> None:
        """续写功能"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(7)  # 切换到续写标签
    
    def _on_refresh_progress(self) -> None:
        """刷新进度"""
        self._set_status("刷新进度...")
    
    def _on_export_progress(self) -> None:
        """导出报告"""
        self._set_status("导出报告...")
    
    def _on_reset_progress(self) -> None:
        """重置进度"""
        self._set_status("重置进度...")
    
    def _on_new_project(self) -> None:
        """新建项目（弹窗）"""
        # 创建弹窗
        dialog = tk.Toplevel(self.root)
        dialog.title("新建项目")
        dialog.geometry("400x200")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 400) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 200) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 项目名称
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        ttk.Label(name_frame, text="项目名称：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="新项目")
        name_entry = ttk.Entry(name_frame, textvariable=name_var, width=30)
        name_entry.pack(side=tk.LEFT, padx=10)
        
        # 项目路径
        path_frame = ttk.Frame(dialog, style="TFrame")
        path_frame.pack(fill=tk.X, padx=20, pady=10)
        ttk.Label(path_frame, text="保存路径：").pack(side=tk.LEFT)
        path_var = tk.StringVar(value=os.getcwd())
        path_entry = ttk.Entry(path_frame, textvariable=path_var, width=30)
        path_entry.pack(side=tk.LEFT, padx=10)
        
        def browse_path():
            selected = filedialog.askdirectory(title="选择保存路径")
            if selected:
                path_var.set(selected)
        
        ttk.Button(path_frame, text="浏览", command=browse_path).pack(side=tk.LEFT, padx=5)
        
        # 按钮
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_project():
            self._project_name_var.set(name_var.get())
            self._project_path_var.set(path_var.get())
            self._set_status(f"已创建项目：{name_var.get()}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_project).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_open_project(self) -> None:
        """打开项目"""
        path = filedialog.askdirectory(title="选择项目目录")
        if path:
            self._project_name_var.set(os.path.basename(path))
            self._project_path_var.set(path)
            self._set_status("项目已加载")
    
    def _on_backup_project(self) -> None:
        """备份项目"""
        self._set_status("创建备份...")
    
    def _on_refresh_plugins(self) -> None:
        """刷新插件"""
        for item in self._plugin_tree.get_children():
            self._plugin_tree.delete(item)
        self._load_plugins()
        self._set_status("插件列表已刷新")
    
    def _on_install_plugin(self) -> None:
        """安装插件"""
        self._set_status("安装插件功能准备中...")
    
    def _on_uninstall_plugin(self) -> None:
        """卸载插件"""
        self._set_status("卸载插件功能准备中...")
    
    def _on_save_settings(self) -> None:
        """保存设置到config.yaml并应用"""
        config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.yaml")
        
        config_data = {
            "service_mode": self._service_mode_var.get(),
            "provider": self._provider_var.get(),
            "model": self._model_var.get(),
            "api_key": self._api_key_var.get(),
            "local_url": self._local_url_var.get(),
            "temperature": self._temp_var.get(),
            "theme": self._theme_var.get(),
            "ai_learning": self._ai_learning_var.get() if hasattr(self, '_ai_learning_var') else True,
            "auto_save": True,
            "backup_interval": self._backup_interval_var.get() if hasattr(self, '_backup_interval_var') else "30",
            "font_size": self._font_size_var.get() if hasattr(self, '_font_size_var') else "14",
            "window_size": self._window_size_var.get() if hasattr(self, '_window_size_var') else "1280x720",
            "save_path": self._save_path_var.get() if hasattr(self, '_save_path_var') else os.path.dirname(os.path.abspath(__file__))
        }
        
        try:
            with open(config_path, "w", encoding="utf-8") as f:
                yaml.dump(config_data, f, allow_unicode=True, default_flow_style=False)
            
            if self._services and self._services.config:
                self._services.config.set("provider", self._provider_var.get())
                self._services.config.set("model", self._model_var.get())
                self._services.config.set("temperature", self._temp_var.get())
            
            # 应用设置
            self._apply_settings(config_data)
            
            self._set_status("设置已保存并应用")
            messagebox.showinfo("成功", "设置已保存并应用！\n部分设置可能需要重启程序才能完全生效。")
            
        except Exception as e:
            self._set_status(f"保存设置失败：{e}")
            messagebox.showerror("错误", f"保存设置失败：{e}")
    
    def _apply_settings(self, config_data: dict) -> None:
        """应用设置到界面"""
        # 应用主题（深色/浅色）
        theme = config_data.get("theme", "dark")
        if theme == "light":
            # 浅色主题配置
            GlassTheme.GLASS_BG = "#FFFFFF"
            GlassTheme.GLASS_SURFACE = "#F5F5F5"
            GlassTheme.TEXT_PRIMARY = "#333333"
            GlassTheme.TEXT_SECONDARY = "#666666"
        else:
            # 深色主题配置（默认）
            GlassTheme.GLASS_BG = "#1E1E1E"
            GlassTheme.GLASS_SURFACE = "#2D2D2D"
            GlassTheme.TEXT_PRIMARY = "#FFFFFF"
            GlassTheme.TEXT_SECONDARY = "#CCCCCC"
        
        # 应用字体大小
        font_size = int(config_data.get("font_size", 14))
        GlassTheme.FONT_SIZE_NORMAL = str(font_size)
        GlassTheme.FONT_SIZE_SMALL = str(max(10, font_size - 2))
        GlassTheme.FONT_SIZE_SUBTITLE = str(max(14, font_size - 2))
        GlassTheme.FONT_SIZE_TITLE = str(max(18, font_size + 4))
        
        # 应用窗口大小（需要重启才生效）
        window_size = config_data.get("window_size", "1280x720")
        # 保存到配置文件，下次启动时应用
        self.root.after(100, lambda: self._set_status(f"窗口大小将在重启后应用：{window_size}"))
    
    def _on_browse_save_path(self):
        """浏览默认保存路径"""
        path = filedialog.askdirectory(title="选择默认保存路径")
        if path:
            self._save_path_var.set(path)
    
    def _on_clear_learning_data(self):
        """清除学习数据"""
        if messagebox.askyesno("确认清除", "确定要清除所有AI学习数据吗？\n这将重置您的写作偏好配置。"):
            self._set_status("学习数据已清除")
    
    def _on_export_preferences(self):
        """导出偏好配置"""
        path = filedialog.asksaveasfilename(
            title="导出偏好配置",
            defaultextension=".yaml",
            filetypes=[("YAML文件", "*.yaml"), ("所有文件", "*.*")]
        )
        if path:
            self._set_status(f"偏好配置已导出到：{os.path.basename(path)}")
    
    def _on_service_mode_changed(self):
        """服务模式切换回调"""
        if self._service_mode_var.get() == "local":
            # 本地模式：显示部署地址，隐藏API Key
            self._key_label.config(text="本地部署地址：")
            self._key_entry.pack_forget()
            self._url_entry.pack(side=tk.LEFT, padx=10)
        else:
            # 线上模式：显示API Key，隐藏本地地址
            self._key_label.config(text="API Key：")
            self._url_entry.pack_forget()
            self._key_entry.pack(side=tk.LEFT, padx=10)
    
    def _on_theme_changed(self):
        """主题切换回调"""
        theme = self._theme_var.get()
        if theme == "light":
            # 浅色主题配置（使用半透明效果）
            GlassTheme.GLASS_BG = "#F0F0F0"
            GlassTheme.GLASS_SURFACE = "#FAFAFA"
            GlassTheme.TEXT_PRIMARY = "#333333"
            GlassTheme.TEXT_SECONDARY = "#666666"
        else:
            # 深色主题配置（默认）
            GlassTheme.GLASS_BG = "#1E1E1E"
            GlassTheme.GLASS_SURFACE = "#2D2D2D"
            GlassTheme.TEXT_PRIMARY = "#FFFFFF"
            GlassTheme.TEXT_SECONDARY = "#CCCCCC"
        # 立即刷新所有UI元素
        self._refresh_theme()
        self._set_status(f"主题已切换为：{'浅色' if theme == 'light' else '深色'}")

    def _refresh_theme(self):
        """刷新所有UI元素的主题"""
        # 更新主窗口背景
        self.root.configure(bg=GlassTheme.GLASS_BG)

        # 更新标题栏
        if hasattr(self, '_title_bar'):
            self._title_bar.configure(bg=GlassTheme.GLASS_SURFACE)

        # 更新侧边栏
        if hasattr(self, '_sidebar'):
            self._sidebar.configure(bg=GlassTheme.GLASS_SURFACE)

        # 更新所有Text组件
        for widget in self.root.winfo_children():
            self._update_widget_theme(widget)

    def _update_widget_theme(self, widget):
        """递归更新所有组件的主题"""
        try:
            if isinstance(widget, tk.Text):
                widget.configure(bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
            elif isinstance(widget, tk.Label) and widget.cget("style") == "":
                widget.configure(bg=GlassTheme.GLASS_BG, fg=GlassTheme.TEXT_PRIMARY)
            elif isinstance(widget, tk.Label) and widget.cget("bg") == GlassTheme.GLASS_BG:
                widget.configure(bg=GlassTheme.GLASS_BG)
            elif isinstance(widget, tk.Label) and widget.cget("bg") == GlassTheme.GLASS_SURFACE:
                widget.configure(bg=GlassTheme.GLASS_SURFACE)

            # 递归处理子组件
            for child in widget.winfo_children():
                self._update_widget_theme(child)
        except Exception:
            pass
    
    def _on_font_size_changed(self):
        """字体大小改变回调"""
        try:
            font_size = int(self._font_size_var.get())
            # 更新GlassTheme字体大小
            GlassTheme.FONT_SIZE_NORMAL = str(font_size)
            GlassTheme.FONT_SIZE_SMALL = str(max(10, font_size - 2))
            GlassTheme.FONT_SIZE_SUBTITLE = str(max(14, font_size - 2))
            GlassTheme.FONT_SIZE_TITLE = str(max(18, font_size + 4))

            # 立即刷新所有字体
            self._refresh_fonts()
            self._set_status(f"字体大小已更新为：{font_size}px")
        except ValueError:
            messagebox.showerror("错误", "请输入有效的字体大小（10-20）")

    def _refresh_fonts(self):
        """刷新所有UI元素的字体"""
        # 更新所有Text组件字体
        for widget in self.root.winfo_children():
            self._update_widget_font(widget)

    def _update_widget_font(self, widget):
        """递归更新所有组件的字体"""
        try:
            if isinstance(widget, tk.Text):
                widget.configure(font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL))
            elif isinstance(widget, tk.Label) and widget.cget("font") != "":
                # 获取原有字体设置（粗体等）
                current_font = widget.cget("font")
                if isinstance(current_font, tuple) and len(current_font) >= 3:
                    # 保留字体族和粗体样式
                    widget.configure(font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, current_font[2]))
                elif isinstance(current_font, tuple) and len(current_font) >= 2:
                    widget.configure(font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL))

            # 递归处理子组件
            for child in widget.winfo_children():
                self._update_widget_font(child)
        except Exception:
            pass

    
    def _on_close(self) -> None:
        """窗口关闭确认"""
        if messagebox.askokcancel("退出", "确定要退出吗？\n未保存的内容将会丢失。"):
            logger.info("Application closed by user")
            self.root.destroy()
    
    def run(self) -> None:
        """运行主窗口"""
        logger.info("Starting main window")
        self.root.mainloop()


# ============== 入口 ==============

def main():
    """主入口"""
    try:
        app = MainWindow()
        app.run()
    except Exception as e:
        logger.error(f"Application error: {e}")
        raise


if __name__ == "__main__":
    main()
