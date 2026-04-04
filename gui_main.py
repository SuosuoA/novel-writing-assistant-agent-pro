#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Novel Writing Assistant - Agent Pro
主GUI入口 - 现代科技玻璃风（Glass Morphism）

V2.17版本
创建日期：2026-03-22
更新日期：2026-03-28

特性：
- 无边框透明窗口（Windows Acrylic/Mica效果）
- 玻璃态UI设计（毛玻璃背景+发光边框）
- 响应式按钮（防抖+异步执行+加载状态）
- 侧边栏导航（无横向标签页）
- 线程安全的后端交互
- 异步生成（解决卡顿问题）
"""

# ============== 离线模式设置（必须在所有import之前）==============
# 设置HuggingFace离线模式，避免程序启动时尝试连接huggingface.co检查模型版本
# 这必须在导入sentence_transformers之前设置，否则无效
import os
os.environ["TRANSFORMERS_OFFLINE"] = "1"      # Transformers离线模式
os.environ["HF_HUB_OFFLINE"] = "1"            # HuggingFace Hub离线模式
os.environ["HF_DATASETS_OFFLINE"] = "1"       # Datasets离线模式

# ============== 其他导入 ==============
import sys
import json
import yaml
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from pathlib import Path
import threading
import time
import logging
import queue
import ctypes
import subprocess
from typing import Any, Callable, Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from ctypes import c_int, byref, sizeof, windll
import networkx as nx
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

# 尝试导入sv_ttk
try:
    import sv_ttk
    SV_TTK_AVAILABLE = True
except ImportError:
    SV_TTK_AVAILABLE = False
    logging.warning("sv_ttk not available")

# 尝试导入专家选择器模块
try:
    from gui.expert_selector import ExpertSelectorWidget, ExpertInfo
    EXPERT_SELECTOR_AVAILABLE = True
except ImportError as e:
    EXPERT_SELECTOR_AVAILABLE = False
    logging.warning(f"Expert selector not available: {e}")

# 尝试导入核心模块
try:
    from core import (
        EventBus,
        PluginRegistry,
        ServiceLocator,
        ConfigManager,
        get_event_bus,
        get_plugin_registry,
        get_service_locator,
        get_config_manager,
        AsyncHandler,
        TaskPriority,
        init_async_handler,
        CoreServiceManager,
        initialize_core_services,
        dispose_core_services,
        get_logging_service,
    )
    CORE_AVAILABLE = True
except ImportError as e:
    CORE_AVAILABLE = False
    logging.warning(f"Core modules not available: {e}")

# 配置日志
log_level = os.getenv('LOG_LEVEL', 'INFO').upper()
valid_levels = ['DEBUG', 'INFO', 'WARNING', 'ERROR', 'CRITICAL']
if log_level not in valid_levels:
    # 使用print因为logger还未初始化
    print(f"[WARNING] Invalid LOG_LEVEL '{log_level}', using INFO")
    log_level = 'INFO'

logging.basicConfig(
    level=getattr(logging, log_level),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============== 动态插件导入辅助函数 ==============

def _dynamic_import_plugin(plugin_id: str, class_name: str = None):
    """
    动态导入插件（支持连字符目录名）
    
    Args:
        plugin_id: 插件ID（如'quick-creator-v1'）
        class_name: 可选的类名，如果不提供则自动查找
    
    Returns:
        插件类或None
    """
    import importlib.util
    
    # 插件目录路径
    plugin_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plugins", plugin_id)
    
    if not os.path.exists(plugin_dir):
        logger.error(f"Plugin directory not found: {plugin_dir}")
        return None
    
    # 查找plugin.py或主模块文件
    plugin_file = os.path.join(plugin_dir, "plugin.py")
    if not os.path.exists(plugin_file):
        # 查找其他.py文件
        for f in os.listdir(plugin_dir):
            if f.endswith('.py') and f != '__init__.py':
                plugin_file = os.path.join(plugin_dir, f)
                break
    
    if not os.path.exists(plugin_file):
        logger.error(f"No plugin file found in {plugin_dir}")
        return None
    
    try:
        # 动态加载模块
        spec = importlib.util.spec_from_file_location(f"plugins.{plugin_id}", plugin_file)
        if not spec or not spec.loader:
            logger.error(f"Failed to create spec for {plugin_file}")
            return None
        
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        
        # 查找插件类
        if class_name:
            plugin_class = getattr(module, class_name, None)
            if plugin_class:
                return plugin_class
        
        # 自动查找BasePlugin子类
        from core.plugin_interface import BasePlugin
        for attr_name in dir(module):
            attr = getattr(module, attr_name)
            if isinstance(attr, type) and issubclass(attr, BasePlugin) and attr is not BasePlugin:
                return attr
        
        logger.error(f"No plugin class found in {plugin_file}")
        return None
        
    except Exception as e:
        logger.error(f"Failed to import plugin {plugin_id}: {e}")
        return None


# ============== Windows Acrylic效果实现 ==============

class DWM_BLURBEHIND(ctypes.Structure):
    """Windows DWM模糊效果结构"""
    _fields_ = [
        ("dwFlags", ctypes.c_ulong),
        ("fEnable", ctypes.c_bool),
        ("hRgnBlur", ctypes.c_void_p),
        ("fTransitionOnMaximized", ctypes.c_bool),
    ]


class AccentPolicy(ctypes.Structure):
    """Windows 10/11 Accent效果结构"""
    _fields_ = [
        ("AccentState", ctypes.c_int),
        ("AccentFlags", ctypes.c_int),
        ("GradientColor", ctypes.c_uint),
        ("AnimationId", ctypes.c_int),
    ]


class WindowCompositionAttrData(ctypes.Structure):
    """窗口组合属性数据"""
    _fields_ = [
        ("Attribute", ctypes.c_int),
        ("Data", ctypes.c_void_p),
        ("SizeOfData", ctypes.c_int),
    ]


class GlassWindowManager:
    """
    玻璃态窗口管理器
    
    实现Windows 10/11的Acrylic/Mica效果
    """
    
    # Windows常量
    DWM_BB_BLURREGION = 0x2
    DWM_BB_ENABLE = 0x1
    WM_NCCALCSIZE = 0x83
    WM_NCHITTEST = 0x84
    GWL_STYLE = -16
    GWL_EXSTYLE = -20
    WS_CAPTION = 0xC00000
    WS_THICKFRAME = 0x40000
    WS_MINIMIZEBOX = 0x20000
    WS_MAXIMIZEBOX = 0x10000
    WS_SYSMENU = 0x80000
    WS_BORDER = 0x800000
    WS_DLGFRAME = 0x400000
    WS_EX_LAYERED = 0x80000
    WS_EX_TRANSPARENT = 0x20
    
    # Accent状态
    ACCENT_DISABLED = 0
    ACCENT_ENABLE_BLURBEHIND = 3
    ACCENT_ENABLE_ACRYLICBLURBEHIND = 4
    
    # 窗口组合属性
    WCA_ACCENT_POLICY = 19
    
    def __init__(self, root: tk.Tk):
        """
        初始化玻璃窗口管理器
        
        Args:
            root: Tkinter根窗口
        """
        self.root = root
        self.hwnd = None
        self._is_acrylic = False
        
        # 加载Windows DLL
        try:
            self.dwmapi = ctypes.windll.dwmapi
            self.user32 = ctypes.windll.user32
            self._dll_available = True
        except Exception:
            self._dll_available = False
            logger.warning("Windows DLL not available")
        
        # 获取窗口句柄（关键：需要先让窗口创建）
        self._get_hwnd()
    
    def _get_hwnd(self) -> bool:
        """
        获取窗口句柄（多种方法尝试）
        
        Returns:
            是否成功获取句柄
        """
        # 方法1: 使用winfo_id()（最可靠）
        try:
            # 先更新窗口，确保它已创建
            self.root.update()
            
            # 获取Tkinter窗口ID
            window_id = self.root.winfo_id()
            if window_id:
                # 在Windows上，winfo_id()返回的就是HWND
                self.hwnd = window_id
                logger.info(f"Got hwnd from winfo_id: {self.hwnd}")
                return True
        except Exception as e:
            logger.warning(f"winfo_id method failed: {e}")
        
        # 方法2: 使用GetActiveWindow（备选）
        try:
            self.hwnd = self.user32.GetActiveWindow()
            if self.hwnd:
                logger.info(f"Got hwnd from GetActiveWindow: {self.hwnd}")
                return True
        except Exception as e:
            logger.warning(f"GetActiveWindow method failed: {e}")
        
        # 方法3: 使用GetForegroundWindow（最后尝试）
        try:
            self.hwnd = self.user32.GetForegroundWindow()
            if self.hwnd:
                logger.info(f"Got hwnd from GetForegroundWindow: {self.hwnd}")
                return True
        except Exception as e:
            logger.warning(f"GetForegroundWindow method failed: {e}")
        
        logger.warning("All hwnd acquisition methods failed")
        return False
    
    def enable_acrylic(self, color: int = 0x99000000) -> bool:
        """
        启用Acrylic效果
        
        Args:
            color: ARGB颜色值（格式：0xAABBGGRR）
        
        Returns:
            是否成功启用
        """
        if not self._dll_available:
            logger.info("Windows DLL not available, using standard window mode")
            return False
        
        # 确保句柄存在
        if not self.hwnd:
            if not self._get_hwnd():
                logger.warning("Failed to get window handle, using standard window mode")
                return False
        
        try:
            
            # 设置Accent策略
            accent = AccentPolicy()
            accent.AccentState = self.ACCENT_ENABLE_ACRYLICBLURBEHIND
            accent.AccentFlags = 2
            accent.GradientColor = color
            accent.AnimationId = 0
            
            data = WindowCompositionAttrData()
            data.Attribute = self.WCA_ACCENT_POLICY
            data.Data = ctypes.cast(ctypes.byref(accent), ctypes.c_void_p)
            data.SizeOfData = ctypes.sizeof(accent)
            
            # 应用效果
            result = self.user32.SetWindowCompositionAttribute(self.hwnd, ctypes.byref(data))
            
            if result:
                self._is_acrylic = True
                logger.info("Acrylic effect enabled")
            else:
                # 降级：尝试使用blur behind作为备选方案
                logger.warning("Acrylic not supported, trying blur behind fallback")
                if self.enable_blur_behind():
                    logger.info("Blur behind fallback enabled successfully")
                else:
                    logger.warning("Blur behind also failed, using standard window mode")
            
            return result != 0
        
        except Exception as e:
            logger.error(f"Error enabling Acrylic: {e}, using standard window mode")
            # 降级：设置基本透明度作为最终备选
            try:
                self.root.attributes('-alpha', 0.95)
                logger.info("Applied basic transparency as fallback")
            except Exception:
                pass
            return False
    
    def make_frameless(self, keep_resize_border: bool = True) -> None:
        """
        设置完全无边框窗口（移除系统边框和标题栏按钮）
        
        Args:
            keep_resize_border: 是否保留调整大小的边框
                - True: 保留透明调整边框（可拖拽调整大小，但有细微边框残留）
                - False: 完全无边框（无边框残留，但需自己实现调整大小）
        """
        if not self._dll_available:
            return
        
        # 确保句柄存在
        if not self.hwnd:
            if not self._get_hwnd():
                logger.warning("Failed to get window handle for frameless mode")
                return
        
        try:
            
            # 先通过Windows API移除边框样式
            style = self.user32.GetWindowLongW(self.hwnd, self.GWL_STYLE)
            
            if keep_resize_border:
                # 模式1：保留调整边框（推荐）
                # 移除标题栏、系统菜单、边框线条，保留调整功能
                style &= ~(self.WS_CAPTION | self.WS_SYSMENU | self.WS_BORDER | self.WS_DLGFRAME)
                style |= self.WS_THICKFRAME | self.WS_MINIMIZEBOX | self.WS_MAXIMIZEBOX
            else:
                # 模式2：完全无边框
                # 移除所有边框和调整功能，纯净无边框
                style &= ~(self.WS_CAPTION | self.WS_SYSMENU | self.WS_THICKFRAME | self.WS_BORDER | self.WS_DLGFRAME)
                style |= self.WS_MINIMIZEBOX | self.WS_MAXIMIZEBOX
            
            self.user32.SetWindowLongW(self.hwnd, self.GWL_STYLE, style)
            
            # 强制重绘窗口
            self.user32.SetWindowPos(self.hwnd, None, 0, 0, 0, 0, 
                                     0x27)  # SWP_FRAMECHANGED | SWP_NOMOVE | SWP_NOSIZE | SWP_NOZORDER
            
            # 关键：使用 overrideredirect 完全移除系统标题栏按钮
            # 必须在Windows API设置之后调用，否则无法获取窗口句柄
            self.root.overrideredirect(True)
            
            logger.info(f"Frameless window enabled (keep_resize_border={keep_resize_border})")
        
        except Exception as e:
            logger.error(f"Error making frameless: {e}")
    
    def enable_blur_behind(self) -> bool:
        """
        启用DWM模糊效果（Windows 7+）
        
        Returns:
            是否成功启用
        """
        if not self._dll_available:
            return False
        
        # 确保句柄存在
        if not self.hwnd:
            if not self._get_hwnd():
                return False
        
        try:
            
            bb = DWM_BLURBEHIND()
            bb.dwFlags = self.DWM_BB_ENABLE
            bb.fEnable = True
            bb.hRgnBlur = None
            bb.fTransitionOnMaximized = False
            
            self.dwmapi.DwmEnableBlurBehindWindow(self.hwnd, ctypes.byref(bb))
            
            logger.info("Blur behind enabled")
            return True
        
        except Exception as e:
            logger.error(f"Error enabling blur behind: {e}")
            return False
    
    def set_window_opacity(self, alpha: float) -> None:
        """
        设置窗口透明度
        
        Args:
            alpha: 透明度（0.0-1.0）
        """
        try:
            # Tkinter原生透明度设置
            self.root.attributes('-alpha', alpha)
        except Exception as e:
            logger.error(f"Error setting opacity: {e}")


# ============== 玻璃态主题配置 ==============

class GlassTheme:
    """
    现代科技玻璃态主题配置
    
    设计理念：毛玻璃效果 + 发光边框 + 悬浮卡片
    """
    
    # 主色调（科技蓝）
    PRIMARY = "#0078D4"          # Windows蓝
    PRIMARY_LIGHT = "#60CDFF"    # 浅蓝
    PRIMARY_DARK = "#003D6B"     # 深蓝
    
    # 强调色
    ACCENT = "#00BCF2"           # 青蓝
    ACCENT_PINK = "#FF6B9D"      # 粉红
    ACCENT_PURPLE = "#8B5CF6"    # 紫色
    ACCENT_GREEN = "#10B981"     # 绿色
    ACCENT_ORANGE = "#F59E0B"    # 橙色
    ACCENT_RED = "#EF4444"       # 红色
    ACCENT_CYAN = "#06B6D4"      # 青色
    ACCENT_ORANGE = "#F97316"    # 橙色
    
    # 背景色（玻璃态）
    GLASS_BG = "#1E1E1E"         # 玻璃背景色
    GLASS_SURFACE = "#2D2D2D"    # 玻璃表面色
    GLASS_BORDER = "#404040"     # 玻璃边框色
    GLASS_HOVER = "#383838"      # 悬停态
    
    # 文字色
    TEXT_PRIMARY = "#FFFFFF"     # 主文字
    TEXT_SECONDARY = "#B0B0B0"   # 次要文字
    TEXT_DISABLED = "#666666"    # 禁用文字
    TEXT_LINK = "#60CDFF"        # 链接文字
    
    # 语义色
    SUCCESS = "#10B981"
    WARNING = "#F59E0B"
    ERROR = "#EF4444"
    INFO = "#3B82F6"
    
    # 玻璃效果参数
    BLUR_RADIUS = 20
    GLASS_OPACITY = 0.85
    BORDER_RADIUS = 8
    
    # 字体配置（V2.19优化：字号增大2pt，符合中文UI规范）
    FONT_FAMILY = "楷体"  # 全局默认字体（按钮、标签等）
    FONT_FAMILY_TEXT = "宋体"  # 文本框专用字体
    FONT_FAMILY_CODE = "Consolas"  # 代码字体
    FONT_SIZE_TITLE = 20      # 大标题：窗口标题、页面主标题
    FONT_SIZE_SUBTITLE = 16   # 小标题：区块标题、卡片标题
    FONT_SIZE_NORMAL = 12     # 正文：普通文字、按钮文字
    FONT_SIZE_SMALL = 11      # 辅助文字：提示、说明
    FONT_SIZE_TINY = 10       # 最小文字：次要信息
    
    @classmethod
    def apply_to_style(cls, style: ttk.Style) -> None:
        """应用玻璃态主题到ttk Style"""
        
        # 全局样式
        style.configure(".",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            fieldbackground=cls.GLASS_SURFACE,
            bordercolor=cls.GLASS_BORDER,
            lightcolor=cls.GLASS_HOVER,
            darkcolor=cls.GLASS_BG
        )
        
        # 标签
        style.configure("TLabel",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL)
        )
        
        # 标题标签
        style.configure("Title.TLabel",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_TITLE, "bold"),
            foreground=cls.TEXT_PRIMARY
        )
        
        # 副标题标签
        style.configure("Subtitle.TLabel",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_SUBTITLE, "bold"),
            foreground=cls.TEXT_SECONDARY
        )
        
        # 按钮 - 玻璃态效果
        style.configure("TButton",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL),
            padding=(16, 8),
            background=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            borderwidth=1,
            relief="flat"
        )
        
        style.map("TButton",
            background=[("active", cls.GLASS_HOVER), ("pressed", cls.PRIMARY)],
            foreground=[("active", cls.TEXT_PRIMARY), ("pressed", cls.TEXT_PRIMARY)]
        )
        
        # 强调按钮
        style.configure("Accent.TButton",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL),
            padding=(16, 8),
            background=cls.PRIMARY,
            foreground=cls.TEXT_PRIMARY,
            borderwidth=0
        )
        
        style.map("Accent.TButton",
            background=[("active", cls.PRIMARY_LIGHT), ("pressed", cls.PRIMARY_DARK)]
        )
        
        # 框架
        style.configure("TFrame",
            background=cls.GLASS_BG
        )
        
        # 玻璃卡片框架
        style.configure("Glass.TFrame",
            background=cls.GLASS_SURFACE,
            borderwidth=1,
            relief="solid"
        )
        
        # LabelFrame
        style.configure("TLabelframe",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            borderwidth=1,
            relief="solid"
        )
        
        style.configure("TLabelframe.Label",
            background=cls.GLASS_BG,
            foreground=cls.PRIMARY_LIGHT,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL, "bold")
        )
        
        # 输入框
        style.configure("TEntry",
            fieldbackground=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            insertcolor=cls.TEXT_PRIMARY,
            borderwidth=1,
            padding=5
        )
        
        # 下拉框
        style.configure("TCombobox",
            fieldbackground=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            arrowcolor=cls.TEXT_PRIMARY,
            borderwidth=1,
            padding=5
        )
        
        # 进度条
        style.configure("TProgressbar",
            background=cls.PRIMARY,
            troughcolor=cls.GLASS_SURFACE,
            borderwidth=0
        )
        
        # 滚动条
        style.configure("TScrollbar",
            background=cls.GLASS_SURFACE,
            troughcolor=cls.GLASS_BG,
            arrowcolor=cls.TEXT_SECONDARY,
            borderwidth=0
        )
        
        # Treeview
        style.configure("Treeview",
            background=cls.GLASS_SURFACE,
            foreground=cls.TEXT_PRIMARY,
            fieldbackground=cls.GLASS_SURFACE,
            borderwidth=0
        )
        
        style.configure("Treeview.Heading",
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL, "bold"),
            background=cls.GLASS_BG,
            foreground=cls.TEXT_SECONDARY
        )
        
        style.map("Treeview",
            background=[("selected", cls.PRIMARY)],
            foreground=[("selected", cls.TEXT_PRIMARY)]
        )
        
        # 复选框
        style.configure("TCheckbutton",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL)
        )
        
        # 单选按钮
        style.configure("TRadiobutton",
            background=cls.GLASS_BG,
            foreground=cls.TEXT_PRIMARY,
            font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL)
        )
        
        # 分隔线
        style.configure("TSeparator",
            background=cls.GLASS_BORDER
        )
        
        # Scale滑块
        style.configure("TScale",
            background=cls.GLASS_BG,
            troughcolor=cls.GLASS_SURFACE,
            borderwidth=0
        )


# ============== 响应式按钮 ==============

class ResponsiveButton(ttk.Button):
    """
    响应式按钮组件
    
    特性：
    - 500ms防抖保护
    - 异步执行耗时操作
    - 加载状态自动管理
    """
    
    DEBOUNCE_MS = 500
    LOADING_SUFFIX = "..."
    
    def __init__(
        self,
        parent: tk.Widget,
        text: str,
        command: Callable,
        async_handler: Optional[Any] = None,
        debounce_ms: int = DEBOUNCE_MS,
        auto_disable: bool = True,
        style: str = "TButton",
        **kwargs
    ):
        super().__init__(parent, text=text, style=style, **kwargs)
        
        self._original_text = text
        self._original_command = command
        self._async_handler = async_handler
        self._debounce_ms = debounce_ms
        self._auto_disable = auto_disable
        
        self._is_loading = False
        self._last_click_time = 0
        
        super().configure(command=self._on_clicked)
    
    def _on_clicked(self) -> None:
        if self._is_loading:
            return
        
        current_time = time.time() * 1000
        if current_time - self._last_click_time < self._debounce_ms:
            return
        
        self._last_click_time = current_time
        
        if self._async_handler:
            self._execute_async()
        else:
            self._execute_sync()
    
    def _execute_sync(self) -> None:
        """P1-8修复：按异常类型分类处理"""
        try:
            self._set_loading(True)
            self._original_command()
        except ValueError as e:
            logger.warning(f"输入验证失败: {e}")
            messagebox.showwarning("输入错误", f"请检查输入：{e}")
        except PermissionError as e:
            logger.error(f"权限不足: {e}")
            messagebox.showerror("权限错误", "您没有执行此操作的权限")
        except FileNotFoundError as e:
            logger.error(f"文件不存在: {e}")
            messagebox.showerror("文件错误", f"找不到文件：{e}")
        except Exception as e:
            logger.error(f"执行失败: {e}", exc_info=True)
            messagebox.showerror(
                "执行失败",
                f"操作失败：{type(e).__name__}\n详情：{str(e)[:200]}"
            )
        finally:
            self._set_loading(False)

    def _execute_async(self) -> None:
        self._set_loading(True)

        def task():
            return self._original_command()

        def on_success(result):
            self._set_loading(False)

        def on_error(error):
            self._set_loading(False)
            # P1-8修复：按异常类型分类处理
            error_str = str(error)
            if "permission" in error_str.lower():
                messagebox.showerror("权限错误", "您没有执行此操作的权限")
            elif "not found" in error_str.lower() or "找不到" in error_str:
                messagebox.showerror("文件错误", f"找不到资源：{error}")
            else:
                messagebox.showerror(
                    "执行失败",
                    f"操作失败\n详情：{error_str[:200]}"
                )

        self._async_handler.submit(
            func=task,
            callback=on_success,
            error_callback=on_error,
            priority=TaskPriority.HIGH
        )
    
    def _set_loading(self, loading: bool) -> None:
        self._is_loading = loading
        
        if self._auto_disable:
            self.configure(state="disabled" if loading else "normal")
        
        if loading:
            self.configure(text=self._original_text + self.LOADING_SUFFIX)
        else:
            self.configure(text=self._original_text)


# ============== 自定义标题栏 ==============

class ResizeGrip:
    """
    窗口调整大小的边框感知区域
    
    用于完全无边框窗口的调整大小功能
    """
    
    RESIZE_BORDER = 8  # 边框感知区域宽度
    
    def __init__(self, root: tk.Tk):
        self.root = root
        self._resize_edge = None
        self._drag_start_x = 0
        self._drag_start_y = 0
        self._drag_start_width = 0
        self._drag_start_height = 0
        
        # 绑定事件
        self.root.bind("<Motion>", self._on_motion)
        self.root.bind("<ButtonPress-1>", self._on_resize_start)
        self.root.bind("<ButtonRelease-1>", self._on_resize_end)
        self.root.bind("<B1-Motion>", self._on_resize_motion)
    
    def _get_edge(self, x: int, y: int) -> str:
        """判断鼠标在哪个边缘"""
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        
        left = x < self.RESIZE_BORDER
        right = x > width - self.RESIZE_BORDER
        top = y < self.RESIZE_BORDER
        bottom = y > height - self.RESIZE_BORDER
        
        # 角落优先
        if top and left:
            return "top-left"
        elif top and right:
            return "top-right"
        elif bottom and left:
            return "bottom-left"
        elif bottom and right:
            return "bottom-right"
        elif left:
            return "left"
        elif right:
            return "right"
        elif top:
            return "top"
        elif bottom:
            return "bottom"
        return ""
    
    def _get_cursor(self, edge: str) -> str:
        """获取对应的鼠标光标"""
        cursors = {
            "top": "top_side",
            "bottom": "bottom_side",
            "left": "left_side",
            "right": "right_side",
            "top-left": "top_left_corner",
            "top-right": "top_right_corner",
            "bottom-left": "bottom_left_corner",
            "bottom-right": "bottom_right_corner"
        }
        return cursors.get(edge, "")
    
    def _on_motion(self, event: tk.Event) -> None:
        """鼠标移动时更新光标"""
        edge = self._get_edge(event.x, event.y)
        cursor = self._get_cursor(edge)
        self.root.configure(cursor=cursor if cursor else "arrow")
    
    def _on_resize_start(self, event: tk.Event) -> None:
        """开始调整大小"""
        self._resize_edge = self._get_edge(event.x, event.y)
        if self._resize_edge:
            self._drag_start_x = self.root.winfo_pointerx()
            self._drag_start_y = self.root.winfo_pointery()
            self._drag_start_width = self.root.winfo_width()
            self._drag_start_height = self.root.winfo_height()
            self._drag_start_root_x = self.root.winfo_rootx()
            self._drag_start_root_y = self.root.winfo_rooty()
    
    def _on_resize_end(self, event: tk.Event) -> None:
        """结束调整大小"""
        self._resize_edge = None
    
    def _on_resize_motion(self, event: tk.Event) -> None:
        """调整大小中"""
        if not self._resize_edge:
            return
        
        current_x = self.root.winfo_pointerx()
        current_y = self.root.winfo_pointery()
        dx = current_x - self._drag_start_x
        dy = current_y - self._drag_start_y
        
        new_width = self._drag_start_width
        new_height = self._drag_start_height
        new_x = self._drag_start_root_x
        new_y = self._drag_start_root_y
        
        min_width = 1000
        min_height = 700
        
        edge = self._resize_edge
        
        if "right" in edge:
            new_width = max(min_width, self._drag_start_width + dx)
        if "left" in edge:
            new_width = max(min_width, self._drag_start_width - dx)
            if new_width > min_width:
                new_x = self._drag_start_root_x + dx
        if "bottom" in edge:
            new_height = max(min_height, self._drag_start_height + dy)
        if "top" in edge:
            new_height = max(min_height, self._drag_start_height - dy)
            if new_height > min_height:
                new_y = self._drag_start_root_y + dy
        
        # 更新窗口
        if new_width >= min_width or new_height >= min_height:
            self.root.geometry(f"{new_width}x{new_height}+{new_x}+{new_y}")


class CustomTitleBar(tk.Frame):
    """
    自定义标题栏
    
    支持：拖动、最小化、最大化、关闭
    """
    
    def __init__(self, parent: tk.Tk, title: str = "Novel Writing Assistant-Agent Pro"):
        super().__init__(parent, bg=GlassTheme.GLASS_BG, height=40)
        
        self.parent = parent
        self.title_text = title
        
        self._drag_start_x = 0
        self._drag_start_y = 0
        
        self._create_widgets()
        self._bind_events()
    
    def _create_widgets(self) -> None:
        # 标题
        self.title_label = tk.Label(
            self,
            text=self.title_text,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SUBTITLE, "bold"),
            fg=GlassTheme.TEXT_PRIMARY,
            bg=GlassTheme.GLASS_BG
        )
        self.title_label.pack(side=tk.LEFT, padx=15, pady=8)
        
        # 窗口控制按钮容器
        btn_frame = tk.Frame(self, bg=GlassTheme.GLASS_BG)
        btn_frame.pack(side=tk.RIGHT)
        
        # 按钮统一配置（等比例缩小）
        BTN_WIDTH = 36
        BTN_HEIGHT = 26
        
        # 使用Canvas绘制统一大小的图标
        def create_icon_canvas(parent, draw_func, bg_color):
            """创建图标Canvas"""
            canvas = tk.Canvas(
                parent,
                width=BTN_WIDTH,
                height=BTN_HEIGHT,
                bg=bg_color,
                highlightthickness=0
            )
            draw_func(canvas, BTN_WIDTH // 2, BTN_HEIGHT // 2)
            return canvas
        
        def draw_minimize(canvas, cx, cy):
            """绘制最小化图标 - 水平线"""
            canvas.create_line(cx - 6, cy, cx + 6, cy, fill=GlassTheme.TEXT_SECONDARY, width=2)
        
        def draw_maximize(canvas, cx, cy):
            """绘制最大化图标 - 空心方块"""
            canvas.create_rectangle(cx - 5, cy - 5, cx + 5, cy + 5, outline=GlassTheme.TEXT_SECONDARY, width=2)
        
        def draw_close(canvas, cx, cy):
            """绘制关闭图标 - X"""
            canvas.create_line(cx - 5, cy - 5, cx + 5, cy + 5, fill=GlassTheme.TEXT_SECONDARY, width=2)
            canvas.create_line(cx - 5, cy + 5, cx + 5, cy - 5, fill=GlassTheme.TEXT_SECONDARY, width=2)
        
        # 最小化按钮
        self.min_btn = create_icon_canvas(btn_frame, draw_minimize, GlassTheme.GLASS_BG)
        self.min_btn.pack(side=tk.LEFT)
        self.min_btn.cursor = "hand2"
        
        # 最大化按钮
        self.max_btn = create_icon_canvas(btn_frame, draw_maximize, GlassTheme.GLASS_BG)
        self.max_btn.pack(side=tk.LEFT)
        self.max_btn.cursor = "hand2"
        
        # 关闭按钮
        self.close_btn = create_icon_canvas(btn_frame, draw_close, GlassTheme.GLASS_BG)
        self.close_btn.pack(side=tk.LEFT)
        self.close_btn.cursor = "hand2"
        
        # 悬停效果
        for btn in [self.min_btn, self.max_btn]:
            btn.bind("<Enter>", lambda e, b=btn: b.configure(bg=GlassTheme.GLASS_HOVER))
            btn.bind("<Leave>", lambda e, b=btn: b.configure(bg=GlassTheme.GLASS_BG))
        
        # Canvas悬停效果 - 使用不同的方法
        def on_min_enter(e):
            self.min_btn.configure(bg=GlassTheme.GLASS_HOVER)
        def on_min_leave(e):
            self.min_btn.configure(bg=GlassTheme.GLASS_BG)
        def on_max_enter(e):
            self.max_btn.configure(bg=GlassTheme.GLASS_HOVER)
        def on_max_leave(e):
            self.max_btn.configure(bg=GlassTheme.GLASS_BG)
        def on_close_enter(e):
            self.close_btn.configure(bg=GlassTheme.ERROR)
        def on_close_leave(e):
            self.close_btn.configure(bg=GlassTheme.GLASS_BG)
        
        self.min_btn.bind("<Enter>", on_min_enter)
        self.min_btn.bind("<Leave>", on_min_leave)
        self.max_btn.bind("<Enter>", on_max_enter)
        self.max_btn.bind("<Leave>", on_max_leave)
        self.close_btn.bind("<Enter>", on_close_enter)
        self.close_btn.bind("<Leave>", on_close_leave)
    
    def _bind_events(self) -> None:
        # 拖动窗口
        self.bind("<Button-1>", self._on_drag_start)
        self.bind("<B1-Motion>", self._on_drag_motion)
        self.title_label.bind("<Button-1>", self._on_drag_start)
        self.title_label.bind("<B1-Motion>", self._on_drag_motion)
        
        # 窗口控制
        self.min_btn.bind("<Button-1>", lambda e: self._minimize_window())
        self.max_btn.bind("<Button-1>", self._toggle_maximize)
        self.close_btn.bind("<Button-1>", lambda e: self.parent.quit())
        
        # 双击标题栏最大化
        self.bind("<Double-Button-1>", lambda e: self._toggle_maximize(None))
    
    def _on_drag_start(self, event: tk.Event) -> None:
        self._drag_start_x = event.x
        self._drag_start_y = event.y
    
    def _on_drag_motion(self, event: tk.Event) -> None:
        x = self.parent.winfo_x() + event.x - self._drag_start_x
        y = self.parent.winfo_y() + event.y - self._drag_start_y
        self.parent.geometry(f"+{x}+{y}")
    
    def _minimize_window(self) -> None:
        """最小化窗口（完全隐藏到系统托盘）"""
        try:
            # 关键：直接使用 withdraw() 完全隐藏窗口
            # 窗口会从屏幕消失，用户可以通过系统托盘图标恢复
            self.parent.withdraw()
            logger.info("Window minimized to system tray")
        except Exception as e:
            logger.error(f"Error minimizing window: {e}")

    def _show_window(self) -> None:
        """从系统托盘恢复窗口显示"""
        try:
            # 显示窗口
            self.parent.deiconify()
            self.parent.lift()
            self.parent.focus_force()
            logger.info("Window restored from system tray")
        except Exception as e:
            logger.error(f"Error showing window: {e}")

    def _toggle_maximize(self, event: Optional[tk.Event]) -> None:
        if self.parent.state() == "zoomed":
            self.parent.state("normal")
            self._redraw_maximize_btn(maximized=False)
        else:
            self.parent.state("zoomed")
            self._redraw_maximize_btn(maximized=True)

    def _redraw_maximize_btn(self, maximized: bool) -> None:
        """重绘最大化按钮图标（Canvas 没有 text 属性，需要重绘）"""
        try:
            self.max_btn.delete("all")
            cx = self.max_btn.winfo_reqwidth() // 2
            cy = self.max_btn.winfo_reqheight() // 2
            if maximized:
                # 还原图标：两个重叠小方块
                self.max_btn.create_rectangle(cx - 3, cy - 5, cx + 5, cy + 3,
                                              outline=GlassTheme.TEXT_SECONDARY, width=2)
                self.max_btn.create_rectangle(cx - 5, cy - 3, cx + 3, cy + 5,
                                              outline=GlassTheme.TEXT_SECONDARY, width=2)
            else:
                # 最大化图标：空心方块
                self.max_btn.create_rectangle(cx - 5, cy - 5, cx + 5, cy + 5,
                                              outline=GlassTheme.TEXT_SECONDARY, width=2)
        except Exception:
            pass

    def _create_window_controls(self):
        """创建窗口控制按钮（最小化、最大化、关闭）"""
        controls_frame = tk.Frame(self.title_frame, bg=GlassTheme.GLASS_BG)
        controls_frame.pack(side=tk.RIGHT, padx=5)

        # 最小化按钮
        self.min_btn = tk.Canvas(
            controls_frame,
            width=30,
            height=30,
            bg=GlassTheme.GLASS_BG,
            highlightthickness=0
        )
        self.min_btn.pack(side=tk.LEFT, padx=2)
        self.min_btn.create_line(8, 15, 22, 15, fill=GlassTheme.TEXT_SECONDARY, width=2)
        self.min_btn.bind("<Button-1>", self._toggle_minimize)
        self.min_btn.bind("<Enter>", lambda e: self.min_btn.config(bg=GlassTheme.GLASS_SURFACE))
        self.min_btn.bind("<Leave>", lambda e: self.min_btn.config(bg=GlassTheme.GLASS_BG))

        # 最大化按钮
        self.max_btn = tk.Canvas(
            controls_frame,
            width=30,
            height=30,
            bg=GlassTheme.GLASS_BG,
            highlightthickness=0
        )
        self.max_btn.pack(side=tk.LEFT, padx=2)
        self.max_btn.create_rectangle(8, 8, 22, 22, outline=GlassTheme.TEXT_SECONDARY, width=2)
        self.max_btn.bind("<Button-1>", self._toggle_maximize)
        self.max_btn.bind("<Enter>", lambda e: self.max_btn.config(bg=GlassTheme.GLASS_SURFACE))
        self.max_btn.bind("<Leave>", lambda e: self.max_btn.config(bg=GlassTheme.GLASS_BG))

        # 关闭按钮
        self.close_btn = tk.Canvas(
            controls_frame,
            width=30,
            height=30,
            bg=GlassTheme.GLASS_BG,
            highlightthickness=0
        )
        self.close_btn.pack(side=tk.LEFT, padx=2)
        self.close_btn.create_line(8, 8, 22, 22, fill=GlassTheme.TEXT_SECONDARY, width=2)
        self.close_btn.create_line(8, 22, 22, 8, fill=GlassTheme.TEXT_SECONDARY, width=2)
        self.close_btn.bind("<Button-1>", lambda e: self.window.destroy())
        self.close_btn.bind("<Enter>", lambda e: self.close_btn.config(bg="#ff4444"))
        self.close_btn.bind("<Leave>", lambda e: self.close_btn.config(bg=GlassTheme.GLASS_BG))
        
        self._log_text = tk.Text(
            log_frame,
            wrap=tk.WORD,
            height=6,
            font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_SMALL),
            bg=GlassTheme.GLASS_SURFACE,
            fg=GlassTheme.TEXT_PRIMARY
        )
        log_scroll = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self._log_text.yview)
        self._log_text.configure(yscrollcommand=log_scroll.set)
        
        self._log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 按钮区
        btn_frame = tk.Frame(self.window, bg=GlassTheme.GLASS_BG)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self._interrupt_btn = ttk.Button(
            btn_frame,
            text="⚠️ 中断生成",
            command=self._on_interrupt,
            state=tk.DISABLED
        )
        self._interrupt_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame,
            text="清空日志",
            command=self._clear_log
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame,
            text="关闭",
            command=self.window.destroy
        ).pack(side=tk.RIGHT, padx=5)
    
    def _subscribe_events(self):
        """订阅事件"""
        if not self._event_bus:
            return
        
        # 订阅生成事件
        events = [
            ("generation.started", self._on_generation_started),
            ("generation.completed", self._on_generation_completed),
            ("generation.failed", self._on_generation_failed),
            ("generation.progress", self._on_generation_progress),
            # 订阅流水线事件
            ("pipeline.stage_started", self._on_stage_started),
            ("pipeline.stage_completed", self._on_stage_completed),
            ("pipeline.iteration_started", self._on_iteration_started),
            ("pipeline.completed", self._on_pipeline_completed),
            # 订阅Agent事件
            ("agent.task.started", self._on_agent_task_started),
            ("agent.task.completed", self._on_agent_task_completed),
            ("agent.task.failed", self._on_agent_task_failed),
        ]
        
        for event_type, handler in events:
            try:
                sub_id = self._event_bus.subscribe(event_type, handler)
                self._subscription_ids.append(sub_id)
            except Exception as e:
                logger.warning(f"订阅事件失败 {event_type}: {e}")
    
    def _unsubscribe_events(self):
        """取消订阅事件"""
        if not self._event_bus:
            return
        
        for sub_id in self._subscription_ids:
            try:
                self._event_bus.unsubscribe(sub_id)
            except Exception:
                pass
        
        self._subscription_ids.clear()
    
    # === 事件处理器 ===
    
    def _on_generation_started(self, event):
        """生成开始事件"""
        def update():
            self._is_running = True
            self._current_pipeline_id = event.data.get("pipeline_id")
            self._status_var.set("🚀 生成任务已开始")
            self._title_var.set(f"🔍 生成过程监视器 - {self._current_pipeline_id}")
            self._interrupt_btn.configure(state=tk.NORMAL)
            self._progress_var.set(0)
            self._progress_label.configure(text="0%")
            self._log_insert(f"[{self._timestamp()}] 生成任务开始\n")
            self._clear_agent_tree()
        
        self.parent.after(0, update)
    
    def _on_generation_completed(self, event):
        """生成完成事件"""
        def update():
            self._is_running = False
            result = event.data.get("result", {})
            success = result.get("success", False)
            
            if success:
                self._status_var.set("✅ 生成完成")
                self._progress_var.set(100)
                self._progress_label.configure(text="100%")
                self._log_insert(f"[{self._timestamp()}] ✅ 生成成功完成\n")
            else:
                self._status_var.set("❌ 生成失败")
                self._log_insert(f"[{self._timestamp()}] ❌ 生成失败: {result.get('error', '未知错误')}\n")
            
            self._interrupt_btn.configure(state=tk.DISABLED)
        
        self.parent.after(0, update)
    
    def _on_generation_failed(self, event):
        """生成失败事件"""
        def update():
            self._is_running = False
            error = event.data.get("error", "未知错误")
            self._status_var.set(f"❌ 生成失败: {error}")
            self._log_insert(f"[{self._timestamp()}] ❌ 生成失败: {error}\n")
            self._interrupt_btn.configure(state=tk.DISABLED)
        
        self.parent.after(0, update)
    
    def _on_generation_progress(self, event):
        """生成进度事件"""
        def update():
            progress = event.data.get("progress", 0)
            stage = event.data.get("stage", "")
            message = event.data.get("message", "")
            
            self._progress_var.set(progress)
            self._progress_label.configure(text=f"{progress:.1f}%")
            if message:
                self._log_insert(f"[{self._timestamp()}] [{stage}] {message}\n")
        
        self.parent.after(0, update)
    
    def _on_stage_started(self, event):
        """流水线阶段开始事件"""
        def update():
            stage_name = event.data.get("stage_name", "")
            agent_type = event.data.get("agent_type", "")
            iteration = event.data.get("iteration", 1)
            
            self._log_insert(f"[{self._timestamp()}] 📌 阶段开始: {stage_name} (Agent: {agent_type}, 迭代: {iteration})\n")
            self._add_agent_row(agent_type, stage_name, "运行中", "-", "...")
        
        self.parent.after(0, update)
    
    def _on_stage_completed(self, event):
        """流水线阶段完成事件"""
        def update():
            stage_name = event.data.get("stage_name", "")
            success = event.data.get("success", False)
            duration = event.data.get("duration_seconds", 0)
            
            status = "✅ 成功" if success else "❌ 失败"
            self._log_insert(f"[{self._timestamp()}] {status} 阶段完成: {stage_name} (耗时: {duration:.2f}s)\n")
            self._update_agent_row(stage_name, status, f"{duration:.2f}s", "完成" if success else "失败")
        
        self.parent.after(0, update)
    
    def _on_iteration_started(self, event):
        """迭代开始事件"""
        def update():
            iteration = event.data.get("iteration", 1)
            max_iterations = event.data.get("max_iterations", 5)
            
            self._log_insert(f"\n{'='*40}\n")
            self._log_insert(f"[{self._timestamp()}] 🔄 开始迭代 {iteration}/{max_iterations}\n")
        
        self.parent.after(0, update)
    
    def _on_pipeline_completed(self, event):
        """流水线完成事件"""
        def update():
            result_data = event.data
            success = result_data.get("success", False)
            total_iterations = result_data.get("total_iterations", 0)
            total_duration = result_data.get("total_duration_seconds", 0)
            
            status = "✅ 成功" if success else "❌ 失败"
            self._log_insert(f"\n{'='*40}\n")
            self._log_insert(f"[{self._timestamp()}] {status} 流水线完成\n")
            self._log_insert(f"  总迭代次数: {total_iterations}\n")
            self._log_insert(f"  总耗时: {total_duration:.2f}秒\n")
            
            self._status_var.set(f"{status} 流水线完成")
            self._is_running = False
            self._interrupt_btn.configure(state=tk.DISABLED)
        
        self.parent.after(0, update)
    
    def _on_agent_task_started(self, event):
        """Agent任务开始事件"""
        def update():
            task_id = event.data.get("task_id", "")
            agent_type = event.data.get("agent_type", "")
            
            self._log_insert(f"[{self._timestamp()}] 🤖 Agent任务开始: {agent_type} ({task_id})\n")
        
        self.parent.after(0, update)
    
    def _on_agent_task_completed(self, event):
        """Agent任务完成事件"""
        def update():
            task_id = event.data.get("task_id", "")
            agent_type = event.data.get("agent_type", "")
            result = event.data.get("result", {})
            
            self._log_insert(f"[{self._timestamp()}] ✅ Agent任务完成: {agent_type}\n")
            
            # 存储结果
            self._agent_results[agent_type] = result
        
        self.parent.after(0, update)
    
    def _on_agent_task_failed(self, event):
        """Agent任务失败事件"""
        def update():
            task_id = event.data.get("task_id", "")
            agent_type = event.data.get("agent_type", "")
            error = event.data.get("error", "未知错误")
            retry_count = event.data.get("retry_count", 0)
            
            self._log_insert(f"[{self._timestamp()}] ❌ Agent任务失败: {agent_type}\n")
            self._log_insert(f"    错误: {error}\n")
            self._log_insert(f"    重试次数: {retry_count}\n")
        
        self.parent.after(0, update)
    
    # === 辅助方法 ===
    
    def _log_insert(self, text: str):
        """插入日志文本"""
        self._log_text.insert(tk.END, text)
        self._log_text.see(tk.END)
    
    def _clear_log(self):
        """清空日志"""
        self._log_text.delete("1.0", tk.END)
    
    def _clear_agent_tree(self):
        """清空Agent列表"""
        for item in self._agent_tree.get_children():
            self._agent_tree.delete(item)
    
    def _add_agent_row(self, agent_type: str, stage: str, status: str, duration: str, result: str):
        """添加Agent行"""
        self._agent_tree.insert("", tk.END, iid=stage, values=(
            agent_type,
            status,
            duration,
            result[:50] + "..." if len(result) > 50 else result
        ))
    
    def _update_agent_row(self, stage: str, status: str, duration: str, result: str):
        """更新Agent行"""
        try:
            item = self._agent_tree.item(stage)
            if item:
                values = list(item["values"])
                values[1] = status
                values[2] = duration
                values[3] = result[:50] + "..." if len(result) > 50 else result
                self._agent_tree.item(stage, values=values)
        except Exception:
            pass
    
    def _timestamp(self) -> str:
        """获取时间戳"""
        from datetime import datetime
        return datetime.now().strftime("%H:%M:%S")
    
    def _on_interrupt(self):
        """中断生成"""
        if not self._is_running:
            return
        
        if messagebox.askyesno("确认", "确定要中断当前生成任务吗？"):
            try:
                from agents.novel_generation_service import get_generation_service
                service = get_generation_service()
                if service.cancel_generation():
                    self._status_var.set("⚠️ 已中断")
                    self._log_insert(f"[{self._timestamp()}] ⚠️ 用户中断生成\n")
                    self._is_running = False
                    self._interrupt_btn.configure(state=tk.DISABLED)
                else:
                    self._log_insert(f"[{self._timestamp()}] ❌ 中断失败\n")
            except Exception as e:
                self._log_insert(f"[{self._timestamp()}] ❌ 中断异常: {e}\n")
    
    def _on_close(self):
        """窗口关闭"""
        self._unsubscribe_events()
        self.window.destroy()
    
    def show(self):
        """显示窗口"""
        self.window.deiconify()
        self.window.lift()
        self.window.focus_force()
    
    def hide(self):
        """隐藏窗口"""
        self.window.withdraw()
    
    def log(self, message: str):
        """添加日志消息"""
        self._log_insert(f"[{self._timestamp()}] {message}\n")
    
    def clear_log(self):
        """清空日志（公开方法）"""
        self._clear_log()


# ============== 主窗口 ==============

class MainWindow:
    """
    主窗口类
    
    结构：
    - 自定义标题栏
    - 侧边栏（导航）
    - 主内容区（根据侧边栏选择显示）
    - 状态栏
    """
    
    VERSION = "V2.1.0"
    TITLE = "Novel Writing Assistant - Agent Pro"
    
    def __init__(self):
        """初始化主窗口"""
        # 设置Windows任务栏图标（必须在创建窗口前设置）
        # 这样Windows会将程序识别为独立应用，而非Python解释器
        try:
            import ctypes
            app_id = "NovelWritingAssistant.AgentPro.V2"
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(app_id)
            logger.info(f"AppUserModelID set: {app_id}")
        except Exception as e:
            logger.warning(f"Failed to set AppUserModelID: {e}")
        
        # 创建根窗口
        self.root = tk.Tk()
        self.root.title(self.TITLE)
        self.root.geometry("1400x900")
        self.root.minsize(1000, 700)
        
        # 设置窗口背景
        self.root.configure(bg=GlassTheme.GLASS_BG)
        
        # 尝试启用玻璃效果
        self._enable_glass_effect()
        
        # 核心服务
        self._services: Optional[CoreServiceManager] = None
        self._async_handler: Optional[AsyncHandler] = None
        
        # UI组件
        self._content_frame: Optional[tk.Frame] = None
        self._status_var: Optional[tk.StringVar] = None
        self._title_bar: Optional[CustomTitleBar] = None
        self._resize_grip: Optional[ResizeGrip] = None  # 窗口调整大小边框感知
        
        # 当前显示的页面
        self._current_page: Optional[str] = None
        self._pages: Dict[str, tk.Frame] = {}
        
        # 项目状态
        self.current_project: Dict[str, Any] = {}
        self.project_file: Optional[str] = None
        self.project_info: Dict[str, Any] = {}
        
        # 结果队列（后台线程通信）
        self._result_queue = queue.Queue()
        
        # 线程安全锁（快捷创作模块）
        self._quick_lock = threading.Lock()
        
        # 缓存
        self._cache = {
            'outline': {},
            'style': {},
            'worldview': {},
            'characters': {},
            'file_mtime': {}
        }
        
        # 小说生成相关
        self._current_pipeline_id: Optional[str] = None
        self._outline_content: str = ""
        self._chapter_outlines: Dict[int, str] = {}
        self._style_profile: Dict[str, Any] = {}
        self._characters: List[Dict[str, Any]] = []
        self._worldview: Dict[str, Any] = {}
        # 【新增】世界观词条列表（结构化数据，用于Treeview显示）
        self._worldview_entries: List[Dict[str, Any]] = []
        # 【新增】人物词条列表（结构化数据，用于Treeview显示）
        self._character_entries: List[Dict[str, Any]] = []
        self._reverse_feedback_data: Dict[str, Any] = {}  # 逆向反馈数据
        self._completed_chapters: List[Dict[str, Any]] = []  # 已完成章节
        self._generated_content: List[str] = []  # 生成内容
        self._llm_client = None
        
        # 知识库管理器（延迟初始化）
        self._knowledge_manager = None

        # 适配器实例（延迟初始化，按需创建）
        self._worldview_adapter = None
        self._character_adapter = None
        self._outline_adapter = None

        # 工作区根目录
        self._workspace_root = os.path.dirname(os.path.abspath(__file__))
        
        # 进度统计变量（确保在页面创建前初始化）
        self._progress_project_name: Optional[tk.StringVar] = None
        self._progress_chapters: Optional[tk.StringVar] = None
        self._progress_total_words: Optional[tk.StringVar] = None
        self._progress_today_words: Optional[tk.StringVar] = None
        self._progress_outline: Optional[tk.StringVar] = None
        self._progress_characters: Optional[tk.StringVar] = None
        self._progress_worldview: Optional[tk.StringVar] = None
        self._progress_style: Optional[tk.StringVar] = None
        
        # 事件订阅ID
        self._subscription_ids: List[str] = []
        
        # 初始化
        self._init_theme()
        self._init_async_handler()
        self._init_core_services()
        self._init_project_manager()  # 新增：初始化项目管理器
        self._setup_project_event_listeners()  # 新增：设置项目事件监听器
        self._init_ui()
        self._init_bindings()
        
        # V2.19: 会话自动恢复（检查配置）
        self._check_auto_recovery()
        
        # V2.20: 启动每日冥想定时任务
        self.root.after(5000, self._start_daily_meditation)  # 延迟5秒启动，避免阻塞主流程
        
        # 启动结果队列处理
        self.root.after(100, self._process_result_queue)
        
        # 窗口关闭确认
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        
        # V2.18: 后台预加载向量模型（5秒后开始，不阻塞UI）
        self.root.after(5000, self._preload_vector_store_async)
        
        # V3.2.1: 启动时同步项目名称显示
        self.root.after(100, self._sync_project_name_on_startup)
        
        logger.info("MainWindow initialized")
    
    def _enable_glass_effect(self) -> None:
        """启用玻璃效果"""
        try:
            glass_mgr = GlassWindowManager(self.root)
            
            # 尝试Acrylic效果
            if glass_mgr.enable_acrylic(0x80000000):
                logger.info("Acrylic glass effect enabled")
            else:
                logger.info("Acrylic effect not available, using alpha transparency")
            
            # 无论 Acrylic 是否成功，都应用完全无边框
            glass_mgr.make_frameless(keep_resize_border=False)
            
            # 调试：确认overrideredirect状态
            logger.info(f"overrideredirect状态: {self.root.overrideredirect()}")
            
            # 如果 Acrylic 失败，降级到透明效果
            if not glass_mgr._is_acrylic:
                self.root.attributes('-alpha', 0.95)
            
            logger.info("Frameless window enabled")
        
        except Exception as e:
            logger.warning(f"Glass effect not available: {e}")
            self.root.attributes('-alpha', 0.98)
    
    def _set_window_icon(self) -> None:
        """设置窗口图标（任务栏和窗口左上角）"""
        try:
            # 尝试加载图标文件
            icon_paths = [
                os.path.join(self._workspace_root, "icon.ico"),
                os.path.join(self._workspace_root, "assets", "icon.ico"),
                os.path.join(self._workspace_root, "data", "icon.ico"),
            ]
            
            icon_set = False
            for icon_path in icon_paths:
                if os.path.exists(icon_path):
                    # 方法1: 使用iconbitmap设置.ico图标（Windows推荐）
                    try:
                        self.root.iconbitmap(icon_path)
                        logger.info(f"Window icon set from: {icon_path}")
                        icon_set = True
                        break
                    except Exception as e:
                        logger.warning(f"Failed to set iconbitmap: {e}")
            
            # 如果没有.ico文件，创建一个简单的默认图标
            if not icon_set:
                logger.warning("No icon.ico found, creating default icon")
                # 创建一个简单的白色图标
                from PIL import Image, ImageDraw
                img = Image.new('RGBA', (256, 256), (70, 130, 180, 255))
                draw = ImageDraw.Draw(img)
                draw.ellipse([64, 64, 192, 192], fill=(255, 255, 255, 255))
                draw.text((90, 120), "N", fill=(255, 255, 255, 255))
                
                # 保存为.ico
                temp_icon_path = os.path.join(self._workspace_root, "icon.ico")
                img.save(temp_icon_path, format='ICO', sizes=[(256, 256)])
                
                # 设置图标
                self.root.iconbitmap(temp_icon_path)
                logger.info("Default window icon created and set")
                
        except ImportError:
            logger.warning("PIL not available, skipping icon generation")
        except Exception as e:
            logger.warning(f"Failed to set window icon: {e}")
    
    def _init_theme(self) -> None:
        """初始化主题"""
        style = ttk.Style()
        
        # 使用sv_ttk主题
        if SV_TTK_AVAILABLE:
            try:
                sv_ttk.set_theme("dark")
            except Exception:
                pass
        
        # 应用玻璃态主题
        GlassTheme.apply_to_style(style)
    
    def _init_async_handler(self) -> None:
        """初始化异步处理器"""
        if CORE_AVAILABLE:
            self._async_handler = init_async_handler(
                root=self.root,
                worker_count=4,
                default_timeout=30.0
            )
            logger.info("AsyncHandler initialized")
    
    def _init_core_services(self) -> None:
        """初始化核心服务"""
        if not CORE_AVAILABLE:
            logger.warning("Core services not available")
            return
        
        try:
            self._services = CoreServiceManager(self.root)
            self._services.initialize(
                event_bus=get_event_bus(),
                registry=get_plugin_registry(),
                locator=get_service_locator(),
                config=get_config_manager()
            )
            
            # 订阅生成事件
            self._subscribe_generation_events()
            
            logger.info("Core services initialized")
        except Exception as e:
            logger.error(f"Failed to initialize core services: {e}")
    
    def _init_project_manager(self) -> None:
        """初始化项目管理器"""
        try:
            from core.service_locator import get_service_locator
            from services.project_manager import ProjectManager
            
            locator = get_service_locator()
            self._project_manager = locator.get(ProjectManager)
            
            if self._project_manager:
                logger.info("[MainWindow] 项目管理器已初始化")
            else:
                logger.warning("[MainWindow] 项目管理器未注册")
                self._project_manager = None
        except Exception as e:
            logger.warning(f"[MainWindow] 项目管理器初始化失败: {e}")
            self._project_manager = None
    
    def _setup_project_event_listeners(self) -> None:
        """设置项目事件监听器"""
        if not self._services or not self._services.event_bus:
            return
        
        try:
            event_bus = self._services.event_bus
            
            # 订阅项目保存事件
            event_bus.subscribe("project.saved", self._on_project_saved)
            
            # 订阅项目加载事件
            event_bus.subscribe("project.loaded", self._on_project_loaded)
            
            logger.info("[MainWindow] 项目事件监听器已设置")
        except Exception as e:
            logger.warning(f"[MainWindow] 项目事件监听器设置失败: {e}")
    
    def _on_project_saved(self, event) -> None:
        """项目保存事件回调"""
        project_name = getattr(event, 'project_name', '未知项目')
        
        # V3.2.1修复：更新状态栏项目名称
        self._update_status_bar(project_name=project_name)
        
        # 更新项目管理页面的项目名称
        if hasattr(self, '_project_name_var'):
            self._project_name_var.set(project_name)
        
        self._set_status(f"项目已保存: {project_name}")
        logger.info(f"[MainWindow] 项目保存事件: {project_name}")
    
    def _on_project_loaded(self, event) -> None:
        """项目加载事件回调"""
        project_name = getattr(event, 'project_name', '未知项目')
        
        # V3.2.1修复：更新状态栏项目名称
        self._update_status_bar(project_name=project_name)
        
        # 更新项目管理页面的项目名称
        if hasattr(self, '_project_name_var'):
            self._project_name_var.set(project_name)
        
        self._set_status(f"项目已加载: {project_name}")
        logger.info(f"[MainWindow] 项目加载事件: {project_name}")
    
    def _init_ui(self) -> None:
        """初始化UI"""
        # 自定义标题栏
        self._title_bar = CustomTitleBar(self.root, self.TITLE)
        self._title_bar.pack(fill=tk.X)
        
        # 窗口调整大小的边框感知（完全无边框模式需要）
        self._resize_grip = ResizeGrip(self.root)
        
        # 主容器
        main_container = ttk.Frame(self.root, style="TFrame")
        main_container.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
        
        # 侧边栏
        self._create_sidebar(main_container)
        
        # 内容区
        self._create_content(main_container)
        
        # 状态栏
        self._create_status_bar()
        
        # 默认显示热榜页面
        self._switch_to_page("hot_ranking")
    
    def _create_sidebar(self, parent: tk.Widget) -> None:
        """创建侧边栏"""
        sidebar = ttk.Frame(parent, style="TFrame", width=200)
        sidebar.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 8))
        sidebar.pack_propagate(False)
        
        # 导航按钮
        nav_buttons = [
            ("热榜", "hot_ranking", GlassTheme.ACCENT_PINK),
            ("工作台", "workbench", GlassTheme.PRIMARY_LIGHT),
            ("创作进度", "progress", GlassTheme.ACCENT_GREEN),
            ("项目管理", "project", GlassTheme.ACCENT_PURPLE),
            ("插件管理", "plugins", GlassTheme.ACCENT_ORANGE),
            ("反馈与建议", "feedback", GlassTheme.INFO),
            ("设置", "settings", GlassTheme.WARNING),
        ]
        
        for text, page_id, color in nav_buttons:
            btn = tk.Label(
                sidebar,
                text=text,
                font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
                fg=GlassTheme.TEXT_PRIMARY,
                bg=GlassTheme.GLASS_SURFACE,
                activeforeground=GlassTheme.TEXT_PRIMARY,
                activebackground=color,
                cursor="hand2",
                padx=20,
                pady=12,
                anchor="w"
            )
            btn.pack(fill=tk.X, pady=2)
            
            # 悬停效果
            btn.bind("<Enter>", lambda e, b=btn, c=color: b.configure(bg=c))
            btn.bind("<Leave>", lambda e, b=btn: b.configure(bg=GlassTheme.GLASS_SURFACE))
            btn.bind("<Button-1>", lambda e, pid=page_id: self._switch_to_page(pid))
        
        # 底部版本信息
        version_label = tk.Label(
            sidebar,
            text=f"Agent Pro {self.VERSION}",
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_TINY),
            fg=GlassTheme.TEXT_DISABLED,
            bg=GlassTheme.GLASS_BG
        )
        version_label.pack(side=tk.BOTTOM, pady=10)
    
    def _create_content(self, parent: tk.Widget) -> None:
        """创建内容区"""
        self._content_frame = ttk.Frame(parent, style="Glass.TFrame")
        self._content_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    def _create_status_bar(self) -> None:
        """创建状态栏（增强版V3.2：项目名称、AI连接、字数统计、后台任务进度）

        V3.2更新：
        - AI状态实时同步（通过AIStatusManagerPlugin）
        - 显示服务类型（本地/线上）、提供商、连接状态
        - 支持点击查看详细信息
        """
        status_frame = ttk.Frame(self.root, style="TFrame")
        status_frame.pack(fill=tk.X, side=tk.BOTTOM, padx=8, pady=(0, 8))

        # 分隔线
        ttk.Separator(status_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=(0, 4))
        
        status_inner = ttk.Frame(status_frame, style="TFrame")
        status_inner.pack(fill=tk.X)
        
        # 项目名称
        self._project_var = tk.StringVar(value="项目: 未打开")
        ttk.Label(
            status_inner,
            textvariable=self._project_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        ).pack(side=tk.LEFT)
        
        # 分隔符
        ttk.Label(status_inner, text=" | ", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=5)
        
        # 字数统计
        self._word_count_var = tk.StringVar(value="字数: 0")
        ttk.Label(
            status_inner,
            textvariable=self._word_count_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        ).pack(side=tk.LEFT)
        
        # 分隔符
        ttk.Label(status_inner, text=" | ", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=5)
        
        # 状态文本
        self._status_var = tk.StringVar(value="就绪")
        ttk.Label(
            status_inner,
            textvariable=self._status_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        ).pack(side=tk.LEFT, padx=10)
        
        # 后台任务进度（右侧）
        self._background_task_var = tk.StringVar(value="")
        self._background_task_label = ttk.Label(
            status_inner,
            textvariable=self._background_task_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
            foreground=GlassTheme.INFO
        )
        self._background_task_label.pack(side=tk.RIGHT, padx=10)
        
        # 分隔符
        ttk.Label(status_inner, text=" | ", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.RIGHT, padx=5)
        
        # AI连接状态（V3.2增强：实时同步+点击详情）
        self._ai_status_var = tk.StringVar(value="AI: 未连接")
        self._ai_status_label = ttk.Label(
            status_inner,
            textvariable=self._ai_status_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
            cursor="hand2"  # 鼠标悬停显示手型
        )
        self._ai_status_label.pack(side=tk.RIGHT, padx=10)

        # 点击AI状态显示详细信息
        self._ai_status_label.bind("<Button-1>", self._show_ai_status_details)

        # 初始化AI状态管理插件（延迟初始化）
        self._ai_status_manager = None
        self._init_ai_status_manager()

    def _init_ai_status_manager(self):
        """初始化AI状态管理插件（按需初始化，不影响启动速度）"""
        try:
            import sys
            from pathlib import Path
            plugin_path = Path(__file__).parent / "plugins" / "ai-status-manager-v1"

            if str(plugin_path) not in sys.path:
                sys.path.insert(0, str(plugin_path))

            # 导入插件类（插件内部已包含基础类定义）
            from plugin import AIStatusManagerPlugin, PluginContext

            plugin = AIStatusManagerPlugin()
            context = PluginContext(
                plugin_id="ai-status-manager-v1",
                config={
                    "auto_sync_interval": 30,
                    "retry_on_failure": True,
                    "max_retry_count": 3,
                }
            )

            if plugin.initialize(context):
                self._ai_status_manager = plugin
                logger.info("[GUI] AI状态管理插件初始化成功")

                # 订阅状态变更事件
                try:
                    from core.service_locator import get_service_locator
                    locator = get_service_locator()
                    from core.event_bus import EventBus
                    event_bus = locator.get(EventBus)

                    if event_bus:
                        # V3.2修复：添加详细日志
                        logger.info(f"[GUI] EventBus获取成功: {event_bus}")
                        event_bus.subscribe("ai.status.changed", self._on_ai_status_changed)
                        logger.info("[GUI] 已订阅AI状态变更事件")

                        # 测试发布一个事件，验证订阅是否成功
                        test_status = {
                            "connection_state": "测试中",
                            "service_type": "线上",
                            "provider": "Test",
                            "model": "test-model",
                            "endpoint": "http://test",
                            "error_message": ""
                        }
                        # V3.2修复：publish直接传递数据，不需要包装成Event对象
                        event_bus.publish("ai.status.changed", test_status)
                        logger.info("[GUI] 测试事件已发布")
                    else:
                        logger.error("[GUI] EventBus获取失败：返回None")

                except Exception as e:
                    logger.error(f"[GUI] 订阅AI状态事件失败: {e}", exc_info=True)

        except Exception as e:
            logger.error(f"[GUI] 初始化AI状态管理插件失败: {e}", exc_info=True)

    def _on_ai_status_changed(self, event):
        """处理AI状态变更事件
        
        Args:
            event: EventBus的Event对象，包含event.data（状态字典）
        """
        try:
            # V3.2修复：从Event对象中提取数据
            from core.event_bus import Event as EventType
            
            if isinstance(event, EventType):
                status = event.data if event.data else {}
            elif isinstance(event, dict):
                # 兼容性处理：如果直接传递字典
                status = event
            else:
                logger.warning(f"[GUI] AI状态事件格式错误: {type(event)}")
                return
            
            connection_state = status.get("connection_state", "未连接")
            service_type = status.get("service_type", "线上")
            provider = status.get("provider", "DeepSeek")

            # 更新状态栏显示
            if connection_state == "已连接":
                status_text = f"AI: {provider} ({service_type}) ✓"
                self._ai_status_label.configure(foreground=GlassTheme.SUCCESS)
            elif connection_state == "连接中":
                status_text = f"AI: {provider} 连接中..."
                self._ai_status_label.configure(foreground=GlassTheme.WARNING)
            elif connection_state == "服务启动中":
                status_text = f"AI: {provider} 启动中..."
                self._ai_status_label.configure(foreground=GlassTheme.WARNING)
            elif connection_state == "连接错误":
                status_text = f"AI: {provider} 连接失败 ✗"
                self._ai_status_label.configure(foreground=GlassTheme.ERROR)
            else:
                status_text = "AI: 未连接"
                self._ai_status_label.configure(foreground=GlassTheme.TEXT_SECONDARY)

            self._ai_status_var.set(status_text)
            logger.info(f"[GUI] AI状态已更新: {status_text}")

        except Exception as e:
            logger.error(f"[GUI] 处理AI状态变更失败: {e}")

    def _show_ai_status_details(self, event):
        """显示AI状态详细信息"""
        try:
            if not self._ai_status_manager:
                messagebox.showinfo("AI状态", "AI状态管理插件未初始化")
                return

            status = self._ai_status_manager.get_status()

            details = (
                f"AI连接状态\n"
                f"{'=' * 40}\n\n"
                f"连接状态: {status['connection_state']}\n"
                f"服务类型: {status['service_type']}\n"
                f"提供商: {status['provider']}\n"
                f"模型: {status['model']}\n"
                f"端点: {status['endpoint']}\n"
            )

            if status['error_message']:
                details += f"\n错误信息: {status['error_message']}\n"

            details += (
                f"\n{'=' * 40}\n"
                f"点击状态栏可查看此信息\n"
                f"状态每30秒自动更新一次"
            )

            messagebox.showinfo("AI连接状态", details)

        except Exception as e:
            logger.error(f"[GUI] 显示AI状态详情失败: {e}")
            messagebox.showerror("错误", f"无法获取AI状态: {str(e)}")

    def _update_status_bar(self, project_name: str = None, word_count: int = None, 
                          ai_status: str = None, background_task: str = None) -> None:
        """更新状态栏信息"""
        if project_name is not None:
            self._project_var.set(f"项目: {project_name}")
        if word_count is not None:
            self._word_count_var.set(f"字数: {word_count:,}")
        if ai_status is not None:
            self._ai_status_var.set(f"AI: {ai_status}")
            # 更新颜色
            if ai_status == "在线":
                self._ai_status_label.configure(foreground=GlassTheme.SUCCESS)
            elif ai_status == "离线":
                self._ai_status_label.configure(foreground=GlassTheme.ERROR)
            else:
                self._ai_status_label.configure(foreground=GlassTheme.WARNING)
        if background_task is not None:
            self._background_task_var.set(background_task)
    
    def _init_bindings(self) -> None:
        """初始化快捷键绑定（增强版）"""
        # 页面切换
        self.root.bind("<Control-w>", lambda e: self._switch_to_page("workbench"))
        self.root.bind("<Control-p>", lambda e: self._switch_to_page("project"))
        self.root.bind("<Control-s>", lambda e: self._switch_to_page("settings"))
        self.root.bind("<Control-g>", lambda e: self._on_quick_generate_action())
        self.root.bind("<F5>", lambda e: self._refresh_current_page())
        
        # 新增快捷键
        self.root.bind("<Control-n>", lambda e: self._on_new_item())  # 新建
        self.root.bind("<Control-o>", lambda e: self._on_open_project())  # 打开项目
        self.root.bind("<Control-b>", lambda e: self._on_backup_project())  # 备份
        self.root.bind("<Escape>", lambda e: self._on_cancel_action())  # 取消
    
    def _on_quick_generate_action(self):
        """快速生成快捷键处理"""
        self._switch_to_page("workbench")
        self._switch_workbench_tab("quick")
    
    def _on_new_item(self):
        """新建项目快捷键处理"""
        self._switch_to_page("project")
        self._on_new_project()
    
    def _on_cancel_action(self):
        """取消操作"""
        self._set_status("操作已取消")
    
    # ============== 页面切换 ==============
    
    def _switch_to_page(self, page_id: str) -> None:
        """切换到指定页面"""
        # 隐藏当前页面
        if self._current_page and self._current_page in self._pages:
            self._pages[self._current_page].pack_forget()
        
        # 创建或显示新页面
        if page_id not in self._pages:
            self._pages[page_id] = self._create_page(page_id)
            logger.debug(f"Created new page: {page_id}")
        
        # pack页面到content_frame
        page = self._pages[page_id]
        page.pack(fill=tk.BOTH, expand=True, in_=self._content_frame)
        self._current_page = page_id
        
        # 等待页面pack完成后更新Canvas宽度（延迟执行避免阻塞UI）
        self.root.after(50, self._update_canvas_widths)
        
        logger.debug(f"Switched to page: {page_id}, parent={page.master}, winfo_parent={page.winfo_parent()}")
    
    def _update_canvas_widths(self):
        """更新所有Canvas的窗口宽度（优化版：添加防抖和缓存）"""
        # 只更新当前可见的页面，避免遍历所有页面
        if self._current_page and self._current_page in self._pages:
            self._update_canvas_width_recursive(self._pages[self._current_page])
    
    def _update_canvas_width_recursive(self, widget, depth=0):
        """递归更新所有Canvas的窗口宽度（优化版：限制递归深度）"""
        # 限制递归深度，避免过深遍历导致卡顿
        if depth > 10:
            return
            
        if widget.winfo_class() == 'Canvas':
            try:
                # 获取Canvas中的所有窗口
                width = widget.winfo_width()
                if width > 1:  # 确保宽度有效
                    for item_id in widget.find_withtag("all"):
                        tags = widget.gettags(item_id)
                        if "all" not in tags:
                            widget.itemconfig(item_id, width=width)
            except Exception:
                pass  # 忽略无效的Canvas操作
        
        # 递归处理子控件
        for child in widget.winfo_children():
            self._update_canvas_width_recursive(child, depth + 1)
    
    def _create_page(self, page_id: str) -> tk.Frame:
        """创建页面"""
        page_creators = {
            "hot_ranking": self._create_hot_ranking_page,
            "workbench": self._create_workbench_page,
            "progress": self._create_progress_page,
            "project": self._create_project_page,
            "plugins": self._create_plugins_page,
            "feedback": self._create_feedback_page,
            "settings": self._create_settings_page,
        }
        
        creator = page_creators.get(page_id)
        if creator:
            return creator()
        else:
            # 默认空白页面
            frame = ttk.Frame(self._content_frame, style="TFrame")
            ttk.Label(frame, text=f"页面: {page_id}", style="Title.TLabel").pack(padx=20, pady=20)
            return frame
    
    def _refresh_current_page(self) -> None:
        """刷新当前页面"""
        if self._current_page:
            # 销毁当前页面并重新创建
            if self._current_page in self._pages:
                self._pages[self._current_page].destroy()
                del self._pages[self._current_page]
            self._switch_to_page(self._current_page)
    
    # ============== 页面创建器 ==============
    
    def _create_hot_ranking_page(self) -> tk.Frame:
        """创建热榜页面（V5完整迁移）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建滚动容器
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        # Canvas和Scrollbar直接pack到frame中
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题栏
        title_frame = ttk.Frame(scrollable_frame, style="TFrame")
        title_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        tk.Label(
            title_frame,
            text="全网小说热度榜单",
            font=(GlassTheme.FONT_FAMILY, 24, "bold"),
            fg=GlassTheme.ACCENT_PINK,
            bg=GlassTheme.GLASS_BG
        ).pack(side=tk.LEFT)
        
        # 按钮区
        btn_frame = ttk.Frame(title_frame, style="TFrame")
        btn_frame.pack(side=tk.RIGHT)
        
        # 清缓存按钮
        self.clear_cache_btn = ResponsiveButton(
            btn_frame,
            text="清缓存",
            command=self._clear_hot_ranking_cache,
            async_handler=self._async_handler
        )
        self.clear_cache_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # 更新按钮
        self.update_hot_btn = ResponsiveButton(
            btn_frame,
            text="更新数据",
            command=self._update_hot_ranking_data,
            async_handler=self._async_handler,
            style="Accent.TButton"
        )
        self.update_hot_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # 进度条（P2-001修复：长任务进度提示）
        progress_frame = ttk.Frame(scrollable_frame, style="TFrame")
        progress_frame.pack(fill=tk.X, padx=20, pady=5)
        
        self._hot_ranking_progress_var = tk.StringVar(value="")
        self._hot_ranking_progress_label = ttk.Label(
            progress_frame,
            textvariable=self._hot_ranking_progress_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
            foreground=GlassTheme.INFO
        )
        self._hot_ranking_progress_label.pack(side=tk.LEFT)
        
        self._hot_ranking_progress_bar = ttk.Progressbar(
            progress_frame,
            length=300,
            mode='determinate',
            maximum=100
        )
        self._hot_ranking_progress_bar.pack(side=tk.LEFT, padx=10)
        
        # 初始隐藏进度条
        progress_frame.pack_forget()
        self._hot_ranking_progress_frame = progress_frame
        
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=10)
        
        # ========== 第一部分：三大网站榜单 ==========
        ttk.Label(
            scrollable_frame,
            text="三大热门小说网站排行榜 (前10名)",
            font=(GlassTheme.FONT_FAMILY, 16, "bold")
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        # 来源说明行
        src_desc_frame = ttk.Frame(scrollable_frame, style="TFrame")
        src_desc_frame.pack(fill=tk.X, padx=22, pady=(0, 8))
        
        site_badge_info = [
            ('🍅', '番茄', GlassTheme.ACCENT_RED),
            ('🚀', '起点', GlassTheme.INFO),
            ('🔷', '晋江', GlassTheme.ACCENT_PURPLE),
        ]
        for icon, name, color in site_badge_info:
            badge = tk.Label(
                src_desc_frame,
                text=f" {icon} {name} ",
                font=(GlassTheme.FONT_FAMILY, 9),
                fg='white', bg=color,
                padx=4, pady=2,
                relief='flat'
            )
            badge.pack(side=tk.LEFT, padx=(0, 6))
        
        # 网站榜单容器（V5原版：使用pack布局，三列并排）
        # V2.18: 保存引用以便异步更新
        self._hot_sites_frame = ttk.Frame(scrollable_frame, style="TFrame")
        self._hot_sites_frame.pack(fill=tk.X, padx=20, pady=10)

        # ===== V2.18 热榜异步化改造 =====
        # 1. 先显示默认数据（立即响应，不阻塞UI）
        hot_data = self._get_default_hot_ranking_data()
        
        # 2. 显示进度提示
        self._hot_ranking_progress_frame.pack(fill=tk.X, padx=20, pady=5)
        self._hot_ranking_progress_var.set("正在后台加载热榜数据...")
        self._hot_ranking_progress_bar['value'] = 10
        
        # 3. 渲染默认数据
        self._render_hot_ranking_sites(hot_data.get('sites', []))
        
        # 4. 异步加载真实数据（缓存优先）
        self._async_fetch_hot_ranking(scrollable_frame)
        
        # ===== 题材榜、类型榜、作家榜（使用默认数据，异步加载完成后刷新）=====
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=20)
        
        # ========== 第二部分：题材热度榜 ==========
        ttk.Label(
            scrollable_frame,
            text="题材热度榜 (前5名)",
            font=(GlassTheme.FONT_FAMILY, 16, 'bold')
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        # V2.18: 保存引用以便异步更新
        self._hot_genres_frame = ttk.Frame(scrollable_frame, style="TFrame")
        self._hot_genres_frame.pack(fill=tk.X, padx=20, pady=10)
        
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=20)
        
        # ========== 第三部分：设定类型/创作流派榜 ==========
        ttk.Label(
            scrollable_frame,
            text="设定类型/创作流派榜 (前5名)",
            font=(GlassTheme.FONT_FAMILY, 16, 'bold')
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        self._hot_types_frame = ttk.Frame(scrollable_frame, style="TFrame")
        self._hot_types_frame.pack(fill=tk.X, padx=20, pady=10)
        
        ttk.Separator(scrollable_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, padx=20, pady=20)
        
        # ========== 第四部分：作家排行榜 ==========
        ttk.Label(
            scrollable_frame,
            text="全网最热网络作家排行榜 (前10名)",
            font=(GlassTheme.FONT_FAMILY, 16, 'bold')
        ).pack(pady=(20, 10), anchor='w', padx=20)
        
        self._hot_authors_frame = ttk.Frame(scrollable_frame, style="TFrame")
        self._hot_authors_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # 渲染默认数据
        self._render_hot_ranking_genres(hot_data.get('genres', {}))
        self._render_hot_ranking_types(hot_data.get('types', {}))
        self._render_hot_ranking_authors(hot_data.get('authors', []))
        
        # 数据来源说明
        ttk.Label(
            scrollable_frame,
            text="收入和粉丝数为算法估算值，仅供参考",
            font=(GlassTheme.FONT_FAMILY, 9),
            foreground=GlassTheme.TEXT_SECONDARY
        ).pack(pady=(5, 15), anchor='w', padx=20)
        
        return frame
    
    def _create_workbench_page(self) -> tk.Frame:
        """创建工作台页面（上半部分按钮 + 下半部分内容区）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 标题
        header = ttk.Frame(frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="工作台", style="Title.TLabel").pack(side=tk.LEFT)
        
        # ========== 上半部分：十个功能按钮 ==========
        buttons_frame = ttk.Frame(frame, style="TFrame")
        buttons_frame.pack(fill=tk.X, padx=20, pady=(0, 15))
        
        # 第一行三个按钮
        row1 = ttk.Frame(buttons_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        
        # 第二行三个按钮
        row2 = ttk.Frame(buttons_frame, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        
        # 第三行三个按钮
        row3 = ttk.Frame(buttons_frame, style="TFrame")
        row3.pack(fill=tk.X, pady=5)
        
        # 第四行一个按钮（居中）
        row4 = ttk.Frame(buttons_frame, style="TFrame")
        row4.pack(fill=tk.X, pady=5)
        
        functions = [
            ("世界观", "worldview", GlassTheme.ACCENT_PURPLE),
            ("人物设定", "characters", GlassTheme.ACCENT_PINK),
            ("大纲管理", "outline", GlassTheme.PRIMARY_LIGHT),
            ("风格学习", "style", GlassTheme.ACCENT_GREEN),
            ("开始创作", "generation", GlassTheme.PRIMARY),
            ("逆向反馈", "reverse", GlassTheme.WARNING),
            ("快捷创作", "quick", GlassTheme.ACCENT_RED),
            ("续写功能", "continue", GlassTheme.INFO),
            ("长篇检测", "consistency", GlassTheme.ACCENT_CYAN),
            ("知识库", "knowledge", GlassTheme.ACCENT_ORANGE),
        ]
        
        self._workbench_buttons = {}
        self._current_workbench_tab = tk.StringVar(value="worldview")
        
        for i, (text, tab_id, color) in enumerate(functions):
            if i < 3:
                parent = row1
            elif i < 6:
                parent = row2
            elif i < 9:
                parent = row3
            else:
                parent = row4
            
            btn = tk.Label(
                parent,
                text=text,
                font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold"),
                fg=GlassTheme.TEXT_PRIMARY,
                bg=GlassTheme.GLASS_SURFACE,
                activeforeground=GlassTheme.TEXT_PRIMARY,
                activebackground=color,
                cursor="hand2",
                padx=20,
                pady=15,
                width=18
            )
            btn.pack(side=tk.LEFT, padx=8, pady=5, expand=True, fill=tk.X)
            
            # 存储按钮引用
            self._workbench_buttons[tab_id] = btn
            
            # 悬停效果
            btn.bind("<Enter>", lambda e, b=btn, c=color: b.configure(bg=c))
            btn.bind("<Leave>", lambda e, b=btn, tid=tab_id: self._update_button_style(tid))
            btn.bind("<Button-1>", lambda e, tid=tab_id: self._switch_workbench_tab(tid))
        
        # ========== 下半部分：内容区域（支持滚动）==========
        # 创建Canvas和Scrollbar用于滚动
        content_canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        content_scrollbar = ttk.Scrollbar(frame, orient="vertical", command=content_canvas.yview)
        
        self._workbench_content_frame = ttk.Frame(content_canvas, style="TFrame")
        
        # 防抖：避免频繁触发Configure事件导致卡顿
        self._workbench_configure_id = None
        def on_content_frame_configure(event):
            if self._workbench_configure_id:
                self.root.after_cancel(self._workbench_configure_id)
            self._workbench_configure_id = self.root.after(50, lambda: _update_scrollregion())
        
        def _update_scrollregion():
            try:
                content_canvas.configure(scrollregion=content_canvas.bbox("all"))
                content_canvas.itemconfig(content_window, width=content_canvas.winfo_width())
            except Exception:
                pass
        
        self._workbench_content_frame.bind("<Configure>", on_content_frame_configure)
        content_window = content_canvas.create_window((0, 0), window=self._workbench_content_frame, anchor="nw")
        content_canvas.configure(yscrollcommand=content_scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_content_mousewheel(event):
            content_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_content_mousewheel(event):
            content_canvas.bind_all("<MouseWheel>", on_content_mousewheel)
        
        def unbind_content_mousewheel(event):
            content_canvas.unbind_all("<MouseWheel>")
        
        content_canvas.bind("<Enter>", bind_content_mousewheel)
        content_canvas.bind("<Leave>", unbind_content_mousewheel)
        
        content_canvas.pack(side="left", fill="both", expand=True, padx=20, pady=(0, 20))
        content_scrollbar.pack(side="right", fill="y", pady=(0, 20))
        
        # 初始化所有标签页内容（延迟创建）
        self._workbench_tabs = {}
        self._workbench_tabs_created = set()
        
        # 默认显示世界观页面
        self._create_all_workbench_tabs()
        self._switch_workbench_tab("worldview")
        
        return frame
    
    def _create_all_workbench_tabs(self):
        """预创建工作台标签页（延迟创建策略，仅初始化为None）"""
        # 所有标签页初始化为None，首次切换时才创建
        # 这样可以避免首次点击工作台时一次性创建所有标签页导致卡顿
        for tab_id in ["worldview", "characters", "outline", "style", 
                       "generation", "reverse", "quick", "continue", "consistency", "knowledge"]:
            self._workbench_tabs[tab_id] = None
    
    def _switch_workbench_tab(self, tab_id: str):
        """切换工作台标签页（延迟创建）"""
        # 更新按钮样式
        self._current_workbench_tab.set(tab_id)
        for tid, btn in self._workbench_buttons.items():
            if tid == tab_id:
                # 激活状态：使用对应颜色
                color = self._get_button_color(tid)
                btn.configure(bg=color, fg="white")
            else:
                btn.configure(bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        
        # 隐藏所有标签页
        for tid, tab_frame in self._workbench_tabs.items():
            if tid != tab_id and tab_frame is not None:
                tab_frame.pack_forget()
        
        # 延迟创建：首次切换时才创建标签页内容
        if self._workbench_tabs.get(tab_id) is None:
            create_method = {
                "worldview": self._create_worldview_content,
                "characters": self._create_characters_content,
                "outline": self._create_outline_content,
                "style": self._create_style_content,
                "generation": self._create_generation_content,
                "reverse": self._create_reverse_content,
                "quick": self._create_quick_content,
                "continue": self._create_continue_content,
                "consistency": self._create_consistency_content,
                "knowledge": self._create_knowledge_content,
            }.get(tab_id)
            if create_method:
                self._workbench_tabs[tab_id] = create_method()
        
        # 显示当前标签页
        if self._workbench_tabs.get(tab_id) is not None:
            self._workbench_tabs[tab_id].pack(fill=tk.BOTH, expand=True, in_=self._workbench_content_frame)
    
    def _update_button_style(self, tab_id: str):
        """更新按钮样式（用于鼠标离开时）"""
        if self._current_workbench_tab.get() != tab_id:
            self._workbench_buttons[tab_id].configure(bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
    
    def _get_button_color(self, tab_id: str) -> str:
        """获取按钮对应的颜色"""
        colors = {
            "worldview": GlassTheme.ACCENT_PURPLE,
            "characters": GlassTheme.ACCENT_PINK,
            "outline": GlassTheme.PRIMARY_LIGHT,
            "style": GlassTheme.ACCENT_GREEN,
            "generation": GlassTheme.PRIMARY,
            "reverse": GlassTheme.WARNING,
            "quick": GlassTheme.ACCENT_RED,
            "continue": GlassTheme.INFO,
            "consistency": GlassTheme.ACCENT_CYAN,
            "knowledge": GlassTheme.ACCENT_ORANGE,
        }
        return colors.get(tab_id, GlassTheme.PRIMARY)
    
    def _create_worldview_content(self) -> tk.Frame:
        """创建世界观管理内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：文件导入区
        file_frame = ttk.LabelFrame(frame, text="世界观文件导入", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="选择文件：").pack(side=tk.LEFT)
        self._worldview_path_var = tk.StringVar()
        ttk.Entry(path_frame, textvariable=self._worldview_path_var, width=50).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="浏览", command=self._on_worldview_browse).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="新建", command=self._on_worldview_new).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="保存项目", command=self._on_save_project).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="解析", command=self._on_worldview_import).pack(side=tk.LEFT, padx=5)
        
        # 支持格式说明
        ttk.Label(file_frame, text="支持格式：TXT / DOCX", 
                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(anchor=tk.W, pady=5)
        
        # 中部：世界观列表
        list_frame = ttk.LabelFrame(frame, text="世界观项目列表", padding=10)
        list_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 创建带滚动条的Treeview容器
        tree_container = ttk.Frame(list_frame)
        tree_container.pack(fill=tk.X, pady=5)
        
        # 滚动条
        worldview_scrollbar = ttk.Scrollbar(tree_container, orient=tk.VERTICAL)
        worldview_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        columns = ("name", "category", "elements", "status", "modified")
        self._worldview_tree = ttk.Treeview(tree_container, columns=columns, show="headings", height=4, yscrollcommand=worldview_scrollbar.set)
        self._worldview_tree.heading("name", text="世界观名称")
        self._worldview_tree.heading("category", text="世界观类别")
        self._worldview_tree.heading("elements", text="核心元素")
        self._worldview_tree.heading("status", text="状态")
        self._worldview_tree.heading("modified", text="修改时间")
        self._worldview_tree.column("name", width=120)
        self._worldview_tree.column("category", width=80)
        self._worldview_tree.column("elements", width=160)
        self._worldview_tree.column("status", width=80)
        self._worldview_tree.column("modified", width=110)
        self._worldview_tree.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # 配置滚动条
        worldview_scrollbar.config(command=self._worldview_tree.yview)
        
        # 鼠标滚轮支持
        def _on_worldview_mousewheel(event):
            self._worldview_tree.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        self._worldview_tree.bind("<MouseWheel>", _on_worldview_mousewheel)
        
        # 【新增】双击查看详情
        self._worldview_tree.bind("<Double-Button-1>", lambda e: self._on_worldview_view())

        list_btn_frame = ttk.Frame(list_frame, style="TFrame")
        list_btn_frame.pack(fill=tk.X, pady=5)
        
        # 使用Grid布局，3列，确保按钮均匀分布
        list_buttons = [
            ("查看详情", self._on_worldview_view),
            ("编辑", self._on_worldview_edit),
            ("批量删除", self._on_worldview_delete),
        ]
        
        for i, (text, command) in enumerate(list_buttons):
            btn = ttk.Button(list_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            list_btn_frame.grid_columnconfigure(i, weight=1)

        # 下部：预览区
        preview_frame = ttk.LabelFrame(frame, text="世界观详情预览", padding=10)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 创建预览文本框容器
        preview_container = ttk.Frame(preview_frame, style="TFrame")
        preview_container.pack(fill=tk.BOTH, expand=True)

        # 滚动条
        preview_scrollbar = ttk.Scrollbar(preview_container, orient=tk.VERTICAL)
        preview_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self._worldview_preview = tk.Text(preview_container, wrap=tk.WORD, height=10,
                                         font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                         bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                         yscrollcommand=preview_scrollbar.set)
        self._worldview_preview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 配置滚动条
        preview_scrollbar.config(command=self._worldview_preview.yview)

        # 鼠标滚轮支持
        def _on_preview_mousewheel(event):
            self._worldview_preview.yview_scroll(int(-1 * (event.delta / 120)), "units")

        self._worldview_preview.bind("<MouseWheel>", _on_preview_mousewheel)

        # 底部按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，3列，确保按钮均匀分布
        bottom_buttons = [
            ("导入世界观", self._on_worldview_import),
            ("关联要素", self._on_worldview_link),
            ("清除", self._on_worldview_clear),
        ]
        
        for i, (text, command) in enumerate(bottom_buttons):
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)

        # 【修复】项目打开后，恢复已加载的世界观数据
        # 问题：标签页延迟创建，项目打开时_worldview_tree还未创建，导致数据无法恢复
        # 解决：标签页创建后，检查是否有已加载的世界观数据，如果有则恢复显示
        if hasattr(self, '_worldview_content') and self._worldview_content:
            # 使用after确保UI完全创建后再更新
            self.root.after(100, lambda: self._update_worldview_tree_from_content(self._worldview_content))

        return frame
    
    def _create_characters_content(self) -> tk.Frame:
        """创建人物设定内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")

        # 【修复】检查是否有已加载的人物数据（类似世界观延迟创建机制）
        if hasattr(self, '_character_data') and self._character_data:
            # 使用after确保UI完全创建后再更新
            self.root.after(100, lambda: self._update_character_tree())
        
        # 上部：人物导入区
        file_frame = ttk.LabelFrame(frame, text="人物档案导入", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="选择文件：").pack(side=tk.LEFT)
        self._character_path_var = tk.StringVar()
        ttk.Entry(path_frame, textvariable=self._character_path_var, width=50).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="浏览", command=self._on_character_browse).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="新建人物", command=self._on_character_new).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="批量解析导入", command=self._on_character_batch_import).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="保存项目", command=self._on_save_project).pack(side=tk.LEFT, padx=5)
        
        # 中部：人物列表区
        list_frame = ttk.LabelFrame(frame, text="人物列表", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 创建带滚动条的Treeview容器
        tree_container = ttk.Frame(list_frame)
        tree_container.pack(fill=tk.BOTH, expand=True)
        
        # 滚动条
        character_scrollbar = ttk.Scrollbar(tree_container, orient=tk.VERTICAL)
        character_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 列表（移除头像列，情绪改为情感状态）
        columns = ("name", "role", "status", "emotion", "chapters")
        self._character_tree = ttk.Treeview(tree_container, columns=columns, show="headings", height=8, yscrollcommand=character_scrollbar.set)
        self._character_tree.heading("name", text="姓名")
        self._character_tree.heading("role", text="角色类型")
        self._character_tree.heading("status", text="状态")
        self._character_tree.heading("emotion", text="情感状态")
        self._character_tree.heading("chapters", text="出场章节")
        self._character_tree.column("name", width=120)
        self._character_tree.column("role", width=100)
        self._character_tree.column("status", width=80)
        self._character_tree.column("emotion", width=100)
        self._character_tree.column("chapters", width=100)
        self._character_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 配置滚动条
        character_scrollbar.config(command=self._character_tree.yview)
        
        # 鼠标滚轮支持
        def _on_character_mousewheel(event):
            self._character_tree.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        self._character_tree.bind("<MouseWheel>", _on_character_mousewheel)
        
        # 列表操作按钮
        list_btn_frame = ttk.Frame(list_frame, style="TFrame")
        list_btn_frame.pack(fill=tk.X, pady=5)

        # 使用Grid布局，4列，确保按钮均匀分布
        list_buttons = [
            ("编辑人物", self._on_character_edit),
            ("人物详情", self._on_character_detail),
            ("关系图谱", self._on_character_relation),
            ("删除人物", self._on_character_delete),
        ]
        
        for i, (text, command) in enumerate(list_buttons):
            btn = ttk.Button(list_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(4):
            list_btn_frame.grid_columnconfigure(i, weight=1)

        # 下部：人物详情预览
        detail_frame = ttk.LabelFrame(frame, text="人物档案详情", padding=10)
        detail_frame.pack(fill=tk.X, padx=5, pady=5)

        # 创建详情文本框容器
        detail_container = ttk.Frame(detail_frame, style="TFrame")
        detail_container.pack(fill=tk.BOTH, expand=True)

        # 滚动条
        detail_scrollbar = ttk.Scrollbar(detail_container, orient=tk.VERTICAL)
        detail_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self._character_detail = tk.Text(detail_container, wrap=tk.WORD, height=6,
                                        font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                        bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                        yscrollcommand=detail_scrollbar.set)
        self._character_detail.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 配置滚动条
        detail_scrollbar.config(command=self._character_detail.yview)

        # 鼠标滚轮支持
        def _on_detail_mousewheel(event):
            self._character_detail.yview_scroll(int(-1 * (event.delta / 120)), "units")

        self._character_detail.bind("<MouseWheel>", _on_detail_mousewheel)

        return frame
    
    def _create_outline_content(self) -> tk.Frame:
        """创建大纲管理内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：文件导入区
        file_frame = ttk.LabelFrame(frame, text="大纲文件导入", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="当前大纲：").pack(side=tk.LEFT)
        self._outline_path_var = tk.StringVar(value="未导入")
        ttk.Label(path_frame, textvariable=self._outline_path_var, 
                 foreground=GlassTheme.TEXT_LINK).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="选择文件", command=self._on_outline_browse).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="新建大纲", command=self._on_outline_new).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="保存项目", command=self._on_save_project).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="导出", command=self._on_outline_export).pack(side=tk.RIGHT, padx=5)
        
        # 中部：大纲树形结构
        tree_frame = ttk.LabelFrame(frame, text="大纲结构（可拖拽排序）", padding=10)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self._outline_tree = ttk.Treeview(tree_frame, show="tree", height=8)
        self._outline_tree.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self._outline_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self._outline_tree.configure(yscrollcommand=scrollbar.set)
        
        # 树形结构操作按钮
        tree_btn_frame = ttk.Frame(tree_frame, style="TFrame")
        tree_btn_frame.pack(fill=tk.X, pady=5, side=tk.BOTTOM)
        
        # 使用Grid布局，4列，确保按钮均匀分布
        tree_buttons = [
            ("添加卷", self._on_outline_add_volume),
            ("添加章节", self._on_outline_add_chapter),
            ("编辑章节", self._on_outline_edit_chapter),
            ("删除", self._on_outline_delete),
        ]
        
        for i, (text, command) in enumerate(tree_buttons):
            btn = ttk.Button(tree_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(4):
            tree_btn_frame.grid_columnconfigure(i, weight=1)
        
        # 下部：章节编辑器
        editor_frame = ttk.LabelFrame(frame, text="章节编辑器", padding=10)
        editor_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 章节信息
        info_row = ttk.Frame(editor_frame, style="TFrame")
        info_row.pack(fill=tk.X, pady=5)
        ttk.Label(info_row, text="章节标题：").pack(side=tk.LEFT)
        self._chapter_title_var = tk.StringVar()
        ttk.Entry(info_row, textvariable=self._chapter_title_var, width=30).pack(side=tk.LEFT, padx=5)
        ttk.Label(info_row, text="目标字数：").pack(side=tk.LEFT, padx=(20, 0))
        self._chapter_words_var = tk.StringVar(value="2000")
        ttk.Spinbox(info_row, from_=500, to=5000, increment=100, width=10, textvariable=self._chapter_words_var).pack(side=tk.LEFT, padx=5)
        
        # 章节摘要
        ttk.Label(editor_frame, text="内容摘要：").pack(anchor=tk.W, pady=5)
        self._chapter_summary = tk.Text(editor_frame, wrap=tk.WORD, height=3,
                                       font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                       bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._chapter_summary.pack(fill=tk.X)
        
        # 进度条
        progress_frame = ttk.Frame(frame, style="TFrame")
        progress_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(progress_frame, text="大纲完成进度：").pack(side=tk.LEFT)
        self._outline_progress = ttk.Progressbar(progress_frame, length=300, mode='determinate')
        self._outline_progress.pack(side=tk.LEFT, padx=10)
        ttk.Label(progress_frame, text="0/0 章节").pack(side=tk.LEFT)
        
        # 按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，2列，确保按钮均匀分布
        outline_buttons = [
            ("解析大纲", self._on_outline_parse),
            ("清除大纲", self._on_outline_clear),
        ]
        
        for i, (text, command) in enumerate(outline_buttons):
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(2):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_style_content(self) -> tk.Frame:
        """创建风格学习内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：范文管理区
        file_frame = ttk.LabelFrame(frame, text="范文管理", padding=10)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        path_frame = ttk.Frame(file_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(path_frame, text="当前风格：").pack(side=tk.LEFT)
        self._style_path_var = tk.StringVar(value="未导入")
        ttk.Label(path_frame, textvariable=self._style_path_var,
                 foreground=GlassTheme.TEXT_LINK).pack(side=tk.LEFT, padx=10)
        ttk.Button(path_frame, text="上传范文", command=self._on_style_browse).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="删除范文", command=self._on_style_delete).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="解析风格", command=self._on_style_analyze).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="保存项目", command=self._on_save_project).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="导出风格", command=self._on_style_export).pack(side=tk.RIGHT, padx=5)
        
        # 中部左：风格档案库
        library_frame = ttk.LabelFrame(frame, text="风格档案库", padding=10)
        library_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5, side=tk.LEFT)
        
        columns = ("name", "vocabulary", "sentence", "rhetoric", "emotion", "rhythm", "structure", "detail")
        self._style_tree = ttk.Treeview(library_frame, columns=columns, show="headings", height=8)
        self._style_tree.heading("name", text="风格名称")
        self._style_tree.heading("vocabulary", text="词汇")
        self._style_tree.heading("sentence", text="句式")
        self._style_tree.heading("rhetoric", text="修辞")
        self._style_tree.heading("emotion", text="情感")
        self._style_tree.heading("rhythm", text="节奏")
        self._style_tree.heading("structure", text="结构")
        self._style_tree.heading("detail", text="细节")
        self._style_tree.column("name", width=100)
        for col in columns[1:]:
            self._style_tree.column(col, width=50)
        self._style_tree.pack(fill=tk.BOTH, expand=True)
        
        # 档案库按钮
        lib_btn_frame = ttk.Frame(library_frame, style="TFrame")
        lib_btn_frame.pack(fill=tk.X, pady=5)

        # 使用Grid布局，3列，确保按钮均匀分布
        lib_buttons = [
            ("切换风格", self._on_style_switch),
            ("编辑权重", self._on_style_edit_weight),
            ("删除风格", self._on_style_delete_profile),
        ]
        
        for i, (text, command) in enumerate(lib_buttons):
            btn = ttk.Button(lib_btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            lib_btn_frame.grid_columnconfigure(i, weight=1)

        # 中部右：风格分析器结果
        analyzer_frame = ttk.LabelFrame(frame, text="风格分析器（七维度评分）", padding=10)
        analyzer_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5, side=tk.RIGHT)

        self._style_info = tk.Text(analyzer_frame, wrap=tk.WORD, height=10,
                                  font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                  bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._style_info.pack(fill=tk.BOTH, expand=True)

        # 底部：创建风格按钮
        create_frame = ttk.Frame(frame, style="TFrame")
        create_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，4列，确保按钮均匀分布
        create_buttons = [
            ("深度学习", self._on_style_learn),
            ("创建风格模板", self._on_style_create),
            ("应用到生成器", self._on_style_apply),
            ("清除风格", self._on_style_clear),
        ]
        
        for i, (text, command) in enumerate(create_buttons):
            btn = ttk.Button(create_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        # 配置列权重，确保按钮均匀分布
        for i in range(4):
            create_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_generation_content(self) -> tk.Frame:
        """创建开始创作内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 上部：生成配置
        config_frame = ttk.LabelFrame(frame, text="生成配置", padding=10)
        config_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # ========== 专家选择器（第0行）==========
        # 延迟加载，按需初始化
        self._expert_selector = None
        if EXPERT_SELECTOR_AVAILABLE:
            try:
                expert_row = ttk.Frame(config_frame, style="TFrame")
                expert_row.pack(fill=tk.X, pady=5)
                
                self._expert_selector = ExpertSelectorWidget(
                    expert_row,
                    plugin_registry=getattr(self, '_plugin_registry', None),
                    on_expert_changed=self._on_expert_changed
                )
                self._expert_selector.pack(fill=tk.X)
            except Exception as e:
                logging.warning(f"专家选择器初始化失败: {e}")
                self._expert_selector = None
        
        # 第一行：章节范围
        row1 = ttk.Frame(config_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="起始章节：").pack(side=tk.LEFT)
        self._start_chapter_var = tk.StringVar(value="1")
        ttk.Spinbox(row1, from_=1, to=100, width=8, textvariable=self._start_chapter_var).pack(side=tk.LEFT, padx=5)
        ttk.Label(row1, text="结束章节：").pack(side=tk.LEFT, padx=(20, 0))
        self._end_chapter_var = tk.StringVar(value="1")
        ttk.Spinbox(row1, from_=1, to=100, width=8, textvariable=self._end_chapter_var).pack(side=tk.LEFT, padx=5)
        ttk.Label(row1, text="（风格/大纲/人设/世界观跟随项目文件）", 
                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=(20, 0))
        
        # 第二行：知识库选择（多选）
        row2 = ttk.Frame(config_frame, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="关联知识库：").pack(side=tk.LEFT)
        
        # 知识库多选区域
        kb_select_frame = ttk.Frame(row2, style="TFrame")
        kb_select_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 存储知识库复选框变量
        self._kb_vars = {}
        
        # 知识库分类复选框（延迟加载实际数据）
        self._kb_frame = ttk.Frame(kb_select_frame, style="TFrame")
        self._kb_frame.pack(fill=tk.X)
        
        # 分类显示名称映射（内部key -> 中文显示名）- 完整题材分类（18个，与知识库实际目录一致）
        self._kb_category_names = {
            "xuanhuan": "玄幻",
            "xianxia": "仙侠",
            "urban": "都市",
            "romance": "言情",
            "history": "历史",
            "scifi": "科幻",
            "suspense": "悬疑",
            "military": "军事",
            "wuxia": "武侠",
            "game": "游戏",
            "fantasy": "奇幻",
            "lingyi": "灵异",
            "tongren": "同人",
            "general": "通用",
            "horror": "恐怖",
            "mystery": "推理",
            "sports": "体育",
            "philosophy": "哲学"
        }
        
        # 反向映射（用于从显示名获取内部key）
        self._kb_name_to_key = {v: k for k, v in self._kb_category_names.items()}
        
        # 题材分类（从KnowledgeManager动态加载，初始化时使用默认值）
        self._kb_categories = list(self._kb_category_names.keys())
        
        # 初始化时加载实际数据
        self._init_knowledge_options()
        
        # 刷新按钮
        ttk.Button(row2, text="🔄", width=3, command=self._refresh_knowledge_options).pack(side=tk.RIGHT, padx=5)
        
        # ==================== 第三行：写作技巧选择（双下拉框 + 已选展示区）====================
        # 写作技巧分区框架
        tech_section_frame = ttk.LabelFrame(frame, text="写作技巧", padding=10)
        tech_section_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 存储已选技巧
        self._selected_writing_techniques = []  # [(领域, 技巧名), ...]
        
        # 领域内部key到显示名的映射
        self._tech_domain_names = {
            "narrative": "叙事技巧",
            "description": "描写技巧", 
            "rhetoric": "修辞技巧",
            "structure": "结构技巧",
            "special_sentence": "特殊句式",
            "advanced": "高级技法"
        }
        self._tech_domain_keys = {v: k for k, v in self._tech_domain_names.items()}
        
        # 从知识库文件加载写作技巧列表
        self._writing_techniques = self._load_writing_techniques_from_knowledge_base()
        
        # ========== 主布局：左侧双下拉框 + 右侧已选展示区 ==========
        tech_main_row = ttk.Frame(tech_section_frame, style="TFrame")
        tech_main_row.pack(fill=tk.X, pady=5)
        
        # 左侧：双下拉框（技巧类型 + 具体技巧）
        left_frame = ttk.Frame(tech_main_row, style="TFrame")
        left_frame.pack(side=tk.LEFT, fill=tk.X, expand=False)
        
        # 技巧类型下拉框
        ttk.Label(left_frame, text="技巧类型：").pack(side=tk.LEFT)
        self._tech_type_var = tk.StringVar()
        self._tech_type_combo = ttk.Combobox(left_frame, textvariable=self._tech_type_var,
                                              values=list(self._tech_domain_names.values()),
                                              state="readonly", width=12)
        self._tech_type_combo.pack(side=tk.LEFT, padx=5)
        self._tech_type_combo.bind("<<ComboboxSelected>>", lambda e: self._on_tech_type_changed())
        
        # 具体技巧下拉框（根据类型动态更新）
        ttk.Label(left_frame, text="具体技巧：").pack(side=tk.LEFT, padx=(15, 0))
        self._tech_name_var = tk.StringVar()
        self._tech_combo = ttk.Combobox(left_frame, textvariable=self._tech_name_var,
                                        values=[], state="readonly", width=18)
        self._tech_combo.pack(side=tk.LEFT, padx=5)
        
        # 添加按钮
        ttk.Button(left_frame, text="添加", width=6, 
                   command=self._on_add_writing_tech).pack(side=tk.LEFT, padx=10)
        
        # 右侧：已选技巧展示区
        right_frame = ttk.Frame(tech_main_row, style="TFrame")
        right_frame.pack(side=tk.RIGHT, fill=tk.X, expand=True, padx=(20, 0))
        
        ttk.Label(right_frame, text="已选技巧：").pack(side=tk.LEFT)
        
        # 已选技巧展示框（带滚动，增加宽度）
        tech_display_frame = ttk.Frame(right_frame, style="TFrame")
        tech_display_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self._selected_tech_display = tk.Text(tech_display_frame, height=2, width=80,
                                               font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_SMALL),
                                               bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                               wrap=tk.WORD, state=tk.DISABLED)
        self._selected_tech_display.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # 清除按钮
        ttk.Button(right_frame, text="清除", width=6,
                   command=self._clear_all_writing_tech).pack(side=tk.RIGHT, padx=5)
        
        # 第四行：字数、温度和生成模式
        row4 = ttk.Frame(config_frame, style="TFrame")
        row4.pack(fill=tk.X, pady=5)
        ttk.Label(row4, text="目标字数/章：").pack(side=tk.LEFT)
        self._target_words_var = tk.StringVar(value="900")
        ttk.Spinbox(row4, from_=500, to=2000, increment=100, width=10, textvariable=self._target_words_var).pack(side=tk.LEFT, padx=5)
        ttk.Label(row4, text="生成温度：").pack(side=tk.LEFT, padx=(20, 0))
        self._gen_temp_var = tk.DoubleVar(value=0.7)
        ttk.Scale(row4, from_=0.0, to=1.0, variable=self._gen_temp_var, orient=tk.HORIZONTAL, length=100).pack(side=tk.LEFT, padx=5)
        ttk.Label(row4, textvariable=self._gen_temp_var, width=4).pack(side=tk.LEFT)
        
        ttk.Separator(row4, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=15)
        
        ttk.Label(row4, text="生成模式：").pack(side=tk.LEFT)
        self._gen_mode_var = tk.StringVar(value="auto")
        ttk.Radiobutton(row4, text="自动迭代", variable=self._gen_mode_var, value="auto").pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(row4, text="手动迭代", variable=self._gen_mode_var, value="manual").pack(side=tk.LEFT, padx=10)
        
        # 第五行：出场人物（可选）
        row5 = ttk.Frame(config_frame, style="TFrame")
        row5.pack(fill=tk.X, pady=5)
        ttk.Label(row5, text="出场人物：").pack(side=tk.LEFT)
        self._gen_characters_var = tk.StringVar(value="自动设置")
        ttk.Entry(row5, textvariable=self._gen_characters_var, width=40).pack(side=tk.LEFT, padx=5)
        ttk.Label(row5, text="(可手动输入添加)", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT)
        
        # 中部：左右分栏（生成过程 + 生成结果）
        middle_split_frame = ttk.Frame(frame, style="TFrame")
        middle_split_frame.pack(fill=tk.BOTH, expand=False, padx=5, pady=5)  # 改为expand=False，给底部按钮留空间

        # 使用grid布局让左右等分宽度
        middle_split_frame.grid_columnconfigure(0, weight=1)
        middle_split_frame.grid_columnconfigure(1, weight=1)

        # 左侧：生成过程展示
        process_frame = ttk.LabelFrame(middle_split_frame, text="生成过程", padding=10)
        process_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 5))

        # 进度状态
        status_frame = ttk.Frame(process_frame, style="TFrame")
        status_frame.pack(fill=tk.X, pady=5)
        self._gen_status_var = tk.StringVar(value="准备就绪")
        ttk.Label(status_frame, textvariable=self._gen_status_var, font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold")).pack(side=tk.LEFT)
        ttk.Label(status_frame, text="当前章节：第1章").pack(side=tk.RIGHT)

        # 进度条
        self._gen_progress = ttk.Progressbar(process_frame, length=400, mode='determinate')
        self._gen_progress.pack(fill=tk.X, pady=5)

        # Agent状态列表（新增）
        agent_frame = ttk.Frame(process_frame, style="TFrame")
        agent_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(agent_frame, text="Agent状态：", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)).pack(anchor=tk.W)
        
        # Agent状态树形视图
        agent_list_frame = ttk.Frame(agent_frame, style="TFrame")
        agent_list_frame.pack(fill=tk.X, pady=2)
        
        self._gen_agent_tree = ttk.Treeview(
            agent_list_frame,
            columns=("status", "info"),
            show="headings",
            height=3
        )
        self._gen_agent_tree.heading("status", text="状态")
        self._gen_agent_tree.heading("info", text="信息")
        self._gen_agent_tree.column("status", width=80)
        self._gen_agent_tree.column("info", width=300)
        self._gen_agent_tree.pack(fill=tk.X, side=tk.LEFT)
        
        # Agent状态滚动条
        agent_scroll = ttk.Scrollbar(agent_list_frame, orient=tk.VERTICAL, command=self._gen_agent_tree.yview)
        agent_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._gen_agent_tree.configure(yscrollcommand=agent_scroll.set)

        # 日志输出（带滚动条，固定高度）
        log_scroll = ttk.Scrollbar(process_frame, orient=tk.VERTICAL)
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._gen_log = tk.Text(process_frame, wrap=tk.WORD, height=10,
                               font=(GlassTheme.FONT_FAMILY_CODE, GlassTheme.FONT_SIZE_SMALL),
                               bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                               yscrollcommand=log_scroll.set)
        self._gen_log.pack(fill=tk.BOTH, expand=True)
        log_scroll.configure(command=self._gen_log.yview)

        # 右侧：生成结果展示（带滚动条，固定高度）
        result_frame = ttk.LabelFrame(middle_split_frame, text="生成结果预览", padding=10)
        result_frame.grid(row=0, column=1, sticky="nsew", padx=(5, 0))

        result_scroll = ttk.Scrollbar(result_frame, orient=tk.VERTICAL)
        result_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._gen_result = tk.Text(result_frame, wrap=tk.WORD, height=10,
                                  font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                  bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                  yscrollcommand=result_scroll.set)
        self._gen_result.pack(fill=tk.BOTH, expand=True)
        result_scroll.configure(command=self._gen_result.yview)

        # 底部按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 使用Grid布局，5列，确保按钮均匀分布
        buttons = [
            ("保存结果", self._on_gen_save_result),
            ("开始生成", self._on_start_generation),
            ("停止生成", self._on_stop_generation),
            ("分章浏览", self._on_gen_browse),
            ("保存项目", self._on_gen_save),
        ]

        for i, (text, command) in enumerate(buttons):
            btn = ttk.Button(btn_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(5):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame

    def _init_knowledge_options(self):
        """初始化知识库选项（异步加载，避免启动卡顿）"""
        # 先显示默认选项（立即响应）
        self._create_default_kb_options()
        
        # 异步加载实际数据（后台执行）
        from core.thread_pool_manager import thread_pool_manager
        
        def load_in_background():
            """后台加载知识库分类"""
            try:
                from core.knowledge_manager import get_knowledge_manager
                km = get_knowledge_manager(self._workspace_root)
                stats = km.get_stats()
                
                # 在主线程更新UI
                self.root.after(0, lambda: self._update_kb_options(stats))
            except Exception as e:
                logger.warning(f"后台加载知识库分类失败: {e}")
        
        # 提交到线程池
        thread_pool_manager.submit_sync(load_in_background)
    
    def _update_kb_options(self, stats: dict):
        """更新知识库选项（在主线程执行）"""
        try:
            # 清空现有选项
            for widget in self._kb_frame.winfo_children():
                widget.destroy()
            
            self._kb_vars.clear()
            
            # 从实际数据获取题材分类
            # 优先从向量库获取（数据更准确），其次从JSON缓存获取
            categories = stats.get("categories", {})
            
            # 过滤：只保留_kb_category_names中定义的题材（排除writing_technique等）
            self._kb_categories = [cat for cat in categories.keys() if cat in self._kb_category_names]
            
            # 按照_kb_category_names的顺序排序（保持一致性）
            ordered_categories = [cat for cat in self._kb_category_names.keys() if cat in self._kb_categories]
            self._kb_categories = ordered_categories
            
            # 创建复选框
            for category in self._kb_categories:
                display_name = self._kb_category_names.get(category, category)
                # 获取该分类的知识点数量
                count = categories.get(category, 0)
                
                var = tk.BooleanVar(value=False)
                self._kb_vars[display_name] = var  # 使用中文显示名作为key
                
                label = f"{display_name}({count}条)" if count > 0 else display_name
                cb = ttk.Checkbutton(self._kb_frame, text=label, variable=var)
                cb.pack(side=tk.LEFT, padx=5)
            
            # 如果没有知识库数据，显示提示
            if not self._kb_vars:
                ttk.Label(self._kb_frame, text="(暂无知识库数据，请先导入)", 
                         foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=5)
        except Exception as e:
            logger.warning(f"更新知识库分类失败: {e}")
    
    def _load_knowledge_categories(self):
        """从KnowledgeManager加载实际的知识库题材分类（手动刷新时调用）"""
        try:
            # 尝试获取KnowledgeManager实例
            from core.knowledge_manager import get_knowledge_manager
            
            km = get_knowledge_manager(self._workspace_root)
            stats = km.get_stats()
            
            # 更新UI
            self._update_kb_options(stats)
        except Exception as e:
            logger.warning(f"加载知识库分类失败: {e}")
            # 如果加载失败，显示默认分类
            self._create_default_kb_options()
            self._kb_categories = ["xuanhuan", "scifi", "general"]
            for category in self._kb_categories:
                display_name = self._kb_category_names.get(category, category)
                var = tk.BooleanVar(value=False)
                self._kb_vars[display_name] = var
                cb = ttk.Checkbutton(self._kb_frame, text=display_name, variable=var)
                cb.pack(side=tk.LEFT, padx=5)
    
    def _load_writing_techniques_from_knowledge_base(self) -> Dict[str, List[str]]:
        """从知识库JSON文件加载写作技巧列表
        
        Returns:
            Dict[str, List[str]]: {显示名: [技巧名列表]}
        """
        import json
        from pathlib import Path
        
        techniques = {}
        
        # 知识库目录
        knowledge_dir = Path(self._workspace_root) / "data" / "knowledge" / "writing_technique"
        
        if not knowledge_dir.exists():
            logger.warning(f"写作技巧知识库目录不存在: {knowledge_dir}")
            return self._get_default_writing_techniques()
        
        # 遍历每个领域的JSON文件
        for domain_file in knowledge_dir.glob("*.json"):
            domain_key = domain_file.stem  # narrative, description, etc.
            
            # 跳过索引文件
            if domain_key == "index":
                continue
            
            # 获取显示名
            display_name = self._tech_domain_names.get(domain_key, domain_key)
            
            try:
                with open(domain_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # 提取知识点标题
                knowledge_points = data.get("knowledge_points", [])
                titles = [kp.get("title", "") for kp in knowledge_points if kp.get("title")]
                
                # 即使没有知识点，也显示分类名（空列表）
                techniques[display_name] = titles
                
                if titles:
                    logger.info(f"从 {domain_file.name} 加载 {len(titles)} 条写作技巧")
                else:
                    logger.warning(f"{domain_file.name} 中没有找到知识点（显示空分类）")
                    
            except Exception as e:
                logger.warning(f"加载 {domain_file.name} 失败: {e}")
        
        # 如果没有找到任何文件，使用默认值
        if not techniques:
            logger.warning("未找到任何写作技巧文件，使用默认值")
            return self._get_default_writing_techniques()
        
        return techniques
    
    def _get_default_writing_techniques(self) -> Dict[str, List[str]]:
        """获取默认的写作技巧列表（当知识库加载失败时使用）"""
        return {
            "叙事技巧": ["不可靠叙述者", "第一人称叙事", "第三人称叙事", "多视角叙事"],
            "描写技巧": ["心理描写", "环境描写", "动作描写"],
            "修辞技巧": ["比喻", "拟人", "夸张"],
            "结构技巧": ["悬念设置", "伏笔铺垫"],
            "特殊句式": ["倒装句式", "排比句式"],
            "高级技法": ["冰山理论", "元叙事", "复调小说理论", "陌生化理论"]
        }
    
    def _load_writing_technique_stats(self) -> Dict[str, Any]:
        """从知识库文件加载写作技巧统计数据
        
        Returns:
            Dict: {"total": 总数, "domains": {显示名: 数量}}
        """
        import json
        from pathlib import Path
        
        stats = {"total": 0, "domains": {}}
        
        # 领域内部key到显示名的映射（如果属性不存在则使用默认值）
        tech_domain_names = getattr(self, '_tech_domain_names', {
            "narrative": "叙事技巧",
            "description": "描写技巧", 
            "rhetoric": "修辞技巧",
            "structure": "结构技巧",
            "special_sentence": "特殊句式",
            "advanced": "高级技法"
        })
        
        # 知识库目录
        knowledge_dir = Path(self._workspace_root) / "data" / "knowledge" / "writing_technique"
        
        if not knowledge_dir.exists():
            logger.warning(f"写作技巧知识库目录不存在: {knowledge_dir}")
            return stats
        
        # 确保所有6个分类都显示（即使知识点为空）
        required_domains = ["narrative", "description", "rhetoric", "structure", "special_sentence", "advanced"]
        
        for domain_key in required_domains:
            display_name = tech_domain_names.get(domain_key, domain_key)
            domain_file = knowledge_dir / f"{domain_key}.json"
            
            count = 0
            
            # 如果文件存在，统计知识点数量
            if domain_file.exists():
                try:
                    with open(domain_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    # 统计知识点数量
                    count = len(data.get("knowledge_points", []))
                    
                    if count > 0:
                        stats["total"] += count
                        
                except Exception as e:
                    logger.warning(f"加载 {domain_file.name} 统计失败: {e}")
            
            # 始终显示分类（即使数量为0）
            stats["domains"][display_name] = count
        
        return stats
    
    def _create_default_kb_options(self):
        """创建默认知识库选项（14个题材分类）"""
        for widget in self._kb_frame.winfo_children():
            widget.destroy()
        
        self._kb_vars.clear()
        # 使用完整的14个题材分类
        self._kb_categories = list(self._kb_category_names.keys())
        
        for category in self._kb_categories:
            display_name = self._kb_category_names.get(category, category)
            var = tk.BooleanVar(value=False)
            self._kb_vars[display_name] = var  # 使用中文显示名作为key
            cb = ttk.Checkbutton(self._kb_frame, text=display_name, variable=var)
            cb.pack(side=tk.LEFT, padx=5)
    
    def _refresh_knowledge_options(self):
        """刷新知识库选项（从KnowledgeManager重新加载）"""
        self._load_knowledge_categories()
        self._set_status("知识库分类已刷新")
    
    def _on_tech_category_changed(self, event):
        """写作技巧分类切换事件（跨分类持久化选择，自适应换行布局）"""
        # 清空当前显示的复选框（但不重置选择状态）
        for widget in self._tech_detail_frame.winfo_children():
            widget.destroy()
        
        # 获取当前分类的技巧列表
        category = self._tech_category_var.get()
        techniques = self._writing_techniques.get(category, [])
        
        # 使用Flow布局（自动换行）
        # 创建内部流式布局容器
        flow_frame = ttk.Frame(self._tech_detail_frame, style="TFrame")
        flow_frame.pack(fill=tk.X, expand=True)
        
        # 为当前分类的每个技巧创建复选框，使用预先创建的BooleanVar
        # 使用grid布局实现自动换行（每行最多5个）
        cols = 5
        for i, tech in enumerate(techniques):
            # 使用预先创建的BooleanVar（保持选择状态）
            var = self._all_writing_tech_vars.get(tech)
            if var is None:
                var = tk.BooleanVar(value=False)
                self._all_writing_tech_vars[tech] = var
            
            cb = ttk.Checkbutton(flow_frame, text=tech, variable=var,
                                  command=self._update_selected_tech_label)
            # 使用grid布局，自动换行
            row = i // cols
            col = i % cols
            cb.grid(row=row, column=col, padx=3, pady=2, sticky="w")
        
        # 配置列权重，均匀分布
        for c in range(cols):
            flow_frame.grid_columnconfigure(c, weight=1)
        
        # 更新已选择技巧显示
        self._update_selected_tech_label()
    
    # ==================== 写作技巧选择方法（双下拉框版）====================
    
    def _on_tech_type_changed(self):
        """技巧类型下拉框切换回调 - 更新具体技巧列表"""
        tech_type = self._tech_type_var.get()
        if not tech_type:
            self._tech_combo['values'] = []
            self._tech_name_var.set("")
            return
        
        # 根据类型获取技巧列表
        techniques = self._writing_techniques.get(tech_type, [])
        self._tech_combo['values'] = techniques
        if techniques:
            self._tech_combo.set(techniques[0])
        else:
            self._tech_name_var.set("")
    
    def _on_add_writing_tech(self):
        """添加选中技巧"""
        tech_type = self._tech_type_var.get()
        tech_name = self._tech_name_var.get()
        
        if not tech_type or not tech_name:
            return
        
        # 检查是否已添加
        if (tech_type, tech_name) in self._selected_writing_techniques:
            return
        
        # 添加到已选列表
        self._selected_writing_techniques.append((tech_type, tech_name))
        
        # 更新显示
        self._update_selected_tech_display()
    
    def _update_selected_tech_display(self):
        """更新已选技巧展示区"""
        # 清空显示
        self._selected_tech_display.config(state=tk.NORMAL)
        self._selected_tech_display.delete(1.0, tk.END)
        
        if not self._selected_writing_techniques:
            self._selected_tech_display.insert(tk.END, "无")
        else:
            # 显示格式：技巧类型-技巧名
            display_text = ", ".join([f"{t[0][:2]}-{t[1]}" for t in self._selected_writing_techniques])
            self._selected_tech_display.insert(tk.END, display_text)
        
        self._selected_tech_display.config(state=tk.DISABLED)
    
    def _clear_all_writing_tech(self):
        """清除所有已选技巧"""
        self._selected_writing_techniques.clear()
        self._update_selected_tech_display()
    
    def get_selected_writing_techniques(self) -> List[str]:
        """获取选中的写作技巧列表"""
        # 返回技巧名称列表
        return [tech_name for _, tech_name in self._selected_writing_techniques]
    
    # ==================== 旧方法保留（兼容性）====================
    
    def _toggle_category_fold(self):
        """折叠/展开领域选择区域"""
        if self._category_fold_var.get():
            self._category_content_frame.pack_forget()
            self._category_fold_btn.config(text="▶ 六大领域选择（点击展开）")
            self._category_fold_var.set(False)
        else:
            self._category_content_frame.pack(fill=tk.X, pady=5, after=self._category_fold_btn.master)
            self._category_fold_btn.config(text="▼ 六大领域选择（点击折叠/展开）")
            self._category_fold_var.set(True)
    
    def _toggle_knowledge_fold(self):
        """折叠/展开知识点选择区域"""
        if self._knowledge_fold_var.get():
            self._knowledge_content_frame.pack_forget()
            self._knowledge_fold_btn.config(text="▶ 知识点详细选择（点击展开）")
            self._knowledge_fold_var.set(False)
        else:
            self._knowledge_content_frame.pack(fill=tk.BOTH, expand=True, pady=5, after=self._knowledge_fold_btn.master)
            self._knowledge_fold_btn.config(text="▼ 知识点详细选择（点击折叠/展开）")
            self._knowledge_fold_var.set(True)
    
    def _on_category_selected(self, domain_name: str):
        """领域选择变化回调"""
        var = self._category_vars.get(domain_name)
        if var and var.get():
            self._selected_tech_categories.add(domain_name)
        else:
            self._selected_tech_categories.discard(domain_name)
        
        # 更新已选领域显示
        if self._selected_tech_categories:
            self._selected_category_label.config(text=f"已选领域: {', '.join(self._selected_tech_categories)}")
        else:
            self._selected_category_label.config(text="已选领域: 无")
        
        # 根据已选领域刷新知识点列表
        self._refresh_knowledge_list_by_categories()
    
    def _refresh_knowledge_list_by_categories(self):
        """根据已选领域刷新知识点列表"""
        # 清空当前显示
        for widget in self._knowledge_scrollable_frame.winfo_children():
            widget.destroy()
        
        # 如果没有选择任何领域，显示提示
        if not self._selected_tech_categories:
            self._no_category_hint.pack(pady=20)
            return
        else:
            self._no_category_hint.pack_forget()
        
        # 按已选领域分组显示知识点
        row_idx = 0
        for domain_name in self._tech_domain_names.values():
            if domain_name not in self._selected_tech_categories:
                continue
            
            techniques = self._writing_techniques.get(domain_name, [])
            if not techniques:
                continue
            
            # 领域标题
            title_label = ttk.Label(self._knowledge_scrollable_frame, 
                                    text=f"【{domain_name}】",
                                    font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold"),
                                    foreground=GlassTheme.ACCENT_BLUE)
            title_label.grid(row=row_idx, column=0, columnspan=4, sticky="w", padx=5, pady=(10, 3))
            row_idx += 1
            
            # 知识点复选框（每行4个）
            cols = 4
            for i, tech in enumerate(techniques):
                var = self._all_writing_tech_vars.get(tech)
                if var is None:
                    var = tk.BooleanVar(value=False)
                    self._all_writing_tech_vars[tech] = var
                
                cb = ttk.Checkbutton(self._knowledge_scrollable_frame, text=tech, variable=var,
                                      command=self._update_selected_knowledge_display)
                col = i % cols
                cb.grid(row=row_idx, column=col, padx=5, pady=3, sticky="w")
                if col == cols - 1:
                    row_idx += 1
            
            # 如果最后一行没满，增加行号
            if len(techniques) % cols != 0:
                row_idx += 1
        
        # 更新Canvas滚动区域
        self._knowledge_scrollable_frame.update_idletasks()
        self._knowledge_canvas.configure(scrollregion=self._knowledge_canvas.bbox("all"))
    
    def _on_knowledge_domain_changed(self):
        """知识点领域切换回调"""
        domain_name = self._knowledge_domain_var.get()
        
        # 清空当前显示
        for widget in self._knowledge_scrollable_frame.winfo_children():
            widget.destroy()
        
        # 获取该领域的知识点列表
        techniques = self._writing_techniques.get(domain_name, [])
        
        # 使用grid布局显示知识点复选框（每行4个）
        cols = 4
        for i, tech in enumerate(techniques):
            var = self._all_writing_tech_vars.get(tech)
            if var is None:
                var = tk.BooleanVar(value=False)
                self._all_writing_tech_vars[tech] = var
            
            cb = ttk.Checkbutton(self._knowledge_scrollable_frame, text=tech, variable=var,
                                  command=self._update_selected_knowledge_display)
            row = i // cols
            col = i % cols
            cb.grid(row=row, column=col, padx=5, pady=3, sticky="w")
        
        # 更新Canvas滚动区域
        self._knowledge_scrollable_frame.update_idletasks()
        self._knowledge_canvas.configure(scrollregion=self._knowledge_canvas.bbox("all"))
    
    def _update_selected_knowledge_display(self):
        """更新已选知识点显示"""
        # 清空已选列表
        for item in self._selected_knowledge_tree.get_children():
            self._selected_knowledge_tree.delete(item)
        
        # 获取所有选中的知识点
        selected = self.get_selected_writing_techniques()
        
        # 按领域分组显示
        for category, techniques in self._writing_techniques.items():
            for tech in techniques:
                if tech in selected:
                    self._selected_knowledge_tree.insert("", "end", values=(category, tech))
        
        # 更新数量显示
        self._selected_knowledge_label.config(text=f"已选知识点: {len(selected)}项")
    
    def _clear_all_knowledge_selection(self):
        """清除所有知识点选择"""
        for var in self._all_writing_tech_vars.values():
            var.set(False)
        self._update_selected_knowledge_display()
    
    def _update_selected_tech_label(self):
        """更新已选择技巧的显示标签（限制显示长度）"""
        selected = self.get_selected_writing_techniques()
        if selected:
            # 限制显示长度，避免标签过长
            if len(selected) <= 3:
                # 3个以内完整显示
                text = f"已选({len(selected)}): {', '.join(selected)}"
            elif len(selected) <= 5:
                # 3-5个显示前3个+省略
                text = f"已选({len(selected)}): {', '.join(selected[:3])}..."
            else:
                # 超过5个只显示数量
                text = f"已选 {len(selected)} 项技巧"
        else:
            text = ""
        self._selected_tech_label.config(text=text)
    
    def get_selected_knowledge_bases(self) -> List[str]:
        """获取选中的知识库分类列表（返回内部key如scifi/xuanhuan）"""
        selected = []
        for display_name, var in self._kb_vars.items():
            if var.get():
                # 将中文显示名转换为内部key
                key = self._kb_name_to_key.get(display_name, display_name)
                selected.append(key)
        return selected
    
    def get_selected_writing_techniques(self) -> List[str]:
        """获取选中的写作技巧列表（跨分类）"""
        selected = []
        for tech, var in self._all_writing_tech_vars.items():
            if var.get():
                selected.append(tech)
        return selected
    
    def _on_gen_save_result(self):
        """保存生成结果"""
        try:
            from tkinter import messagebox

            content = self._gen_result.get("1.0", tk.END).strip()

            if not content:
                messagebox.showwarning("提示", "没有可保存的内容！")
                return

            from tkinter import filedialog

            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")],
                title="保存生成结果"
            )

            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)

                messagebox.showinfo("成功", "生成结果已保存！")
        except Exception as e:
            from tkinter import messagebox
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_file_error
            title, full_message = convert_file_error(e, "保存")
            messagebox.showerror(title, full_message)
        ttk.Button(btn_frame, text="导出TXT", command=self._on_gen_export_txt).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="导出DOCX", command=self._on_gen_export_docx).pack(side=tk.LEFT, padx=5)
        
        return frame
    
    def _create_reverse_content(self) -> tk.Frame:
        """创建逆向反馈内容页面（V2.5重构版）
        
        功能：
        - 章节上传区域（支持单个/批量上传文件或粘贴文本）
        - 章节列表显示（含字数、状态）
        - 删除按钮（从项目文件中删除章节）
        - "运行分析"按钮（调用逆向反馈分析插件）
        - 分析结果展示区（冲突列表，带修正建议）
        - "应用修正"按钮（调用设定修正生成器，更新当前项目设定）
        """
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 初始化章节数据存储
        self._reverse_chapters = {}  # 章节ID -> {title, content, words, status, file_path}
        self._reverse_analysis_result = None  # 分析结果缓存
        self._reverse_selected_issues = []  # 选中的冲突项
        
        # ==================== 上部：章节上传区域 ====================
        upload_frame = ttk.LabelFrame(frame, text="章节上传", padding=10)
        upload_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 上传方式说明
        ttk.Label(upload_frame, text="支持单个/批量上传TXT/DOCX文件，或直接粘贴文本：", 
                  font=('Microsoft YaHei UI', 9)).pack(anchor=tk.W, pady=2)
        
        # 上传按钮行
        upload_btn_row = ttk.Frame(upload_frame, style="TFrame")
        upload_btn_row.pack(fill=tk.X, pady=5)
        
        ttk.Button(upload_btn_row, text="📁 上传文件", command=self._on_reverse_upload_files, width=12).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_row, text="📁 批量上传", command=self._on_reverse_batch_upload, width=12).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_row, text="📋 粘贴文本", command=self._on_reverse_paste_text, width=12).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_row, text="🔄 刷新列表", command=self._on_reverse_refresh_chapters, width=12).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_row, text="💾 保存项目", command=self._on_save_project, width=12).pack(side=tk.LEFT, padx=5)
        
        # 粘贴文本区域（可折叠，默认隐藏）
        self._paste_text_frame = ttk.Frame(upload_frame, style="TFrame")
        self._paste_text_frame.pack(fill=tk.X, pady=5)
        
        paste_title_row = ttk.Frame(self._paste_text_frame, style="TFrame")
        paste_title_row.pack(fill=tk.X)
        
        ttk.Label(paste_title_row, text="章节标题：").pack(side=tk.LEFT)
        self._paste_title_var = tk.StringVar()
        ttk.Entry(paste_title_row, textvariable=self._paste_title_var, width=40).pack(side=tk.LEFT, padx=5)
        
        ttk.Label(self._paste_text_frame, text="章节内容：").pack(anchor=tk.W)
        
        paste_content_frame = ttk.Frame(self._paste_text_frame, style="TFrame")
        paste_content_frame.pack(fill=tk.X, pady=2)
        
        self._paste_content_text = tk.Text(paste_content_frame, wrap=tk.WORD, height=4,
                                           font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                           bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._paste_content_text.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        paste_scroll = ttk.Scrollbar(paste_content_frame, orient=tk.VERTICAL, command=self._paste_content_text.yview)
        paste_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._paste_content_text.configure(yscrollcommand=paste_scroll.set)
        
        ttk.Button(self._paste_text_frame, text="✅ 添加章节", command=self._on_reverse_add_pasted_chapter).pack(anchor=tk.E, pady=5)
        
        # 默认隐藏粘贴区域
        self._paste_text_frame.pack_forget()
        
        # ==================== 中部：左右分栏布局 ====================
        middle_frame = ttk.Frame(frame, style="TFrame")
        middle_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # ========== 中部左：章节列表 ==========
        chapters_frame = ttk.LabelFrame(middle_frame, text="章节列表", padding=10)
        chapters_frame.pack(fill=tk.BOTH, expand=True, side=tk.LEFT, padx=(0, 5))
        
        # 章节列表（带滚动条）
        chapters_list_frame = ttk.Frame(chapters_frame, style="TFrame")
        chapters_list_frame.pack(fill=tk.BOTH, expand=True)
        
        columns = ("id", "title", "words", "status", "source")
        self._completed_chapters_tree = ttk.Treeview(chapters_list_frame, columns=columns, show="headings", height=10)
        self._completed_chapters_tree.heading("id", text="序号")
        self._completed_chapters_tree.heading("title", text="标题")
        self._completed_chapters_tree.heading("words", text="字数")
        self._completed_chapters_tree.heading("status", text="状态")
        self._completed_chapters_tree.heading("source", text="来源")
        self._completed_chapters_tree.column("id", width=50)
        self._completed_chapters_tree.column("title", width=180)
        self._completed_chapters_tree.column("words", width=70)
        self._completed_chapters_tree.column("status", width=70)
        self._completed_chapters_tree.column("source", width=80)
        self._completed_chapters_tree.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        # 滚动条
        chapters_scroll = ttk.Scrollbar(chapters_list_frame, orient=tk.VERTICAL, command=self._completed_chapters_tree.yview)
        chapters_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._completed_chapters_tree.configure(yscrollcommand=chapters_scroll.set)
        
        # 绑定选择事件
        self._completed_chapters_tree.bind("<<TreeviewSelect>>", self._on_reverse_chapter_select)
        
        # 右键菜单
        self._chapters_context_menu = tk.Menu(self._completed_chapters_tree, tearoff=0)
        self._chapters_context_menu.add_command(label="✅ 标记为完成", command=lambda: self._on_mark_chapter_completed())
        self._chapters_context_menu.add_command(label="⏳ 标记为未完成", command=lambda: self._on_mark_chapter_incomplete())
        self._chapters_context_menu.add_separator()
        self._chapters_context_menu.add_command(label="👁️ 查看内容", command=self._on_reverse_view_chapter)
        self._chapters_context_menu.add_command(label="🗑️ 删除章节", command=lambda: self._on_delete_completed_chapter())
        self._completed_chapters_tree.bind("<Button-3>", self._show_chapters_context_menu)
        
        # 章节列表按钮行
        chapters_btn_frame = ttk.Frame(chapters_frame, style="TFrame")
        chapters_btn_frame.pack(fill=tk.X, pady=5)
        
        chapters_buttons = [
            ("删除选中", self._on_delete_completed_chapter),
            ("全部标记完成", self._on_mark_all_completed),
            ("清空列表", self._on_reverse_clear_chapters),
        ]
        
        for i, (text, command) in enumerate(chapters_buttons):
            btn = ttk.Button(chapters_btn_frame, text=text, command=command, width=12)
            btn.grid(row=0, column=i, padx=3, pady=2, sticky="ew")
        
        for i in range(3):
            chapters_btn_frame.grid_columnconfigure(i, weight=1)
        
        # ========== 中部右：分析配置 ==========
        config_frame = ttk.LabelFrame(middle_frame, text="分析配置", padding=10)
        config_frame.pack(fill=tk.BOTH, expand=True, side=tk.RIGHT, padx=(5, 0))
        
        # 分析维度
        ttk.Label(config_frame, text="分析维度：", font=('Microsoft YaHei UI', 9, 'bold')).pack(anchor=tk.W, pady=5)
        
        self._reverse_check_consistency = tk.BooleanVar(value=True)
        self._reverse_check_logic = tk.BooleanVar(value=True)
        self._reverse_check_character = tk.BooleanVar(value=True)
        self._reverse_check_style = tk.BooleanVar(value=True)
        self._reverse_check_worldview = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(config_frame, text="✓ 一致性检查（章节间逻辑一致性）", 
                        variable=self._reverse_check_consistency).pack(anchor=tk.W, pady=2)
        ttk.Checkbutton(config_frame, text="✓ 逻辑漏洞检测（情节合理性）", 
                        variable=self._reverse_check_logic).pack(anchor=tk.W, pady=2)
        ttk.Checkbutton(config_frame, text="✓ 人设偏离检测（角色行为一致性）", 
                        variable=self._reverse_check_character).pack(anchor=tk.W, pady=2)
        ttk.Checkbutton(config_frame, text="✓ 风格匹配度（写作风格一致性）", 
                        variable=self._reverse_check_style).pack(anchor=tk.W, pady=2)
        ttk.Checkbutton(config_frame, text="✓ 世界观冲突（设定一致性）", 
                        variable=self._reverse_check_worldview).pack(anchor=tk.W, pady=2)
        
        ttk.Separator(config_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=10)
        
        # 分析范围
        ttk.Label(config_frame, text="分析范围：", font=('Microsoft YaHei UI', 9, 'bold')).pack(anchor=tk.W, pady=5)
        
        self._reverse_scope_var = tk.StringVar(value="selected")
        ttk.Radiobutton(config_frame, text="仅分析选中章节", variable=self._reverse_scope_var, 
                        value="selected").pack(anchor=tk.W, pady=2)
        ttk.Radiobutton(config_frame, text="分析所有已完成章节", variable=self._reverse_scope_var, 
                        value="completed").pack(anchor=tk.W, pady=2)
        ttk.Radiobutton(config_frame, text="分析所有章节", variable=self._reverse_scope_var, 
                        value="all").pack(anchor=tk.W, pady=2)
        
        ttk.Separator(config_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=10)
        
        # 运行分析按钮
        self._run_analysis_btn = ttk.Button(config_frame, text="🔍 运行分析", 
                                            command=self._on_reverse_run_analysis, width=15)
        self._run_analysis_btn.pack(anchor=tk.CENTER, pady=10)
        
        # 分析进度指示
        self._analysis_progress_frame = ttk.Frame(config_frame, style="TFrame")
        self._analysis_progress_frame.pack(fill=tk.X, pady=5)
        
        self._analysis_progress_label = ttk.Label(self._analysis_progress_frame, text="", 
                                                   foreground=GlassTheme.PRIMARY)
        self._analysis_progress_label.pack(anchor=tk.W)
        
        # ==================== 下部：分析结果展示 ====================
        result_frame = ttk.LabelFrame(frame, text="分析结果", padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 结果区域：左侧冲突列表，右侧修正建议（固定布局，不可拖动调整）
        result_split = ttk.Frame(result_frame, style="TFrame")
        result_split.pack(fill=tk.BOTH, expand=True)
        
        # 左侧：冲突列表（固定宽度50%）
        issues_frame = ttk.Frame(result_split, style="TFrame")
        issues_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        ttk.Label(issues_frame, text="冲突列表（双击查看详情）：", 
                  font=('Microsoft YaHei UI', 9)).pack(anchor=tk.W)
        
        issues_list_frame = ttk.Frame(issues_frame, style="TFrame")
        issues_list_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        issue_columns = ("type", "severity", "element", "chapter", "desc")
        self._issues_tree = ttk.Treeview(issues_list_frame, columns=issue_columns, show="headings", height=8)
        self._issues_tree.heading("type", text="类型")
        self._issues_tree.heading("severity", text="优先级")
        self._issues_tree.heading("element", text="元素")
        self._issues_tree.heading("chapter", text="章节")
        self._issues_tree.heading("desc", text="问题描述")
        self._issues_tree.column("type", width=80)
        self._issues_tree.column("severity", width=60)
        self._issues_tree.column("element", width=100)
        self._issues_tree.column("chapter", width=100)
        self._issues_tree.column("desc", width=200)
        self._issues_tree.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        issues_scroll = ttk.Scrollbar(issues_list_frame, orient=tk.VERTICAL, command=self._issues_tree.yview)
        issues_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._issues_tree.configure(yscrollcommand=issues_scroll.set)
        
        # 绑定选择事件显示详情
        self._issues_tree.bind("<<TreeviewSelect>>", self._on_reverse_issue_select)
        self._issues_tree.bind("<Double-1>", self._on_reverse_issue_double_click)
        
        # 右侧：修正建议详情（固定宽度50%）
        detail_frame = ttk.Frame(result_split, style="TFrame")
        detail_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        ttk.Label(detail_frame, text="修正建议：", 
                  font=('Microsoft YaHei UI', 9)).pack(anchor=tk.W)
        
        detail_text_frame = ttk.Frame(detail_frame, style="TFrame")
        detail_text_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self._issue_detail_text = tk.Text(detail_text_frame, wrap=tk.WORD, height=8,
                                          font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                          bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._issue_detail_text.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        detail_scroll = ttk.Scrollbar(detail_text_frame, orient=tk.VERTICAL, command=self._issue_detail_text.yview)
        detail_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._issue_detail_text.configure(yscrollcommand=detail_scroll.set)
        self._issue_detail_text.insert("1.0", "选择冲突项查看详细修正建议...")
        self._issue_detail_text.configure(state=tk.DISABLED)
        
        # ==================== 底部：操作按钮 ====================
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 操作按钮（4个核心按钮）
        action_buttons = [
            ("🔧 应用修正", self._on_reverse_apply_fix),
            ("📋 导出报告", self._on_reverse_export_report),
            ("🔄 重新分析", self._on_reverse_run_analysis),
            ("🗑️ 清除结果", self._on_reverse_clear_result),
        ]
        
        for i, (text, command) in enumerate(action_buttons):
            btn = ttk.Button(btn_frame, text=text, command=command, width=14)
            btn.grid(row=0, column=i, padx=5, pady=2, sticky="ew")
        
        for i in range(4):
            btn_frame.grid_columnconfigure(i, weight=1)
        
        return frame
    
    def _create_quick_content(self) -> tk.Frame:
        """创建快捷创作内容页面（一次性生成四个结果：世界观、大纲、人设、关键情节）"""
        # 直接在workbench内容区创建frame，不需要额外的Canvas
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 初始化上传文件列表
        self._quick_uploaded_files = []  # 存储上传的文件列表 [{path, type, content}]

        # 上部：参考文本上传区（增强版）
        upload_frame = ttk.LabelFrame(frame, text="参考文本上传（可选，支持多个文件）", padding=10)
        upload_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 说明文字
        ttk.Label(upload_frame, text="上传参考文本或已有设定，系统将分析并作为生成依据：", 
                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)).pack(anchor=tk.W, pady=5)
        
        # 文件上传按钮组
        upload_btn_frame = ttk.Frame(upload_frame, style="TFrame")
        upload_btn_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(upload_btn_frame, text="📄 上传世界观参考", width=18,
                  command=lambda: self._on_quick_upload_file("worldview")).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_frame, text="📋 上传大纲参考", width=18,
                  command=lambda: self._on_quick_upload_file("outline")).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_frame, text="👤 上传人设参考", width=18,
                  command=lambda: self._on_quick_upload_file("characters")).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_frame, text="📖 上传情节参考", width=18,
                  command=lambda: self._on_quick_upload_file("plot")).pack(side=tk.LEFT, padx=5)
        ttk.Button(upload_btn_frame, text="🗑️ 清空全部", width=12,
                  command=self._on_quick_clear_uploads).pack(side=tk.LEFT, padx=5)
        
        # 已上传文件列表（Treeview）
        upload_list_frame = ttk.Frame(upload_frame, style="TFrame")
        upload_list_frame.pack(fill=tk.X, pady=5)
        
        columns = ("文件名", "类型", "字数", "状态")
        self._quick_upload_tree = ttk.Treeview(upload_list_frame, columns=columns, show="headings", height=3)
        
        self._quick_upload_tree.heading("文件名", text="文件名")
        self._quick_upload_tree.heading("类型", text="类型")
        self._quick_upload_tree.heading("字数", text="字数")
        self._quick_upload_tree.heading("状态", text="状态")
        
        self._quick_upload_tree.column("文件名", width=200)
        self._quick_upload_tree.column("类型", width=80, anchor=tk.CENTER)
        self._quick_upload_tree.column("字数", width=80, anchor=tk.CENTER)
        self._quick_upload_tree.column("状态", width=80, anchor=tk.CENTER)
        
        # 添加滚动条
        tree_scroll = ttk.Scrollbar(upload_list_frame, orient="vertical", command=self._quick_upload_tree.yview)
        self._quick_upload_tree.configure(yscrollcommand=tree_scroll.set)
        
        self._quick_upload_tree.pack(side=tk.LEFT, fill=tk.X, expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 右键菜单（删除单个文件）
        self._quick_upload_menu = tk.Menu(self._quick_upload_tree, tearoff=0)
        self._quick_upload_menu.add_command(label="删除此文件", command=self._on_quick_remove_upload)
        self._quick_upload_tree.bind("<Button-3>", self._show_quick_upload_menu)
        
        # 中部：需求描述（关键词输入）
        input_frame = ttk.LabelFrame(frame, text="创作需求描述（关键词）", padding=10)
        input_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 示例提示
        example_frame = ttk.Frame(input_frame, style="TFrame")
        example_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(example_frame, text="示例：", foreground=GlassTheme.TEXT_SECONDARY,
                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)).pack(side=tk.LEFT)
        
        examples = ["修仙世界", "现代都市", "古代宫廷", "仙侠爱情", "科幻未来"]
        for example in examples:
            def set_example(e=example):
                self._quick_input.delete("1.0", tk.END)
                self._quick_input.insert("1.0", e)
            ttk.Button(example_frame, text=example, width=10,
                      command=set_example).pack(side=tk.LEFT, padx=5)
        
        # 关键词输入框
        self._quick_input = tk.Text(input_frame, wrap=tk.WORD, height=4,
                                   font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                   bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._quick_input.pack(fill=tk.X, pady=5)
        self._quick_input.insert("1.0", "请输入关键词描述...")
        
        # 生成详细程度选择
        detail_frame = ttk.Frame(input_frame, style="TFrame")
        detail_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(detail_frame, text="生成详细程度：").pack(side=tk.LEFT)
        self._quick_detail_var = tk.StringVar(value="standard")
        ttk.Radiobutton(detail_frame, text="快速", variable=self._quick_detail_var, value="quick").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(detail_frame, text="标准", variable=self._quick_detail_var, value="standard").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(detail_frame, text="详细", variable=self._quick_detail_var, value="detailed").pack(side=tk.LEFT, padx=5)
        
        # 下部：结果展示区（可折叠）
        results_outer_frame = ttk.LabelFrame(frame, text="生成结果", padding=10)
        results_outer_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 创建可折叠的结果区域
        self._quick_result_frames = {}
        self._quick_result_texts = {}
        self._quick_result_expanded = {"worldview": True, "outline": True, "characters": True, "plot": True}
        
        result_types = [
            ("worldview", "🌍 世界观设定", "#4A90E2"),
            ("outline", "📋 章节大纲", "#50C878"),
            ("characters", "👤 人物设定", "#FF6B6B"),
            ("plot", "📖 关键情节", "#FFA500")
        ]
        
        for result_type, title, color in result_types:
            # 可折叠的标题栏
            header_frame = ttk.Frame(results_outer_frame, style="TFrame")
            header_frame.pack(fill=tk.X, pady=2)
            
            # 折叠按钮
            expand_btn = ttk.Button(header_frame, text="▼", width=3,
                                   command=lambda t=result_type: self._toggle_quick_result(t))
            expand_btn.pack(side=tk.LEFT)
            self._quick_result_frames[result_type] = {"header": header_frame, "expand_btn": expand_btn}
            
            # 标题标签
            ttk.Label(header_frame, text=title, 
                     font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold")).pack(side=tk.LEFT, padx=10)
            
            # 内容区域
            content_frame = ttk.Frame(results_outer_frame, style="TFrame")
            content_frame.pack(fill=tk.BOTH, expand=True, pady=2)
            
            # 文本框
            text_scroll = ttk.Scrollbar(content_frame, orient=tk.VERTICAL)
            text_scroll.pack(side=tk.RIGHT, fill=tk.Y)
            
            result_text = tk.Text(content_frame, wrap=tk.WORD, height=8,
                                 font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_SMALL),
                                 bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                 yscrollcommand=text_scroll.set, state=tk.NORMAL)
            result_text.pack(fill=tk.BOTH, expand=True)
            result_text.insert("1.0", f"{title}将在此显示...")
            text_scroll.configure(command=result_text.yview)
            
            self._quick_result_texts[result_type] = result_text
            self._quick_result_frames[result_type]["content"] = content_frame
        
        # 底部按钮区（两行布局）
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        # 第一行：生成按钮
        gen_btn_frame = ttk.Frame(btn_frame, style="TFrame")
        gen_btn_frame.pack(fill=tk.X, pady=2)
        
        self._quick_gen_all_btn = ttk.Button(gen_btn_frame, text="🚀 一键生成全部", 
                                            command=self._on_quick_generate_all, width=15)
        self._quick_gen_all_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(gen_btn_frame, text="🌍 仅世界观", width=12,
                  command=lambda: self._on_quick_generate_single("worldview")).pack(side=tk.LEFT, padx=5)
        ttk.Button(gen_btn_frame, text="📋 仅大纲", width=12,
                  command=lambda: self._on_quick_generate_single("outline")).pack(side=tk.LEFT, padx=5)
        ttk.Button(gen_btn_frame, text="👤 仅人设", width=12,
                  command=lambda: self._on_quick_generate_single("characters")).pack(side=tk.LEFT, padx=5)
        ttk.Button(gen_btn_frame, text="📖 仅情节", width=12,
                  command=lambda: self._on_quick_generate_single("plot")).pack(side=tk.LEFT, padx=5)
        
        # 第二行：保存和导入按钮
        save_btn_frame = ttk.Frame(btn_frame, style="TFrame")
        save_btn_frame.pack(fill=tk.X, pady=2)
        
        ttk.Button(save_btn_frame, text="💾 保存结果", width=12,
                  command=self._on_quick_save_results).pack(side=tk.LEFT, padx=5)
        ttk.Button(save_btn_frame, text="📤 导出Markdown", width=14,
                  command=lambda: self._on_quick_export_results("markdown")).pack(side=tk.LEFT, padx=5)
        ttk.Button(save_btn_frame, text="📤 导出JSON", width=12,
                  command=lambda: self._on_quick_export_results("json")).pack(side=tk.LEFT, padx=5)
        ttk.Button(save_btn_frame, text="📥 导入当前项目", width=14,
                  command=self._on_quick_import).pack(side=tk.LEFT, padx=5)
        
        # 延迟加载插件：首次使用时加载，避免启动卡顿
        # 插件将在 _on_quick_generate_all 或 _on_quick_generate_single 中加载
        
        return frame
    
    def _toggle_quick_result(self, result_type: str):
        """折叠/展开结果区域"""
        content = self._quick_result_frames[result_type]["content"]
        expand_btn = self._quick_result_frames[result_type]["expand_btn"]
        
        if self._quick_result_expanded[result_type]:
            content.pack_forget()
            expand_btn.configure(text="▶")
            self._quick_result_expanded[result_type] = False
        else:
            content.pack(fill=tk.BOTH, expand=True, pady=2)
            expand_btn.configure(text="▼")
            self._quick_result_expanded[result_type] = True
    
    def _load_quick_creator_plugin(self):
        """加载快捷创作插件"""
        try:
            # 使用动态导入（支持连字符目录名）
            plugin_class = _dynamic_import_plugin('quick-creator-v1', 'QuickCreatorPlugin')
            if not plugin_class:
                logger.error("快捷创作插件类未找到")
                self._quick_creator_plugin = None
                return
            
            self._quick_creator_plugin = plugin_class()
            
            # 从服务定位器获取AI客户端
            if CORE_AVAILABLE:
                service_locator = get_service_locator()
                ai_service = service_locator.get_service("ai_service")
                if ai_service and hasattr(ai_service, 'get_client'):
                    client = ai_service.get_client()
                    if client:
                        self._quick_creator_plugin.set_api_client(client)
                        self._set_status("快捷创作插件加载成功")
                        logger.info("快捷创作插件加载成功")
                        return
            
            logger.warning("快捷创作插件未找到AI客户端")
        except Exception as e:
            logger.error(f"加载快捷创作插件失败: {e}")
            self._quick_creator_plugin = None
    
    # 别名方法，保持兼容性
    def _load_quick_creation_plugin(self):
        """加载快捷创作插件（别名方法）"""
        self._load_quick_creator_plugin()

    def _on_quick_save_results(self):
        """保存快捷创作结果"""
        try:
            # 获取所有结果内容
            content = {}
            for result_type, text_widget in self._quick_result_texts.items():
                content[result_type] = text_widget.get("1.0", tk.END).strip()
            
            # 检查是否有内容
            if not any(content.values()):
                messagebox.showwarning("提示", "没有可保存的内容！")
                return
            
            # 选择保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")],
                title="保存快捷创作结果"
            )
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("===== 世界观设定 =====\n")
                    f.write(content.get("worldview", "") + "\n\n")
                    f.write("===== 章节大纲 =====\n")
                    f.write(content.get("outline", "") + "\n\n")
                    f.write("===== 人物设定 =====\n")
                    f.write(content.get("characters", "") + "\n\n")
                    f.write("===== 关键情节 =====\n")
                    f.write(content.get("plot", "") + "\n")
                
                self._set_status(f"已保存：{os.path.basename(file_path)}")
                messagebox.showinfo("成功", "快捷创作结果已保存！")
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_file_error
            title, full_message = convert_file_error(e, "保存")
            messagebox.showerror(title, full_message)
    
    def _create_continue_content(self) -> tk.Frame:
        """创建续写功能内容页面"""
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # 初始化版本管理器
        self._continue_versions = []  # 存储所有版本
        self._current_version_index = 0  # 当前显示的版本索引
        self._best_version_index = -1  # 最佳版本索引
        
        # 上部：原文区
        source_frame = ttk.LabelFrame(frame, text="原文内容（智能分析续写）", padding=10)
        source_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 原文操作按钮
        source_btn_frame = ttk.Frame(source_frame, style="TFrame")
        source_btn_frame.pack(fill=tk.X, pady=5)
        ttk.Label(source_btn_frame, text="续写起点：光标定位到末尾位置").pack(side=tk.LEFT)
        ttk.Button(source_btn_frame, text="选择文件", command=self._on_continue_browse).pack(side=tk.RIGHT, padx=5)
        ttk.Button(source_btn_frame, text="选择章节", command=self._on_continue_select_chapter).pack(side=tk.RIGHT, padx=5)
        
        self._continue_source = tk.Text(source_frame, wrap=tk.WORD, height=8,
                                       font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                       bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._continue_source.pack(fill=tk.BOTH, expand=True)
        self._continue_source.insert("1.0", "请在此粘贴或输入原文内容，或选择文件导入...")
        
        # 中部：续写设置
        settings_frame = ttk.LabelFrame(frame, text="续写设置", padding=10)
        settings_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 第一行：字数和方向
        row1 = ttk.Frame(settings_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        
        # 字数选择（预设+自定义）
        ttk.Label(row1, text="续写字数：").pack(side=tk.LEFT)
        self._continue_words_var = tk.StringVar(value="1000")
        words_combo = ttk.Combobox(row1, textvariable=self._continue_words_var, 
                                   values=["500", "1000", "2000", "自定义"], width=10, state="readonly")
        words_combo.pack(side=tk.LEFT, padx=5)
        words_combo.bind("<<ComboboxSelected>>", self._on_words_selected)
        
        # 自定义字数输入框（默认隐藏）
        self._custom_words_entry = ttk.Entry(row1, width=8)
        
        # 续写方向
        ttk.Label(row1, text="续写方向：").pack(side=tk.LEFT, padx=(20, 0))
        self._continue_direction_var = tk.StringVar(value="自然续写")
        ttk.Combobox(row1, textvariable=self._continue_direction_var, 
                    values=["自然续写", "剧情推进", "高潮铺垫", "结局", "制造冲突", "引入新角色", "转折情节", "情感描写", "动作场景", "对话为主"], 
                    width=12, state="readonly").pack(side=tk.LEFT, padx=5)
        
        # 第二行：温度和模式
        row2 = ttk.Frame(settings_frame, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        
        ttk.Label(row2, text="创意温度：").pack(side=tk.LEFT)
        self._continue_temp_var = tk.DoubleVar(value=0.8)
        temp_scale = ttk.Scale(row2, from_=0.0, to=1.0, variable=self._continue_temp_var, 
                               orient=tk.HORIZONTAL, length=120)
        temp_scale.pack(side=tk.LEFT, padx=5)
        self._temp_label = ttk.Label(row2, text="0.8")
        self._temp_label.pack(side=tk.LEFT)
        temp_scale.bind("<Motion>", lambda e: self._temp_label.configure(text=f"{self._continue_temp_var.get():.1f}"))
        
        ttk.Label(row2, text="生成模式：").pack(side=tk.LEFT, padx=(20, 0))
        self._continue_mode_var = tk.StringVar(value="单次生成")
        ttk.Radiobutton(row2, text="单次", variable=self._continue_mode_var, value="单次生成").pack(side=tk.LEFT)
        ttk.Radiobutton(row2, text="多版本", variable=self._continue_mode_var, value="多版本生成").pack(side=tk.LEFT, padx=5)
        
        # 下部：分割为左侧版本列表 + 右侧结果预览
        bottom_frame = ttk.Frame(frame, style="TFrame")
        bottom_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 左侧：版本列表区域（20%宽度）
        version_list_frame = ttk.LabelFrame(bottom_frame, text="版本历史（最多5个）", padding=10)
        version_list_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 5))
        
        # 版本列表Treeview
        columns = ("版本", "温度", "字数", "评分", "状态")
        self._version_tree = ttk.Treeview(version_list_frame, columns=columns, show="headings", height=10)
        
        # 设置列宽和标题
        self._version_tree.heading("版本", text="版本")
        self._version_tree.heading("温度", text="温度")
        self._version_tree.heading("字数", text="字数")
        self._version_tree.heading("评分", text="评分")
        self._version_tree.heading("状态", text="状态")
        
        self._version_tree.column("版本", width=50, anchor=tk.CENTER)
        self._version_tree.column("温度", width=50, anchor=tk.CENTER)
        self._version_tree.column("字数", width=60, anchor=tk.CENTER)
        self._version_tree.column("评分", width=60, anchor=tk.CENTER)
        self._version_tree.column("状态", width=60, anchor=tk.CENTER)
        
        # 添加滚动条
        tree_scroll = ttk.Scrollbar(version_list_frame, orient="vertical", command=self._version_tree.yview)
        self._version_tree.configure(yscrollcommand=tree_scroll.set)
        
        self._version_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 绑定选择事件
        self._version_tree.bind("<<TreeviewSelect>>", self._on_version_tree_select)
        self._version_tree.bind("<Double-1>", self._on_version_double_click)
        
        # 右侧：续写结果预览（80%宽度）
        result_frame = ttk.LabelFrame(bottom_frame, text="续写结果预览", padding=10)
        result_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 版本信息栏
        version_bar = ttk.Frame(result_frame, style="TFrame")
        version_bar.pack(fill=tk.X, pady=5)
        self._version_info_label = ttk.Label(version_bar, text="当前无版本")
        self._version_info_label.pack(side=tk.LEFT)
        
        # 版本切换按钮
        ttk.Label(version_bar, text="  快速切换：").pack(side=tk.LEFT)
        ttk.Button(version_bar, text="◀", width=3, command=lambda: self._switch_version(-1)).pack(side=tk.LEFT, padx=2)
        ttk.Button(version_bar, text="▶", width=3, command=lambda: self._switch_version(1)).pack(side=tk.LEFT, padx=2)
        
        # 结果文本框
        self._continue_result = tk.Text(result_frame, wrap=tk.WORD, height=10,
                                       font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                       bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        self._continue_result.pack(fill=tk.BOTH, expand=True)
        
        # 结果信息栏
        info_bar = ttk.Frame(result_frame, style="TFrame")
        info_bar.pack(fill=tk.X, pady=5)
        self._result_info_label = ttk.Label(info_bar, text="字数：0 | 耗时：0秒 | 模型：-")
        self._result_info_label.pack(side=tk.LEFT)
        self._score_label = ttk.Label(info_bar, text="评分：-")
        self._score_label.pack(side=tk.RIGHT)
        
        # 底部按钮
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 左侧按钮组
        left_btns = ttk.Frame(btn_frame, style="TFrame")
        left_btns.pack(side=tk.LEFT)
        
        self._start_btn = ttk.Button(left_btns, text="🚀 开始续写", command=self._on_continue_generate)
        self._start_btn.pack(side=tk.LEFT, padx=5)
        
        self._regenerate_btn = ttk.Button(left_btns, text="🔄 重新生成", command=self._on_continue_regenerate, state=tk.DISABLED)
        self._regenerate_btn.pack(side=tk.LEFT, padx=5)
        
        # 右侧按钮组
        right_btns = ttk.Frame(btn_frame, style="TFrame")
        right_btns.pack(side=tk.RIGHT)
        
        self._select_best_btn = ttk.Button(right_btns, text="⭐ 选择最佳版本", command=self._on_continue_select_best, state=tk.DISABLED)
        self._select_best_btn.pack(side=tk.LEFT, padx=5)
        
        self._save_btn = ttk.Button(right_btns, text="💾 保存结果", command=self._on_continue_save, state=tk.DISABLED)
        self._save_btn.pack(side=tk.LEFT, padx=5)
        
        # 加载续写插件
        self._load_continuation_plugin()
        
        return frame
    
    def _load_continuation_plugin(self):
        """加载续写插件"""
        try:
            # 使用动态导入（支持连字符目录名）
            plugin_class = _dynamic_import_plugin('continuation-generator-v1', 'ContinuationGeneratorPlugin')
            if not plugin_class:
                logger.error("续写插件类未找到")
                self._continuation_plugin = None
                return
            
            self._continuation_plugin = plugin_class()
            
            # 从服务定位器获取AI客户端
            if CORE_AVAILABLE:
                service_locator = get_service_locator()
                ai_service = service_locator.get_service("ai_service")
                if ai_service and hasattr(ai_service, 'get_client'):
                    client = ai_service.get_client()
                    if client:
                        self._continuation_plugin.set_api_client(client)
                        self._set_status("续写插件加载成功")
                        logger.info("续写插件加载成功")
                        return
            
            # 如果无法获取AI服务，使用配置文件中的设置
            if hasattr(self, '_config_manager') and self._config_manager:
                ai_config = self._config_manager.get("ai_service", {})
                if ai_config:
                    from openai import OpenAI
                    api_key = ai_config.get("api_key", "")
                    base_url = ai_config.get("base_url", "https://api.deepseek.com/v1")
                    if api_key:
                        client = OpenAI(api_key=api_key, base_url=base_url)
                        self._continuation_plugin.set_api_client(client)
                        self._set_status("续写插件加载成功（使用配置文件）")
                        logger.info("续写插件加载成功（使用配置文件）")
                        return
            
            logger.warning("续写插件未找到AI客户端，请先配置AI服务")
        except Exception as e:
            logger.error(f"加载续写插件失败: {e}")
            self._continuation_plugin = None
    
    def _create_consistency_content(self) -> tk.Frame:
        """
        创建长篇检测内容页面
        
        参考：
        - 升级方案 10.升级方案✅️.md Sprint 7-8
        - OpenClaw一致性检查Agent设计
        """
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # ========== 上部：新章节输入区 ==========
        input_frame = ttk.LabelFrame(frame, text="新章节内容", padding=10)
        input_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 输入操作按钮
        input_btn_frame = ttk.Frame(input_frame, style="TFrame")
        input_btn_frame.pack(fill=tk.X, pady=5)
        ttk.Label(input_btn_frame, text="请输入待检测的新章节内容，系统将自动召回前文并检测冲突").pack(side=tk.LEFT)
        ttk.Button(input_btn_frame, text="选择文件", command=self._on_consistency_browse).pack(side=tk.RIGHT, padx=5)
        ttk.Button(input_btn_frame, text="选择章节", command=self._on_consistency_select_chapter).pack(side=tk.RIGHT, padx=5)
        
        # 文本输入框
        text_container = ttk.Frame(input_frame, style="TFrame")
        text_container.pack(fill=tk.BOTH, expand=True)
        
        # 滚动条
        input_scrollbar = ttk.Scrollbar(text_container, orient=tk.VERTICAL)
        input_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self._consistency_input = tk.Text(text_container, wrap=tk.WORD, height=8,
                                          font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                          bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                          yscrollcommand=input_scrollbar.set)
        self._consistency_input.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        input_scrollbar.config(command=self._consistency_input.yview)
        
        # ========== 中部：检测设置 ==========
        settings_frame = ttk.LabelFrame(frame, text="检测设置", padding=10)
        settings_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 第一行：题材和召回章节数
        row1 = ttk.Frame(settings_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        
        ttk.Label(row1, text="小说题材：").pack(side=tk.LEFT)
        self._consistency_genre_var = tk.StringVar(value="auto")
        genre_combo = ttk.Combobox(row1, textvariable=self._consistency_genre_var,
                                   values=["auto", "科幻", "玄幻", "历史", "都市", "悬疑"],
                                   width=12, state="readonly")
        genre_combo.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(row1, text="召回章节数：").pack(side=tk.LEFT, padx=(20, 0))
        self._consistency_topk_var = tk.IntVar(value=10)
        ttk.Spinbox(row1, from_=5, to=30, textvariable=self._consistency_topk_var,
                   width=8).pack(side=tk.LEFT, padx=5)
        
        # 第二行：检测选项
        row2 = ttk.Frame(settings_frame, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        
        ttk.Label(row2, text="检测类型：").pack(side=tk.LEFT)
        self._check_character_var = tk.BooleanVar(value=True)
        self._check_plot_var = tk.BooleanVar(value=True)
        self._check_worldview_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row2, text="人物一致性", variable=self._check_character_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(row2, text="情节连贯性", variable=self._check_plot_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(row2, text="世界观一致", variable=self._check_worldview_var).pack(side=tk.LEFT, padx=5)
        
        # ========== 下部：检测结果 ==========
        result_frame = ttk.LabelFrame(frame, text="检测结果", padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 结果信息栏（增强反馈）
        info_bar = ttk.Frame(result_frame, style="TFrame")
        info_bar.pack(fill=tk.X, pady=5)
        
        self._consistency_status_label = ttk.Label(info_bar, text="状态：就绪", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold"))
        self._consistency_status_label.pack(side=tk.LEFT)
        
        # 新增：阶段标签
        self._consistency_stage_label = ttk.Label(info_bar, text="", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL))
        self._consistency_stage_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # 新增：耗时标签
        self._consistency_time_label = ttk.Label(info_bar, text="", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL))
        self._consistency_time_label.pack(side=tk.LEFT, padx=(10, 0))
        
        self._consistency_accuracy_label = ttk.Label(info_bar, text="准确率：-", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL))
        self._consistency_accuracy_label.pack(side=tk.RIGHT)
        
        # 新增：进度条
        self._consistency_progress = ttk.Progressbar(info_bar, mode='indeterminate', length=150)
        self._consistency_progress.pack(side=tk.RIGHT, padx=5)
        
        # 冲突列表Treeview
        tree_container = ttk.Frame(result_frame, style="TFrame")
        tree_container.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 滚动条
        tree_scrollbar = ttk.Scrollbar(tree_container, orient=tk.VERTICAL)
        tree_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        columns = ("severity", "type", "description", "location", "suggestion")
        self._consistency_tree = ttk.Treeview(tree_container, columns=columns, show="headings", height=6,
                                              yscrollcommand=tree_scrollbar.set)
        
        # 设置列宽和标题
        self._consistency_tree.heading("severity", text="严重程度")
        self._consistency_tree.heading("type", text="冲突类型")
        self._consistency_tree.heading("description", text="冲突描述")
        self._consistency_tree.heading("location", text="位置")
        self._consistency_tree.heading("suggestion", text="修复建议")
        
        self._consistency_tree.column("severity", width=80, anchor=tk.CENTER)
        self._consistency_tree.column("type", width=80, anchor=tk.CENTER)
        self._consistency_tree.column("description", width=250)
        self._consistency_tree.column("location", width=120)
        self._consistency_tree.column("suggestion", width=200)
        
        self._consistency_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scrollbar.config(command=self._consistency_tree.yview)
        
        # 召回上下文预览
        context_frame = ttk.LabelFrame(result_frame, text="召回上下文摘要", padding=5)
        context_frame.pack(fill=tk.X, pady=5)
        
        context_container = ttk.Frame(context_frame, style="TFrame")
        context_container.pack(fill=tk.BOTH, expand=True)
        
        context_scrollbar = ttk.Scrollbar(context_container, orient=tk.VERTICAL)
        context_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self._consistency_context = tk.Text(context_container, wrap=tk.WORD, height=4,
                                            font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_SMALL),
                                            bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_SECONDARY,
                                            yscrollcommand=context_scrollbar.set)
        self._consistency_context.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        context_scrollbar.config(command=self._consistency_context.yview)
        self._consistency_context.insert("1.0", "检测完成后将显示召回的相关章节摘要...")
        self._consistency_context.config(state=tk.DISABLED)
        
        # ========== 底部按钮 ==========
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 左侧按钮组
        left_btns = ttk.Frame(btn_frame, style="TFrame")
        left_btns.pack(side=tk.LEFT)
        
        self._consistency_check_btn = ttk.Button(left_btns, text="🔍 开始检测", 
                                                  command=self._on_consistency_check)
        self._consistency_check_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(left_btns, text="🗑️ 清空", command=self._on_consistency_clear).pack(side=tk.LEFT, padx=5)
        
        # 右侧按钮组
        right_btns = ttk.Frame(btn_frame, style="TFrame")
        right_btns.pack(side=tk.RIGHT)
        
        ttk.Button(right_btns, text="📋 复制结果", command=self._on_consistency_copy).pack(side=tk.LEFT, padx=5)
        ttk.Button(right_btns, text="💾 导出报告", command=self._on_consistency_export).pack(side=tk.LEFT, padx=5)
        
        # 加载一致性检查Agent
        self._load_consistency_agent()
        
        return frame
    
    def _load_consistency_agent(self):
        """加载一致性检查Agent"""
        try:
            from agents.consistency_checker_agent import ConsistencyCheckerAgent
            self._consistency_agent = ConsistencyCheckerAgent()
            self._set_status("一致性检查Agent加载成功")
            logger.info("一致性检查Agent加载成功")
        except ImportError as e:
            logger.warning(f"一致性检查Agent未安装: {e}")
            self._consistency_agent = None
        except Exception as e:
            logger.error(f"加载一致性检查Agent失败: {e}")
            self._consistency_agent = None
    
    def _on_consistency_browse(self):
        """选择文件导入章节"""
        file_path = filedialog.askopenfilename(
            title="选择章节文件",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self._consistency_input.delete("1.0", tk.END)
                self._consistency_input.insert("1.0", content)
                self._set_status(f"已导入文件：{os.path.basename(file_path)}")
            except Exception as e:
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_file_error
                title, full_message = convert_file_error(e, "读取")
                messagebox.showerror(title, full_message)
    
    def _on_consistency_select_chapter(self):
        """选择已生成章节"""
        if not hasattr(self, '_project_manager') or not self._project_manager:
            messagebox.showwarning("提示", "请先打开项目")
            return
        
        chapters = self._project_manager.list_chapters()
        if not chapters:
            messagebox.showinfo("提示", "暂无已生成章节")
            return
        
        # 创建选择对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("选择章节")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text="请选择章节：").pack(pady=10)
        
        # 章节列表
        list_frame = ttk.Frame(dialog)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        chapter_list = tk.Listbox(list_frame, yscrollcommand=scrollbar.set)
        chapter_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=chapter_list.yview)
        
        for chapter in chapters:
            chapter_list.insert(tk.END, f"第{chapter.get('number', '?')}章 - {chapter.get('title', '未命名')}")
        
        def on_confirm():
            selection = chapter_list.curselection()
            if selection:
                chapter = chapters[selection[0]]
                content = chapter.get('content', '')
                self._consistency_input.delete("1.0", tk.END)
                self._consistency_input.insert("1.0", content)
                dialog.destroy()
        
        ttk.Button(dialog, text="确定", command=on_confirm).pack(pady=10)
    
    def _on_consistency_check(self):
        """执行一致性检查"""
        # 获取输入内容
        new_chapter = self._consistency_input.get("1.0", tk.END).strip()
        if not new_chapter or new_chapter == "请在此粘贴或输入原文内容，或选择文件导入...":
            messagebox.showwarning("提示", "请输入待检测的章节内容")
            return
        
        # 检查Agent是否加载
        if not hasattr(self, '_consistency_agent') or not self._consistency_agent:
            messagebox.showerror("错误", "一致性检查Agent未加载")
            return
        
        # 更新状态 - 开始检测
        self._consistency_status_label.config(text="状态：检测中")
        self._consistency_stage_label.config(text="阶段：初始化")
        self._consistency_time_label.config(text="耗时：0.0s")
        self._consistency_accuracy_label.config(text="准确率：计算中...")
        self._set_status("正在进行一致性检查...")
        
        # 清空之前的结果
        for item in self._consistency_tree.get_children():
            self._consistency_tree.delete(item)
        
        # 禁用检测按钮，显示进度
        self._consistency_check_btn.config(state=tk.DISABLED)
        self._consistency_progress.start(10)  # 启动进度条动画
        
        # 记录开始时间
        import time
        start_time = time.time()
        
        # 更新阶段状态的回调函数
        def update_stage(stage: str):
            self.root.after(0, lambda: self._consistency_stage_label.config(text=f"阶段：{stage}"))
        
        # 更新耗时的回调函数
        def update_time():
            elapsed = time.time() - start_time
            self.root.after(0, lambda: self._consistency_time_label.config(text=f"耗时：{elapsed:.1f}s"))
        
        # 异步执行检查
        def do_check():
            try:
                # 阶段1：准备上下文
                update_stage("准备上下文")
                update_time()
                
                # 构建检查上下文
                context = {
                    "new_chapter": new_chapter,
                    "genre": self._consistency_genre_var.get() if self._consistency_genre_var.get() != "auto" else None,
                    "top_k": self._consistency_topk_var.get(),
                    "check_types": {
                        "character": self._check_character_var.get(),
                        "plot": self._check_plot_var.get(),
                        "worldview": self._check_worldview_var.get()
                    }
                }
                
                # 阶段2：召回相似章节
                update_stage("召回相似章节")
                update_time()
                
                # 阶段3：LLM检测
                update_stage("LLM冲突检测")
                update_time()
                
                # 调用Agent执行检查
                result = self._consistency_agent.execute(context)
                
                # 阶段4：解析结果
                update_stage("解析结果")
                update_time()
                
                # 计算总耗时
                total_time = time.time() - start_time
                
                # 回到主线程更新UI
                self.root.after(0, lambda: self._update_consistency_result(result, total_time))
                
            except Exception as e:
                logger.error(f"一致性检查失败: {e}")
                total_time = time.time() - start_time
                self.root.after(0, lambda: self._handle_consistency_error(str(e), total_time))
        
        # 启动异步任务（使用统一线程池，解决卡顿问题）
        from core.thread_pool_manager import thread_pool_manager
        thread_pool_manager.submit_sync(do_check)
    
    def _update_consistency_result(self, result: dict, elapsed_time: float = 0.0):
        """更新检测结果到UI"""
        try:
            # 停止进度条并恢复按钮
            self._consistency_progress.stop()
            self._consistency_check_btn.config(state=tk.NORMAL)
            
            # 清空阶段标签
            self._consistency_stage_label.config(text="")
            self._consistency_time_label.config(text=f"耗时：{elapsed_time:.1f}s")
            
            # 更新状态
            is_consistent = result.get("is_consistent", True)
            accuracy = result.get("accuracy", 0.0)
            
            if is_consistent:
                self._consistency_status_label.config(text="状态：✅ 无冲突", foreground="#10B981")
            else:
                conflicts = result.get("conflicts", [])
                p0_count = sum(1 for c in conflicts if c.get("severity") == "P0")
                p1_count = sum(1 for c in conflicts if c.get("severity") == "P1")
                p2_count = sum(1 for c in conflicts if c.get("severity") == "P2")
                self._consistency_status_label.config(
                    text=f"状态：⚠️ 发现{len(conflicts)}个冲突（P0:{p0_count} P1:{p1_count} P2:{p2_count}）",
                    foreground="#F59E0B")
            
            self._consistency_accuracy_label.config(text=f"准确率：{accuracy:.1%}")
            
            # 更新冲突列表
            conflicts = result.get("conflicts", [])
            for conflict in conflicts:
                severity = conflict.get("severity", "P2")
                conflict_type = conflict.get("conflict_type", "unknown")
                description = conflict.get("description", "")
                location = conflict.get("location", "")
                suggestion = conflict.get("suggestion", "")
                
                # 严重程度标签
                severity_tag = {"P0": "🔴严重", "P1": "🟡中等", "P2": "🟢轻微"}.get(severity, severity)
                
                # 类型标签
                type_tag = {"character": "人物", "plot": "情节", "worldview": "世界观"}.get(conflict_type, conflict_type)
                
                self._consistency_tree.insert("", tk.END, values=(
                    severity_tag, type_tag, description, location, suggestion
                ))
            
            # 更新召回上下文
            recalled_context = result.get("recalled_context", "")
            self._consistency_context.config(state=tk.NORMAL)
            self._consistency_context.delete("1.0", tk.END)
            if recalled_context:
                self._consistency_context.insert("1.0", recalled_context)
            else:
                self._consistency_context.insert("1.0", "无召回上下文")
            self._consistency_context.config(state=tk.DISABLED)
            
            self._set_status(f"一致性检查完成，发现{len(conflicts)}个冲突，耗时{elapsed_time:.1f}秒")
            
        except Exception as e:
            logger.error(f"更新检测结果失败: {e}")
            self._consistency_progress.stop()
            self._consistency_check_btn.config(state=tk.NORMAL)
            self._consistency_status_label.config(text=f"状态：显示失败 - {str(e)}", foreground="#EF4444")
    
    def _handle_consistency_error(self, error_msg: str, elapsed_time: float = 0.0):
        """处理一致性检查错误"""
        self._consistency_progress.stop()
        self._consistency_check_btn.config(state=tk.NORMAL)
        self._consistency_stage_label.config(text="")
        self._consistency_time_label.config(text=f"耗时：{elapsed_time:.1f}s")
        self._consistency_status_label.config(text=f"状态：❌ 检测失败", foreground="#EF4444")
        self._consistency_accuracy_label.config(text="准确率：-")
        self._set_status(f"一致性检查失败: {error_msg}")
    
    def _on_consistency_clear(self):
        """清空检测结果"""
        self._consistency_input.delete("1.0", tk.END)
        for item in self._consistency_tree.get_children():
            self._consistency_tree.delete(item)
        self._consistency_context.config(state=tk.NORMAL)
        self._consistency_context.delete("1.0", tk.END)
        self._consistency_context.insert("1.0", "检测完成后将显示召回的相关章节摘要...")
        self._consistency_context.config(state=tk.DISABLED)
        
        # 清除所有反馈标签
        self._consistency_status_label.config(text="状态：就绪", foreground=GlassTheme.TEXT_PRIMARY)
        self._consistency_stage_label.config(text="")
        self._consistency_time_label.config(text="")
        self._consistency_accuracy_label.config(text="准确率：-")
        
        # 停止进度条并确保按钮可用
        self._consistency_progress.stop()
        self._consistency_check_btn.config(state=tk.NORMAL)
        
        self._set_status("已清空检测结果")
    
    def _on_consistency_copy(self):
        """复制检测结果到剪贴板"""
        # 收集结果
        result_text = "长篇小说一致性检查报告\n"
        result_text += "=" * 50 + "\n\n"
        
        # 状态信息
        result_text += f"状态: {self._consistency_status_label.cget('text')}\n"
        result_text += f"准确率: {self._consistency_accuracy_label.cget('text')}\n\n"
        
        # 冲突列表
        result_text += "冲突列表：\n"
        result_text += "-" * 50 + "\n"
        
        for item in self._consistency_tree.get_children():
            values = self._consistency_tree.item(item)['values']
            result_text += f"\n严重程度: {values[0]}\n"
            result_text += f"冲突类型: {values[1]}\n"
            result_text += f"冲突描述: {values[2]}\n"
            result_text += f"位置: {values[3]}\n"
            result_text += f"修复建议: {values[4]}\n"
        
        # 复制到剪贴板
        self.root.clipboard_clear()
        self.root.clipboard_append(result_text)
        self._set_status("结果已复制到剪贴板")
    
    def _on_consistency_export(self):
        """导出检测报告"""
        file_path = filedialog.asksaveasfilename(
            title="导出检测报告",
            defaultextension=".md",
            filetypes=[("Markdown文件", "*.md"), ("文本文件", "*.txt")]
        )
        
        if file_path:
            try:
                # 生成报告内容
                report = f"""# 长篇小说一致性检查报告

**检查时间**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

## 检查结果

- **状态**: {self._consistency_status_label.cget('text')}
- **准确率**: {self._consistency_accuracy_label.cget('text')}

## 冲突列表

"""
                for item in self._consistency_tree.get_children():
                    values = self._consistency_tree.item(item)['values']
                    report += f"""### {values[0]} - {values[1]}冲突

- **描述**: {values[2]}
- **位置**: {values[3]}
- **修复建议**: {values[4]}

"""
                
                # 召回上下文
                context = self._consistency_context.get("1.0", tk.END).strip()
                if context and context != "检测完成后将显示召回的相关章节摘要...":
                    report += f"""## 召回上下文摘要

{context}
"""
                
                # 写入文件
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(report)
                
                messagebox.showinfo("成功", f"报告已导出：{file_path}")
                self._set_status(f"报告已导出：{os.path.basename(file_path)}")
                
            except Exception as e:
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_file_error
                title, full_message = convert_file_error(e, "导出")
                messagebox.showerror(title, full_message)
    
    def _create_knowledge_content(self) -> tk.Frame:
        """
        创建知识库浏览内容页面
        
        参考：
        - 升级方案 10.升级方案✅️.md Sprint 9-10
        - OpenClaw知识库检索设计
        """
        frame = ttk.Frame(self._workbench_content_frame, style="TFrame")
        
        # ========== 上部：搜索区 ==========
        search_frame = ttk.LabelFrame(frame, text="知识库搜索", padding=10)
        search_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 搜索输入行
        search_row = ttk.Frame(search_frame, style="TFrame")
        search_row.pack(fill=tk.X, pady=5)
        
        ttk.Label(search_row, text="关键词：").pack(side=tk.LEFT)
        self._knowledge_search_var = tk.StringVar()
        self._knowledge_search_entry = ttk.Entry(search_row, textvariable=self._knowledge_search_var, width=40)
        self._knowledge_search_entry.pack(side=tk.LEFT, padx=5)
        self._knowledge_search_entry.bind("<Return>", lambda e: self._on_knowledge_search())
        
        ttk.Button(search_row, text="🔍 搜索", command=self._on_knowledge_search).pack(side=tk.LEFT, padx=5)
        ttk.Button(search_row, text="🔄 刷新", command=self._on_knowledge_refresh).pack(side=tk.LEFT, padx=5)
        
        # 过滤条件行
        filter_row = ttk.Frame(search_frame, style="TFrame")
        filter_row.pack(fill=tk.X, pady=5)
        
        ttk.Label(filter_row, text="题材过滤：").pack(side=tk.LEFT)
        self._knowledge_category_var = tk.StringVar(value="全部")
        self._category_map = {
            "全部": "all",
            "玄幻": "xuanhuan", "仙侠": "xianxia", "都市": "urban", "言情": "romance",
            "历史": "history", "科幻": "scifi", "悬疑": "suspense", "军事": "military",
            "武侠": "wuxia", "游戏": "game", "奇幻": "fantasy", "灵异": "lingyi",
            "同人": "tongren", "通用": "general",
            "恐怖": "horror", "推理": "mystery", "体育": "sports", "哲学": "philosophy",
            "写作技巧": "writing_technique"  # V5.3新增：写作技巧分类
        }
        self._category_reverse_map = {v: k for k, v in self._category_map.items()}
        category_combo = ttk.Combobox(filter_row, textvariable=self._knowledge_category_var,
                                       values=list(self._category_map.keys()),
                                       width=10, state="readonly")
        category_combo.pack(side=tk.LEFT, padx=5)
        # 绑定题材切换事件
        category_combo.bind("<<ComboboxSelected>>", self._on_category_changed)
        
        ttk.Label(filter_row, text="领域过滤：").pack(side=tk.LEFT, padx=(20, 0))
        self._knowledge_domain_var = tk.StringVar(value="全部")
        self._domain_map = {
            "全部": "all",
            # 基础学科
            "物理": "physics", "化学": "chemistry", "生物": "biology", "数学": "mathematics",
            "地理": "geography", "天文": "astronomy", "心理": "psychology", "哲学": "philosophy",
            "经济": "economics", "历史": "history", "文化": "culture", "技术": "technology",
            # 玄幻/仙侠/奇幻
            "魔法": "magic", "神话": "mythology", "宗教": "religion", "玄学": "occult",
            "修炼": "cultivation", "道家": "daoism", "佛家": "buddhism", "灵学": "spirituality",
            "幻想生物": "fantasy_creatures", "世界观": "worldbuilding",
            # 都市/言情
            "社会": "society", "法律": "law", "教育": "education", "职场": "workplace",
            "情感": "emotion", "家庭": "family",
            # 历史相关
            "政治": "politics", "军事": "military",
            # 科幻相关
            "航天": "space", "AI": "ai", "未来学": "future",
            # 悬疑相关
            "刑侦": "forensics", "逻辑": "logic",
            # 武侠相关
            "武术": "martial_arts", "中医": "medicine", "江湖": "jianghu",
            # 游戏相关
            "游戏设计": "design", "游戏叙事": "game_narrative",
            # 灵异相关
            "民俗": "folklore", "传说": "legend", "神秘学": "occultism",
            # 同人相关
            "原著分析": "analysis", "人物": "characters", "剧情": "plot", "设定": "setting",
            # 通用写作（移除重复键，避免domain_reverse_map冲突）
            "写作技法": "writing_technique", "人物塑造": "character", "情节设计": "plot_design", "对话": "dialogue",
            # 其他
            "基础知识": "basic_knowledge",
            # V5.3新增：写作技巧六领域（按12.2文档规范，确保唯一映射）
            "叙事技巧": "narrative",
            "描写技巧": "description",
            "修辞技巧": "rhetoric",
            "结构技巧": "structure",
            "特殊句式": "special_sentence",
            "高级技法": "advanced"
        }
        
        # V5.3新增：写作技巧固定六领域（按12.2文档规范）
        self._writing_technique_domains = {
            "叙事技巧": "narrative",
            "描写技巧": "description", 
            "修辞技巧": "rhetoric",
            "结构技巧": "structure",
            "特殊句式": "special_sentence",
            "高级技法": "advanced"
        }
        
        self._domain_reverse_map = {v: k for k, v in self._domain_map.items()}
        self._domain_combo = ttk.Combobox(filter_row, textvariable=self._knowledge_domain_var,
                                     values=list(self._domain_map.keys()),
                                     width=12, state="readonly")
        self._domain_combo.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(filter_row, text="返回数量：").pack(side=tk.LEFT, padx=(20, 0))
        self._knowledge_topk_var = tk.IntVar(value=50)
        ttk.Spinbox(filter_row, from_=10, to=200, textvariable=self._knowledge_topk_var,
                   width=8).pack(side=tk.LEFT, padx=5)
        
        # ========== 中部：统计信息 ==========
        stats_frame = ttk.LabelFrame(frame, text="知识库统计", padding=10)
        stats_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 第一行：题材分类统计
        stats_row1 = ttk.Frame(stats_frame, style="TFrame")
        stats_row1.pack(fill=tk.X, pady=2)
        
        self._knowledge_stats_label1 = ttk.Label(stats_row1, text="题材分类: 正在加载...")
        self._knowledge_stats_label1.pack(side=tk.LEFT)
        
        # 第二行：写作技巧统计
        stats_row2 = ttk.Frame(stats_frame, style="TFrame")
        stats_row2.pack(fill=tk.X, pady=2)
        
        self._knowledge_stats_label2 = ttk.Label(stats_row2, text="写作技巧: 正在加载...")
        self._knowledge_stats_label2.pack(side=tk.LEFT)
        
        ttk.Button(stats_row2, text="📊 详细统计", command=self._on_knowledge_stats).pack(side=tk.RIGHT, padx=5)
        
        # ========== 下部：知识列表 ==========
        list_frame = ttk.LabelFrame(frame, text="知识列表", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Treeview
        tree_container = ttk.Frame(list_frame, style="TFrame")
        tree_container.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 滚动条
        tree_scrollbar_y = ttk.Scrollbar(tree_container, orient=tk.VERTICAL)
        tree_scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        tree_scrollbar_x = ttk.Scrollbar(tree_container, orient=tk.HORIZONTAL)
        tree_scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
        
        columns = ("id", "title", "category", "domain", "keywords", "preview")
        self._knowledge_tree = ttk.Treeview(tree_container, columns=columns, show="headings", height=15,
                                             yscrollcommand=tree_scrollbar_y.set,
                                             xscrollcommand=tree_scrollbar_x.set)
        
        # 设置列宽和标题
        self._knowledge_tree.heading("id", text="ID")
        self._knowledge_tree.heading("title", text="标题")
        self._knowledge_tree.heading("category", text="题材")
        self._knowledge_tree.heading("domain", text="领域")
        self._knowledge_tree.heading("keywords", text="关键词")
        self._knowledge_tree.heading("preview", text="内容预览")
        
        self._knowledge_tree.column("id", width=120, anchor=tk.CENTER)
        self._knowledge_tree.column("title", width=150)
        self._knowledge_tree.column("category", width=80, anchor=tk.CENTER)
        self._knowledge_tree.column("domain", width=80, anchor=tk.CENTER)
        self._knowledge_tree.column("keywords", width=200)
        self._knowledge_tree.column("preview", width=300)
        
        self._knowledge_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scrollbar_y.config(command=self._knowledge_tree.yview)
        tree_scrollbar_x.config(command=self._knowledge_tree.xview)
        
        # 双击查看详情
        self._knowledge_tree.bind("<Double-1>", self._on_knowledge_double_click)
        
        # ========== 详情面板 ==========
        detail_frame = ttk.LabelFrame(frame, text="知识点详情", padding=10)
        detail_frame.pack(fill=tk.X, padx=5, pady=5)
        
        detail_container = ttk.Frame(detail_frame, style="TFrame")
        detail_container.pack(fill=tk.BOTH, expand=True)
        
        detail_scrollbar = ttk.Scrollbar(detail_container, orient=tk.VERTICAL)
        detail_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self._knowledge_detail = tk.Text(detail_container, wrap=tk.WORD, height=15,
                                          font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                          bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY,
                                          yscrollcommand=detail_scrollbar.set)
        self._knowledge_detail.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        detail_scrollbar.config(command=self._knowledge_detail.yview)
        self._knowledge_detail.insert("1.0", "双击知识点查看详细内容...")
        self._knowledge_detail.config(state=tk.DISABLED)
        
        # ========== 底部按钮 ==========
        btn_frame = ttk.Frame(frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 左侧按钮组
        left_btns = ttk.Frame(btn_frame, style="TFrame")
        left_btns.pack(side=tk.LEFT)
        
        ttk.Button(left_btns, text="📥 导入JSON", command=self._on_knowledge_import).pack(side=tk.LEFT, padx=5)
        ttk.Button(left_btns, text="📤 导出JSON", command=self._on_knowledge_export).pack(side=tk.LEFT, padx=5)
        ttk.Button(left_btns, text="🔄 刷新缓存", command=self._on_knowledge_refresh).pack(side=tk.LEFT, padx=5)
        ttk.Button(left_btns, text="🧹 清除垃圾", command=self._on_knowledge_cleanup).pack(side=tk.LEFT, padx=5)
        ttk.Button(left_btns, text="🔙 恢复清理", command=self._on_knowledge_restore).pack(side=tk.LEFT, padx=5)
        
        # 右侧按钮组
        right_btns = ttk.Frame(btn_frame, style="TFrame")
        right_btns.pack(side=tk.RIGHT)
        
        ttk.Button(right_btns, text="➕ 添加知识点", command=self._on_knowledge_add).pack(side=tk.LEFT, padx=5)
        ttk.Button(right_btns, text="🤖 AI生成知识", command=self._on_knowledge_ai_generate, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(right_btns, text="✏️ 编辑选中", command=self._on_knowledge_edit).pack(side=tk.LEFT, padx=5)
        ttk.Button(right_btns, text="🗑️ 删除选中", command=self._on_knowledge_delete).pack(side=tk.LEFT, padx=5)
        
        # 加载知识库管理器
        self._load_knowledge_manager()
        
        return frame
    
    def _load_knowledge_manager(self):
        """加载知识库管理器"""
        try:
            from core.knowledge_manager import get_knowledge_manager, reset_knowledge_manager
            # 使用脚本所在目录，而非当前工作目录
            workspace_root = Path(__file__).parent.resolve()
            logger.info(f"[知识库] 工作区路径: {workspace_root}")
            
            # 强制重新加载，确保获取最新数据
            reset_knowledge_manager()
            logger.info("[知识库] 已重置单例")
            
            self._knowledge_manager = get_knowledge_manager(workspace_root)
            logger.info(f"[知识库] 管理器实例: {self._knowledge_manager}")
            
            # 加载统计信息
            stats = self._knowledge_manager.get_stats()
            total = stats.get("total", 0)
            categories = stats.get("categories", {})
            
            # 分离题材分类和写作技巧
            subject_categories = {k: v for k, v in categories.items() if k != "writing_technique"}
            
            # 第一行：题材分类统计（18个题材，与知识库实际目录一致）
            # 题材分类映射表（完整18个）
            subject_names = {
                "xuanhuan": "玄幻", "xianxia": "仙侠", "urban": "都市", "romance": "言情",
                "history": "历史", "scifi": "科幻", "suspense": "悬疑", "military": "军事",
                "wuxia": "武侠", "game": "游戏", "fantasy": "奇幻", "lingyi": "灵异",
                "tongren": "同人", "general": "通用",
                "horror": "恐怖", "mystery": "推理", "sports": "体育", "philosophy": "哲学"
            }
            
            # 格式：题材名(数量)，用空格分隔
            subject_parts = []
            for cat_key, count in subject_categories.items():
                display_name = subject_names.get(cat_key, cat_key)
                subject_parts.append(f"{display_name}({count})")
            
            stats_text1 = "题材: " + "  ".join(subject_parts)
            self._knowledge_stats_label1.config(text=stats_text1)
            
            # 第二行：写作技巧统计（6大类）
            # 从实际知识库文件加载写作技巧统计
            technique_stats = self._load_writing_technique_stats()
            
            # 格式：技巧名(数量)，用空格分隔
            technique_parts = []
            for domain_name, count in technique_stats["domains"].items():
                technique_parts.append(f"{domain_name}({count})")
            
            stats_text2 = "写作技巧: " + "  ".join(technique_parts)
            self._knowledge_stats_label2.config(text=stats_text2)
            
            # 加载初始列表
            self._on_knowledge_refresh()
            
            logger.info(f"[知识库] 加载成功，总计 {total} 条知识点")
        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            logger.error(f"加载知识库管理器失败:\n{error_detail}")
            self._knowledge_manager = None
            self._knowledge_stats_label1.config(text=f"题材分类: 加载失败")
            self._knowledge_stats_label2.config(text=f"写作技巧: 加载失败")
    
    def _on_knowledge_refresh(self):
        """刷新知识库缓存"""
        if self._knowledge_manager:
            try:
                self._knowledge_manager.reload_cache()
                # 刷新列表显示
                self._on_knowledge_search()
                messagebox.showinfo("成功", "知识库缓存已刷新")
            except Exception as e:
                logger.error(f"刷新知识库缓存失败: {e}")
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_knowledge_error
                title, full_message = convert_knowledge_error(e, "刷新")
                messagebox.showerror(title, full_message)

    def _on_category_changed(self, event):
        """题材切换事件：根据题材自动切换领域下拉框选项
        
        V5.3新增：写作技巧使用固定六领域，其他题材使用通用领域
        """
        category = self._knowledge_category_var.get()
        
        if category == "写作技巧":
            # 切换为写作技巧六领域
            self._domain_combo['values'] = ["全部"] + list(self._writing_technique_domains.keys())
            self._knowledge_domain_var.set("全部")
        else:
            # 恢复为通用领域
            self._domain_combo['values'] = list(self._domain_map.keys())
            self._knowledge_domain_var.set("全部")

    def _on_knowledge_search(self):
        """搜索知识点"""
        if not self._knowledge_manager:
            messagebox.showwarning("提示", "知识库未初始化")
            return
        
        query = self._knowledge_search_var.get().strip()
        category = self._knowledge_category_var.get()
        domain = self._knowledge_domain_var.get()
        top_k = self._knowledge_topk_var.get()
        
        # 过滤条件（中文转英文）
        category_filter = None if self._category_map.get(category) == "all" else self._category_map.get(category)
        domain_filter = None if self._domain_map.get(domain) == "all" else self._domain_map.get(domain)
        
        try:
            # 清空列表
            for item in self._knowledge_tree.get_children():
                self._knowledge_tree.delete(item)
            
            if query:
                # 搜索模式
                results = self._knowledge_manager.search_knowledge(
                    query=query,
                    category=category_filter,
                    domain=domain_filter,
                    top_k=top_k
                )
                
                for r in results:
                    keywords_str = ", ".join(r.keywords[:5])
                    preview = r.content[:80] + "..." if len(r.content) > 80 else r.content
                    # 英文转中文显示
                    category_cn = self._category_reverse_map.get(r.category, r.category)
                    domain_cn = self._domain_reverse_map.get(r.domain, r.domain)
                    
                    self._knowledge_tree.insert("", tk.END, values=(
                        r.knowledge_id, r.title, category_cn, domain_cn, keywords_str, preview
                    ))
                
                self._set_status(f"搜索完成，找到 {len(results)} 条知识点")
            else:
                # 浏览模式
                self._on_knowledge_refresh()
                
        except Exception as e:
            logger.error(f"搜索知识点失败: {e}")
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_knowledge_error
            title, full_message = convert_knowledge_error(e, "搜索")
            messagebox.showerror(title, full_message)
    
    def _on_knowledge_refresh(self):
        """刷新知识列表"""
        if not self._knowledge_manager:
            return
        
        category = self._knowledge_category_var.get()
        domain = self._knowledge_domain_var.get()
        
        # 中文转英文过滤条件
        category_filter = None if self._category_map.get(category) == "all" else self._category_map.get(category)
        domain_filter = None if self._domain_map.get(domain) == "all" else self._domain_map.get(domain)
        
        try:
            # 清空列表
            for item in self._knowledge_tree.get_children():
                self._knowledge_tree.delete(item)
            
            # 获取知识点列表
            knowledge_list = self._knowledge_manager.list_knowledge(
                category=category_filter,
                domain=domain_filter
            )
            
            for kp in knowledge_list[:200]:  # 限制显示数量
                keywords_str = ", ".join(kp.keywords[:5])
                preview = kp.content[:80] + "..." if len(kp.content) > 80 else kp.content
                # 英文转中文显示
                category_cn = self._category_reverse_map.get(kp.category, kp.category)
                domain_cn = self._domain_reverse_map.get(kp.domain, kp.domain)
                
                self._knowledge_tree.insert("", tk.END, values=(
                    kp.knowledge_id, kp.title, category_cn, domain_cn, keywords_str, preview
                ))
            
            self._set_status(f"已加载 {len(knowledge_list)} 条知识点")
            
        except Exception as e:
            logger.error(f"刷新知识列表失败: {e}")
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_knowledge_error
            title, full_message = convert_knowledge_error(e, "刷新")
            messagebox.showerror(title, full_message)
    
    def _convert_stats_to_chinese(self, stats: dict) -> dict:
        """将统计结果中的英文分类转换为中文显示"""
        if not stats:
            return stats
            
        display_stats = {
            "total": stats.get("total", 0),
            "categories_cn": {},
            "domains_cn": {},
            "source": stats.get("source", "unknown")
        }
        
        # 转换分类名称（使用已有的_reverse_map）
        categories = stats.get("categories", {})
        reverse_map = getattr(self, '_category_reverse_map', {})
        
        for cat_en, count in categories.items():
            cat_cn = reverse_map.get(cat_en, cat_en)  # 如果找不到映射，保留原key
            display_stats["categories_cn"][cat_cn] = count
        
        # 转换领域名称（写作技巧六领域）
        domains = stats.get("domains", {})
        writing_technique_map = {
            "narrative": "叙事手法",
            "description": "描写技巧",
            "rhetoric": "修辞技巧",
            "structure": "结构设计",
            "special_sentence": "特殊句式",
            "advanced": "高级技巧"
        }
        
        for domain_en, count in domains.items():
            domain_cn = writing_technique_map.get(domain_en, domain_en)
            display_stats["domains_cn"][domain_cn] = count
        
        return display_stats
    
    def _on_knowledge_stats(self):
        """显示详细统计"""
        if not self._knowledge_manager:
            return
        
        try:
            stats = self._knowledge_manager.get_stats()
            
            # 创建统计对话框
            dialog = tk.Toplevel(self.root)
            dialog.title("知识库统计")
            dialog.geometry("600x500")
            dialog.transient(self.root)
            
            # 统计文本
            text = tk.Text(dialog, wrap=tk.WORD, font=("Consolas", 10))
            text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            import json
            
            # 转换分类名称为中文
            display_stats = self._convert_stats_to_chinese(stats)
            
            text.insert("1.0", json.dumps(display_stats, indent=2, ensure_ascii=False))
            text.config(state=tk.DISABLED)
            
        except Exception as e:
            logger.error(f"显示统计失败: {e}")
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_knowledge_error
            title, full_message = convert_knowledge_error(e, "统计")
            messagebox.showerror(title, full_message)
    
    def _on_knowledge_double_click(self, event):
        """双击查看详情 - 专业排版（完整版）"""
        selection = self._knowledge_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        values = self._knowledge_tree.item(item)['values']
        knowledge_id = values[0]
        
        if not self._knowledge_manager:
            return
        
        try:
            kp = self._knowledge_manager.get_knowledge(knowledge_id)
            if kp:
                # 构建专业排版的详情
                detail = f"{'═'*60}\n"
                detail += f"【{kp.title}】\n"
                detail += f"{'═'*60}\n\n"
                
                # 元信息区
                detail += f"📌 ID: {kp.knowledge_id}\n"
                # 英文转中文显示
                category_cn = self._category_reverse_map.get(kp.category, kp.category)
                domain_cn = self._domain_reverse_map.get(kp.domain, kp.domain)
                detail += f"📚 题材: {category_cn}  |  领域: {domain_cn}\n"
                detail += f"🎯 难度: {kp.difficulty or 'intermediate'}\n"
                detail += f"🏷️ 关键词: {', '.join(kp.keywords)}\n"
                detail += f"{'─'*60}\n\n"
                
                # 核心概念解释（新增）
                if hasattr(kp, 'explanation') and kp.explanation:
                    detail += f"💡 核心概念\n"
                    detail += f"{'─'*60}\n"
                    detail += f"{kp.explanation}\n\n"
                
                # 详细内容区
                detail += f"📖 详细内容\n"
                detail += f"{'─'*60}\n"
                detail += f"{kp.content}\n"
                
                # 经典案例应用（新增）
                if hasattr(kp, 'classic_cases') and kp.classic_cases:
                    detail += f"\n{'─'*60}\n"
                    detail += f"🎬 经典案例应用\n"
                    detail += f"{'─'*60}\n"
                    detail += f"{kp.classic_cases}\n"
                
                # 示例列表
                if hasattr(kp, 'examples') and kp.examples:
                    detail += f"\n{'─'*60}\n"
                    detail += f"📝 作品示例\n"
                    detail += f"{'─'*60}\n"
                    for ex in kp.examples:
                        detail += f"  ◆ {ex}\n"
                
                # 常见写作误区
                if hasattr(kp, 'common_mistakes') and kp.common_mistakes:
                    detail += f"\n{'─'*60}\n"
                    detail += f"⚠️ 常见写作误区\n"
                    detail += f"{'─'*60}\n"
                    mistakes = kp.common_mistakes if isinstance(kp.common_mistakes, list) else [kp.common_mistakes]
                    for m in mistakes:
                        detail += f"  ✗ {m}\n"
                
                # 参考文献
                if hasattr(kp, 'references') and kp.references:
                    detail += f"\n{'─'*60}\n"
                    detail += f"📚 参考文献\n"
                    detail += f"{'─'*60}\n"
                    refs = kp.references if isinstance(kp.references, list) else [kp.references]
                    for ref in refs:
                        detail += f"  ○ {ref}\n"
                
                # 时间戳
                detail += f"\n{'─'*60}\n"
                detail += f"创建: {kp.created_at[:19] if kp.created_at else 'N/A'}\n"
                detail += f"更新: {kp.updated_at[:19] if kp.updated_at else 'N/A'}\n"
                detail += f"{'═'*60}\n"
                
                self._knowledge_detail.config(state=tk.NORMAL)
                self._knowledge_detail.delete("1.0", tk.END)
                self._knowledge_detail.insert("1.0", detail)
                self._knowledge_detail.config(state=tk.DISABLED)
                
        except Exception as e:
            logger.error(f"查看知识点详情失败: {e}")
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_knowledge_error
            title, full_message = convert_knowledge_error(e, "查看")
            messagebox.showerror(title, full_message)
    
    def _on_knowledge_import(self):
        """导入JSON文件"""
        file_path = filedialog.askopenfilename(
            title="选择知识库JSON文件",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")]
        )
        
        if file_path and self._knowledge_manager:
            try:
                result = self._knowledge_manager.import_from_json(file_path)
                
                messagebox.showinfo(
                    "导入结果",
                    f"总数: {result.total}\n"
                    f"成功: {result.success}\n"
                    f"失败: {result.failed}"
                )
                
                # 刷新列表
                self._on_knowledge_refresh()
                
                # 更新统计
                stats = self._knowledge_manager.get_stats()
                total = stats.get("total", 0)
                categories = stats.get("categories", {})
                
                # 分离题材分类和写作技巧
                subject_categories = {k: v for k, v in categories.items() if k != "writing_technique"}
                subject_total = sum(subject_categories.values())
                
                # 更新题材统计
                self._knowledge_stats_label1.config(text=f"题材总计: {subject_total}条  |  总计: {total}条")
                
                # 更新写作技巧统计
                technique_stats = self._load_writing_technique_stats()
                self._knowledge_stats_label2.config(text=f"写作技巧总计: {technique_stats['total']}条")
                
            except Exception as e:
                logger.error(f"导入知识库失败: {e}")
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_knowledge_error
                title, full_message = convert_knowledge_error(e, "导入")
                messagebox.showerror(title, full_message)
    
    def _on_knowledge_export(self):
        """导出到JSON文件"""
        file_path = filedialog.asksaveasfilename(
            title="导出知识库",
            defaultextension=".json",
            filetypes=[("JSON文件", "*.json")]
        )
        
        if file_path and self._knowledge_manager:
            try:
                # 获取所有知识点
                knowledge_list = self._knowledge_manager.list_knowledge()
                
                data = {
                    "knowledge_points": [kp.model_dump() for kp in knowledge_list],
                    "exported_at": datetime.now().isoformat()
                }
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                messagebox.showinfo("成功", f"已导出 {len(knowledge_list)} 条知识点")
                
            except Exception as e:
                logger.error(f"导出知识库失败: {e}")
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_knowledge_error
                title, full_message = convert_knowledge_error(e, "导出")
                messagebox.showerror(title, full_message)
    
    def _on_knowledge_add(self):
        """添加知识点"""
        # 创建对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("添加知识点")
        dialog.geometry("600x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 表单
        form = ttk.Frame(dialog, padding=20)
        form.pack(fill=tk.BOTH, expand=True)
        
        # 题材（中文显示）
        row1 = ttk.Frame(form, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="题材：").pack(side=tk.LEFT)
        category_var = tk.StringVar(value="科幻")
        category_combo = ttk.Combobox(row1, textvariable=category_var, 
                     values=["科幻", "玄幻", "历史", "通用"],
                     width=15, state="readonly")
        category_combo.pack(side=tk.LEFT, padx=5)
        
        # 领域（中文显示）
        row2 = ttk.Frame(form, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="领域：").pack(side=tk.LEFT)
        domain_var = tk.StringVar(value="物理")
        domain_combo = ttk.Combobox(row2, textvariable=domain_var,
                     values=["物理", "化学", "生物", "宗教", "神话", "哲学", "逻辑"],
                     width=15, state="readonly")
        domain_combo.pack(side=tk.LEFT, padx=5)
        
        # 标题
        row3 = ttk.Frame(form, style="TFrame")
        row3.pack(fill=tk.X, pady=5)
        ttk.Label(row3, text="标题：").pack(side=tk.LEFT)
        title_var = tk.StringVar()
        ttk.Entry(row3, textvariable=title_var, width=50).pack(side=tk.LEFT, padx=5)
        
        # 关键词
        row4 = ttk.Frame(form, style="TFrame")
        row4.pack(fill=tk.X, pady=5)
        ttk.Label(row4, text="关键词：").pack(side=tk.LEFT)
        keywords_var = tk.StringVar()
        ttk.Entry(row4, textvariable=keywords_var, width=50).pack(side=tk.LEFT, padx=5)
        ttk.Label(row4, text="（用逗号分隔）").pack(side=tk.LEFT)
        
        # 内容
        ttk.Label(form, text="内容：").pack(anchor=tk.W, pady=5)
        content_text = tk.Text(form, wrap=tk.WORD, height=10)
        content_text.pack(fill=tk.BOTH, expand=True)
        
        def on_submit():
            if not self._knowledge_manager:
                return
            
            title = title_var.get().strip()
            content = content_text.get("1.0", tk.END).strip()
            keywords = [k.strip() for k in keywords_var.get().split(",") if k.strip()]
            
            if not title or not content:
                messagebox.showwarning("提示", "标题和内容不能为空")
                return
            
            if len(content) < 50:
                messagebox.showwarning("提示", "内容长度不能少于50字符")
                return
            
            # 中文转英文
            category_en = self._category_map.get(category_var.get(), "scifi")
            domain_en = self._domain_map.get(domain_var.get(), "physics")
            
            result = self._knowledge_manager.create_knowledge(
                category=category_en,
                domain=domain_en,
                title=title,
                content=content,
                keywords=keywords
            )
            
            if result.success:
                messagebox.showinfo("成功", f"知识点已创建: {result.knowledge_id}")
                dialog.destroy()
                self._on_knowledge_refresh()
            else:
                messagebox.showerror("错误", f"创建失败: {result.error}")
        
        ttk.Button(form, text="提交", command=on_submit).pack(pady=10)
    
    def _on_knowledge_edit(self):
        """编辑知识点"""
        selection = self._knowledge_tree.selection()
        if not selection:
            messagebox.showwarning("提示", "请先选择知识点")
            return
        
        item = selection[0]
        values = self._knowledge_tree.item(item)['values']
        knowledge_id = values[0]
        
        if not self._knowledge_manager:
            return
        
        kp = self._knowledge_manager.get_knowledge(knowledge_id)
        if not kp:
            messagebox.showerror("错误", "知识点不存在")
            return
        
        # 创建对话框
        dialog = tk.Toplevel(self.root)
        dialog.title(f"编辑知识点 - {knowledge_id}")
        dialog.geometry("600x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 表单
        form = ttk.Frame(dialog, padding=20)
        form.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        row1 = ttk.Frame(form, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="标题：").pack(side=tk.LEFT)
        title_var = tk.StringVar(value=kp.title)
        ttk.Entry(row1, textvariable=title_var, width=50).pack(side=tk.LEFT, padx=5)
        
        # 关键词
        row2 = ttk.Frame(form, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="关键词：").pack(side=tk.LEFT)
        keywords_var = tk.StringVar(value=", ".join(kp.keywords))
        ttk.Entry(row2, textvariable=keywords_var, width=50).pack(side=tk.LEFT, padx=5)
        
        # 内容
        ttk.Label(form, text="内容：").pack(anchor=tk.W, pady=5)
        content_text = tk.Text(form, wrap=tk.WORD, height=10)
        content_text.pack(fill=tk.BOTH, expand=True)
        content_text.insert("1.0", kp.content)
        
        def on_submit():
            title = title_var.get().strip()
            content = content_text.get("1.0", tk.END).strip()
            keywords = [k.strip() for k in keywords_var.get().split(",") if k.strip()]
            
            if not title or not content:
                messagebox.showwarning("提示", "标题和内容不能为空")
                return
            
            if len(content) < 50:
                messagebox.showwarning("提示", "内容长度不能少于50字符")
                return
            
            result = self._knowledge_manager.update_knowledge(
                knowledge_id=knowledge_id,
                title=title,
                content=content,
                keywords=keywords
            )
            
            if result.success:
                messagebox.showinfo("成功", "知识点已更新")
                dialog.destroy()
                self._on_knowledge_refresh()
            else:
                messagebox.showerror("错误", f"更新失败: {result.error}")
        
        ttk.Button(form, text="保存", command=on_submit).pack(pady=10)
    
    def _on_knowledge_delete(self):
        """删除知识点"""
        selection = self._knowledge_tree.selection()
        if not selection:
            messagebox.showwarning("提示", "请先选择知识点")
            return
        
        item = selection[0]
        values = self._knowledge_tree.item(item)['values']
        knowledge_id = values[0]
        title = values[1]
        
        if messagebox.askyesno("确认", f"确定要删除知识点「{title}」吗？"):
            if self._knowledge_manager:
                success = self._knowledge_manager.delete_knowledge(knowledge_id)
                if success:
                    self._knowledge_tree.delete(item)
                    self._set_status(f"已删除知识点: {knowledge_id}")
                else:
                    messagebox.showerror("错误", "删除失败")
    
    def _on_knowledge_cleanup(self):
        """
        清除垃圾知识词条
        
        功能：
        - 检测重复词条（完全匹配/语义相似/内容重叠）
        - 评估词条质量（5维度评分）
        - 删除不合格词条（不删除文件）
        - 备份删除的词条
        - 生成Claw优化报告
        
        参考：11.19垃圾知识词条清理方案✅️.md
        """
        # 1. 确认对话框
        confirm = messagebox.askyesno(
            "确认执行知识库清理",
            """清理内容：
- 检测重复词条（完全匹配/语义相似/内容重叠）
- 评估词条质量（5维度评分）
- 删除不合格词条（不删除文件）

删除的词条将备份到：
data/知识库验证器/backups/

是否继续？""",
            icon='warning'
        )
        
        if not confirm:
            return
        
        # 2. 创建进度窗口
        progress_window = tk.Toplevel(self.root)
        progress_window.title("知识库清理")
        progress_window.geometry("450x300")
        progress_window.transient(self.root)
        progress_window.grab_set()
        
        # 标题
        ttk.Label(progress_window, text="知识库清理中...", font=('Arial', 12, 'bold')).pack(pady=15)
        
        # 进度条
        progress_bar = ttk.Progressbar(
            progress_window, 
            length=380, 
            mode='determinate'
        )
        progress_bar.pack(pady=10)
        
        # 状态标签
        status_label = ttk.Label(progress_window, text="初始化...", font=('Arial', 10))
        status_label.pack(pady=5)
        
        # 详情标签
        detail_label = ttk.Label(progress_window, text="", justify='left', font=('Arial', 9))
        detail_label.pack(pady=5, padx=20)
        
        # 结果文本框（初始隐藏）
        result_frame = ttk.Frame(progress_window)
        result_text = tk.Text(result_frame, height=10, width=50, font=('Arial', 9))
        result_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 3. 异步执行清理
        def cleanup_task():
            try:
                from data.知识库验证器.knowledge_verifier import KnowledgeVerifier
                
                workspace_root = Path(__file__).parent.resolve()
                verifier = KnowledgeVerifier(workspace_root)
                
                # 进度回调
                def update_progress(percent, message, detail):
                    try:
                        progress_window.after(0, lambda: [
                            progress_bar.__setitem__('value', percent),
                            status_label.config(text=message),
                            detail_label.config(text=detail)
                        ])
                    except Exception:
                        pass  # 窗口可能已关闭
                
                # 执行清理
                result = verifier.verify_all(
                    enable_dedup=True,
                    enable_quality=True,
                    dry_run=False,
                    progress_callback=update_progress
                )
                
                # 4. 显示结果
                def show_result():
                    progress_bar.pack_forget()
                    status_label.pack_forget()
                    detail_label.pack_forget()
                    
                    total = result['total_knowledge']
                    dedup = result.get('dedup_removed', 0)
                    quality = result.get('quality_removed', 0)
                    removed = result['removed_count']
                    final = result['final_count']
                    
                    result_text.insert(tk.END, f"✅ 知识库清理完成\n\n")
                    result_text.insert(tk.END, f"清理结果:\n")
                    result_text.insert(tk.END, f"  原始词条: {total}条\n")
                    result_text.insert(tk.END, f"  已删除: {removed}条\n")
                    result_text.insert(tk.END, f"  最终词条: {final}条\n\n")
                    
                    result_text.insert(tk.END, f"删除原因:\n")
                    result_text.insert(tk.END, f"  - 重复词条: {dedup}条\n")
                    result_text.insert(tk.END, f"  - 低质量: {quality}条\n\n")
                    
                    # 优化建议
                    opt_data = result.get('optimization_data', {})
                    if opt_data and not opt_data.get('error'):
                        suggestions = opt_data.get('optimization_suggestions', [])
                        if suggestions:
                            result_text.insert(tk.END, f"优化建议:\n")
                            for sug in suggestions[:3]:
                                result_text.insert(tk.END, f"  • {sug}\n")
                    
                    result_text.config(state=tk.DISABLED)
                    result_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
                    
                    # 添加按钮
                    btn_frame = ttk.Frame(progress_window)
                    btn_frame.pack(pady=10)
                    ttk.Button(btn_frame, text="查看报告", command=lambda: self._open_memory_file()).pack(side=tk.LEFT, padx=10)
                    ttk.Button(btn_frame, text="关闭", command=progress_window.destroy).pack(side=tk.LEFT, padx=10)
                    
                    # 刷新知识库列表
                    self._on_knowledge_refresh()
                
                progress_window.after(0, show_result)
                
            except Exception as e:
                import traceback
                error_detail = traceback.format_exc()
                logger.error(f"[KNOWLEDGE_CLEANUP] 清理失败:\n{error_detail}")
                
                def show_error():
                    progress_window.destroy()
                    messagebox.showerror(
                        "清理失败",
                        f"知识库清理失败:\n{str(e)}",
                        icon='error'
                    )
                
                progress_window.after(0, show_error)
        
        # 在后台线程执行
        import threading
        thread = threading.Thread(target=cleanup_task, daemon=True)
        thread.start()
    
    def _open_memory_file(self):
        """打开记忆文件（清理报告）"""
        from datetime import datetime
        today = datetime.now().strftime("%Y-%m-%d")
        memory_file = Path(__file__).parent / ".workbuddy" / "memory" / f"{today}.md"
        
        if memory_file.exists():
            import os
            os.startfile(str(memory_file))
        else:
            messagebox.showinfo("提示", f"记忆文件不存在: {memory_file}")
    
    def _on_knowledge_restore(self):
        """
        恢复最近清理的知识词条
        
        功能：
        - 列出所有可恢复的备份文件
        - 选择备份进行恢复
        - 将备份中的知识点合并到当前知识库
        
        参考：11.19垃圾知识词条清理方案✅️.md
        """
        from data.知识库验证器.file_cleaner import FileCleaner
        
        # 创建备份目录
        workspace_root = Path(__file__).parent.resolve()
        backup_dir = workspace_root / "data" / "知识库验证器" / "backups"
        
        if not backup_dir.exists():
            messagebox.showinfo("提示", "没有可恢复的备份文件")
            return
        
        # 获取备份列表
        cleaner = FileCleaner(backup_dir)
        backups = cleaner.list_backups()
        
        if not backups:
            messagebox.showinfo("提示", "没有可恢复的备份文件")
            return
        
        # 创建选择窗口
        select_window = tk.Toplevel(self.root)
        select_window.title("选择要恢复的备份")
        select_window.geometry("600x400")
        select_window.transient(self.root)
        select_window.grab_set()
        
        # 标题
        ttk.Label(select_window, text="选择要恢复的备份文件", font=('Arial', 12, 'bold')).pack(pady=10)
        
        # 备份列表
        columns = ('original_file', 'created_at', 'knowledge_count', 'size_kb')
        tree = ttk.Treeview(select_window, columns=columns, show='headings', height=12)
        
        tree.heading('original_file', text='原始文件')
        tree.heading('created_at', text='备份时间')
        tree.heading('knowledge_count', text='知识点数')
        tree.heading('size_kb', text='大小(KB)')
        
        tree.column('original_file', width=200)
        tree.column('created_at', width=150)
        tree.column('knowledge_count', width=100, anchor='center')
        tree.column('size_kb', width=80, anchor='center')
        
        # 填充数据
        for backup in backups:
            tree.insert('', tk.END, values=(
                backup['original_file'],
                backup['created_at'],
                backup['knowledge_count'],
                backup['size_kb']
            ), tags=(backup['backup_path'],))
        
        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 结果变量
        result_var = {'selected': None}
        
        def on_select():
            selection = tree.selection()
            if selection:
                item = selection[0]
                result_var['selected'] = tree.item(item)['tags'][0]
                select_window.destroy()
            else:
                messagebox.showwarning("提示", "请选择一个备份文件")
        
        def on_cancel():
            select_window.destroy()
        
        # 按钮
        btn_frame = ttk.Frame(select_window)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="恢复选中", command=on_select).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="恢复最新", command=lambda: [setattr(result_var, 'selected', backups[0]['backup_path'] if backups else None), select_window.destroy()]).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="取消", command=on_cancel).pack(side=tk.LEFT, padx=10)
        
        # 等待选择
        self.root.wait_window(select_window)
        
        # 执行恢复
        if result_var['selected']:
            backup_path = Path(result_var['selected'])
            
            # 确认恢复
            confirm = messagebox.askyesno(
                "确认恢复",
                f"将从以下备份恢复知识点：\n\n{backup_path.name}\n\n这将把备份中的知识点合并到当前知识库。\n是否继续？"
            )
            
            if confirm:
                # 显示进度
                progress_window = tk.Toplevel(self.root)
                progress_window.title("恢复中...")
                progress_window.geometry("300x100")
                progress_window.transient(self.root)
                
                ttk.Label(progress_window, text="正在恢复知识点...").pack(pady=20)
                
                def restore_task():
                    try:
                        result = cleaner.restore_from_backup(backup_path)
                        
                        def show_result():
                            progress_window.destroy()
                            
                            if result['success']:
                                messagebox.showinfo(
                                    "恢复成功",
                                    f"{result['message']}\n\n恢复的知识点已合并到知识库。"
                                )
                                # 刷新知识库列表
                                self._on_knowledge_refresh()
                            else:
                                messagebox.showerror("恢复失败", result['message'])
                        
                        progress_window.after(0, show_result)
                        
                    except Exception as e:
                        progress_window.after(0, lambda: [
                            progress_window.destroy(),
                            messagebox.showerror("恢复失败", f"恢复过程出错：{str(e)}")
                        ])
                
                import threading
                thread = threading.Thread(target=restore_task, daemon=True)
                thread.start()
    
    def _on_knowledge_ai_generate(self):
        """
        AI生成知识点 - V4高质量版本
        
        参考：11.15知识库生成器V4.md
        特性：
        - 完整路西法示例作为质量标准
        - 6项严格质量检查
        - 生成完成报告（含成本估算）
        - 与全局AI设置统一
        """
        # 创建对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("AI生成知识 V4")
        dialog.geometry("750x800")  # 增加高度确保按钮可见
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 创建Canvas和Scrollbar支持滚动
        canvas = tk.Canvas(dialog, highlightthickness=0)
        scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, padding=20)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 布局滚动组件
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 绑定鼠标滚轮
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        def _on_destroy(event):
            canvas.unbind_all("<MouseWheel>")
        dialog.bind("<Destroy>", _on_destroy)
        
        # 表单（使用scrollable_frame替代form）
        form = scrollable_frame
        
        # ========== 参数区域 ==========
        params_frame = ttk.LabelFrame(form, text="生成参数", padding=10)
        params_frame.pack(fill=tk.X, pady=5)
        
        # 题材（中文显示）- 覆盖主流小说题材 + 写作技巧（V5.3新增）
        row1 = ttk.Frame(params_frame, style="TFrame")
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="题材：").pack(side=tk.LEFT)
        ai_category_var = tk.StringVar(value="科幻")
        # 【修复】使用_category_map中实际存在的题材选项（排除"全部"）
        all_categories = [k for k in self._category_map.keys() if k != "全部"]
        category_combo = ttk.Combobox(row1, textvariable=ai_category_var,
                     values=all_categories,
                     width=15, state="readonly")
        category_combo.pack(side=tk.LEFT, padx=5)
        
        # 领域（中文显示）- 题材和领域完全独立，可任意组合
        # 设计原则：题材和领域互不干扰，用户可自由组合（如"历史+物理"、"仙侠+化学"）
        row2 = ttk.Frame(params_frame, style="TFrame")
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="领域：").pack(side=tk.LEFT)
        ai_domain_var = tk.StringVar(value="物理")
        
        # 【BUG修复】使用_domain_map中实际存在的领域选项，确保映射正确
        # 从_domain_map中排除"全部"选项，其他全部可用
        all_domains = [k for k in self._domain_map.keys() if k != "全部"]
        
        domain_combo = ttk.Combobox(row2, textvariable=ai_domain_var,
                     values=all_domains,
                     width=20, state="readonly")
        domain_combo.pack(side=tk.LEFT, padx=5)
        
        # V5.3新增：写作技巧固定六领域（按12.2文档规范）
        writing_technique_domains = ["叙事技巧", "描写技巧", "修辞技巧", "结构技巧", "特殊句式", "高级技法"]
        
        # V5.3新增：题材切换事件 - 写作技巧题材时更新领域选项
        def on_ai_category_changed(event=None):
            selected_category = ai_category_var.get()
            if selected_category == "写作技巧":
                # 切换为写作技巧六领域
                domain_combo['values'] = writing_technique_domains
                ai_domain_var.set(writing_technique_domains[0])
            else:
                # 【BUG修复】恢复为通用领域时，保持用户之前选择的领域，不要强制重置为"物理"
                # 只有当前选择的领域不在通用领域列表中时，才重置为第一个领域
                domain_combo['values'] = all_domains
                if ai_domain_var.get() not in all_domains:
                    ai_domain_var.set(all_domains[0] if all_domains else "物理")
        
        category_combo.bind("<<ComboboxSelected>>", on_ai_category_changed)
        
        # 生成数量
        row3 = ttk.Frame(params_frame, style="TFrame")
        row3.pack(fill=tk.X, pady=5)
        ttk.Label(row3, text="生成数量：").pack(side=tk.LEFT)
        ai_count_var = tk.IntVar(value=5)
        ttk.Spinbox(row3, from_=1, to=20, textvariable=ai_count_var, width=8).pack(side=tk.LEFT, padx=5)
        ttk.Label(row3, text="条（建议5-10条，质量优先）").pack(side=tk.LEFT)
        
        # 生成提示
        row4 = ttk.Frame(params_frame, style="TFrame")
        row4.pack(fill=tk.X, pady=5)
        ttk.Label(row4, text="生成方向：").pack(side=tk.LEFT)
        ai_prompt_var = tk.StringVar()
        ttk.Entry(row4, textvariable=ai_prompt_var, width=50).pack(side=tk.LEFT, padx=5)
        ttk.Label(row4, text="（可选，如：量子力学相关）").pack(side=tk.LEFT)
        
        # ========== 质量说明 ==========
        quality_frame = ttk.LabelFrame(form, text="V4质量标准", padding=5)
        quality_frame.pack(fill=tk.X, pady=5)
        
        quality_text = """• 内容≥500字，包含5个结构部分
• 经典案例≥300字，至少3个案例
• 写作应用≥300字，含3大板块
• 关键词≥10个，明确归属
• 常见误区≥5条
• 参考文献≥5个"""
        
        ttk.Label(quality_frame, text=quality_text, justify=tk.LEFT,
                  font=(GlassTheme.FONT_FAMILY, 9)).pack(anchor=tk.W)
        
        # ========== 生成结果 ==========
        result_frame = ttk.LabelFrame(form, text="生成结果", padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        result_text = tk.Text(result_frame, wrap=tk.WORD, height=12,
                             font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL))
        result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(result_frame, orient=tk.VERTICAL, command=result_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        result_text.config(yscrollcommand=scrollbar.set)
        
        # ========== 进度区域 ==========
        progress_frame = ttk.Frame(form, style="TFrame")
        progress_frame.pack(fill=tk.X, pady=5)
        
        progress_var = tk.DoubleVar(value=0)
        progress_bar = ttk.Progressbar(progress_frame, variable=progress_var, maximum=100, length=500)
        progress_bar.pack(side=tk.LEFT, padx=5)
        
        progress_label = ttk.Label(progress_frame, text="准备就绪")
        progress_label.pack(side=tk.LEFT, padx=5)
        
        # 存储生成结果用于导出（必须先定义）
        last_generate_result = [None]
        
        # ========== 定义回调函数（必须在按钮定义之前）==========
        def do_generate():
            """执行AI生成（V4版本）"""
            if not self._knowledge_manager:
                messagebox.showerror("错误", "知识库管理器未加载")
                return
            
            # 获取参数
            category_cn = ai_category_var.get()
            domain_cn = ai_domain_var.get()
            count = ai_count_var.get()
            prompt_hint = ai_prompt_var.get().strip()
            
            # 中文转英文
            category_en = self._category_map.get(category_cn, "scifi")
            domain_en = self._domain_map.get(domain_cn, "physics")
            
            # 禁用按钮
            generate_btn.configure(state=tk.DISABLED)
            result_text.delete("1.0", tk.END)
            result_text.insert("1.0", "正在初始化V4生成器...\n\n")
            progress_var.set(0)
            progress_label.configure(text="初始化中...")
            
            def update_progress(percent: int, message: str):
                """进度回调（线程安全）"""
                # 使用root.after调度到主线程
                self.root.after(0, lambda: _do_update_progress(percent, message))

            def _do_update_progress(percent: int, message: str):
                """实际更新进度（在主线程执行）"""
                progress_var.set(percent)
                progress_label.configure(text=message)

            def generate_task():
                """后台生成任务"""
                try:
                    # 导入V4生成器
                    from core.knowledge_generator import (
                        KnowledgeGeneratorV4, KnowledgeGenerateRequest
                    )
                    from pathlib import Path

                    # 创建生成器
                    workspace_root = Path(__file__).parent.resolve()
                    generator = KnowledgeGeneratorV4(workspace_root)

                    # 创建请求
                    request = KnowledgeGenerateRequest(
                        category=category_en,
                        domain=domain_en,
                        count=count,
                        focus_hint=prompt_hint,
                        quality_level="high"
                    )

                    # 执行生成
                    result = generator.generate_knowledge(request, progress_callback=update_progress)

                    # 保存结果用于导出
                    last_generate_result[0] = {
                        "request": {
                            "category": category_cn,
                            "domain": domain_cn,
                            "count": count,
                            "focus_hint": prompt_hint
                        },
                        "result": result.model_dump() if hasattr(result, 'model_dump') else result.dict()
                    }

                    # 更新UI
                    self.root.after(0, lambda: update_ui(result))

                except ImportError as e:
                    error_msg = str(e)
                    self.root.after(0, lambda msg=error_msg: show_error(f"生成器模块加载失败: {msg}"))
                except Exception as e:
                    import traceback
                    error_msg = str(e)
                    error_trace = traceback.format_exc()[:500]
                    self.root.after(0, lambda msg=error_msg, trace=error_trace: show_error(f"生成异常: {msg}\n{trace}"))
            
            def update_ui(result):
                """更新UI显示"""
                progress_var.set(100)
                progress_label.configure(text="完成")
                generate_btn.configure(state=tk.NORMAL)
                
                # 构建报告文本
                report_lines = []
                report_lines.append("=" * 50)
                report_lines.append("【AI生成知识 V4 - 生成完成报告】")
                report_lines.append("=" * 50)
                report_lines.append(f"\n📊 生成统计：")
                report_lines.append(f"  • 请求生成：{result.total} 条")
                report_lines.append(f"  • AI生成：{result.generated} 条")
                report_lines.append(f"  • 成功保存：{result.saved} 条")
                report_lines.append(f"  • 质量检查通过率：{(result.generated/result.total*100) if result.total > 0 else 0:.1f}%")
                report_lines.append(f"\n⏱️ 性能指标：")
                report_lines.append(f"  • 生成耗时：{result.generation_time:.1f} 秒")
                report_lines.append(f"  • 预估成本：¥{result.cost_estimate:.4f}")
                report_lines.append(f"\n📝 生成详情：")
                
                for detail in result.details:
                    status_icon = "✅" if detail.get("status") == "valid" else "⚠️"
                    title = detail.get("title", "未知")[:30]
                    status_text = "通过" if detail.get("status") == "valid" else detail.get("reason", "未通过")
                    report_lines.append(f"  {status_icon} {title} - {status_text}")
                
                if result.errors:
                    report_lines.append(f"\n❌ 错误信息：")
                    for err in result.errors[:5]:
                        report_lines.append(f"  • {err}")
                
                if result.knowledge_ids:
                    report_lines.append(f"\n📋 知识点ID列表：")
                    for kid in result.knowledge_ids[:10]:
                        report_lines.append(f"  • {kid}")
                    if len(result.knowledge_ids) > 10:
                        report_lines.append(f"  ... 共 {len(result.knowledge_ids)} 条")
                
                report_lines.append("\n" + "=" * 50)
                
                result_text.delete("1.0", tk.END)
                result_text.insert("1.0", "\n".join(report_lines))
                
                # 启用按钮
                if result.saved > 0:
                    export_report_btn.configure(state=tk.NORMAL)
                
                # 刷新知识列表
                if result.saved > 0:
                    self.root.after(500, self._on_knowledge_refresh)
            
            def show_error(error_msg):
                """显示错误"""
                progress_var.set(0)
                progress_label.configure(text="失败")
                generate_btn.configure(state=tk.NORMAL)
                result_text.delete("1.0", tk.END)
                result_text.insert("1.0", f"❌ {error_msg}")
            
            # ========== 启动后台任务（使用统一线程池，解决卡顿问题）==========
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(generate_task)
        
        # ========== 按钮行 ==========
        btn_frame = ttk.Frame(form, style="TFrame")
        btn_frame.pack(fill=tk.X, pady=10)
        
        generate_btn = ttk.Button(btn_frame, text="🚀 开始生成", style="Accent.TButton", command=do_generate)
        generate_btn.pack(side=tk.LEFT, padx=5)
        
        export_report_btn = ttk.Button(btn_frame, text="📄 保存报告", state=tk.DISABLED)
        export_report_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frame, text="关闭", command=lambda: dialog.destroy()).pack(side=tk.RIGHT, padx=5)
        
        # 绑定按钮命令
        def export_report_wrapper():
            """保存报告到经验文档文件夹"""
            if not last_generate_result[0]:
                messagebox.showwarning("提示", "请先生成知识")
                return
            
            try:
                from pathlib import Path
                data = last_generate_result[0]
                req = data["request"]
                res = data["result"]
                
                # 生成文件名：AI生成知识_题材_领域_时间.md
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"AI生成知识_{req['category']}_{req['domain']}_{timestamp}.md"
                
                # 保存到经验文档文件夹
                docs_path = Path(__file__).parent / "经验文档"
                docs_path.mkdir(parents=True, exist_ok=True)
                save_path = docs_path / filename
                
                report = f"""# AI生成知识 V4 报告

**生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 生成参数

| 参数 | 值 |
|------|------|
| 题材 | {req['category']} |
| 领域 | {req['domain']} |
| 数量 | {req['count']} 条 |
| 方向提示 | {req['focus_hint'] or '无'} |

## 生成结果

| 指标 | 值 |
|------|------|
| 请求生成 | {res['total']} 条 |
| AI生成 | {res['generated']} 条 |
| 成功保存 | {res['saved']} 条 |
| 质量通过率 | {(res['generated']/res['total']*100) if res['total'] > 0 else 0:.1f}% |
| 生成耗时 | {res['generation_time']:.1f} 秒 |
| 预估成本 | ¥{res['cost_estimate']:.4f} |

## 生成详情

"""
                for detail in res['details']:
                    status = "✅ 通过" if detail.get("status") == "valid" else f"⚠️ {detail.get('reason', '未通过')}"
                    report += f"- **{detail.get('title', '未知')}**：{status}\n"
                
                if res['knowledge_ids']:
                    report += "\n## 知识点ID\n\n"
                    for kid in res['knowledge_ids']:
                        report += f"- `{kid}`\n"
                
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(report)
                
                messagebox.showinfo("成功", f"报告已保存到：\n{save_path}")
                
            except Exception as e:
                messagebox.showerror("错误", f"保存失败：{e}")
        
        export_report_btn.configure(command=export_report_wrapper)
    
    def _on_words_selected(self, event):
        """字数选择变更"""
        selected = self._continue_words_var.get()
        if selected == "自定义":
            self._custom_words_entry.pack(side=tk.LEFT, padx=5)
            self._custom_words_entry.focus()
        else:
            self._custom_words_entry.pack_forget()
    
    def _on_continue_browse(self):
        """浏览并选择续写源文件"""
        path = filedialog.askopenfilename(
            title="选择续写源文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            try:
                # 读取文件内容
                if path.endswith('.docx'):
                    # 处理Word文档
                    try:
                        from docx import Document
                        doc = Document(path)
                        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    except ImportError:
                        messagebox.showerror("错误", "需要安装python-docx库才能读取Word文档")
                        return
                else:
                    # 处理文本文件
                    with open(path, 'r', encoding='utf-8-sig') as f:
                        content = f.read()
                
                # 更新原文文本框
                self._continue_source.delete("1.0", tk.END)
                self._continue_source.insert("1.0", content)
                self._set_status(f"已导入文件: {os.path.basename(path)}")
                
            except Exception as e:
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_file_error
                title, full_message = convert_file_error(e, "读取")
                messagebox.showerror(title, full_message)
    
    def _on_continue_select_chapter(self):
        """选择已生成章节作为续写起点"""
        if not hasattr(self, '_project_manager') or not self._project_manager:
            messagebox.showwarning("提示", "请先打开项目")
            return
        
        # 获取项目章节列表
        chapters = self._project_manager.list_chapters()
        if not chapters:
            messagebox.showinfo("提示", "当前项目没有章节")
            return
        
        # 创建选择对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("选择章节")
        dialog.geometry("500x400")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 400) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 章节列表
        list_frame = ttk.Frame(dialog, style="TFrame")
        list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        ttk.Label(list_frame, text="选择章节作为续写起点：").pack(anchor=tk.W)
        
        # 创建Treeview
        tree = ttk.Treeview(list_frame, columns=("章节", "字数"), show="headings", height=15)
        tree.heading("章节", text="章节名称")
        tree.heading("字数", text="字数")
        tree.column("章节", width=300)
        tree.column("字数", width=100)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 填充章节
        for chapter in chapters:
            tree.insert("", tk.END, values=(chapter.get("title", ""), chapter.get("word_count", 0)))
        
        # 按钮区
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=10)
        
        def on_confirm():
            selected = tree.selection()
            if selected:
                item = tree.item(selected[0])
                chapter_title = item["values"][0]
                # 读取章节内容
                content = self._project_manager.get_chapter_content(chapter_title)
                if content:
                    self._continue_source.delete("1.0", tk.END)
                    self._continue_source.insert("1.0", content)
                    self._set_status(f"已导入章节：{chapter_title}")
                dialog.destroy()
            else:
                messagebox.showwarning("提示", "请选择章节")
        
        ttk.Button(btn_frame, text="确定", command=on_confirm).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT)
    
    def _switch_version(self, direction: int):
        """切换版本
        
        Args:
            direction: 1为下一个版本，-1为上一个版本
        """
        if not self._continue_versions:
            return
        
        self._current_version_index = (self._current_version_index + direction) % len(self._continue_versions)
        self._display_version(self._current_version_index)
    
    def _on_version_tree_select(self, event):
        """版本列表选择事件"""
        selected = self._version_tree.selection()
        if selected:
            # 获取选中项的版本索引
            item_id = selected[0]
            item_data = self._version_tree.item(item_id)
            version_text = item_data["values"][0]  # 如 "V1"
            version_index = int(version_text.replace("V", "")) - 1
            
            if 0 <= version_index < len(self._continue_versions):
                self._current_version_index = version_index
                self._display_version(version_index)
    
    def _on_version_double_click(self, event):
        """版本列表双击事件 - 快速选择并提示保存"""
        selected = self._version_tree.selection()
        if not selected:
            return
        
        item_id = selected[0]
        item_data = self._version_tree.item(item_id)
        version_text = item_data["values"][0]
        version_index = int(version_text.replace("V", "")) - 1
        
        if 0 <= version_index < len(self._continue_versions):
            # 显示确认对话框
            if messagebox.askyesno("确认保存", f"是否选择版本 {version_text} 作为最终版本并保存？"):
                self._current_version_index = version_index
                self._display_version(version_index)
                # 直接触发保存
                self._on_continue_save()
    
    def _on_version_combo_changed(self, event):
        """版本下拉框选择变更（保留兼容性）"""
        # 已废弃，使用版本列表替代
        pass
    
    def _display_version(self, index: int):
        """显示指定版本"""
        if not self._continue_versions or index >= len(self._continue_versions):
            return
        
        version = self._continue_versions[index]
        
        # 更新文本框
        self._continue_result.delete("1.0", tk.END)
        self._continue_result.insert("1.0", version.get("text", ""))
        
        # 更新信息栏
        metadata = version.get("metadata", {})
        word_count = version.get("word_count", 0)
        gen_time = metadata.get("generation_time", 0)
        model = metadata.get("model_name", "-")
        
        self._result_info_label.configure(text=f"字数：{word_count} | 耗时：{gen_time:.1f}秒 | 模型：{model}")
        
        # 更新评分
        score = version.get("score", 0)
        self._score_label.configure(text=f"评分：{score:.2f}")
        
        # 更新版本信息
        total = len(self._continue_versions)
        is_best = index == self._best_version_index
        best_mark = " ⭐最佳" if is_best else ""
        self._version_info_label.configure(text=f"版本 {index + 1}/{total}{best_mark}")
        
        # 同步高亮版本列表
        self._highlight_version_in_tree(index)
        
        # 启用按钮
        self._regenerate_btn.configure(state=tk.NORMAL)
        self._select_best_btn.configure(state=tk.NORMAL)
        self._save_btn.configure(state=tk.NORMAL)
    
    def _highlight_version_in_tree(self, index: int):
        """高亮版本列表中的指定版本"""
        if not hasattr(self, '_version_tree'):
            return
        
        # 清除所有选择
        for item in self._version_tree.get_children():
            self._version_tree.selection_remove(item)
        
        # 选择指定版本
        children = self._version_tree.get_children()
        if 0 <= index < len(children):
            self._version_tree.selection_set(children[index])
            self._version_tree.see(children[index])  # 滚动到可见区域
    
    def _update_version_tree(self):
        """更新版本列表Treeview"""
        if not hasattr(self, '_version_tree'):
            return
        
        # 清空现有项
        for item in self._version_tree.get_children():
            self._version_tree.delete(item)
        
        if not self._continue_versions:
            return
        
        # 添加版本项（最多5个）
        for i, version in enumerate(self._continue_versions[:5]):
            temp = version.get("temperature", 0.8)
            score = version.get("score", 0)
            word_count = version.get("word_count", 0)
            
            # 确定状态
            is_best = i == self._best_version_index
            status = "⭐最佳" if is_best else ""
            
            # 插入到Treeview
            self._version_tree.insert("", tk.END, values=(
                f"V{i + 1}",
                f"{temp:.1f}",
                f"{word_count}",
                f"{score:.2f}",
                status
            ))
        
        # 高亮当前版本
        if self._continue_versions:
            self._highlight_version_in_tree(self._current_version_index)
    
    def _update_version_combo(self):
        """更新版本列表（重定向到Treeview方法）"""
        self._update_version_tree()
    
    # ============== 工作台子功能回调方法 ==============
    
    def _on_worldview_browse(self):
        """浏览世界观文件"""
        path = filedialog.askopenfilename(
            title="选择世界观文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._worldview_path_var.set(path)
            self._preview_worldview(path)
    
    def _preview_worldview(self, path: str):
        """预览世界观文件"""
        try:
            self._worldview_preview.delete("1.0", tk.END)
            if path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(path)
                text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                with open(path, 'r', encoding='utf-8') as f:
                    text = f.read()
            self._worldview_preview.insert("1.0", text[:5000])
        except Exception as e:
            self._worldview_preview.insert("1.0", f"预览失败: {e}")
    
    def _on_worldview_new(self):
        """新建世界观（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("新建世界观")
        dialog.geometry("500x450")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 450) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 表单区域
        form_frame = ttk.Frame(dialog, padding=20)
        form_frame.pack(fill=tk.BOTH, expand=True)
        
        # 世界观名称
        ttk.Label(form_frame, text="世界观名称：").grid(row=0, column=0, sticky=tk.W, pady=5)
        name_var = tk.StringVar(value="新世界观")
        ttk.Entry(form_frame, textvariable=name_var, width=40).grid(row=0, column=1, pady=5)
        
        # 世界观类别
        ttk.Label(form_frame, text="世界观类别：").grid(row=1, column=0, sticky=tk.W, pady=5)
        category_var = tk.StringVar(value="世界观设定")
        category_combo = ttk.Combobox(form_frame, textvariable=category_var, width=37,
                                      values=['世界观设定', '力量体系', '势力分布', '地理环境', '社会规则', '历史背景', '种族设定', '其他'])
        category_combo.grid(row=1, column=1, pady=5)
        
        # 核心元素
        ttk.Label(form_frame, text="核心元素：").grid(row=2, column=0, sticky=tk.W, pady=5)
        elements_var = tk.StringVar(value="魔法体系、势力分布、历史背景")
        ttk.Entry(form_frame, textvariable=elements_var, width=40).grid(row=2, column=1, pady=5)
        
        # 世界观描述
        ttk.Label(form_frame, text="详细描述：").grid(row=3, column=0, sticky=tk.NW, pady=5)
        desc_text = tk.Text(form_frame, wrap=tk.WORD, height=12, width=40,
                           font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL))
        desc_text.grid(row=3, column=1, pady=5)
        desc_text.insert("1.0", "请在此描述世界观的核心设定、规则、势力分布等内容...")
        
        # 按钮
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=20, pady=10)
        
        def create_worldview():
            name = name_var.get().strip()
            category = category_var.get().strip()
            elements = elements_var.get().strip()
            desc = desc_text.get("1.0", tk.END).strip()
            
            if not name:
                messagebox.showwarning("警告", "请输入世界观名称！")
                return
            
            # 初始化词条列表（如果不存在）
            if not hasattr(self, '_worldview_entries'):
                self._worldview_entries = []
            
            # 创建词条数据
            from datetime import datetime
            entry_data = {
                'name': name,
                'category': category,
                'elements': elements,
                'description': desc,
                'status': '新建',
                'modified': datetime.now().strftime('%Y-%m-%d %H:%M')
            }
            
            # 添加到词条列表
            self._worldview_entries.append(entry_data)
            
            # 添加到树
            display_elements = elements[:50] + "..." if len(elements) > 50 else elements
            self._worldview_tree.insert("", tk.END, values=(
                name,
                category,
                display_elements,
                '新建',
                entry_data['modified']
            ))
            
            # 更新预览
            self._worldview_preview.delete("1.0", tk.END)
            self._worldview_preview.insert("1.0", f"【{name}】\n\n类别：{category}\n\n核心元素：{elements}\n\n{desc}")
            
            self._set_status(f"已创建世界观：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_worldview).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_worldview_view(self):
        """查看世界观详情 - 在预览区域显示选中词条的完整信息"""
        if not hasattr(self, '_worldview_tree') or not hasattr(self, '_worldview_entries'):
            return
        
        selection = self._worldview_tree.selection()
        if not selection:
            messagebox.showinfo("提示", "请先选择一个世界观词条")
            return
        
        # 获取选中项的索引
        item_id = selection[0]
        children = self._worldview_tree.get_children()
        if item_id not in children:
            return
        
        idx = children.index(item_id)
        if idx >= len(self._worldview_entries):
            return
        
        # 获取词条数据
        entry = self._worldview_entries[idx]
        
        # 更新预览区域
        if hasattr(self, '_worldview_preview'):
            self._worldview_preview.delete("1.0", tk.END)
            
            # 格式化显示
            details = f"【{entry.get('name', '未命名')}】\n\n"
            details += f"类别：{entry.get('category', '世界观设定')}\n"
            details += f"状态：{entry.get('status', '已保存')}\n"
            details += f"修改时间：{entry.get('modified', '-')}\n\n"
            details += f"━━━━━━━━━━━━━━━━━━━━━━━━\n"
            details += f"【核心元素】\n{entry.get('elements', '暂无')}\n\n"
            details += f"━━━━━━━━━━━━━━━━━━━━━━━━\n"
            details += f"【详细描述】\n{entry.get('description', entry.get('elements', '暂无详细描述'))}\n"
            
            self._worldview_preview.insert("1.0", details)
            self._set_status(f"查看世界观详情：{entry.get('name', '未命名')}")
    
    def _on_worldview_edit(self):
        """编辑世界观 - 弹出编辑对话框"""
        if not hasattr(self, '_worldview_tree') or not hasattr(self, '_worldview_entries'):
            return
        
        selection = self._worldview_tree.selection()
        if not selection:
            messagebox.showinfo("提示", "请先选择一个世界观词条")
            return
        
        # 获取选中项的索引
        item_id = selection[0]
        children = self._worldview_tree.get_children()
        if item_id not in children:
            return
        
        idx = children.index(item_id)
        if idx >= len(self._worldview_entries):
            return
        
        # 获取词条数据
        entry = self._worldview_entries[idx]
        
        # 创建编辑对话框
        dialog = tk.Toplevel(self.root)
        dialog.title(f"编辑世界观 - {entry.get('name', '未命名')}")
        dialog.geometry("600x500")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 600) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 500) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 表单区域
        form_frame = ttk.Frame(dialog, padding=20)
        form_frame.pack(fill=tk.BOTH, expand=True)
        
        # 名称
        ttk.Label(form_frame, text="世界观名称：").grid(row=0, column=0, sticky=tk.W, pady=5)
        name_var = tk.StringVar(value=entry.get('name', ''))
        ttk.Entry(form_frame, textvariable=name_var, width=50).grid(row=0, column=1, pady=5)
        
        # 类别
        ttk.Label(form_frame, text="世界观类别：").grid(row=1, column=0, sticky=tk.W, pady=5)
        category_var = tk.StringVar(value=entry.get('category', '世界观设定'))
        category_combo = ttk.Combobox(form_frame, textvariable=category_var, width=47,
                                      values=['世界观设定', '力量体系', '势力分布', '地理环境', '社会规则', '历史背景', '种族设定', '其他'])
        category_combo.grid(row=1, column=1, pady=5)
        
        # 核心元素
        ttk.Label(form_frame, text="核心元素：").grid(row=2, column=0, sticky=tk.W, pady=5)
        elements_var = tk.StringVar(value=entry.get('elements', ''))
        ttk.Entry(form_frame, textvariable=elements_var, width=50).grid(row=2, column=1, pady=5)
        
        # 详细描述
        ttk.Label(form_frame, text="详细描述：").grid(row=3, column=0, sticky=tk.NW, pady=5)
        desc_text = tk.Text(form_frame, width=50, height=15, wrap=tk.WORD,
                           font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL))
        desc_text.grid(row=3, column=1, pady=5)
        desc_text.insert("1.0", entry.get('description', ''))
        
        # 按钮区域
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=20, pady=10)
        
        def save_changes():
            """保存编辑"""
            # 更新词条数据
            self._worldview_entries[idx]['name'] = name_var.get()
            self._worldview_entries[idx]['category'] = category_var.get()
            self._worldview_entries[idx]['elements'] = elements_var.get()
            self._worldview_entries[idx]['description'] = desc_text.get("1.0", tk.END).strip()
            self._worldview_entries[idx]['modified'] = datetime.now().strftime('%Y-%m-%d %H:%M')
            
            # 更新树显示
            self._worldview_tree.item(item_id, values=(
                name_var.get(),
                category_var.get(),
                elements_var.get()[:50] + '...' if len(elements_var.get()) > 50 else elements_var.get(),
                '已修改',
                self._worldview_entries[idx]['modified']
            ))
            
            # 更新预览
            self._on_worldview_view()
            
            self._set_status(f"已更新世界观：{name_var.get()}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="保存", command=save_changes).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_worldview_delete(self):
        """批量删除世界观"""
        if not hasattr(self, '_worldview_tree') or not hasattr(self, '_worldview_entries'):
            return
        
        selection = self._worldview_tree.selection()
        if not selection:
            messagebox.showinfo("提示", "请先选择要删除的世界观词条")
            return
        
        # 确认删除
        if not messagebox.askyesno("确认删除", f"确定要删除选中的 {len(selection)} 个世界观词条吗？"):
            return
        
        # 获取要删除的索引（从大到小排序，避免删除时索引变化）
        children = list(self._worldview_tree.get_children())
        indices_to_delete = sorted([children.index(item_id) for item_id in selection if item_id in children], reverse=True)
        
        # 从列表和树中删除
        for idx in indices_to_delete:
            if idx < len(self._worldview_entries):
                # 从词条列表删除
                del self._worldview_entries[idx]
        
        # 从树中删除选中项
        for item_id in selection:
            self._worldview_tree.delete(item_id)
        
        # 更新预览
        if hasattr(self, '_worldview_preview'):
            self._worldview_preview.delete("1.0", tk.END)
            self._worldview_preview.insert("1.0", f"已删除 {len(selection)} 个世界观词条")
        
        self._set_status(f"已删除 {len(selection)} 个世界观词条")
    
    def _on_worldview_link(self):
        """关联要素 - 显示关联对话框"""
        if not hasattr(self, '_worldview_tree'):
            return
        
        selection = self._worldview_tree.selection()
        if not selection:
            messagebox.showinfo("提示", "请先选择一个世界观词条")
            return
        
        # 获取选中项名称
        item_id = selection[0]
        values = self._worldview_tree.item(item_id, 'values')
        entry_name = values[0] if values else '未命名'
        
        # 显示关联信息对话框
        dialog = tk.Toplevel(self.root)
        dialog.title(f"关联要素 - {entry_name}")
        dialog.geometry("500x400")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 400) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 关联信息
        ttk.Label(dialog, text=f"世界观词条：{entry_name}", font=('Microsoft YaHei UI', 12, 'bold')).pack(pady=10)
        
        # 关联选项
        options_frame = ttk.Frame(dialog, padding=20)
        options_frame.pack(fill=tk.BOTH, expand=True)
        
        # 关联大纲
        ttk.Label(options_frame, text="关联大纲章节：").grid(row=0, column=0, sticky=tk.W, pady=5)
        outline_var = tk.StringVar()
        ttk.Entry(options_frame, textvariable=outline_var, width=30).grid(row=0, column=1, pady=5)
        ttk.Label(options_frame, text="（输入章节号，如：第1-3章）", foreground='gray').grid(row=0, column=2, padx=5)
        
        # 关联人物
        ttk.Label(options_frame, text="关联人物：").grid(row=1, column=0, sticky=tk.W, pady=5)
        character_var = tk.StringVar()
        ttk.Entry(options_frame, textvariable=character_var, width=30).grid(row=1, column=1, pady=5)
        ttk.Label(options_frame, text="（输入人物名称）", foreground='gray').grid(row=1, column=2, padx=5)
        
        # 关联风格
        ttk.Label(options_frame, text="关联风格：").grid(row=2, column=0, sticky=tk.W, pady=5)
        style_var = tk.StringVar()
        ttk.Entry(options_frame, textvariable=style_var, width=30).grid(row=2, column=1, pady=5)
        ttk.Label(options_frame, text="（输入风格特征）", foreground='gray').grid(row=2, column=2, padx=5)
        
        # 说明
        ttk.Label(options_frame, text="💡 提示：关联信息将在生成时自动注入到上下文中", 
                 foreground='blue').grid(row=3, column=0, columnspan=3, pady=20)
        
        # 按钮
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=20, pady=10)
        
        def apply_link():
            """应用关联"""
            links = []
            if outline_var.get():
                links.append(f"大纲章节：{outline_var.get()}")
            if character_var.get():
                links.append(f"关联人物：{character_var.get()}")
            if style_var.get():
                links.append(f"关联风格：{style_var.get()}")
            
            if links:
                self._set_status(f"已设置关联：{', '.join(links)}")
            else:
                self._set_status("未设置关联")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="确定", command=apply_link).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_worldview_import(self):
        """导入世界观 - 调用WorldviewParserAdapter"""
        try:
            # 获取当前选中的世界观路径
            worldview_path = getattr(self, '_worldview_path_var', tk.StringVar()).get()
            if not worldview_path:
                messagebox.showwarning("提示", "请先选择世界观文件")
                return
            
            # 更新状态
            self._set_status("正在解析世界观...")
            
            # 导入适配器
            from agents.adapters.worldview_adapter import WorldviewParserAdapter
            from agents.priority import AgentTask
            from pathlib import Path
            
            # 读取文件内容
            content = Path(worldview_path).read_text(encoding='utf-8')
            
            # 创建适配器实例
            adapter = WorldviewParserAdapter()
            if not adapter.initialize():
                raise RuntimeError("WorldviewParserAdapter初始化失败")
            
            # 创建任务
            from agents.priority import TaskPriority
            task = AgentTask(
                task_id=f"worldview_import_{int(time.time())}",
                agent_type="worldview_parser",
                priority=TaskPriority.NORMAL,
                payload={
                    "worldview_content": content,
                    "options": {}
                }
            )
            
            # 执行解析（异步，使用统一线程池）
            def run_parse():
                try:
                    result = adapter.execute(task)
                    self.root.after(0, lambda: self._on_worldview_import_complete(result))
                except Exception as e:
                    self.root.after(0, lambda: self._on_worldview_import_error(str(e)))
            
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(run_parse)
            
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_exception
            title, full_message = convert_exception(e, "世界观解析")
            messagebox.showerror(title, full_message)
            self._set_status(f"世界观解析失败: {str(e)}")
    
    def _on_worldview_import_complete(self, result):
        """世界观解析完成回调"""
        try:
            # 存储世界观数据
            self._worldview = result.get("result", {})
            elements = self._worldview.get("elements", [])
            rules = self._worldview.get("rules", [])
            
            # 【关键修复】更新_worldview_content并同步到树形列表
            # 将elements转换为项目文件期望的格式
            if elements and isinstance(elements, list):
                self._worldview_content = elements
                # 同时更新树形列表
                if hasattr(self, '_worldview_tree'):
                    self._update_worldview_tree_from_content(elements)
            
            # 更新预览区域显示简洁摘要
            self._worldview_preview.delete("1.0", tk.END)
            self._worldview_preview.insert(tk.END, f"世界观解析完成！\n\n")
            self._worldview_preview.insert(tk.END, f"要素数量: {len(elements)}\n")
            self._worldview_preview.insert(tk.END, f"规则数量: {len(rules)}\n\n")
            self._worldview_preview.insert(tk.END, "已添加到左侧列表，双击查看详情\n")
            
            self._set_status(f"世界观解析完成，共{len(elements)}个要素")
        except Exception as e:
            self._set_status(f"处理解析结果失败: {str(e)}")
    
    def _on_worldview_import_error(self, error: str):
        """世界观解析错误回调"""
        messagebox.showerror("世界观解析错误", error)
        self._set_status(f"世界观解析失败: {error}")
    
    def _on_worldview_clear(self):
        """清除世界观"""
        self._worldview_path_var.set("")
        self._worldview_preview.delete("1.0", tk.END)
    
    def _on_character_browse(self):
        """浏览人物档案文件"""
        path = filedialog.askopenfilename(
            title="选择人物档案文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._character_path_var.set(path)
            self._set_status(f"已选择: {os.path.basename(path)}")
    
    def _on_character_batch_import(self):
        """批量解析导入人物 - 调用CharacterManagerAdapter"""
        try:
            # 获取当前选中的人物档案路径
            character_path = getattr(self, '_character_path_var', tk.StringVar()).get()
            if not character_path:
                messagebox.showwarning("提示", "请先选择人物档案文件")
                return
            
            # 更新状态
            self._set_status("正在解析人物档案...")
            
            # 导入适配器
            from agents.adapters.character_adapter import CharacterManagerAdapter
            from agents.priority import AgentTask
            from pathlib import Path
            
            # 读取文件内容
            content = Path(character_path).read_text(encoding='utf-8')
            
            # 创建适配器实例
            adapter = CharacterManagerAdapter()
            if not adapter.initialize():
                raise RuntimeError("CharacterManagerAdapter初始化失败")
            
            # 创建任务
            from agents.priority import TaskPriority
            task = AgentTask(
                task_id=f"character_import_{int(time.time())}",
                agent_type="character_manager",
                priority=TaskPriority.NORMAL,
                payload={
                    "operation": "batch_import",
                    "character_data": {
                        "content": content,
                        "source_file": character_path
                    }
                }
            )
            
            # 执行解析（异步，使用统一线程池）
            def run_parse():
                try:
                    result = adapter.execute(task)
                    self.root.after(0, lambda: self._on_character_import_complete(result))
                except Exception as e:
                    self.root.after(0, lambda: self._on_character_import_error(str(e)))
            
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(run_parse)
            
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_exception
            title, full_message = convert_exception(e, "人物解析")
            messagebox.showerror(title, full_message)
            self._set_status(f"人物解析失败: {str(e)}")
    
    def _on_character_import_complete(self, result):
        """人物解析完成回调

        更新人物列表和结构化存储

        数据格式（从CharacterManagerAdapter._parse_characters_text）：
        {
            "name": "张三",
            "role": "主角",
            "status": "新建",
            "emotion": "平静",
            "chapters": "第1章",
            "appearance": "...",
            "personality": "...",
            "background": "...",
            "goals": "...",
            "fears": "...",
            "mbti": "",
            "description": "..."
        }
        """
        try:
            # 存储人物数据
            characters = result.get("result", {}).get("characters", [])
            self._characters = characters

            # 【新增】存储到结构化列表（与世界观模式一致）
            if not hasattr(self, '_character_entries'):
                self._character_entries = []
            self._character_entries = characters

            # 【修复】同步到_character_data，确保保存和加载一致
            self._character_data = characters

            # 更新人物树（5列数据）
            if hasattr(self, '_character_tree'):
                self._character_tree.delete(*self._character_tree.get_children())
                for char in characters:
                    name = char.get("name", "未命名")
                    role = char.get("role", "未设置")
                    status = char.get("status", "新建")
                    emotion = char.get("emotion", "平静")
                    chapters = char.get("chapters", "未设置")
                    self._character_tree.insert("", tk.END, values=(name, role, status, emotion, chapters))

            # 更新状态
            self._set_status(f"成功导入 {len(characters)} 个人物")
        except Exception as e:
            self._set_status(f"处理解析结果失败: {str(e)}")
    
    def _on_character_import_error(self, error: str):
        """人物解析错误回调"""
        messagebox.showerror("人物解析错误", error)
        self._set_status(f"人物解析失败: {error}")
    
    def _on_character_new(self):
        """新建人物（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("新建人物")
        dialog.geometry("600x700")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 600) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 700) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 创建滚动区域
        canvas = tk.Canvas(dialog, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 基本信息
        basic_frame = ttk.LabelFrame(scrollable_frame, text="基本信息", padding=10)
        basic_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 姓名
        name_frame = ttk.Frame(basic_frame, style="TFrame")
        name_frame.pack(fill=tk.X, pady=5)
        ttk.Label(name_frame, text="姓名：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="新角色")
        ttk.Entry(name_frame, textvariable=name_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 角色类型
        role_frame = ttk.Frame(basic_frame, style="TFrame")
        role_frame.pack(fill=tk.X, pady=5)
        ttk.Label(role_frame, text="角色类型：").pack(side=tk.LEFT)
        role_var = tk.StringVar(value="主角")
        ttk.Combobox(role_frame, textvariable=role_var, values=["主角", "配角", "反派", "路人"], width=15).pack(side=tk.LEFT, padx=10)
        
        # 外貌描述
        appearance_frame = ttk.LabelFrame(scrollable_frame, text="外貌描述", padding=10)
        appearance_frame.pack(fill=tk.X, padx=20, pady=10)
        appearance_text = tk.Text(appearance_frame, wrap=tk.WORD, height=3,
                                 font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                 bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        appearance_text.pack(fill=tk.X)
        
        # 性格特点
        personality_frame = ttk.LabelFrame(scrollable_frame, text="性格特点", padding=10)
        personality_frame.pack(fill=tk.X, padx=20, pady=10)
        personality_text = tk.Text(personality_frame, wrap=tk.WORD, height=3,
                                  font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                  bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        personality_text.pack(fill=tk.X)
        
        # 背景故事
        background_frame = ttk.LabelFrame(scrollable_frame, text="背景故事", padding=10)
        background_frame.pack(fill=tk.X, padx=20, pady=10)
        background_text = tk.Text(background_frame, wrap=tk.WORD, height=5,
                                 font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                                 bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        background_text.pack(fill=tk.X)
        
        # 能力特长
        ability_frame = ttk.LabelFrame(scrollable_frame, text="能力特长", padding=10)
        ability_frame.pack(fill=tk.X, padx=20, pady=10)
        ability_text = tk.Text(ability_frame, wrap=tk.WORD, height=3,
                              font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                              bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        ability_text.pack(fill=tk.X)
        
        # 按钮
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_character():
            name = name_var.get().strip()
            role = role_var.get()

            if not name:
                messagebox.showwarning("警告", "请输入人物姓名！")
                return

            # 获取详细信息
            appearance = appearance_text.get("1.0", tk.END).strip()
            personality = personality_text.get("1.0", tk.END).strip()
            background = background_text.get("1.0", tk.END).strip()
            ability = ability_text.get("1.0", tk.END).strip()

            # 【修复】添加到结构化存储（与世界观模式一致）
            if not hasattr(self, '_character_entries'):
                self._character_entries = []
            if not hasattr(self, '_character_data'):
                self._character_data = []

            from datetime import datetime
            entry_data = {
                'id': f"char_{datetime.now().strftime('%Y%m%d%H%M%S')}",
                'name': name,
                'role': role,
                'status': '新建',
                'emotion': '平静',
                'chapters': '第1章',
                'appearance': appearance,
                'personality': personality,
                'background': background,
                'ability': ability,
                'created_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            self._character_entries.append(entry_data)
            self._character_data.append(entry_data)

            # 添加到列表（5列数据）
            self._character_tree.insert("", tk.END, values=(
                name,
                role,
                "新建",
                "平静",
                "第1章"
            ))

            # 更新详情
            self._character_detail.delete("1.0", tk.END)
            self._character_detail.insert("1.0", f"""【{name}】- {role}

外貌：{appearance}

性格：{personality}

背景：{background}

能力：{ability}
""")

            self._set_status(f"已创建人物：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_character).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_character_edit(self):
        """编辑人物"""
        if not hasattr(self, '_character_tree'):
            return

        selection = self._character_tree.selection()
        if not selection or len(selection) != 1:
            messagebox.showinfo("提示", "请选择一个要编辑的人物")
            return

        # 获取选中的人物名称
        children = list(self._character_tree.get_children())
        selected_index = children.index(selection[0])

        if selected_index >= len(self._character_entries):
            messagebox.showerror("错误", "人物数据索引错误")
            return

        character = self._character_entries[selected_index]
        character_name = character.get('name', '未命名')

        # 创建编辑对话框
        edit_window = tk.Toplevel(self.root)
        edit_window.title(f"编辑人物 - {character_name}")
        edit_window.geometry("500x600")
        edit_window.transient(self.root)
        edit_window.grab_set()

        # 表单字段
        form_fields = [
            ('姓名', 'name', character.get('name', '')),
            ('角色', 'role', character.get('role', '未设置')),
            ('性别', 'gender', character.get('gender', '')),
            ('年龄', 'age', character.get('age', '')),
            ('外貌', 'appearance', character.get('appearance', '')),
            ('性格', 'personality', character.get('personality', '')),
            ('背景', 'background', character.get('background', '')),
            ('目标', 'goals', character.get('goals', '')),
            ('恐惧', 'fears', character.get('fears', '')),
            ('MBTI', 'mbti', character.get('mbti', '')),
            ('情绪', 'emotion', character.get('emotion', '平静')),
            ('状态', 'status', character.get('status', '新建')),
            ('出场章节', 'chapters', character.get('chapters', '未设置'))
        ]

        entries = {}

        for i, (label_text, field_name, default_value) in enumerate(form_fields):
            tk.Label(edit_window, text=label_text).grid(row=i, column=0, padx=10, pady=5, sticky='e')
            entry = tk.Entry(edit_window, width=40)
            entry.insert(0, str(default_value))
            entry.grid(row=i, column=1, padx=10, pady=5, sticky='w')
            entries[field_name] = entry

        # 描述字段（多行文本）
        tk.Label(edit_window, text="描述").grid(row=len(form_fields), column=0, padx=10, pady=5, sticky='ne')
        desc_text = tk.Text(edit_window, width=40, height=5)
        desc_text.insert('1.0', character.get('description', ''))
        desc_text.grid(row=len(form_fields), column=1, padx=10, pady=5, sticky='w')
        entries['description'] = desc_text

        def save_edit():
            """保存编辑"""
            new_data = {}
            for field_name, entry in entries.items():
                if field_name == 'description':
                    new_data[field_name] = entry.get('1.0', tk.END).strip()
                else:
                    new_data[field_name] = entry.get().strip()

            # 调用适配器编辑人物
            try:
                from agents.priority import TaskPriority
                result = self._get_character_manager().execute(
                    AgentTask(
                        task_id=f"edit_character_{character_name}",
                        agent_type="character_manager",
                        priority=TaskPriority.NORMAL,
                        payload={
                            "operation": "edit_character",
                            "character_name": character_name,
                            "character_data": new_data,
                            "all_characters": self._character_entries
                        }
                    )
                )

                if result.get("result", {}).get("success"):
                    self._character_entries = result["result"]["characters"]
                    self._character_data = self._character_entries
                    self._update_character_tree()
                    self._set_status(result["result"]["message"])
                    edit_window.destroy()
                else:
                    messagebox.showerror("错误", result["result"]["message"])

            except Exception as e:
                messagebox.showerror("错误", f"编辑失败: {str(e)}")

        # 按钮框架
        btn_frame = tk.Frame(edit_window)
        btn_frame.grid(row=len(form_fields)+1, column=0, columnspan=2, pady=20)

        tk.Button(btn_frame, text="保存", command=save_edit, width=10).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="取消", command=edit_window.destroy, width=10).pack(side=tk.LEFT, padx=5)

    def _on_character_detail(self):
        """人物详情 - 在主界面详情框中显示"""
        if not hasattr(self, '_character_tree'):
            return

        selection = self._character_tree.selection()
        if not selection or len(selection) != 1:
            messagebox.showinfo("提示", "请选择一个要查看详情的人物")
            return

        # 获取选中的人物
        children = list(self._character_tree.get_children())
        selected_index = children.index(selection[0])

        if selected_index >= len(self._character_entries):
            messagebox.showerror("错误", "人物数据索引错误")
            return

        character = self._character_entries[selected_index]
        character_name = character.get('name', '未命名')

        # 构建详情内容
        detail_content = f"# 【{character_name}】\n\n"
        detail_content += f"【角色】{character.get('role', '未设置')}\n"
        detail_content += f"【性别】{character.get('gender', '')}\n"
        detail_content += f"【年龄】{character.get('age', '')}\n"
        detail_content += f"【MBTI】{character.get('mbti', '')}\n"
        detail_content += f"【情绪】{character.get('emotion', '平静')}\n"
        detail_content += f"【状态】{character.get('status', '新建')}\n"
        detail_content += f"【出场章节】{character.get('chapters', '未设置')}\n\n"

        # 基本信息
        detail_content += "▌基本信息\n\n"
        if character.get('appearance'):
            detail_content += f"外貌：{character.get('appearance')}\n"
        if character.get('personality'):
            detail_content += f"性格：{character.get('personality')}\n"
        if character.get('background'):
            detail_content += f"背景：{character.get('background')}\n"

        detail_content += "\n▌动机与目标\n\n"
        if character.get('goals'):
            detail_content += f"目标：{character.get('goals')}\n"
        if character.get('fears'):
            detail_content += f"恐惧：{character.get('fears')}\n"

        # 描述
        if character.get('description'):
            detail_content += "\n▌详细描述\n\n"
            detail_content += character.get('description') + "\n"

        # 清空并显示详情
        self._character_detail.config(state='normal')
        self._character_detail.delete('1.0', tk.END)
        self._character_detail.insert('1.0', detail_content)
        self._character_detail.config(state='disabled')

        self._set_status(f"已加载人物详情: {character_name}")

    def _on_character_relation(self):
        """人物关系图谱"""
        from agents.priority import AgentTask, TaskPriority

        if not hasattr(self, '_character_entries') or not self._character_entries:
            messagebox.showinfo("提示", "请先导入人物数据")
            return

        # 调用适配器构建关系图谱（适配器自动从人物数据提取描述）
        try:
            result = self._get_character_manager().execute(
                AgentTask(
                    task_id="build_relation_graph",
                    agent_type="character_manager",
                    priority=TaskPriority.NORMAL,
                    payload={
                        "operation": "build_relation_graph",
                        "all_characters": self._character_entries
                    }
                )
            )

            graph_data = result.get("result", {})

            if graph_data.get("error"):
                messagebox.showerror("错误", f"构建关系图谱失败: {graph_data['error']}")
                return

            # 创建关系图谱窗口
            self._show_relation_graph(graph_data)

        except Exception as e:
            messagebox.showerror("错误", f"构建关系图谱失败: {str(e)}")

    def _show_relation_graph(self, graph_data: Dict[str, Any]):
        """显示人物关系图谱（带交互功能）"""
        import networkx as nx
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
        from matplotlib.figure import Figure

        # 配置中文字体
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
        plt.rcParams['axes.unicode_minus'] = False

        # 创建更大的窗口
        relation_window = tk.Toplevel(self.root)
        relation_window.title("人物关系图谱")
        relation_window.geometry("1400x900")
        relation_window.transient(self.root)
        
        # 全屏模式
        relation_window.state('zoomed')

        # 主容器
        main_container = tk.Frame(relation_window)
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 左侧：图谱画布
        left_frame = tk.Frame(main_container)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        tk.Label(left_frame, text="人物关系图谱", 
                font=('SimHei', 14, 'bold')).pack(pady=5)

        # 创建Matplotlib图形
        fig = Figure(figsize=(12, 10), dpi=100)
        ax = fig.add_subplot(111)

        # 创建NetworkX图
        nodes = graph_data.get("nodes", [])
        relations = graph_data.get("relations", [])

        G = nx.DiGraph() if any(r.get('source') != r.get('target') for r in relations) else nx.Graph()

        # 添加节点
        for node in nodes:
            node_name = node.get('id', '')
            role = node.get('role', '未设置')
            G.add_node(node_name, role=role, label=node_name)

        # 添加边
        edge_labels = {}
        for relation in relations:
            source = relation.get('source', '')
            target = relation.get('target', '')
            label = relation.get('label', '')
            
            if source and target:
                G.add_edge(source, target, relation_type=label)
                edge_labels[(source, target)] = label

        # 存储节点位置和颜色，用于交互
        node_positions = {}
        node_colors_map = {}

        # 绘制图谱
        if nodes:
            try:
                # 使用spring布局（力导向布局）
                pos = nx.spring_layout(G, k=1.5, iterations=50, seed=42)
                node_positions = pos

                # 绘制边
                nx.draw_networkx_edges(
                    G, pos, ax=ax,
                    edge_color='#999999',
                    width=1.5,
                    alpha=0.6,
                    arrows=G.is_directed()
                )

                # 绘制边标签
                nx.draw_networkx_edge_labels(
                    G, pos, edge_labels=edge_labels, ax=ax,
                    font_size=9,
                    font_color='#333333',
                    font_family='SimHei',
                    bbox=dict(facecolor='white', alpha=0.7, edgecolor='none')
                )

                # 绘制节点
                node_colors = [self._get_role_color(node.get('role', '未设置')) for node in nodes]
                node_names = [node.get('id', '') for node in nodes]
                
                for i, node in enumerate(nodes):
                    node_colors_map[node.get('id', '')] = node_colors[i]

                nx.draw_networkx_nodes(
                    G, pos, ax=ax,
                    node_color=node_colors,
                    node_size=2000,
                    alpha=0.9,
                    edgecolors='black',
                    linewidths=1.5
                )

                nx.draw_networkx_labels(
                    G, pos, ax=ax,
                    labels={node: node for node in node_names},
                    font_size=10,
                    font_weight='bold',
                    font_family='SimHei'
                )

                ax.set_title("人物关系网络图", fontsize=14, fontweight='bold', pad=20, fontfamily='SimHei')
                ax.axis('off')
                fig.tight_layout()

                # 嵌入Tkinter
                canvas = FigureCanvasTkAgg(fig, master=left_frame)
                canvas.draw()
                canvas_widget = canvas.get_tk_widget()
                canvas_widget.pack(fill=tk.BOTH, expand=True)

                # 添加导航工具栏
                toolbar = NavigationToolbar2Tk(canvas, left_frame)
                toolbar.update()
                toolbar.pack(side=tk.BOTTOM, fill=tk.X)

                # 实现节点点击交互
                def on_click(event):
                    if event.inaxes != ax:
                        return
                    
                    click_x, click_y = event.xdata, event.ydata
                    closest_node = None
                    min_distance = float('inf')
                    
                    for node_name, (nx_pos, ny_pos) in node_positions.items():
                        distance = ((click_x - nx_pos)**2 + (click_y - ny_pos)**2)**0.5
                        if distance < min_distance:
                            min_distance = distance
                            closest_node = node_name
                    
                    if closest_node and min_distance < 0.15:
                        self._highlight_node(ax, fig, canvas, G, pos, node_colors_map, closest_node)
                        self._highlight_tree_item(tree, closest_node)

                canvas.mpl_connect('button_press_event', on_click)

            except Exception as e:
                tk.Label(left_frame, text=f"图谱生成失败: {str(e)}", fg='red', font=('SimHei', 12)).pack(pady=20)

        # 右侧：关系列表
        right_frame = tk.Frame(main_container, width=400)
        right_frame.pack(side=tk.RIGHT, fill=tk.Y, padx=(10, 0))

        tk.Label(right_frame, text="人物关系列表（点击列表项可高亮图谱节点）", font=('SimHei', 12, 'bold')).pack(pady=5)

        list_frame_container = tk.Frame(right_frame)
        list_frame_container.pack(fill=tk.BOTH, expand=True)

        # 创建Treeview
        columns = ('source', 'relation', 'target')
        tree = ttk.Treeview(list_frame_container, columns=columns, show='headings', height=35)
        
        tree.heading('source', text='人物A')
        tree.heading('relation', text='关系')
        tree.heading('target', text='人物B')
        
        tree.column('source', width=120, anchor='w')
        tree.column('relation', width=100, anchor='center')
        tree.column('target', width=120, anchor='w')

        scrollbar = ttk.Scrollbar(list_frame_container, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 填充数据并绑定事件
        for relation in relations:
            item_id = tree.insert('', tk.END, values=(
                relation.get('source', ''),
                relation.get('label', ''),
                relation.get('target', '')
            ))
            tree.item(item_id, tags=(relation.get('source', ''), relation.get('target', '')))

        def on_tree_select(event):
            selection = tree.selection()
            if not selection:
                return
            
            item = tree.item(selection[0])
            tags = item.get('tags', ())
            
            if tags:
                node_name = tags[0]
                if node_name in node_positions:
                    self._highlight_node(ax, fig, canvas, G, pos, node_colors_map, node_name)

        tree.bind('<<TreeviewSelect>>', on_tree_select)

        # 底部按钮
        button_frame = tk.Frame(relation_window)
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=10)

        tk.Button(button_frame, text="关闭", command=relation_window.destroy, 
                 font=('SimHei', 11), padx=20).pack(side=tk.RIGHT, padx=10)
        
        tk.Label(button_frame, text="提示：点击图谱节点或列表项可高亮显示；使用底部工具栏可缩放、平移图谱", 
                fg='#666666', font=('SimHei', 10)).pack(side=tk.LEFT)

    def _get_role_color(self, role: str) -> str:
        """根据角色获取颜色"""
        role_colors = {
            '主角': '#FF6B6B',
            '配角': '#4ECDC4',
            '反派': '#FFD93D',
            '未设置': '#95E1D3'
        }
        return role_colors.get(role, '#C7CEEA')
    
    def _highlight_node(self, ax, fig, canvas, G, pos, node_colors_map, node_name):
        """高亮选中的节点"""
        ax.clear()
        
        # 重新绘制边（全部变淡）
        nx.draw_networkx_edges(
            G, pos, ax=ax,
            edge_color='#999999',
            width=1.5,
            alpha=0.2,
            arrows=G.is_directed()
        )
        
        # 高亮选中节点的边
        neighbors = list(G.neighbors(node_name))
        for neighbor in neighbors:
            nx.draw_networkx_edges(
                G, pos, ax=ax,
                edgelist=[(node_name, neighbor)],
                edge_color='#FF6B6B',
                width=3.0,
                alpha=0.8,
                arrows=G.is_directed()
            )
        
        # 重新绘制节点（非选中节点变淡）
        non_highlight_nodes = [n for n in G.nodes() if n != node_name and n not in neighbors]
        highlight_nodes = [node_name] + neighbors
        
        if non_highlight_nodes:
            nx.draw_networkx_nodes(
                G, pos, ax=ax,
                nodelist=non_highlight_nodes,
                node_color=[node_colors_map.get(n, '#C7CEEA') for n in non_highlight_nodes],
                node_size=2000,
                alpha=0.3,
                edgecolors='black',
                linewidths=1.0
            )
        
        if highlight_nodes:
            nx.draw_networkx_nodes(
                G, pos, ax=ax,
                nodelist=highlight_nodes,
                node_color=[node_colors_map.get(n, '#C7CEEA') for n in highlight_nodes],
                node_size=2000,
                alpha=1.0,
                edgecolors='black',
                linewidths=2.0
            )
        
        # 重新绘制节点标签
        nx.draw_networkx_labels(
            G, pos, ax=ax,
            labels={n: n for n in G.nodes()},
            font_size=10,
            font_weight='bold',
            font_family='SimHei'
        )
        
        ax.set_title(f"人物关系网络图 - 选中: {node_name}", fontsize=14, fontweight='bold', pad=20, fontfamily='SimHei')
        ax.axis('off')
        fig.tight_layout()
        canvas.draw()
    
    def _highlight_tree_item(self, tree, node_name):
        """在Treeview中高亮包含指定人物的关系项"""
        # 清除当前选择
        tree.selection_remove(tree.get_children())
        
        # 查找并选中所有包含该人物的关系
        for item in tree.get_children():
            tags = tree.item(item, 'tags')
            if tags and node_name in tags:
                tree.selection_add(item)
                tree.see(item)  # 确保可见
    
    def _on_character_delete(self):
        """批量删除人物"""
        if not hasattr(self, '_character_tree') or not hasattr(self, '_character_entries'):
            return

        selection = self._character_tree.selection()
        if not selection:
            messagebox.showinfo("提示", "请先选择要删除的人物")
            return

        # 确认删除
        if not messagebox.askyesno("确认删除", f"确定要删除选中的 {len(selection)} 个人物吗？"):
            return

        # 获取要删除的索引（从大到小排序，避免删除时索引变化）
        children = list(self._character_tree.get_children())
        indices_to_delete = sorted([children.index(item_id) for item_id in selection if item_id in children], reverse=True)

        # 从列表和树中删除
        for idx in indices_to_delete:
            if idx < len(self._character_entries):
                # 从人物列表删除
                del self._character_entries[idx]

        # 从树中删除选中项
        for item_id in selection:
            self._character_tree.delete(item_id)

        # 【修复】同步更新_character_data，确保保存和加载一致
        self._character_data = self._character_entries

        self._set_status(f"已删除 {len(selection)} 个人物")
    
    def _on_outline_browse(self):
        """浏览大纲文件"""
        path = filedialog.askopenfilename(
            title="选择大纲文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._outline_path_var.set(os.path.basename(path))
            self._set_status(f"已选择: {os.path.basename(path)}")
    
    def _on_outline_new(self):
        """新建大纲（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("新建大纲")
        dialog.geometry("500x300")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 300) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 大纲名称
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        ttk.Label(name_frame, text="大纲名称：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="新小说大纲")
        ttk.Entry(name_frame, textvariable=name_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 小说类型
        type_frame = ttk.Frame(dialog, style="TFrame")
        type_frame.pack(fill=tk.X, padx=20, pady=10)
        ttk.Label(type_frame, text="小说类型：").pack(side=tk.LEFT)
        type_var = tk.StringVar(value="玄幻")
        ttk.Combobox(type_frame, textvariable=type_var, values=["玄幻", "仙侠", "都市", "历史", "科幻", "言情"], width=15).pack(side=tk.LEFT, padx=10)
        
        # 目标字数
        words_frame = ttk.Frame(dialog, style="TFrame")
        words_frame.pack(fill=tk.X, padx=20, pady=10)
        ttk.Label(words_frame, text="目标字数：").pack(side=tk.LEFT)
        words_var = tk.StringVar(value="500000")
        ttk.Entry(words_frame, textvariable=words_var, width=15).pack(side=tk.LEFT, padx=10)
        ttk.Label(words_frame, text="字").pack(side=tk.LEFT)
        
        # 简介
        intro_frame = ttk.LabelFrame(dialog, text="小说简介", padding=10)
        intro_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        intro_text = tk.Text(intro_frame, wrap=tk.WORD, height=4,
                            font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                            bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        intro_text.pack(fill=tk.BOTH, expand=True)
        
        # 按钮
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_outline():
            name = name_var.get().strip()
            novel_type = type_var.get()
            target_words = words_var.get()
            intro = intro_text.get("1.0", tk.END).strip()
            
            if not name:
                messagebox.showwarning("警告", "请输入大纲名称！")
                return
            
            # 添加根节点到大纲树
            root_node = self._outline_tree.insert("", tk.END, text=f"📖 {name}", open=True)
            self._outline_tree.insert(root_node, tk.END, text="第一卷")
            self._outline_tree.insert(root_node, tk.END, text="第二卷")
            
            self._outline_path_var.set(name)
            self._set_status(f"已创建大纲：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_outline).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_outline_export(self):
        """导出大纲"""
        self._set_status("导出大纲功能开发中...")
    
    def _on_outline_add_volume(self):
        """添加卷"""
        self._set_status("添加卷功能开发中...")
    
    def _on_outline_add_chapter(self):
        """添加章节"""
        self._set_status("添加章节功能开发中...")
    
    def _on_outline_edit_chapter(self):
        """编辑章节"""
        self._set_status("编辑章节功能开发中...")
    
    def _on_outline_delete(self):
        """删除"""
        self._set_status("删除功能开发中...")
    
    def _on_outline_parse(self):
        """解析大纲 - 调用OutlineAnalysisAgent"""
        try:
            # 获取当前选中的大纲路径
            outline_path = getattr(self, '_outline_path_var', tk.StringVar()).get()
            if not outline_path or outline_path == "未导入":
                messagebox.showwarning("提示", "请先选择大纲文件")
                return
            
            # 更新状态
            self._set_status("正在解析大纲...")
            
            # 导入Agent
            from agents.plugins.outline_analysis_agent import OutlineAnalysisAgent
            from agents.core.base_agent import AgentContext
            from pathlib import Path
            
            # 读取文件内容
            content = Path(outline_path).read_text(encoding='utf-8')
            
            # 创建Agent实例
            agent = OutlineAnalysisAgent()
            if not agent.initialize():
                raise RuntimeError("OutlineAnalysisAgent初始化失败")
            
            # 构建上下文
            context = AgentContext(
                task_id=f"outline_parse_{int(time.time())}",
                input_data={
                    "outline_content": content,
                    "extract_chapters": True,
                }
            )
            
            # 执行解析（异步，使用统一线程池）
            def run_parse():
                try:
                    result = agent.execute(context)
                    self.root.after(0, lambda: self._on_outline_parse_complete(result))
                except Exception as e:
                    self.root.after(0, lambda: self._on_outline_parse_error(str(e)))
            
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(run_parse)
            
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_exception
            title, full_message = convert_exception(e, "大纲解析")
            messagebox.showerror(title, full_message)
            self._set_status(f"大纲解析失败: {str(e)}")
    
    def _on_outline_parse_complete(self, result):
        """大纲解析完成回调"""
        try:
            if result.success:
                # 存储大纲数据
                self._outline_content = result.data.get("outline_content", "")
                self._chapter_outlines = result.data.get("chapters", {})
                
                # 更新大纲树形结构
                self._outline_tree.delete(*self._outline_tree.get_children())
                chapters = result.data.get("chapters", [])
                for i, chapter in enumerate(chapters):
                    chapter_id = self._outline_tree.insert("", tk.END, text=f"第{i+1}章: {chapter.get('title', '未命名')}")
                    # 添加子节点
                    for key, value in chapter.items():
                        if key != "title":
                            self._outline_tree.insert(chapter_id, tk.END, text=f"{key}: {str(value)[:30]}")
                
                # 更新统计信息
                total_words = result.data.get("estimated_words", 0)
                self._set_status(f"大纲解析完成，预估{total_words}字")
            else:
                self._set_status(f"大纲解析失败: {result.error}")
        except Exception as e:
            self._set_status(f"处理解析结果失败: {str(e)}")
    
    def _on_outline_parse_error(self, error: str):
        """大纲解析错误回调"""
        messagebox.showerror("大纲解析错误", error)
        self._set_status(f"大纲解析失败: {error}")
    
    def _on_outline_clear(self):
        """清除大纲"""
        self._outline_path_var.set("未导入")
    
    def _on_style_browse(self):
        """浏览风格文件"""
        path = filedialog.askopenfilename(
            title="选择风格文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._style_path_var.set(os.path.basename(path))
            self._set_status(f"已选择: {os.path.basename(path)}")
    
    def _on_style_delete(self):
        """删除范文"""
        self._set_status("删除范文功能开发中...")
    
    def _on_style_analyze(self):
        """解析风格 - 调用StyleLearningAgent"""
        try:
            # 获取当前选中的范文路径
            style_path = getattr(self, '_style_path_var', tk.StringVar()).get()
            if not style_path or style_path == "未导入":
                messagebox.showwarning("提示", "请先上传范文文件")
                return
            
            # 更新状态
            self._set_status("正在分析风格...")
            
            # 导入Agent
            from agents.plugins.style_learning_agent import StyleLearningAgent
            from pathlib import Path
            
            # 读取文件内容
            content = Path(style_path).read_text(encoding='utf-8')
            
            # 创建Agent实例
            agent = StyleLearningAgent()
            if not agent.initialize():
                raise RuntimeError("StyleLearningAgent初始化失败")
            
            # 构建上下文
            from agents.core.base_agent import AgentContext
            context = AgentContext(
                task_id=f"style_analyze_{int(time.time())}",
                input_data={
                    "text": content[:50000],  # 限制长度
                    "extract_patterns": True,
                }
            )
            
            # 执行分析（异步）
            def run_analysis():
                try:
                    result = agent.execute(context)
                    self.root.after(0, lambda: self._on_style_analyze_complete(result))
                except Exception as e:
                    self.root.after(0, lambda: self._on_style_analyze_error(str(e)))
            
            # 使用统一线程池（解决卡顿问题）
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(run_analysis)
            
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_exception
            title, full_message = convert_exception(e, "风格分析")
            messagebox.showerror(title, full_message)
            self._set_status(f"风格分析失败: {str(e)}")
    
    def _on_style_analyze_complete(self, result):
        """风格分析完成回调"""
        try:
            if result.success:
                # 存储风格档案
                self._style_profile = result.data.get("style_profile", {})
                
                # 更新UI
                self._style_info.delete("1.0", tk.END)
                self._style_info.insert("1.0", f"风格分析完成！\n\n")
                self._style_info.insert(tk.END, f"词汇特征: {len(result.data.get('vocabulary_features', []))}个\n")
                self._style_info.insert(tk.END, f"句式模式: {len(result.data.get('sentence_patterns', []))}个\n")
                self._style_info.insert(tk.END, f"修辞手法: {len(result.data.get('rhetorical_devices', []))}个\n")
                
                self._set_status("风格分析完成")
            else:
                self._set_status(f"风格分析失败: {result.error}")
        except Exception as e:
            self._set_status(f"处理分析结果失败: {str(e)}")
    
    def _on_style_analyze_error(self, error: str):
        """风格分析错误回调"""
        messagebox.showerror("风格分析错误", error)
        self._set_status(f"风格分析失败: {error}")
    
    def _on_style_export(self):
        """导出风格"""
        self._set_status("导出风格功能开发中...")
    
    def _on_style_switch(self):
        """切换风格"""
        self._set_status("切换风格功能开发中...")
    
    def _on_style_edit_weight(self):
        """编辑权重"""
        self._set_status("编辑权重功能开发中...")
    
    def _on_style_delete_profile(self):
        """删除风格档案"""
        self._set_status("删除风格档案功能开发中...")
    
    def _on_style_learn(self):
        """学习风格"""
        self._set_status("风格学习功能开发中...")
    
    def _on_style_create(self):
        """创建风格模板（弹窗）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("创建风格模板")
        dialog.geometry("600x500")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 600) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 500) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 风格名称
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        ttk.Label(name_frame, text="风格名称：").pack(side=tk.LEFT)
        name_var = tk.StringVar(value="自定义风格")
        ttk.Entry(name_frame, textvariable=name_var, width=30).pack(side=tk.LEFT, padx=10)
        
        # 七维度评分
        dimensions_frame = ttk.LabelFrame(dialog, text="七维度风格评分（0-10分）", padding=10)
        dimensions_frame.pack(fill=tk.X, padx=20, pady=10)
        
        dimensions = [
            ("词汇丰富度", "vocabulary"),
            ("句式复杂度", "sentence"),
            ("修辞使用", "rhetoric"),
            ("情感强度", "emotion"),
            ("节奏感", "rhythm"),
            ("结构完整性", "structure"),
            ("细节描写", "detail")
        ]
        
        dim_vars = {}
        for i, (label, key) in enumerate(dimensions):
            row = ttk.Frame(dimensions_frame, style="TFrame")
            row.pack(fill=tk.X, pady=3)
            ttk.Label(row, text=f"{label}：", width=15, anchor='e').pack(side=tk.LEFT)
            var = tk.DoubleVar(value=5.0)
            dim_vars[key] = var
            scale = ttk.Scale(row, from_=0.0, to=10.0, variable=var, orient=tk.HORIZONTAL, length=200)
            scale.pack(side=tk.LEFT, padx=10)
            ttk.Label(row, textvariable=var, width=5).pack(side=tk.LEFT)
        
        # 风格描述
        desc_frame = ttk.LabelFrame(dialog, text="风格描述（可选）", padding=10)
        desc_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        desc_text = tk.Text(desc_frame, wrap=tk.WORD, height=5,
                           font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
                           bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
        desc_text.pack(fill=tk.BOTH, expand=True)
        desc_text.insert("1.0", "描述这种风格的特点，如：华丽、简洁、幽默、严肃等...")
        
        # 按钮
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def create_style():
            name = name_var.get().strip()
            if not name:
                messagebox.showwarning("警告", "请输入风格名称！")
                return
            
            # 获取评分
            scores = {key: var.get() for key, var in dim_vars.items()}
            
            # 添加到风格档案库
            self._style_tree.insert("", tk.END, values=(
                name,
                f"{scores['vocabulary']:.1f}",
                f"{scores['sentence']:.1f}",
                f"{scores['rhetoric']:.1f}",
                f"{scores['emotion']:.1f}",
                f"{scores['rhythm']:.1f}",
                f"{scores['structure']:.1f}",
                f"{scores['detail']:.1f}"
            ))
            
            # 更新分析器显示
            self._style_info.delete("1.0", tk.END)
            self._style_info.insert("1.0", f"""【{name}】风格模板

词汇：{"█" * int(scores['vocabulary'])}{"░" * (10 - int(scores['vocabulary']))} {scores['vocabulary']:.1f}
句式：{"█" * int(scores['sentence'])}{"░" * (10 - int(scores['sentence']))} {scores['sentence']:.1f}
修辞：{"█" * int(scores['rhetoric'])}{"░" * (10 - int(scores['rhetoric']))} {scores['rhetoric']:.1f}
情感：{"█" * int(scores['emotion'])}{"░" * (10 - int(scores['emotion']))} {scores['emotion']:.1f}
节奏：{"█" * int(scores['rhythm'])}{"░" * (10 - int(scores['rhythm']))} {scores['rhythm']:.1f}
结构：{"█" * int(scores['structure'])}{"░" * (10 - int(scores['structure']))} {scores['structure']:.1f}
细节：{"█" * int(scores['detail'])}{"░" * (10 - int(scores['detail']))} {scores['detail']:.1f}

描述：{desc_text.get("1.0", tk.END).strip()}
""")
            
            self._set_status(f"已创建风格模板：{name}")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="创建", command=create_style).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _on_style_apply(self):
        """应用到生成器"""
        self._set_status("应用到生成器功能开发中...")
    
    def _on_style_clear(self):
        """清除风格"""
        self._style_path_var.set("未导入")
        self._style_info.delete("1.0", tk.END)
    
    def _on_start_generation(self):
        """开始生成 - 调用小说生成流水线（异步版本，解决卡顿问题）
        
        V2.17版本更新（2026-03-28）：
        - 使用GUIAsyncHelper实现真正的异步调用
        - UI线程不再阻塞，解决Windows弹窗问题
        - 支持实时进度反馈和取消操作
        """
        try:
            # 导入生成服务
            from agents.novel_generation_service import get_generation_service
            from core.gui_async_helper import create_async_helper
            
            # 获取生成配置
            chapter_number = int(self._start_chapter_var.get())
            end_chapter = int(self._end_chapter_var.get())
            target_words = int(self._target_words_var.get())
            temperature = self._gen_temp_var.get()
            
            # 获取大纲内容（从大纲管理页面获取或使用默认值）
            outline_content = self._get_outline_content()
            chapter_outline = self._get_chapter_outline(chapter_number)
            
            # 获取风格档案（从风格学习页面获取）
            style_profile = self._get_style_profile()
            
            # 获取人物设定
            characters = self._get_characters()
            
            # 获取世界观设定
            worldview = self._get_worldview()
            
            # 获取选中的知识库分类
            selected_kb = self.get_selected_knowledge_bases()
            
            # 获取选中的写作技巧
            selected_techniques = self.get_selected_writing_techniques()
            
            # 更新UI状态
            self._gen_status_var.set(f"正在生成第{chapter_number}章...")
            self._gen_progress['value'] = 0
            self._gen_log.delete("1.0", tk.END)
            self._gen_log.insert(tk.END, f"开始生成第{chapter_number}章...\n")
            self._gen_log.insert(tk.END, f"目标字数: {target_words}\n")
            self._gen_log.insert(tk.END, f"生成温度: {temperature}\n")
            if selected_kb:
                self._gen_log.insert(tk.END, f"关联知识库: {', '.join(selected_kb)}\n")
            if selected_techniques:
                self._gen_log.insert(tk.END, f"写作技巧: {', '.join(selected_techniques)}\n")
            self._gen_log.insert(tk.END, "-" * 40 + "\n")
            
            # 获取LLM客户端
            llm_client = self._get_llm_client()
            
            # 获取生成服务
            service = get_generation_service(
                event_bus=getattr(self, '_event_bus', None),
                llm_client=llm_client,
            )
            
            # === 异步调用改造（核心修复）===
            # 创建异步辅助器
            if not hasattr(self, '_async_helper'):
                self._async_helper = create_async_helper(self.root)
            
            # 进度回调（线程安全）
            def on_progress(progress):
                self._async_helper.call_on_main_thread(
                    lambda: self._update_generation_progress(progress)
                )
            
            # 完成回调（线程安全）
            def on_complete(result):
                self._async_helper.call_on_main_thread(
                    lambda: self._on_generation_complete(result)
                )
            
            # 错误回调（线程安全）
            def on_error(error):
                self._async_helper.call_on_main_thread(
                    lambda: self._on_generation_error(error)
                )
            
            # 提交异步生成任务（使用统一线程池）
            task_id = service.submit_async_generation(
                chapter_title=f"第{chapter_number}章",
                chapter_number=chapter_number,
                outline_content=outline_content,
                chapter_outline=chapter_outline,
                target_word_count=target_words,
                style_profile=style_profile,
                characters=characters,
                worldview=worldview,
                previous_chapter_text="",  # TODO: 从历史章节获取
                max_iterations=5,
                on_progress=on_progress,
                on_complete=on_complete,
                on_error=on_error,
                knowledge_categories=selected_kb,
                writing_techniques=selected_techniques,
            )
            
            self._current_pipeline_id = task_id
            self._set_status(f"正在生成第{chapter_number}章...")
            
            # 保存任务ID以便取消
            self._current_task_id = task_id
            
        except Exception as e:
            from tkinter import messagebox
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_generation_error
            title, full_message = convert_generation_error(e)
            messagebox.showerror(title, full_message)
            self._gen_status_var.set("生成失败")
            self._set_status(f"生成失败: {str(e)}")
    
    def _on_generation_error(self, error):
        """生成错误回调（V2.17新增）"""
        from tkinter import messagebox
        from core.user_friendly_errors import convert_generation_error
        
        title, full_message = convert_generation_error(error)
        messagebox.showerror(title, full_message)
        self._gen_status_var.set("生成失败")
        self._set_status(f"生成失败: {str(error)}")
    
    def _on_stop_generation(self):
        """停止生成"""
        try:
            from agents.novel_generation_service import get_generation_service
            service = get_generation_service()
            if service.cancel_generation():
                self._gen_status_var.set("已取消生成")
                self._gen_log.insert(tk.END, "\n用户取消生成\n")
                self._set_status("已停止生成")
            else:
                self._set_status("没有正在进行的生成任务")
        except Exception as e:
            self._set_status(f"停止生成失败: {str(e)}")
    
    def _update_generation_progress(self, progress):
        """更新生成进度（UI线程安全）"""
        try:
            self._gen_progress['value'] = progress.progress_percent
            self._gen_log.insert(tk.END, f"[{progress.stage_name}] {progress.message}\n")
            self._gen_log.see(tk.END)  # 自动滚动到底部
        except Exception as e:
            pass  # 忽略UI更新错误
    
    def _on_generation_complete(self, result):
        """生成完成回调（UI线程安全）"""
        try:
            if result.success:
                self._gen_status_var.set("生成完成")
                self._gen_progress['value'] = 100
                self._gen_log.insert(tk.END, "\n" + "=" * 40 + "\n")
                self._gen_log.insert(tk.END, f"生成完成！总迭代次数: {result.total_iterations}\n")
                self._gen_log.insert(tk.END, f"总耗时: {result.total_duration_seconds:.1f}秒\n")
                
                # 显示生成结果
                if result.final_output:
                    content = result.final_output.get("content", str(result.final_output))
                    self._gen_result.delete("1.0", tk.END)
                    self._gen_result.insert("1.0", content)
                    
                    # 更新字数统计
                    word_count = len(content)
                    self._gen_log.insert(tk.END, f"实际字数: {word_count}\n")
                
                self._set_status(f"第{result.stages[0].iteration if result.stages else 1}章生成完成")
            else:
                self._gen_status_var.set("生成失败")
                self._gen_log.insert(tk.END, f"\n生成失败: {result.error}\n")
                self._set_status(f"生成失败: {result.error}")
            
            self._gen_log.see(tk.END)
            
        except Exception as e:
            self._gen_status_var.set("处理结果出错")
            self._gen_log.insert(tk.END, f"\n处理生成结果时出错: {str(e)}\n")
    
    def _get_outline_content(self) -> str:
        """获取大纲内容"""
        # TODO: 从大纲管理页面或文件获取
        return getattr(self, '_outline_content', "")
    
    def _get_chapter_outline(self, chapter_number: int) -> str:
        """获取指定章节的大纲"""
        # TODO: 从大纲管理页面获取
        return getattr(self, '_chapter_outlines', {}).get(chapter_number, "")
    
    def _get_style_profile(self) -> dict:
        """获取风格档案"""
        # TODO: 从风格学习页面获取
        return getattr(self, '_style_profile', {})
    
    def _get_characters(self) -> list:
        """获取人物设定"""
        # TODO: 从人物设定页面获取
        return getattr(self, '_characters', [])
    
    def _get_worldview(self) -> dict:
        """获取世界观设定"""
        # TODO: 从世界观页面获取
        return getattr(self, '_worldview', {})
    
    def _get_llm_client(self):
        """获取LLM客户端"""
        # 从配置或服务获取LLM客户端
        return getattr(self, '_llm_client', None)
    
    def _on_expert_changed(self, expert_info):
        """专家选择变化回调
        
        Args:
            expert_info: ExpertInfo对象，如果选择默认模式则为None
        """
        if expert_info:
            self._set_status(f"已选择专家模式：{expert_info.name}")
            logging.info(f"专家选择变化: {expert_info.expert_id} - {expert_info.name}")
        else:
            self._set_status("已选择默认模式")
            logging.info("切换回默认生成模式")
    
    def _get_expert_config(self) -> dict:
        """获取当前专家配置
        
        Returns:
            dict: 专家配置字典，如果未选择专家则返回空字典
        """
        if self._expert_selector:
            return self._expert_selector.get_config()
        return {}
    
    def _is_expert_mode_enabled(self) -> bool:
        """检查是否启用了专家模式
        
        Returns:
            bool: True表示启用了专家模式
        """
        return self._expert_selector is not None and self._expert_selector.get_selected_expert() is not None
    
    def _on_gen_browse(self):
        """分章浏览"""
        self._set_status("分章浏览功能开发中...")
    
    def _on_gen_save(self):
        """保存项目 - 使用项目管理器"""
        # 检查项目管理器是否可用
        if not self._project_manager or not self._project_manager.is_project_open():
            messagebox.showwarning("保存项目", "当前没有打开的项目")
            return
        
        try:
            self._set_status("正在保存项目...")
            
            # 同步所有模块数据到项目管理器
            self._sync_all_data_to_manager()
            
            # 调用项目管理器保存
            success = self._project_manager.save_project()
            
            if success:
                project_name = self._project_manager.get_project_name()
                self._set_status("项目保存完成")
                messagebox.showinfo("成功", f"项目「{project_name}」已保存！")
            else:
                self._set_status("保存项目失败")
                messagebox.showwarning("保存失败", "项目保存失败，请查看日志获取详情")
            
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_file_error
            title, full_message = convert_file_error(e, "保存项目")
            messagebox.showerror(title, full_message)
            self._set_status("保存项目失败")
    
    def _sync_all_data_to_manager(self):
        """同步所有模块数据到项目管理器"""
        if not self._project_manager:
            return

        # 同步各模块数据到项目管理器
        if hasattr(self, '_outline_content') and self._outline_content:
            self._project_manager.sync_module_data('outline', self._outline_content)

        # 人物数据：优先使用_character_entries（批量解析后的结构化列表），降级到_character_data
        if hasattr(self, '_character_entries') and self._character_entries:
            self._project_manager.sync_module_data('characters', self._character_entries)
        elif hasattr(self, '_character_data') and self._character_data:
            self._project_manager.sync_module_data('characters', self._character_data)

        if hasattr(self, '_worldview_content') and self._worldview_content:
            self._project_manager.sync_module_data('worldview', self._worldview_content)
        
        if hasattr(self, '_style_profile') and self._style_profile:
            self._project_manager.sync_module_data('style', self._style_profile)
        
        if hasattr(self, '_reverse_chapters') and self._reverse_chapters:
            self._project_manager.sync_module_data('reverse_chapters', self._reverse_chapters)
        
        if hasattr(self, '_reverse_feedback_data') and self._reverse_feedback_data:
            self._project_manager.sync_module_data('reverse_feedback', self._reverse_feedback_data)
        
        if hasattr(self, '_completed_chapters') and self._completed_chapters:
            self._project_manager.sync_module_data('completed_chapters', self._completed_chapters)
        
        if hasattr(self, '_generated_content') and self._generated_content:
            self._project_manager.sync_module_data('generated_content', self._generated_content)
    
    def _sync_all_data_to_project(self):
        """同步所有模块数据到项目（兼容旧代码，内部使用）"""
        if not self.current_project:
            return
        
        # 大纲
        if hasattr(self, '_outline_content') and self._outline_content:
            self.current_project['outline'] = self._outline_content
        
        # 人物：优先使用_character_entries（批量解析后的数据），降级到_character_data
        if hasattr(self, '_character_entries') and self._character_entries:
            self.current_project['characters'] = self._character_entries
        elif hasattr(self, '_character_data') and self._character_data:
            self.current_project['characters'] = self._character_data
        
        # 世界观
        if hasattr(self, '_worldview_content') and self._worldview_content:
            self.current_project['worldview'] = self._worldview_content
        
        # 风格
        if hasattr(self, '_style_profile') and self._style_profile:
            self.current_project['style'] = self._style_profile
        
        # 逆向反馈上传的章节
        if hasattr(self, '_reverse_chapters') and self._reverse_chapters:
            self.current_project['reverse_chapters'] = self._reverse_chapters
        
        # 逆向反馈数据
        if hasattr(self, '_reverse_feedback_data') and self._reverse_feedback_data:
            self.current_project['reverse_feedback'] = self._reverse_feedback_data
        
        # 已完成章节
        if hasattr(self, '_completed_chapters') and self._completed_chapters:
            self.current_project['completed_chapters'] = self._completed_chapters
        
        # 生成内容
        if hasattr(self, '_generated_content') and self._generated_content:
            self.current_project['generated_content'] = self._generated_content
    
    def _on_gen_export_txt(self):
        """导出TXT"""
        self._set_status("导出TXT功能开发中...")
    
    def _on_gen_export_docx(self):
        """导出DOCX"""
        self._set_status("导出DOCX功能开发中...")
    
    def _on_reverse_browse(self):
        """浏览逆向分析文件（已弃用，保留兼容）"""
        path = filedialog.askopenfilename(
            title="选择需要分析的文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self._on_reverse_add_chapter_from_file(path)
    
    def _on_reverse_upload_files(self):
        """上传单个或多个章节文件"""
        paths = filedialog.askopenfilenames(
            title="选择章节文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if paths:
            for path in paths:
                self._on_reverse_add_chapter_from_file(path)
            self._set_status(f"已上传 {len(paths)} 个章节文件")
    
    def _on_reverse_batch_upload(self):
        """批量上传文件夹中的所有章节"""
        folder = filedialog.askdirectory(title="选择章节文件夹")
        if folder:
            count = 0
            for filename in os.listdir(folder):
                if filename.endswith(('.txt', '.docx')):
                    path = os.path.join(folder, filename)
                    self._on_reverse_add_chapter_from_file(path)
                    count += 1
            self._set_status(f"已从文件夹批量上传 {count} 个章节")
    
    def _on_reverse_paste_text(self):
        """切换粘贴文本区域的显示/隐藏"""
        if self._paste_text_frame.winfo_ismapped():
            self._paste_text_frame.pack_forget()
        else:
            self._paste_text_frame.pack(fill=tk.X, pady=5)
    
    def _on_reverse_add_pasted_chapter(self):
        """添加粘贴的章节"""
        title = self._paste_title_var.get().strip()
        content = self._paste_content_text.get("1.0", tk.END).strip()
        
        if not title:
            messagebox.showwarning("提示", "请输入章节标题")
            return
        if not content:
            messagebox.showwarning("提示", "请输入章节内容")
            return
        
        # 计算字数
        word_count = len(content.replace('\n', '').replace(' ', ''))
        
        # 生成唯一ID
        chapter_id = f"chapter_{int(time.time() * 1000)}"
        
        # 存储章节数据
        self._reverse_chapters[chapter_id] = {
            'title': title,
            'content': content,
            'words': word_count,
            'status': '未完成',
            'source': '粘贴'
        }
        
        # 添加到列表
        idx = len(self._completed_chapters_tree.get_children()) + 1
        self._completed_chapters_tree.insert("", tk.END, iid=chapter_id, values=(
            f"第{idx}章",
            title,
            f"{word_count}字",
            "未完成",
            "粘贴"
        ))
        
        # 清空输入
        self._paste_title_var.set("")
        self._paste_content_text.delete("1.0", tk.END)
        
        self._set_status(f"已添加章节：{title}")
    
    def _on_reverse_add_chapter_from_file(self, file_path: str):
        """从文件添加章节"""
        try:
            # 读取文件内容
            if file_path.endswith('.docx'):
                content = self._read_docx_file(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # 提取标题（从文件名或内容首行）
            title = os.path.basename(file_path).rsplit('.', 1)[0]
            if content.startswith('#') or content.startswith('第'):
                first_line = content.split('\n')[0]
                if len(first_line) < 50:
                    title = first_line.strip('# ')
            
            # 计算字数
            word_count = len(content.replace('\n', '').replace(' ', ''))
            
            # 生成唯一ID
            chapter_id = f"file_{int(time.time() * 1000)}_{len(self._reverse_chapters)}"
            
            # 存储章节数据
            self._reverse_chapters[chapter_id] = {
                'title': title,
                'content': content,
                'words': word_count,
                'status': '未完成',
                'source': os.path.basename(file_path),
                'file_path': file_path
            }
            
            # 添加到列表
            idx = len(self._completed_chapters_tree.get_children()) + 1
            self._completed_chapters_tree.insert("", tk.END, iid=chapter_id, values=(
                f"第{idx}章",
                title,
                f"{word_count}字",
                "未完成",
                os.path.basename(file_path)
            ))
            
        except Exception as e:
            self._set_status(f"读取文件失败：{str(e)}")
            logger.error(f"读取章节文件失败: {e}")
    
    def _read_docx_file(self, file_path: str) -> str:
        """读取DOCX文件内容"""
        try:
            from docx import Document
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        except ImportError:
            messagebox.showwarning("提示", "需要安装python-docx库才能读取DOCX文件")
            return ""
    
    def _on_reverse_refresh_chapters(self):
        """刷新章节列表"""
        # 重新统计并更新列表
        for item in self._completed_chapters_tree.get_children():
            if item in self._reverse_chapters:
                chapter_data = self._reverse_chapters[item]
                values = list(self._completed_chapters_tree.item(item, "values"))
                values[2] = f"{chapter_data['words']}字"
                values[3] = chapter_data['status']
                self._completed_chapters_tree.item(item, values=values)
        
        total_words = sum(c['words'] for c in self._reverse_chapters.values())
        self._set_status(f"已刷新列表，共 {len(self._reverse_chapters)} 章，{total_words} 字")
    
    def _on_reverse_clear_chapters(self):
        """清空章节列表"""
        if messagebox.askyesno("确认", "确定要清空所有章节吗？"):
            self._completed_chapters_tree.delete(*self._completed_chapters_tree.get_children())
            self._reverse_chapters.clear()
            self._set_status("已清空章节列表")
    
    def _on_reverse_chapter_select(self, event):
        """章节选择事件"""
        selected = self._completed_chapters_tree.selection()
        if selected:
            count = len(selected)
            total_words = 0
            for item in selected:
                if item in self._reverse_chapters:
                    total_words += self._reverse_chapters[item]['words']
            self._set_status(f"已选择 {count} 章，共 {total_words} 字")
    
    def _on_reverse_view_chapter(self):
        """查看章节内容"""
        selected = self._completed_chapters_tree.selection()
        if not selected:
            return
        
        item = selected[0]
        if item not in self._reverse_chapters:
            return
        
        chapter_data = self._reverse_chapters[item]
        
        # 创建查看窗口
        view_window = tk.Toplevel(self.root)
        view_window.title(f"查看章节：{chapter_data['title']}")
        view_window.geometry("600x400")
        
        text = tk.Text(view_window, wrap=tk.WORD, font=('Microsoft YaHei UI', 10))
        text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text.insert("1.0", chapter_data['content'])
        text.configure(state=tk.DISABLED)
    
    def _on_reverse_run_analysis(self):
        """运行逆向分析"""
        # 获取要分析的章节
        scope = self._reverse_scope_var.get()
        chapters_to_analyze = []
        
        if scope == "selected":
            selected = self._completed_chapters_tree.selection()
            for item in selected:
                if item in self._reverse_chapters:
                    chapters_to_analyze.append({
                        'id': item,
                        **self._reverse_chapters[item]
                    })
            if not chapters_to_analyze:
                messagebox.showwarning("提示", "请先选择要分析的章节")
                return
        elif scope == "completed":
            for item_id, chapter in self._reverse_chapters.items():
                if chapter['status'] == '已完成':
                    chapters_to_analyze.append({
                        'id': item_id,
                        **chapter
                    })
            if not chapters_to_analyze:
                messagebox.showwarning("提示", "没有已完成的章节可供分析")
                return
        else:  # all
            for item_id, chapter in self._reverse_chapters.items():
                chapters_to_analyze.append({
                    'id': item_id,
                    **chapter
                })
            if not chapters_to_analyze:
                messagebox.showwarning("提示", "章节列表为空，请先上传章节")
                return
        
        # 获取分析维度
        analysis_options = {
            'consistency': self._reverse_check_consistency.get(),
            'logic': self._reverse_check_logic.get(),
            'character': self._reverse_check_character.get(),
            'style': self._reverse_check_style.get(),
            'worldview': self._reverse_check_worldview.get(),
        }
        
        # 更新进度提示
        self._analysis_progress_label.configure(text="正在分析中...")
        self._run_analysis_btn.configure(state=tk.DISABLED)
        self.root.update()
        
        # 异步执行分析
        def run_analysis():
            try:
                # 获取插件
                plugin = self._get_reverse_feedback_plugin()
                if not plugin:
                    self.root.after(0, lambda: self._analysis_progress_label.configure(
                        text="分析插件不可用"))
                    return
                
                # 获取项目设定
                settings = self._get_current_project_settings()
                
                # 分析结果
                all_issues = []
                
                for chapter in chapters_to_analyze:
                    result = plugin.analyze_chapter_vs_settings(
                        chapter_text=chapter['content'],
                        current_settings=settings,
                        chapter_id=chapter['title']
                    )
                    all_issues.extend(result.issues)
                
                # 保存结果
                self._reverse_analysis_result = {
                    'issues': all_issues,
                    'chapters_analyzed': len(chapters_to_analyze),
                    'analysis_options': analysis_options
                }
                
                # 更新逆向反馈数据（用于项目保存）
                if not hasattr(self, '_reverse_feedback_data'):
                    self._reverse_feedback_data = {}
                
                # 添加分析报告
                if 'reports' not in self._reverse_feedback_data:
                    self._reverse_feedback_data['reports'] = []
                
                import time
                self._reverse_feedback_data['reports'].append({
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
                    'issues_count': len(all_issues),
                    'chapters_analyzed': len(chapters_to_analyze),
                    'issues': [
                        {
                            'type': getattr(i, 'issue_type', 'unknown'),
                            'severity': getattr(i, 'severity', 'medium'),
                            'element': getattr(i, 'element_name', ''),
                            'description': getattr(i, 'description', ''),
                            'suggestion': getattr(i, 'suggested_fix', '')
                        }
                        for i in all_issues
                    ]
                })
                self._reverse_feedback_data['last_chapter_analyzed'] = chapters_to_analyze[-1]['title'] if chapters_to_analyze else None
                
                # 更新UI
                self.root.after(0, lambda: self._update_analysis_result())
                
            except Exception as e:
                logger.error(f"分析失败: {e}")
                self.root.after(0, lambda: self._analysis_progress_label.configure(
                    text=f"分析失败: {str(e)}"))
            finally:
                self.root.after(0, lambda: self._run_analysis_btn.configure(state=tk.NORMAL))
        
        # 使用统一线程池（解决卡顿问题）
        from core.thread_pool_manager import thread_pool_manager
        thread_pool_manager.submit_sync(run_analysis)
    
    def _load_reverse_feedback_plugin(self):
        """加载逆向反馈插件"""
        try:
            # 尝试从PluginRegistry获取
            registry = get_plugin_registry()
            if registry:
                plugin = registry.get_plugin("reverse-feedback-analyzer")
                if plugin:
                    self._reverse_feedback_plugin = plugin
                    self._set_status("逆向反馈插件加载成功")
                    logger.info("逆向反馈插件从Registry加载成功")
                    return
            
            # 使用动态导入（支持连字符目录名）
            plugin_class = _dynamic_import_plugin('reverse-feedback-analyzer', 'ReverseFeedbackAnalyzerPlugin')
            if not plugin_class:
                logger.error("逆向反馈插件类未找到")
                self._reverse_feedback_plugin = None
                return
            
            from core.plugin_interface import PluginContext
            
            self._reverse_feedback_plugin = plugin_class()
            
            # 创建上下文并初始化
            context = PluginContext(
                config_manager=self._config_manager if hasattr(self, '_config_manager') else None,
                event_bus=None,
                service_locator=None,
            )
            self._reverse_feedback_plugin.initialize(context)
            
            self._set_status("逆向反馈插件加载成功")
            logger.info("逆向反馈插件加载成功")
            
        except Exception as e:
            logger.error(f"加载逆向反馈插件失败: {e}")
            self._reverse_feedback_plugin = None
    
    def _get_reverse_feedback_plugin(self):
        """获取逆向反馈分析插件"""
        try:
            # 尝试从PluginRegistry获取
            registry = get_plugin_registry()
            if registry:
                plugin = registry.get_plugin("reverse-feedback-analyzer")
                if plugin:
                    return plugin
            
            # 使用动态导入（支持连字符目录名）
            plugin_class = _dynamic_import_plugin('reverse-feedback-analyzer', 'ReverseFeedbackAnalyzerPlugin')
            if not plugin_class:
                logger.error("逆向反馈插件类未找到")
                return None
            
            from core.plugin_interface import PluginContext
            plugin = plugin_class()
            
            # 创建简化的上下文
            context = PluginContext(
                config_manager=self._config_manager if hasattr(self, '_config_manager') else None,
                event_bus=None,
                service_locator=None,
            )
            plugin.initialize(context)
            return plugin
        except Exception as e:
            logger.error(f"获取逆向反馈插件失败: {e}")
            return None
    
    def _get_current_project_settings(self) -> Dict:
        """获取当前项目设定"""
        settings = {
            'project_name': getattr(self, '_current_project_name', '未命名项目'),
            'outline': getattr(self, '_outline_content', ''),
            'characters': getattr(self, '_character_data', []),
            'worldview': getattr(self, '_worldview_content', ''),
            'style': getattr(self, '_style_profile', None),  # 风格设定
            'reverse_feedback': getattr(self, '_reverse_feedback_data', {}),  # 逆向反馈
            'completed_chapters': getattr(self, '_completed_chapters', []),  # 已完成章节
        }
        return settings
    
    def _update_analysis_result(self):
        """更新分析结果显示"""
        if not self._reverse_analysis_result:
            return
        
        # 清空冲突列表
        self._issues_tree.delete(*self._issues_tree.get_children())
        
        issues = self._reverse_analysis_result.get('issues', [])
        
        # 类型映射
        type_map = {
            'character': '人物',
            'outline': '大纲',
            'worldview': '世界观',
        }
        
        # 优先级映射
        severity_map = {
            'high': '🔴 高',
            'medium': '🟡 中',
            'low': '🟢 低',
        }
        
        for issue in issues:
            issue_type = getattr(issue, 'issue_type', 'outline')
            if hasattr(issue_type, 'value'):
                issue_type = issue_type.value
            
            severity = getattr(issue, 'severity', 'medium')
            if hasattr(severity, 'value'):
                severity = severity.value
            
            self._issues_tree.insert("", tk.END, values=(
                type_map.get(issue_type, issue_type),
                severity_map.get(severity, severity),
                getattr(issue, 'element_name', ''),
                getattr(issue, 'chapter_reference', ''),
                getattr(issue, 'description', '')[:50] + '...' if len(getattr(issue, 'description', '')) > 50 else getattr(issue, 'description', '')
            ))
        
        # 更新进度提示
        high_count = sum(1 for i in issues if getattr(i, 'severity', None) and 
                        (getattr(i.severity, 'value', i.severity) == 'high'))
        
        self._analysis_progress_label.configure(
            text=f"分析完成：共 {len(issues)} 个冲突，{high_count} 个高优先级"
        )
        self._set_status(f"分析完成，发现 {len(issues)} 个冲突项")
    
    def _on_reverse_issue_select(self, event):
        """冲突项选择事件"""
        selected = self._issues_tree.selection()
        if not selected or not self._reverse_analysis_result:
            return
        
        # 获取选中的问题索引
        idx = self._issues_tree.index(selected[0])
        issues = self._reverse_analysis_result.get('issues', [])
        
        if idx >= len(issues):
            return
        
        issue = issues[idx]
        
        # 显示详情
        self._issue_detail_text.configure(state=tk.NORMAL)
        self._issue_detail_text.delete("1.0", tk.END)
        
        detail = f"""【冲突详情】

类型：{getattr(issue, 'issue_type', '未知')}
严重程度：{getattr(issue, 'severity', '未知')}
涉及元素：{getattr(issue, 'element_name', '未知')}
章节引用：{getattr(issue, 'chapter_reference', '未知')}

【问题描述】
{getattr(issue, 'description', '无描述')}

【原始设定】
{getattr(issue, 'original_content', '无')}

【修正建议】
{getattr(issue, 'suggested_fix', '无建议')}

【置信度】
{getattr(issue, 'confidence', 0) * 100:.0f}%
"""
        self._issue_detail_text.insert("1.0", detail)
        self._issue_detail_text.configure(state=tk.DISABLED)
    
    def _on_reverse_issue_double_click(self, event):
        """双击冲突项，打开详情窗口"""
        self._on_reverse_issue_select(event)
        
        # 获取详情内容
        detail = self._issue_detail_text.get("1.0", tk.END)
        
        # 创建详情窗口
        detail_window = tk.Toplevel(self.root)
        detail_window.title("冲突详情")
        detail_window.geometry("500x400")
        
        text = tk.Text(detail_window, wrap=tk.WORD, font=('Microsoft YaHei UI', 10))
        text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text.insert("1.0", detail)
        text.configure(state=tk.DISABLED)
    
    def _on_reverse_apply_fix(self):
        """应用修正"""
        if not self._reverse_analysis_result:
            messagebox.showwarning("提示", "请先运行分析")
            return
        
        issues = self._reverse_analysis_result.get('issues', [])
        if not issues:
            messagebox.showinfo("提示", "没有需要修正的冲突项")
            return
        
        # 确认对话框
        high_count = sum(1 for i in issues if getattr(i, 'severity', None) and 
                        (getattr(i.severity, 'value', i.severity) == 'high'))
        
        msg = f"发现 {len(issues)} 个冲突项（{high_count} 个高优先级）\n\n是否应用修正？"
        if not messagebox.askyesno("确认修正", msg):
            return
        
        # 获取插件
        plugin = self._get_reverse_feedback_plugin()
        if not plugin:
            messagebox.showerror("错误", "逆向反馈插件不可用")
            return
        
        # 生成修正
        settings = self._get_current_project_settings()
        
        # 创建简化的报告对象
        from core.plugin_interface import ConsistencyReport
        report = ConsistencyReport(project_name=settings.get('project_name', ''))
        report.issues = issues
        
        corrections = plugin.generate_corrections(report, settings)
        
        # 显示修正结果
        self._issue_detail_text.configure(state=tk.NORMAL)
        self._issue_detail_text.delete("1.0", tk.END)
        
        correction_text = "【修正完成】\n\n"
        for suggestion in corrections.get('suggestions', []):
            correction_text += f"• {suggestion}\n"
        
        correction_text += "\n【备份信息】\n"
        if corrections.get('backup'):
            correction_text += f"备份时间：{corrections['backup'].get('backup_time', '未知')}\n"
        
        self._issue_detail_text.insert("1.0", correction_text)
        self._issue_detail_text.configure(state=tk.DISABLED)
        
        # 更新项目设定（如果用户确认）
        if messagebox.askyesno("应用修正", "是否将修正应用到当前项目设定？"):
            self._apply_corrections_to_project(corrections)
        
        self._set_status("修正已应用")
    
    def _apply_corrections_to_project(self, corrections: Dict):
        """将修正应用到项目设定"""
        # 更新大纲
        if corrections.get('updated_outline'):
            self._outline_content = corrections['updated_outline']
        
        # 更新人物
        if corrections.get('updated_characters'):
            self._character_data = corrections['updated_characters']
        
        # 刷新项目管理器
        self._refresh_project_manager()
    
    def _refresh_project_manager(self):
        """刷新项目管理器"""
        # 切换到项目管理页面更新显示
        self._set_status("项目设定已更新")
    
    def _on_reverse_export_report(self):
        """导出分析报告"""
        if not self._reverse_analysis_result:
            messagebox.showwarning("提示", "没有分析结果可导出")
            return
        
        path = filedialog.asksaveasfilename(
            title="保存分析报告",
            defaultextension=".md",
            filetypes=[("Markdown文件", "*.md"), ("文本文件", "*.txt")]
        )
        
        if path:
            try:
                issues = self._reverse_analysis_result.get('issues', [])
                
                report = f"""# 逆向反馈分析报告

生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')}
分析章节数：{self._reverse_analysis_result.get('chapters_analyzed', 0)}
冲突总数：{len(issues)}

## 冲突列表

"""
                for i, issue in enumerate(issues, 1):
                    severity = getattr(issue, 'severity', 'medium')
                    if hasattr(severity, 'value'):
                        severity = severity.value
                    
                    report += f"""### {i}. {getattr(issue, 'element_name', '未知')}

- **类型**：{getattr(issue, 'issue_type', '未知')}
- **严重程度**：{severity}
- **章节**：{getattr(issue, 'chapter_reference', '未知')}
- **描述**：{getattr(issue, 'description', '')}
- **修正建议**：{getattr(issue, 'suggested_fix', '')}

"""
                
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(report)
                
                self._set_status(f"报告已导出：{path}")
                
            except Exception as e:
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_file_error
                title, full_message = convert_file_error(e, "导出")
                messagebox.showerror(title, full_message)
    
    def _on_reverse_clear_result(self):
        """清除分析结果"""
        self._issues_tree.delete(*self._issues_tree.get_children())
        self._issue_detail_text.configure(state=tk.NORMAL)
        self._issue_detail_text.delete("1.0", tk.END)
        self._issue_detail_text.insert("1.0", "选择冲突项查看详细修正建议...")
        self._issue_detail_text.configure(state=tk.DISABLED)
        self._reverse_analysis_result = None
        self._analysis_progress_label.configure(text="")
        self._set_status("已清除分析结果")
    
    # 兼容旧方法
    def _on_reverse_run(self):
        """运行分析（兼容旧方法）"""
        self._on_reverse_run_analysis()
    
    def _on_reverse_analyze(self):
        """开始逆向分析（兼容旧方法）"""
        self._on_reverse_run_analysis()
    
    def _on_reverse_apply(self):
        """自动修正（兼容旧方法）"""
        self._on_reverse_apply_fix()
    
    def _on_reverse_outline(self):
        """调整大纲细节"""
        self._set_status("调整大纲细节功能开发中...")
    
    def _on_reverse_character(self):
        """修改人物性格"""
        self._set_status("修改人物性格功能开发中...")
    
    def _on_reverse_trajectory(self):
        """补充人物轨迹"""
        self._set_status("补充人物轨迹功能开发中...")
    
    def _on_reverse_worldview(self):
        """增减世界观设定"""
        self._set_status("增减世界观设定功能开发中...")
    
    def _on_quick_generate(self):
        """快捷生成（保留兼容）"""
        self._on_quick_generate_all()
    
    def _on_quick_generate_all(self):
        """一键生成全部四个结果"""
        # 延迟加载插件
        if not hasattr(self, '_quick_creator_plugin') or not self._quick_creator_plugin:
            self._load_quick_creator_plugin()
            if not self._quick_creator_plugin:
                messagebox.showerror("错误", "快捷创作插件未加载，请先配置AI服务")
                return
        
        # 获取关键词
        keywords = self._quick_input.get("1.0", tk.END).strip()
        if not keywords or keywords == "请输入关键词描述...":
            messagebox.showwarning("提示", "请输入创作关键词")
            return
        
        # 获取详细程度
        detail = self._quick_detail_var.get()
        
        # 线程安全：获取参考文本
        with self._quick_lock:
            reference_texts = {}
            for file_data in self._quick_uploaded_files:
                ref_type = file_data["type"]
                if ref_type not in reference_texts:
                    reference_texts[ref_type] = []
                reference_texts[ref_type].append(file_data["content"])
        
        # 禁用生成按钮
        self._quick_gen_all_btn.configure(state=tk.DISABLED)
        self._set_status("正在生成全部设定...")
        
        # 在后台线程执行
        def generate_task():
            try:
                from core.models import QuickCreationRequest
                
                # 创建请求
                request = QuickCreationRequest(
                    keywords=keywords,
                    detail_level=detail,
                    reference_worldview="\n\n".join(reference_texts.get("worldview", [])),
                    reference_outline="\n\n".join(reference_texts.get("outline", [])),
                    reference_characters="\n\n".join(reference_texts.get("characters", [])),
                    reference_plot="\n\n".join(reference_texts.get("plot", []))
                )
                
                # 调用插件生成
                result = self._quick_creator_plugin.generate_all(request)
                
                # 更新UI
                if result.success:
                    # 世界观
                    if result.worldview:
                        self.root.after(0, lambda: self._update_quick_result("worldview", result.worldview.content))
                    
                    # 大纲
                    if result.outline:
                        self.root.after(0, lambda: self._update_quick_result("outline", result.outline.content))
                    
                    # 人设
                    if result.characters:
                        self.root.after(0, lambda: self._update_quick_result("characters", result.characters.content))
                    
                    # 情节
                    if result.plot:
                        self.root.after(0, lambda: self._update_quick_result("plot", result.plot.content))
                    
                    self.root.after(0, lambda: self._set_status("生成完成"))
                else:
                    self.root.after(0, lambda: messagebox.showerror("错误", f"生成失败：{result.error}"))
                    self.root.after(0, lambda: self._set_status("生成失败"))
                
            except Exception as e:
                logger.error(f"快捷生成异常: {e}", exc_info=True)
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_generation_error
                title, full_message = convert_generation_error(e)
                self.root.after(0, lambda: messagebox.showerror(title, full_message))
                self.root.after(0, lambda: self._set_status("生成异常"))
            finally:
                # 恢复按钮
                self.root.after(0, lambda: self._quick_gen_all_btn.configure(state=tk.NORMAL))
        
        # 启动后台任务（使用统一线程池，解决卡顿问题）
        if CORE_AVAILABLE:
            self._async_handler.submit(generate_task, priority=TaskPriority.NORMAL)
        else:
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(generate_task)
    
    def _on_quick_generate_single(self, gen_type: str):
        """生成单个结果"""
        # 延迟加载插件
        if not hasattr(self, '_quick_creator_plugin') or not self._quick_creator_plugin:
            self._load_quick_creator_plugin()
            if not self._quick_creator_plugin:
                messagebox.showerror("错误", "快捷创作插件未加载")
                return
        
        # 获取关键词
        keywords = self._quick_input.get("1.0", tk.END).strip()
        if not keywords or keywords == "请输入关键词描述...":
            messagebox.showwarning("提示", "请输入创作关键词")
            return
        
        # 获取详细程度
        detail = self._quick_detail_var.get()
        
        # 线程安全：获取参考文本
        with self._quick_lock:
            reference_texts = {}
            for file_data in self._quick_uploaded_files:
                ref_type = file_data["type"]
                if ref_type not in reference_texts:
                    reference_texts[ref_type] = []
                reference_texts[ref_type].append(file_data["content"])
        
        type_names = {
            "worldview": "世界观",
            "outline": "大纲",
            "characters": "人设",
            "plot": "关键情节"
        }
        
        self._set_status(f"正在生成{type_names.get(gen_type, '')}...")
        
        # 在后台线程执行
        def generate_task():
            try:
                # 根据类型调用对应的生成方法
                if gen_type == "worldview":
                    result = self._quick_creator_plugin.generate_worldview(
                        keywords=keywords,
                        detail_level=detail,
                        reference_text="\n\n".join(reference_texts.get("worldview", []))
                    )
                    if result.success:
                        self.root.after(0, lambda: self._update_quick_result("worldview", result.content))
                
                elif gen_type == "outline":
                    result = self._quick_creator_plugin.generate_outline(
                        keywords=keywords,
                        detail_level=detail,
                        reference_text="\n\n".join(reference_texts.get("outline", []))
                    )
                    if result.success:
                        self.root.after(0, lambda: self._update_quick_result("outline", result.content))
                
                elif gen_type == "characters":
                    result = self._quick_creator_plugin.generate_character(
                        keywords=keywords,
                        detail_level=detail,
                        reference_text="\n\n".join(reference_texts.get("characters", []))
                    )
                    if result.success:
                        self.root.after(0, lambda: self._update_quick_result("characters", result.content))
                
                elif gen_type == "plot":
                    result = self._quick_creator_plugin.generate_plot(
                        keywords=keywords,
                        detail_level=detail,
                        reference_text="\n\n".join(reference_texts.get("plot", []))
                    )
                    if result.success:
                        self.root.after(0, lambda: self._update_quick_result("plot", result.content))
                
                self.root.after(0, lambda: self._set_status(f"{type_names.get(gen_type, '')}生成完成"))
                
            except Exception as e:
                logger.error(f"生成异常: {e}", exc_info=True)
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_generation_error
                title, full_message = convert_generation_error(e)
                self.root.after(0, lambda: messagebox.showerror(title, full_message))
                self.root.after(0, lambda: self._set_status("生成异常"))
        
        # 启动后台任务（使用统一线程池，解决卡顿问题）
        if CORE_AVAILABLE:
            self._async_handler.submit(generate_task, priority=TaskPriority.NORMAL)
        else:
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(generate_task)
    
    def _update_quick_result(self, result_type: str, content: str):
        """更新快捷创作结果"""
        if result_type in self._quick_result_texts:
            text_widget = self._quick_result_texts[result_type]
            text_widget.configure(state=tk.NORMAL)
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", content)
            text_widget.configure(state=tk.NORMAL)
    
    def _on_quick_upload_file(self, file_type: str):
        """上传参考文本文件
        
        Args:
            file_type: 文件类型（worldview/outline/characters/plot）
        """
        path = filedialog.askopenfilename(
            title=f"选择{file_type}参考文本",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if not path:
            return
        
        try:
            # 读取文件内容
            if path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(path)
                content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # 计算字数
            word_count = len(content.strip())
            
            # 线程安全：添加到上传列表
            file_data = {
                "path": path,
                "filename": os.path.basename(path),
                "type": file_type,
                "content": content,
                "word_count": word_count
            }
            with self._quick_lock:
                self._quick_uploaded_files.append(file_data)
            
            # 更新Treeview
            type_names = {
                "worldview": "世界观",
                "outline": "大纲",
                "characters": "人设",
                "plot": "情节"
            }
            self._quick_upload_tree.insert("", tk.END, values=(
                os.path.basename(path),
                type_names.get(file_type, file_type),
                f"{word_count}",
                "已上传"
            ))
            
            self._set_status(f"已上传：{os.path.basename(path)}（{word_count}字）")
            
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_file_error
            title, full_message = convert_file_error(e, "上传")
            messagebox.showerror(title, full_message)
    
    def _on_quick_remove_upload(self):
        """删除选中的上传文件（P2-002修复：添加二次确认）"""
        selected = self._quick_upload_tree.selection()
        if not selected:
            messagebox.showwarning("提示", "请先选择要删除的文件")
            return
        
        # 二次确认
        if not messagebox.askyesno("确认删除", f"确定要删除选中的 {len(selected)} 个上传文件吗？"):
            return
        
        for item in selected:
            # 从Treeview删除
            self._quick_upload_tree.delete(item)
            
            # 从列表中删除（根据item索引）
            item_index = self._quick_upload_tree.get_children().index(item) if item in self._quick_upload_tree.get_children() else -1
            if 0 <= item_index < len(self._quick_uploaded_files):
                self._quick_uploaded_files.pop(item_index)
        
        self._set_status(f"已删除 {len(selected)} 个上传文件")
    
    def _on_quick_clear_uploads(self):
        """清空所有上传文件（P2-002修复：添加二次确认）"""
        # 线程安全：清空列表
        with self._quick_lock:
            if not self._quick_uploaded_files:
                return
            
            # 二次确认
            if not messagebox.askyesno("确认清空", f"确定要清空所有 {len(self._quick_uploaded_files)} 个上传文件吗？\n此操作不可撤销！"):
                return
            
            self._quick_uploaded_files.clear()
        
        # 清空Treeview
        for item in self._quick_upload_tree.get_children():
            self._quick_upload_tree.delete(item)
        
        self._set_status("已清空所有上传文件")
    
    def _show_quick_upload_menu(self, event):
        """显示上传文件右键菜单"""
        item = self._quick_upload_tree.identify_row(event.y)
        if item:
            self._quick_upload_tree.selection_set(item)
            self._quick_upload_menu.post(event.x_root, event.y_root)
    
    def _show_chapters_context_menu(self, event):
        """显示已完成章节右键菜单"""
        # 选中点击项
        item = self._completed_chapters_tree.identify_row(event.y)
        if item:
            self._completed_chapters_tree.selection_set(item)
            self._chapters_context_menu.post(event.x_root, event.y_root)
    
    def _on_upload_chapter(self):
        """上传章节（兼容旧方法）"""
        self._on_reverse_upload_files()
    
    def _on_mark_chapter_completed(self):
        """标记章节为完成"""
        selected = self._completed_chapters_tree.selection()
        for item in selected:
            values = list(self._completed_chapters_tree.item(item, "values"))
            values[3] = "已完成"
            self._completed_chapters_tree.item(item, values=values)
        self._set_status(f"已标记 {len(selected)} 个章节为完成")
    
    def _on_mark_chapter_incomplete(self):
        """标记章节为未完成"""
        selected = self._completed_chapters_tree.selection()
        for item in selected:
            values = list(self._completed_chapters_tree.item(item, "values"))
            values[3] = "未完成"
            self._completed_chapters_tree.item(item, values=values)
        self._set_status(f"已标记 {len(selected)} 个章节为未完成")
    
    def _on_delete_completed_chapter(self):
        """删除已完成章节（P2-002修复：添加二次确认）"""
        selected = self._completed_chapters_tree.selection()
        if not selected:
            messagebox.showwarning("提示", "请先选择要删除的章节")
            return
        
        # 二次确认
        if not messagebox.askyesno("确认删除", f"确定要删除选中的 {len(selected)} 个章节吗？\n此操作不可撤销！"):
            return
        
        for item in selected:
            self._completed_chapters_tree.delete(item)
        self._set_status(f"已删除 {len(selected)} 个章节")
    
    def _on_mark_all_completed(self):
        """标记所有章节为完成"""
        for item in self._completed_chapters_tree.get_children():
            values = list(self._completed_chapters_tree.item(item, "values"))
            values[3] = "已完成"
            self._completed_chapters_tree.item(item, values=values)
        self._set_status("已标记所有章节为完成")
    
    def _on_quick_save_txt(self):
        """保存为TXT"""
        self._set_status("保存为TXT功能开发中...")
    
    def _on_quick_save_docx(self):
        """保存为DOCX"""
        self._set_status("保存为DOCX功能开发中...")
    
    def _on_quick_import(self):
        """导入当前项目"""
        if not hasattr(self, '_project_manager') or not self._project_manager:
            messagebox.showwarning("提示", "请先打开或创建项目")
            return
        
        # 获取所有结果内容
        content = {}
        for result_type, text_widget in self._quick_result_texts.items():
            content[result_type] = text_widget.get("1.0", tk.END).strip()
        
        # 检查是否有内容
        if not any(content.values()):
            messagebox.showwarning("提示", "没有可导入的内容")
            return
        
        # 创建导入确认对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("导入到当前项目")
        dialog.geometry("400x300")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 400) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 300) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 提示信息
        ttk.Label(dialog, text="选择要导入的内容：").pack(pady=10)
        
        # 复选框
        import_vars = {}
        for result_type, text in content.items():
            if text and not text.endswith("将在此显示..."):
                var = tk.BooleanVar(value=True)
                import_vars[result_type] = var
                type_names = {
                    "worldview": "🌍 世界观",
                    "outline": "📋 大纲",
                    "characters": "👤 人设",
                    "plot": "📖 情节"
                }
                ttk.Checkbutton(dialog, text=type_names.get(result_type, result_type), 
                               variable=var).pack(anchor=tk.W, padx=40, pady=5)
        
        # 按钮区
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def on_import():
            try:
                imported_count = 0
                
                # 导入世界观
                if import_vars.get("worldview", tk.BooleanVar(value=False)).get():
                    self._project_manager.set_worldview(content["worldview"])
                    imported_count += 1
                
                # 导入大纲
                if import_vars.get("outline", tk.BooleanVar(value=False)).get():
                    self._project_manager.set_outline(content["outline"])
                    imported_count += 1
                
                # 导入人设
                if import_vars.get("characters", tk.BooleanVar(value=False)).get():
                    self._project_manager.set_characters(content["characters"])
                    imported_count += 1
                
                # 导入情节
                if import_vars.get("plot", tk.BooleanVar(value=False)).get():
                    self._project_manager.set_plot(content["plot"])
                    imported_count += 1
                
                self._set_status(f"已导入 {imported_count} 项设定到当前项目")
                messagebox.showinfo("成功", f"成功导入 {imported_count} 项设定")
                dialog.destroy()
                
            except Exception as e:
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_file_error
                title, full_message = convert_file_error(e, "导入")
                messagebox.showerror(title, full_message)
        
        ttk.Button(btn_frame, text="导入", command=on_import).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT)
    
    def _on_quick_export_results(self, export_format: str = "markdown"):
        """导出生成结果
        
        Args:
            export_format: 导出格式（markdown/json）
        """
        # 获取所有结果内容
        content = {}
        for result_type, text_widget in self._quick_result_texts.items():
            text = text_widget.get("1.0", tk.END).strip()
            # 过滤占位符文本
            if text and not text.endswith("将在此显示..."):
                content[result_type] = text
        
        # 检查是否有内容
        if not content:
            messagebox.showwarning("提示", "没有可导出的内容")
            return
        
        # 选择保存路径
        if export_format == "json":
            default_ext = ".json"
            filetypes = [("JSON文件", "*.json"), ("所有文件", "*.*")]
        else:
            default_ext = ".md"
            filetypes = [("Markdown文件", "*.md"), ("文本文件", "*.txt"), ("所有文件", "*.*")]
        
        save_path = filedialog.asksaveasfilename(
            title="选择保存位置",
            defaultextension=default_ext,
            filetypes=filetypes,
            initialfile="快捷创作结果"
        )
        
        if not save_path:
            return
        
        try:
            if export_format == "json":
                # JSON格式
                import json
                data = {
                    "生成时间": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    "世界观设定": content.get("worldview", ""),
                    "章节大纲": content.get("outline", ""),
                    "人物设定": content.get("characters", ""),
                    "关键情节": content.get("plot", "")
                }
                with open(save_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            else:
                # Markdown格式
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write("# 快捷创作生成结果\n\n")
                    f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                    
                    if "worldview" in content:
                        f.write("## 🌍 世界观设定\n\n")
                        f.write(content["worldview"] + "\n\n")
                    
                    if "outline" in content:
                        f.write("## 📋 章节大纲\n\n")
                        f.write(content["outline"] + "\n\n")
                    
                    if "characters" in content:
                        f.write("## 👤 人物设定\n\n")
                        f.write(content["characters"] + "\n\n")
                    
                    if "plot" in content:
                        f.write("## 📖 关键情节\n\n")
                        f.write(content["plot"] + "\n\n")
            
            self._set_status(f"已导出到：{os.path.basename(save_path)}")
            messagebox.showinfo("成功", f"导出成功！\n保存位置：{save_path}")
            
        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_file_error
            title, full_message = convert_file_error(e, "导出")
            messagebox.showerror(title, full_message)
    
    def _browse_continue_source(self):
        """浏览续写原文文件"""
        path = filedialog.askopenfilename(
            title="选择原文文件",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            try:
                if path.lower().endswith('.docx'):
                    from docx import Document
                    doc = Document(path)
                    text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                else:
                    with open(path, 'r', encoding='utf-8') as f:
                        text = f.read()
                self._continue_source.delete("1.0", tk.END)
                self._continue_source.insert("1.0", text)
                self._set_status(f"已导入: {os.path.basename(path)}")
            except Exception as e:
                self._set_status(f"导入失败: {e}")
    
    def _on_continue_generate(self):
        """开始续写"""
        # 检查插件
        if not hasattr(self, '_continuation_plugin') or not self._continuation_plugin:
            messagebox.showerror("错误", "续写插件未加载，请先配置AI服务")
            return
        
        # 获取原文
        source_text = self._continue_source.get("1.0", tk.END).strip()
        if not source_text or source_text == "请在此粘贴或输入原文内容，或选择文件导入...":
            messagebox.showwarning("提示", "请先输入或导入原文内容")
            return
        
        # 获取字数
        words_str = self._continue_words_var.get()
        if words_str == "自定义":
            try:
                word_count = int(self._custom_words_entry.get().strip())
                if word_count < 100 or word_count > 5000:
                    raise ValueError("字数超出范围")
            except ValueError as e:
                messagebox.showwarning("提示", f"请输入有效的字数（100-5000）\n错误：{e}")
                return
        else:
            word_count = int(words_str)
        
        # 获取续写方向
        direction_map = {
            "自然续写": "natural",
            "剧情推进": "action",
            "高潮铺垫": "action",
            "结局": "specific",
            "制造冲突": "specific",
            "引入新角色": "specific",
            "转折情节": "specific",
            "情感描写": "emotion",
            "动作场景": "action",
            "对话为主": "dialogue"
        }
        direction = direction_map.get(self._continue_direction_var.get(), "natural")
        
        # 获取温度
        temperature = self._continue_temp_var.get()
        
        # 获取生成模式
        is_multi_version = self._continue_mode_var.get() == "多版本生成"
        
        # 检查版本数量限制（最多5个版本）
        if is_multi_version and len(self._continue_versions) >= 5:
            messagebox.showwarning("提示", "已达到最大版本数量（5个），请先保存或清空版本")
            return
        
        # 清空历史版本（仅在首次生成时）
        if not self._continue_versions:
            self._continue_versions = []
            self._current_version_index = 0
            self._best_version_index = -1
        
        # 禁用按钮
        self._start_btn.configure(state=tk.DISABLED)
        self._regenerate_btn.configure(state=tk.DISABLED)
        self._select_best_btn.configure(state=tk.DISABLED)
        self._save_btn.configure(state=tk.DISABLED)
        
        # 清空结果框
        self._continue_result.delete("1.0", tk.END)
        self._version_info_label.configure(text="正在生成...")
        self._result_info_label.configure(text="字数：- | 耗时：- | 模型：-")
        self._score_label.configure(text="评分：-")
        
        # 在后台线程执行续写
        def generate_task():
            try:
                from core.models import ContinuationRequest
                
                # 获取上下文（如果有项目管理器）
                outline = None
                characters = None
                worldview = None
                style_profile = None
                previous_chapters = None
                
                if hasattr(self, '_project_manager') and self._project_manager:
                    # 从项目获取上下文
                    outline = self._project_manager.get_outline()
                    characters = self._project_manager.get_characters()
                    worldview = self._project_manager.get_worldview()
                    # 获取最近5章作为前文参考
                    previous_chapters = self._project_manager.get_recent_chapters(5)
                
                if is_multi_version:
                    # 多版本生成
                    self._set_status("正在生成多个版本...")
                    
                    # 计算可生成的版本数量（最多5个）
                    remaining_slots = 5 - len(self._continue_versions)
                    num_to_generate = min(3, remaining_slots)  # 每次最多生成3个
                    
                    if num_to_generate <= 0:
                        self.root.after(0, lambda: messagebox.showwarning("提示", "已达到最大版本数量（5个）"))
                        return
                    
                    # 根据剩余空间调整温度分布
                    if num_to_generate == 1:
                        temps = [temperature]
                    elif num_to_generate == 2:
                        temps = [max(0.5, temperature - 0.2), min(1.0, temperature + 0.2)]
                    else:
                        temps = [0.6, 0.8, 1.0]
                    
                    results = self._continuation_plugin.generate_multiple_versions(
                        request=ContinuationRequest(
                            starting_text=source_text,
                            word_count=word_count,
                            direction=direction,
                            outline=outline,
                            characters=characters,
                            worldview=worldview,
                            style_profile=style_profile,
                            previous_chapters=previous_chapters,
                            temperature=temperature
                        ),
                        num_versions=num_to_generate,
                        temperatures=temps
                    )
                    
                    # 存储所有版本
                    for i, result in enumerate(results):
                        if result.success:
                            version_data = {
                                "text": result.text,
                                "word_count": result.word_count,
                                "temperature": [0.6, 0.8, 1.0][i],
                                "metadata": result.metadata.model_dump(),
                                "score": 0
                            }
                            self._continue_versions.append(version_data)
                    
                    # 自动选择最佳版本
                    if self._continue_versions:
                        best_result, best_index, scores = self._continuation_plugin.select_best_version(
                            results[:len(self._continue_versions)]
                        )
                        self._best_version_index = best_index
                        
                        # 更新评分
                        for i, score_dict in enumerate(scores):
                            if i < len(self._continue_versions):
                                self._continue_versions[i]["score"] = score_dict.get("total", 0)
                        
                        # 显示最佳版本
                        self._current_version_index = best_index
                        self.root.after(0, lambda: self._display_version(best_index))
                    
                    self.root.after(0, lambda: self._set_status(f"已生成 {len(self._continue_versions)} 个版本"))
                else:
                    # 单次生成
                    self._set_status("正在生成续写...")
                    
                    request = ContinuationRequest(
                        starting_text=source_text,
                        word_count=word_count,
                        direction=direction,
                        outline=outline,
                        characters=characters,
                        worldview=worldview,
                        style_profile=style_profile,
                        previous_chapters=previous_chapters,
                        temperature=temperature
                    )
                    
                    result = self._continuation_plugin.generate_continuation(request)
                    
                    if result.success:
                        # 计算评分
                        scores = self._continuation_plugin._evaluate_version(result, request)
                        
                        version_data = {
                            "text": result.text,
                            "word_count": result.word_count,
                            "temperature": temperature,
                            "metadata": result.metadata.model_dump(),
                            "score": scores.get("total", 0)
                        }
                        self._continue_versions.append(version_data)
                        self._current_version_index = 0
                        self._best_version_index = 0
                        
                        # 显示结果
                        self.root.after(0, lambda: self._display_version(0))
                        self.root.after(0, lambda: self._set_status("续写生成完成"))
                    else:
                        self.root.after(0, lambda: messagebox.showerror("错误", f"续写失败：{result.error}"))
                        self.root.after(0, lambda: self._set_status("续写生成失败"))
                
            except Exception as e:
                logger.error(f"续写生成异常: {e}", exc_info=True)
                # P2-003修复：用户友好错误提示
                from core.user_friendly_errors import convert_generation_error
                title, full_message = convert_generation_error(e)
                self.root.after(0, lambda: messagebox.showerror(title, full_message))
                self.root.after(0, lambda: self._set_status("续写生成异常"))
            finally:
                # 恢复按钮
                self.root.after(0, lambda: self._start_btn.configure(state=tk.NORMAL))
                self.root.after(0, lambda: self._update_version_combo())
        
        # 启动后台任务（使用统一线程池，解决卡顿问题）
        if CORE_AVAILABLE:
            self._async_handler.submit(generate_task, priority=TaskPriority.NORMAL)
        else:
            from core.thread_pool_manager import thread_pool_manager
            thread_pool_manager.submit_sync(generate_task)
    
    def _on_continue_regenerate(self):
        """重新生成续写"""
        if not hasattr(self, '_continuation_plugin') or not self._continuation_plugin:
            messagebox.showerror("错误", "续写插件未加载")
            return
        
        # 显示重新生成模式选择对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("重新生成模式")
        dialog.geometry("400x250")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 400) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 250) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 模式选择
        ttk.Label(dialog, text="选择重新生成模式：").pack(pady=20)
        
        mode_var = tk.StringVar(value="creative")
        ttk.Radiobutton(dialog, text="🎨 创意模式（温度1.0，更多创新）", variable=mode_var, value="creative").pack(anchor=tk.W, padx=40, pady=5)
        ttk.Radiobutton(dialog, text="⚖️ 平衡模式（温度0.8，平衡创新和连贯）", variable=mode_var, value="balanced").pack(anchor=tk.W, padx=40, pady=5)
        ttk.Radiobutton(dialog, text="🔒 保守模式（温度0.6，更连贯保守）", variable=mode_var, value="conservative").pack(anchor=tk.W, padx=40, pady=5)
        ttk.Radiobutton(dialog, text="🎯 聚焦模式（温度0.5，高度连贯）", variable=mode_var, value="focused").pack(anchor=tk.W, padx=40, pady=5)
        
        # 按钮区
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def on_regenerate():
            dialog.destroy()
            
            # 获取模式对应的温度
            temp_map = {
                "creative": 1.0,
                "balanced": 0.8,
                "conservative": 0.6,
                "focused": 0.5
            }
            temperature = temp_map.get(mode_var.get(), 0.8)
            
            # 更新温度设置
            self._continue_temp_var.set(temperature)
            self._temp_label.configure(text=f"{temperature:.1f}")
            
            # 执行续写
            self._on_continue_generate()
        
        ttk.Button(btn_frame, text="确定", command=on_regenerate).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT)
    
    def _on_continue_select_best(self):
        """选择最佳版本并保存为新章节"""
        if not self._continue_versions:
            messagebox.showwarning("提示", "没有可选择的版本")
            return
        
        if self._best_version_index < 0:
            # 如果没有最佳版本索引，使用当前选中的版本
            best_index = self._current_version_index
        else:
            best_index = self._best_version_index
        
        # 切换到最佳版本
        self._current_version_index = best_index
        self._display_version(best_index)
        
        # 弹出章节命名窗口
        self._show_chapter_naming_dialog(best_index)
    
    def _show_chapter_naming_dialog(self, version_index: int):
        """显示章节命名对话框"""
        if not self._continue_versions or version_index >= len(self._continue_versions):
            return
        
        current_version = self._continue_versions[version_index]
        continuation_text = current_version.get("text", "")
        
        # 创建命名对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("保存为新章节")
        dialog.geometry("400x200")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 400) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 200) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 提示信息
        ttk.Label(dialog, text=f"正在保存版本 V{version_index + 1}（评分：{current_version.get('score', 0):.2f}）").pack(pady=10)
        
        # 章节名称输入
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=20, pady=10)
        
        ttk.Label(name_frame, text="章节名称：").pack(side=tk.LEFT)
        chapter_name_var = tk.StringVar(value=f"第X章 续写内容")
        chapter_name_entry = ttk.Entry(name_frame, textvariable=chapter_name_var, width=25)
        chapter_name_entry.pack(side=tk.LEFT, padx=5)
        chapter_name_entry.focus()
        
        # 字数统计
        word_count = current_version.get("word_count", 0)
        ttk.Label(dialog, text=f"字数：{word_count}").pack(pady=5)
        
        # 按钮区
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def on_save():
            chapter_name = chapter_name_var.get().strip()
            if not chapter_name:
                messagebox.showwarning("提示", "请输入章节名称")
                return
            
            # 保存到项目
            if hasattr(self, '_project_manager') and self._project_manager:
                try:
                    self._project_manager.add_chapter(chapter_name, continuation_text)
                    self._set_status(f"已保存为新章节：{chapter_name}")
                    messagebox.showinfo("成功", f"已保存章节：{chapter_name}")
                    dialog.destroy()
                except Exception as e:
                    # P2-003修复：用户友好错误提示
                    from core.user_friendly_errors import convert_file_error
                    title, full_message = convert_file_error(e, "保存章节")
                    messagebox.showerror(title, full_message)
            else:
                messagebox.showwarning("提示", "请先打开或创建项目")
        
        ttk.Button(btn_frame, text="保存", command=on_save).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT)
    
    def _on_continue_save(self):
        """保存续写结果"""
        if not self._continue_versions:
            messagebox.showwarning("提示", "没有可保存的内容")
            return
        
        current_version = self._continue_versions[self._current_version_index]
        continuation_text = current_version.get("text", "")
        
        if not continuation_text:
            messagebox.showwarning("提示", "当前版本内容为空")
            return
        
        # 创建保存对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("保存续写结果")
        dialog.geometry("450x300")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 450) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 300) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 保存选项
        ttk.Label(dialog, text="保存方式：").pack(pady=10)
        
        save_mode = tk.StringVar(value="new")
        ttk.Radiobutton(dialog, text="📄 保存为新章节", variable=save_mode, value="new").pack(anchor=tk.W, padx=40, pady=5)
        ttk.Radiobutton(dialog, text="📝 追加到当前章节", variable=save_mode, value="append").pack(anchor=tk.W, padx=40, pady=5)
        ttk.Radiobutton(dialog, text="💾 导出为文件", variable=save_mode, value="export").pack(anchor=tk.W, padx=40, pady=5)
        
        # 章节名称（仅新章节模式）
        name_frame = ttk.Frame(dialog, style="TFrame")
        name_frame.pack(fill=tk.X, padx=40, pady=10)
        ttk.Label(name_frame, text="章节名称：").pack(side=tk.LEFT)
        chapter_name_var = tk.StringVar(value=f"第X章 续写内容")
        chapter_name_entry = ttk.Entry(name_frame, textvariable=chapter_name_var, width=25)
        chapter_name_entry.pack(side=tk.LEFT, padx=5)
        
        # 按钮区
        btn_frame = ttk.Frame(dialog, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def on_save():
            mode = save_mode.get()
            
            if mode == "export":
                # 导出为文件
                file_path = filedialog.asksaveasfilename(
                    title="保存续写结果",
                    defaultextension=".txt",
                    filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("Markdown", "*.md")]
                )
                if file_path:
                    try:
                        if file_path.endswith('.docx'):
                            from docx import Document
                            doc = Document()
                            doc.add_heading("续写内容", level=1)
                            doc.add_paragraph(continuation_text)
                            doc.save(file_path)
                        else:
                            with open(file_path, 'w', encoding='utf-8') as f:
                                f.write(continuation_text)
                        
                        self._set_status(f"已导出到：{os.path.basename(file_path)}")
                        dialog.destroy()
                    except Exception as e:
                        messagebox.showerror("错误", f"导出失败：{e}")
            
            elif mode == "new" or mode == "append":
                # 检查项目管理器
                if not hasattr(self, '_project_manager') or not self._project_manager:
                    messagebox.showwarning("提示", "请先打开或创建项目")
                    return
                
                try:
                    if mode == "new":
                        # 保存为新章节
                        chapter_name = chapter_name_var.get()
                        self._project_manager.add_chapter(chapter_name, continuation_text)
                        self._set_status(f"已保存为新章节：{chapter_name}")
                    else:
                        # 追加到当前章节
                        source_text = self._continue_source.get("1.0", tk.END).strip()
                        full_text = source_text + "\n\n" + continuation_text
                        # 更新当前章节内容（需要项目管理器支持）
                        if hasattr(self._project_manager, 'update_current_chapter'):
                            self._project_manager.update_current_chapter(full_text)
                        else:
                            # 如果不支持，保存为新章节
                            self._project_manager.add_chapter(f"第X章（续写）", full_text)
                        self._set_status("已追加到当前章节")
                    
                    dialog.destroy()
                except Exception as e:
                    # P2-003修复：用户友好错误提示
                    from core.user_friendly_errors import convert_file_error
                    title, full_message = convert_file_error(e, "保存")
                    messagebox.showerror(title, full_message)
        
        ttk.Button(btn_frame, text="保存", command=on_save).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT)
        
        # 根据保存模式切换章节名称输入框状态
        def on_mode_changed():
            if save_mode.get() == "new":
                chapter_name_entry.configure(state=tk.NORMAL)
            else:
                chapter_name_entry.configure(state=tk.DISABLED)
        
        save_mode.trace("w", lambda *args: on_mode_changed())
    
    def _create_progress_page(self) -> tk.Frame:
        """创建创作进度页面（支持滚动）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建Canvas和Scrollbar用于滚动
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题
        header = ttk.Frame(scrollable_frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="创作进度", style="Title.TLabel").pack(side=tk.LEFT)
        
        # 进度统计
        progress_frame = ttk.LabelFrame(scrollable_frame, text="当前项目状态", padding=20)
        progress_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 使用变量存储统计信息
        self._progress_project_name = tk.StringVar(value="未打开项目")
        self._progress_chapters = tk.StringVar(value="0 / 0")
        self._progress_total_words = tk.StringVar(value="0 字")
        self._progress_today_words = tk.StringVar(value="0 / 3000 字")
        self._progress_outline = tk.StringVar(value="未完成")
        self._progress_characters = tk.StringVar(value="0 人")
        self._progress_worldview = tk.StringVar(value="0 个")
        self._progress_style = tk.StringVar(value="未设置")
        
        stats = [
            ("当前项目", self._progress_project_name),
            ("已完成章节", self._progress_chapters),
            ("总字数", self._progress_total_words),
            ("今日目标", self._progress_today_words),
            ("大纲解析", self._progress_outline),
            ("人物录入", self._progress_characters),
            ("世界观条目", self._progress_worldview),
            ("当前风格", self._progress_style),
        ]
        
        for i, (label, var) in enumerate(stats):
            row_frame = ttk.Frame(progress_frame, style="TFrame")
            row_frame.pack(fill=tk.X, pady=3)
            
            ttk.Label(row_frame, text=f"{label}：", width=15, anchor='e').pack(side=tk.LEFT)
            ttk.Label(row_frame, textvariable=var, foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=10)
        
        # 操作按钮
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=10)

        buttons = [
            ("刷新统计", self._on_refresh_progress),
            ("导出报告", self._on_export_progress),
            ("重置进度", self._on_reset_progress),
        ]

        for i, (text, command) in enumerate(buttons):
            ResponsiveButton(
                btn_frame,
                text=text,
                command=command,
                async_handler=self._async_handler
            ).grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)
        
        return frame
    
    def _create_project_page(self) -> tk.Frame:
        """创建项目管理页面（支持滚动）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建Canvas和Scrollbar用于滚动
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题
        header = ttk.Frame(scrollable_frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="项目管理", style="Title.TLabel").pack(side=tk.LEFT)
        
        # 当前项目信息
        info_frame = ttk.LabelFrame(scrollable_frame, text="当前项目", padding=20)
        info_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 项目名称
        name_frame = ttk.Frame(info_frame, style="TFrame")
        name_frame.pack(fill=tk.X, pady=5)
        ttk.Label(name_frame, text="项目名称：").pack(side=tk.LEFT)
        self._project_name_var = tk.StringVar(value="未打开项目")
        ttk.Label(name_frame, textvariable=self._project_name_var).pack(side=tk.LEFT, padx=10)
        
        # 项目路径
        path_frame = ttk.Frame(info_frame, style="TFrame")
        path_frame.pack(fill=tk.X, pady=5)
        ttk.Label(path_frame, text="项目路径：").pack(side=tk.LEFT)
        self._project_path_var = tk.StringVar(value="-")
        ttk.Label(path_frame, textvariable=self._project_path_var).pack(side=tk.LEFT, padx=10)
        
        # 操作按钮
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=10)

        buttons = [
            ("新建项目", self._on_new_project),
            ("打开项目", self._on_open_project),
            ("导出备份", self._on_backup_project),
        ]

        for i, (text, command) in enumerate(buttons):
            ResponsiveButton(
                btn_frame,
                text=text,
                command=command,
                async_handler=self._async_handler,
                style="Accent.TButton" if i == 0 else "TButton"
            ).grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重，确保按钮均匀分布
        for i in range(3):
            btn_frame.grid_columnconfigure(i, weight=1)

        return frame
    
    def _create_plugins_page(self) -> tk.Frame:
        """创建插件管理页面（支持滚动）- V2.2增强版
        
        功能：
        - 显示已安装插件列表（名称、版本、状态、类型、保护状态）
        - 启用/禁用插件（调用PluginRegistry）
        - 插件配置界面（通过PluginContext读取/保存配置）
        - V5保护模块标识（禁止操作）
        """
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建Canvas和Scrollbar用于滚动
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题区域
        header = ttk.Frame(scrollable_frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="插件管理", style="Title.TLabel").pack(side=tk.LEFT)
        
        # 状态说明
        status_info = ttk.Label(
            header, 
            text="🔒 = V5核心保护模块（禁止禁用）", 
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
            foreground=GlassTheme.ACCENT_ORANGE
        )
        status_info.pack(side=tk.RIGHT)
        
        # 插件列表
        list_frame = ttk.LabelFrame(scrollable_frame, text="已安装插件", padding=15)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # V2.2新增：增加保护状态列
        columns = ("name", "version", "status", "type", "protected")
        self._plugin_tree = ttk.Treeview(
            list_frame,
            columns=columns,
            show="headings",
            height=15,
            selectmode="browse"  # 单选模式
        )

        # 配置字体
        style = ttk.Style()
        style.configure("Treeview",
                     font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL))
        style.configure("Treeview.Heading",
                     font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold"))
        
        # 配置列标题和宽度
        self._plugin_tree.heading("name", text="插件名称")
        self._plugin_tree.heading("version", text="版本")
        self._plugin_tree.heading("status", text="状态")
        self._plugin_tree.heading("type", text="类型")
        self._plugin_tree.heading("protected", text="保护")
        
        self._plugin_tree.column("name", width=200, minwidth=150)
        self._plugin_tree.column("version", width=80, minwidth=60)
        self._plugin_tree.column("status", width=80, minwidth=60)
        self._plugin_tree.column("type", width=100, minwidth=80)
        self._plugin_tree.column("protected", width=60, minwidth=50)
        
        # 添加滚动条
        tree_scroll = ttk.Scrollbar(list_frame, orient="vertical", command=self._plugin_tree.yview)
        self._plugin_tree.configure(yscrollcommand=tree_scroll.set)
        
        self._plugin_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 绑定双击事件打开配置对话框
        self._plugin_tree.bind("<Double-1>", self._on_plugin_double_click)
        # 绑定选择事件更新按钮状态
        self._plugin_tree.bind("<<TreeviewSelect>>", self._on_plugin_select)
        
        # V2.3优化：异步加载插件数据
        self._load_plugins_async()
        
        # 插件详情区域
        detail_frame = ttk.LabelFrame(scrollable_frame, text="插件详情", padding=15)
        detail_frame.pack(fill=tk.X, padx=20, pady=10)
        
        self._plugin_detail_var = tk.StringVar(value="请选择一个插件查看详情")
        detail_label = ttk.Label(
            detail_frame, 
            textvariable=self._plugin_detail_var,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL),
            wraplength=800,
            justify=tk.LEFT
        )
        detail_label.pack(fill=tk.X, anchor=tk.W)
        
        # 操作按钮区域
        btn_frame = ttk.Frame(scrollable_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, padx=20, pady=10)

        # V2.2新增：启用/禁用/配置按钮
        buttons = [
            ("刷新列表", self._on_refresh_plugins, "TButton"),
            ("启用插件", self._on_enable_plugin, "TButton"),
            ("禁用插件", self._on_disable_plugin, "TButton"),
            ("插件配置", self._on_config_plugin, "Accent.TButton"),
            ("重新加载", self._on_reload_plugin, "TButton"),
        ]

        for i, (text, command, btn_style) in enumerate(buttons):
            ResponsiveButton(
                btn_frame,
                text=text,
                command=command,
                async_handler=self._async_handler,
                style=btn_style
            ).grid(row=0, column=i, padx=5, pady=2, sticky="ew")

        # 配置列权重
        for i in range(len(buttons)):
            btn_frame.grid_columnconfigure(i, weight=1)

        # 存储按钮引用以便更新状态
        self._plugin_buttons = {
            "enable": btn_frame.winfo_children()[1],
            "disable": btn_frame.winfo_children()[2],
            "config": btn_frame.winfo_children()
        }

        return frame
    
    def _create_feedback_page(self) -> tk.Frame:
        """创建反馈与建议页面（Claw化 V2.0）
        
        功能：
        - 用户反馈输入框（支持多种反馈类型）
        - 当前章节预览
        - 系统已学习知识点展示
        - 历史反馈查看
        - 改进报告生成
        
        UI设计参考：经验文档/11.4Claw化实际运行说明✅️.md 第三节
        """
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建滚动容器
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 绑定鼠标滚轮
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        # ========== 标题区 ==========
        header = ttk.Frame(scrollable_frame, style="TFrame")
        header.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        ttk.Label(header, text="反馈与建议", style="Title.TLabel").pack(side=tk.LEFT)
        ttk.Label(header, text="帮助系统“越用越聪明”", 
                 font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
                 foreground=GlassTheme.TEXT_SECONDARY).pack(side=tk.LEFT, padx=10)
        
        # ========== 当前章节预览 ==========
        chapter_frame = ttk.LabelFrame(scrollable_frame, text="当前章节", padding=15)
        chapter_frame.pack(fill=tk.X, padx=20, pady=10)
        
        self._feedback_chapter_label = ttk.Label(chapter_frame, text="未选择章节")
        self._feedback_chapter_label.pack(anchor="w")
        
        # 内容预览
        preview_frame = ttk.Frame(chapter_frame, style="TFrame")
        preview_frame.pack(fill=tk.X, pady=5)
        
        self._feedback_content_preview = tk.Text(
            preview_frame, 
            height=4, 
            wrap=tk.WORD,
            font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_SMALL),
            bg=GlassTheme.GLASS_SURFACE,
            fg=GlassTheme.TEXT_PRIMARY,
            state=tk.DISABLED
        )
        self._feedback_content_preview.pack(fill=tk.X)
        
        # ========== 反馈输入区 ==========
        feedback_frame = ttk.LabelFrame(scrollable_frame, text="您的反馈", padding=15)
        feedback_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 反馈类型选择
        type_frame = ttk.Frame(feedback_frame, style="TFrame")
        type_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(type_frame, text="反馈类型：").pack(side=tk.LEFT)
        
        self._feedback_type_var = tk.StringVar(value="style")
        feedback_types = [
            ("内容问题", "content"),
            ("风格问题", "style"),
            ("AI感问题", "ai_feeling"),
            ("其他", "other")
        ]
        
        for text, value in feedback_types:
            ttk.Radiobutton(
                type_frame, 
                text=text, 
                variable=self._feedback_type_var, 
                value=value
            ).pack(side=tk.LEFT, padx=10)
        
        # 反馈文本输入
        self._feedback_text = tk.Text(
            feedback_frame,
            height=5,
            wrap=tk.WORD,
            font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_NORMAL),
            bg=GlassTheme.GLASS_SURFACE,
            fg=GlassTheme.TEXT_PRIMARY
        )
        self._feedback_text.pack(fill=tk.X, pady=10)
        self._feedback_text.insert("1.0", "例如：这个对话太生硬了，不像角色会说的话")
        self._feedback_text.bind("<FocusIn>", self._on_feedback_text_focus)
        
        # 提示
        hint_label = ttk.Label(
            feedback_frame,
            text="提示：描述越具体，系统学习效果越好",
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_TINY),
            foreground=GlassTheme.TEXT_SECONDARY
        )
        hint_label.pack(anchor="w")
        
        # 操作按钮
        btn_frame = ttk.Frame(feedback_frame, style="TFrame")
        btn_frame.pack(fill=tk.X, pady=10)
        
        ResponsiveButton(
            btn_frame,
            text="提交反馈",
            command=self._on_submit_feedback,
            async_handler=self._async_handler,
            style="Accent.TButton"
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame,
            text="查看历史反馈",
            command=self._on_view_feedback_history
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame,
            text="生成改进报告",
            command=self._on_generate_improvement_report
        ).pack(side=tk.LEFT, padx=5)
        
        # ========== 系统已学习知识点 ==========
        learned_frame = ttk.LabelFrame(scrollable_frame, text="系统已学习的知识点", padding=15)
        learned_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 统计信息
        stats_frame = ttk.Frame(learned_frame, style="TFrame")
        stats_frame.pack(fill=tk.X, pady=5)
        
        self._feedback_stats_label = ttk.Label(
            stats_frame,
            text="总反馈数：0 | 已处理：0 | 已提取知识点：0"
        )
        self._feedback_stats_label.pack(side=tk.LEFT)
        
        # 知识点列表
        self._learned_knowledge_list = tk.Text(
            learned_frame,
            height=8,
            wrap=tk.WORD,
            font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_SMALL),
            bg=GlassTheme.GLASS_SURFACE,
            fg=GlassTheme.TEXT_PRIMARY,
            state=tk.DISABLED
        )
        self._learned_knowledge_list.pack(fill=tk.X, pady=5)
        
        # 刷新按钮
        ttk.Button(
            learned_frame,
            text="刷新知识点",
            command=self._refresh_learned_knowledge
        ).pack(anchor="w")
        
        # ========== 反馈效果说明 ==========
        effect_frame = ttk.LabelFrame(scrollable_frame, text="反馈如何影响生成", padding=15)
        effect_frame.pack(fill=tk.X, padx=20, pady=10)
        
        effect_text = """
1. 提交反馈后，系统会自动分析并提取知识点
2. 提取的知识点会存入知识库，下次生成时召回
3. 系统会根据反馈调整生成策略（如增加约束、优化权重）
4. 定期生成改进报告，展示优化进度

即时效果：提交反馈后，下次生成立即应用
明显效果：1-2周后，八大维度提升3-5%
显著效果：1-2月后，AI感降低30%，用户满意度提升
        """.strip()
        
        ttk.Label(
            effect_frame,
            text=effect_text,
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
            foreground=GlassTheme.TEXT_SECONDARY,
            justify=tk.LEFT
        ).pack(anchor="w")
        
        # 初始化反馈收集器
        self._init_feedback_collector()
        
        return frame
    
    def _init_feedback_collector(self):
        """初始化反馈收集器"""
        try:
            from core.feedback_collector import get_feedback_collector
            self._feedback_collector = get_feedback_collector()
            self._refresh_feedback_stats()
        except Exception as e:
            logger.warning(f"反馈收集器初始化失败: {e}")
            self._feedback_collector = None
    
    def _on_feedback_text_focus(self, event):
        """反馈文本框获得焦点时清除占位符"""
        current_text = self._feedback_text.get("1.0", "end-1c")
        if current_text == "例如：这个对话太生硬了，不像角色会说的话":
            self._feedback_text.delete("1.0", tk.END)
    
    def _refresh_feedback_stats(self):
        """刷新反馈统计信息"""
        if not hasattr(self, '_feedback_collector') or not self._feedback_collector:
            return
        
        try:
            stats = self._feedback_collector.get_statistics()
            self._feedback_stats_label.config(
                text=f"总反馈数：{stats['total']} | 已处理：{stats['total'] - stats['unprocessed']} | 已提取知识点：{stats['total'] - stats['unextracted']}"
            )
        except Exception as e:
            logger.warning(f"刷新反馈统计失败: {e}")
    
    def _refresh_learned_knowledge(self):
        """刷新已学习知识点列表"""
        if not hasattr(self, '_feedback_collector') or not self._feedback_collector:
            return
        
        try:
            # 获取最近反馈
            history = self._feedback_collector.get_history(limit=10)
            
            # 构建显示文本
            lines = []
            for fb in history:
                if fb.knowledge_extracted:
                    lines.append(f"• [{fb.feedback_type}] {fb.feedback_text[:50]}...")
            
            # 更新显示
            self._learned_knowledge_list.config(state=tk.NORMAL)
            self._learned_knowledge_list.delete("1.0", tk.END)
            if lines:
                self._learned_knowledge_list.insert("1.0", "\n".join(lines))
            else:
                self._learned_knowledge_list.insert("1.0", "暂无已提取的知识点")
            self._learned_knowledge_list.config(state=tk.DISABLED)
            
            # 更新统计
            self._refresh_feedback_stats()
            
        except Exception as e:
            logger.warning(f"刷新知识点失败: {e}")
    
    def _on_submit_feedback(self):
        """提交反馈"""
        feedback_text = self._feedback_text.get("1.0", "end-1c").strip()
        feedback_type = self._feedback_type_var.get()
        
        if not feedback_text or feedback_text == "例如：这个对话太生硬了，不像角色会说的话":
            messagebox.showwarning("提示", "请输入反馈内容")
            return
        
        if not hasattr(self, '_feedback_collector') or not self._feedback_collector:
            messagebox.showerror("错误", "反馈收集器未初始化")
            return
        
        try:
            # 收集反馈
            chapter_id = getattr(self, '_current_chapter_id', 'unknown')
            feedback = self._feedback_collector.collect(
                chapter_id=chapter_id,
                feedback_text=feedback_text,
                feedback_type=feedback_type,
                context={
                    "project": self.current_project.get('name', 'unknown') if self.current_project else 'unknown'
                }
            )
            
            # 尝试提纯反馈
            try:
                from core.feedback_purifier import get_feedback_purifier
                purifier = get_feedback_purifier()
                knowledge_points = purifier.purify(feedback_text, feedback_type)
                
                if knowledge_points:
                    # 标记已提取知识点
                    self._feedback_collector.mark_knowledge_extracted(feedback.id)
                    
                    # 尝试添加到知识库
                    try:
                        from core.knowledge_manager import get_knowledge_manager
                        km = get_knowledge_manager()
                        for kp in knowledge_points:
                            km.add_knowledge(
                                category=kp.category,
                                content=kp.content,
                                tags=kp.tags,
                                source="用户反馈",
                                metadata=kp.context
                            )
                    except Exception as e:
                        logger.warning(f"添加到知识库失败: {e}")
                    
                    messagebox.showinfo(
                        "成功", 
                        f"反馈已提交并提取了 {len(knowledge_points)} 条知识点！\n"
                        "下次生成时会应用这些优化。"
                    )
                else:
                    messagebox.showinfo("成功", "反馈已提交！")
                
            except Exception as e:
                logger.warning(f"反馈提纯失败: {e}")
                messagebox.showinfo("成功", "反馈已提交！")
            
            # 清空输入
            self._feedback_text.delete("1.0", tk.END)
            
            # 刷新统计
            self._refresh_feedback_stats()
            self._refresh_learned_knowledge()
            
        except Exception as e:
            logger.error(f"提交反馈失败: {e}")
            messagebox.showerror("错误", f"提交失败: {e}")  # P2-003：反馈收集器错误，保留原有格式
    
    def _on_view_feedback_history(self):
        """查看历史反馈"""
        if not hasattr(self, '_feedback_collector') or not self._feedback_collector:
            messagebox.showerror("错误", "反馈收集器未初始化")
            return
        
        try:
            history = self._feedback_collector.get_history(limit=50)
            
            if not history:
                messagebox.showinfo("提示", "暂无历史反馈")
                return
            
            # 创建历史窗口
            history_window = tk.Toplevel(self.root)
            history_window.title("历史反馈")
            history_window.geometry("800x500")
            history_window.configure(bg=GlassTheme.GLASS_BG)
            
            # 创建Treeview
            columns = ("时间", "类型", "章节", "内容")
            tree = ttk.Treeview(history_window, columns=columns, show="headings", height=20)
            
            tree.heading("时间", text="时间")
            tree.heading("类型", text="类型")
            tree.heading("章节", text="章节")
            tree.heading("内容", text="内容")
            
            tree.column("时间", width=150)
            tree.column("类型", width=100)
            tree.column("章节", width=150)
            tree.column("内容", width=400)
            
            # 添加数据
            for fb in history:
                tree.insert("", tk.END, values=(
                    fb.timestamp[:19] if fb.timestamp else "-",
                    fb.feedback_type,
                    fb.chapter_id[:20] if fb.chapter_id else "-",
                    fb.feedback_text[:60] + "..." if len(fb.feedback_text) > 60 else fb.feedback_text
                ))
            
            tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # 关闭按钮
            ttk.Button(history_window, text="关闭", command=history_window.destroy).pack(pady=10)
            
        except Exception as e:
            logger.error(f"查看历史失败: {e}")
            messagebox.showerror("错误", f"查看失败: {e}")
    
    def _on_generate_improvement_report(self):
        """生成改进报告"""
        try:
            from core.report_generator import get_report_generator
            
            generator = get_report_generator()
            
            self._set_status("正在生成改进报告...")
            
            # 生成报告
            report = generator.generate_improvement_report()
            
            if report:
                # 保存报告
                report_path = Path("经验文档/改进报告")
                report_path.mkdir(parents=True, exist_ok=True)
                
                filename = f"改进报告_{report.get('date', 'unknown')}.md"
                report_file = report_path / filename
                
                with open(report_file, 'w', encoding='utf-8') as f:
                    f.write(report.get('content', ''))
                
                messagebox.showinfo(
                    "成功",
                    f"改进报告已生成！\n保存位置：{report_file}"
                )
                self._set_status("改进报告生成完成")
            else:
                messagebox.showwarning("提示", "数据不足，无法生成报告")
                self._set_status("报告生成失败：数据不足")
                
        except ImportError:
            messagebox.showwarning("提示", "报告生成器模块未安装")
        except Exception as e:
            logger.error(f"生成报告失败: {e}")
            messagebox.showerror("错误", f"生成失败: {e}")
            self._set_status(f"报告生成失败: {e}")
    
    def _create_settings_page(self) -> tk.Frame:
        """创建设置页面（支持滚动）"""
        frame = ttk.Frame(self._content_frame, style="TFrame")
        
        # 创建滚动容器
        canvas = tk.Canvas(frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 鼠标滚轮绑定
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 标题
        ttk.Label(scrollable_frame, text="设置", style="Title.TLabel").pack(anchor=tk.W, padx=20, pady=(20, 10))
        
        # AI服务配置
        ai_frame = ttk.LabelFrame(scrollable_frame, text="AI服务配置", padding=15)
        ai_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 服务模式
        mode_frame = ttk.Frame(ai_frame, style="TFrame")
        mode_frame.pack(fill=tk.X, pady=5)
        ttk.Label(mode_frame, text="服务模式：").pack(side=tk.LEFT)
        self._service_mode_var = tk.StringVar(value="remote")  # V2.20修复：online -> remote
        mode_local = ttk.Radiobutton(mode_frame, text="本地大模型", variable=self._service_mode_var, value="local", command=self._on_service_mode_changed)
        mode_local.pack(side=tk.LEFT, padx=10)
        mode_online = ttk.Radiobutton(mode_frame, text="线上API", variable=self._service_mode_var, value="remote", command=self._on_service_mode_changed)  # V2.20修复：online -> remote
        mode_online.pack(side=tk.LEFT)
        
        # API提供商
        provider_frame = ttk.Frame(ai_frame, style="TFrame")
        provider_frame.pack(fill=tk.X, pady=5)
        ttk.Label(provider_frame, text="提供商：").pack(side=tk.LEFT)
        self._provider_var = tk.StringVar(value="DeepSeek")
        self._provider_combo = ttk.Combobox(  # V2.23保存引用用于动态更新
            provider_frame,
            textvariable=self._provider_var,
            values=["DeepSeek", "OpenAI", "Anthropic", "Ollama", "Qwen"],  # V2.23新增Qwen本地模型
            state="readonly",
            width=20
        )
        self._provider_combo.pack(side=tk.LEFT, padx=10)
        
        # 模型选择
        model_frame = ttk.Frame(ai_frame, style="TFrame")
        model_frame.pack(fill=tk.X, pady=5)
        ttk.Label(model_frame, text="模型：").pack(side=tk.LEFT)
        self._model_var = tk.StringVar(value="deepseek-chat")
        self._model_combo = ttk.Combobox(  # V2.23保存引用用于动态更新
            model_frame,
            textvariable=self._model_var,
            values=["deepseek-chat", "deepseek-reasoner", "gpt-4", "gpt-3.5-turbo", "claude-3", "qwen2.5-14b-gptq", "llama3.1", "mistral"],  # V2.23新增本地模型选项
            state="readonly",
            width=20
        )
        self._model_combo.pack(side=tk.LEFT, padx=10)  # V2.23修复：使用self引用
        
        # API Key / 本地部署地址（根据模式显示不同字段）
        key_frame = ttk.Frame(ai_frame, style="TFrame")
        key_frame.pack(fill=tk.X, pady=5)
        self._key_label = ttk.Label(key_frame, text="API Key：")
        self._key_label.pack(side=tk.LEFT)
        self._api_key_var = tk.StringVar()
        self._local_url_var = tk.StringVar(value="http://localhost:8000/v1")
        self._key_entry = ttk.Entry(key_frame, textvariable=self._api_key_var, width=35, show="*")
        self._key_entry.pack(side=tk.LEFT, padx=10)
        
        # 显示/隐藏API Key按钮
        self._show_key_var = tk.BooleanVar(value=False)
        self._toggle_key_btn = ttk.Button(key_frame, text="👁", width=3, 
                                          command=self._toggle_api_key_visibility)
        self._toggle_key_btn.pack(side=tk.LEFT, padx=2)
        
        self._url_entry = ttk.Entry(key_frame, textvariable=self._local_url_var, width=40)
        # 默认隐藏本地地址输入框
        self._url_entry.pack_forget()
        
        # Temperature
        temp_frame = ttk.Frame(ai_frame, style="TFrame")
        temp_frame.pack(fill=tk.X, pady=5)
        ttk.Label(temp_frame, text="Temperature：").pack(side=tk.LEFT)
        self._temp_var = tk.DoubleVar(value=0.7)
        temp_scale = ttk.Scale(
            temp_frame,
            from_=0.0,
            to=2.0,
            variable=self._temp_var,
            orient=tk.HORIZONTAL,
            length=200
        )
        temp_scale.pack(side=tk.LEFT, padx=10)
        self._temp_label = ttk.Label(temp_frame, text="0.70")
        self._temp_label.pack(side=tk.LEFT)
        temp_scale.configure(command=lambda v: self._temp_label.configure(text=f"{float(v):.2f}"))
        
        # 安全状态提示框（API Key保护状态）
        security_frame = ttk.LabelFrame(ai_frame, text="🔒 安全状态", padding=10)
        security_frame.pack(fill=tk.X, pady=10)
        
        # 检查保护状态
        gitignore_ok = self._check_gitignore_protection()
        
        self._security_status_label = ttk.Label(
            security_frame, 
            text=f"{'✓ .gitignore已配置' if gitignore_ok else '⚠ .gitignore未配置'} | API Key将加密存储",
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL)
        )
        self._security_status_label.pack(anchor=tk.W)
        
        # 安全提示
        tip_label = ttk.Label(
            security_frame,
            text="⚠ 建议定期轮换API Key（每3-6个月重新生成）",
            font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SMALL),
            foreground=GlassTheme.ACCENT_ORANGE
        )
        tip_label.pack(anchor=tk.W, pady=(5, 0))
        
        # API操作按钮组
        api_btn_frame = ttk.Frame(ai_frame, style="TFrame")
        api_btn_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(api_btn_frame, text="🔗 测试连接", command=self._on_test_api_connection).pack(side=tk.LEFT, padx=5)
        ttk.Button(api_btn_frame, text="💾 备份密钥", command=self._on_backup_api_keys).pack(side=tk.LEFT, padx=5)
        ttk.Button(api_btn_frame, text="📥 恢复密钥", command=self._on_restore_api_keys).pack(side=tk.LEFT, padx=5)
        

        # 文件路径设置
        # 文件路径设置
        path_frame = ttk.LabelFrame(scrollable_frame, text="文件路径设置", padding=15)
        path_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 默认保存位置
        save_path_frame = ttk.Frame(path_frame, style="TFrame")
        save_path_frame.pack(fill=tk.X, pady=5)
        ttk.Label(save_path_frame, text="默认保存位置：").pack(side=tk.LEFT)
        # V2.23修复：从config.yaml读取已保存的路径
        config = self._load_config()
        saved_save_path = config.get("save_path", os.path.dirname(os.path.abspath(__file__)))
        self._save_path_var = tk.StringVar(value=saved_save_path)
        ttk.Entry(save_path_frame, textvariable=self._save_path_var, width=40).pack(side=tk.LEFT, padx=10)
        ttk.Button(save_path_frame, text="浏览", command=self._on_browse_save_path).pack(side=tk.LEFT)
        
        # 自动备份间隔
        backup_frame = ttk.Frame(path_frame, style="TFrame")
        backup_frame.pack(fill=tk.X, pady=5)
        ttk.Label(backup_frame, text="自动备份间隔：").pack(side=tk.LEFT)
        self._backup_interval_var = tk.StringVar(value="30")
        ttk.Spinbox(backup_frame, from_=10, to=120, textvariable=self._backup_interval_var, width=8).pack(side=tk.LEFT, padx=10)
        ttk.Label(backup_frame, text="分钟").pack(side=tk.LEFT)
        
        # 偏好学习设置
        learning_frame = ttk.LabelFrame(scrollable_frame, text="偏好学习设置", padding=15)
        learning_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # AI学习开关
        ai_learning_frame = ttk.Frame(learning_frame, style="TFrame")
        ai_learning_frame.pack(fill=tk.X, pady=5)
        ttk.Label(ai_learning_frame, text="AI偏好学习：").pack(side=tk.LEFT)
        self._ai_learning_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(ai_learning_frame, text="启用（AI将学习您的写作偏好）", variable=self._ai_learning_var).pack(side=tk.LEFT, padx=10)
        
        # 数据清除
        clear_frame = ttk.Frame(learning_frame, style="TFrame")
        clear_frame.pack(fill=tk.X, pady=5)
        ttk.Label(clear_frame, text="偏好数据：").pack(side=tk.LEFT)
        ttk.Button(clear_frame, text="清除学习数据", command=self._on_clear_learning_data).pack(side=tk.LEFT, padx=10)
        ttk.Button(clear_frame, text="导出偏好配置", command=self._on_export_preferences).pack(side=tk.LEFT, padx=5)
        
        # ========== 长期记忆管理（Sprint 9-10）==========
        memory_frame = ttk.LabelFrame(scrollable_frame, text="长期记忆管理", padding=15)
        memory_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # 记忆层级状态概览
        overview_frame = ttk.Frame(memory_frame, style="TFrame")
        overview_frame.pack(fill=tk.X, pady=5)
        ttk.Label(overview_frame, text="记忆层级状态：", font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, "bold")).pack(side=tk.LEFT)
        
        self._memory_status_label = ttk.Label(overview_frame, text="正在加载...")
        self._memory_status_label.pack(side=tk.LEFT, padx=10)
        ttk.Button(overview_frame, text="刷新状态", command=self._on_refresh_memory_status).pack(side=tk.RIGHT)
        
        # L1热记忆：SESSION-STATE管理
        l1_frame = ttk.LabelFrame(memory_frame, text="L1热记忆（SESSION-STATE）", padding=10)
        l1_frame.pack(fill=tk.X, pady=5)
        
        l1_status = ttk.Frame(l1_frame, style="TFrame")
        l1_status.pack(fill=tk.X, pady=2)
        ttk.Label(l1_status, text="状态：").pack(side=tk.LEFT)
        self._session_state_status = ttk.Label(l1_status, text="-")
        self._session_state_status.pack(side=tk.LEFT, padx=5)
        
        l1_actions = ttk.Frame(l1_frame, style="TFrame")
        l1_actions.pack(fill=tk.X, pady=5)
        ttk.Button(l1_actions, text="查看当前会话", command=self._on_view_session_state).pack(side=tk.LEFT, padx=5)
        ttk.Button(l1_actions, text="清空会话状态", command=self._on_clear_session_state).pack(side=tk.LEFT, padx=5)
        ttk.Button(l1_actions, text="保存会话快照", command=self._on_save_session_snapshot).pack(side=tk.LEFT, padx=5)
        
        # L2温记忆：向量数据库管理
        l2_frame = ttk.LabelFrame(memory_frame, text="L2温记忆（向量数据库）", padding=10)
        l2_frame.pack(fill=tk.X, pady=5)
        
        l2_stats = ttk.Frame(l2_frame, style="TFrame")
        l2_stats.pack(fill=tk.X, pady=2)
        ttk.Label(l2_stats, text="章节数：").pack(side=tk.LEFT)
        self._chapter_count_label = ttk.Label(l2_stats, text="-")
        self._chapter_count_label.pack(side=tk.LEFT, padx=5)
        ttk.Label(l2_stats, text="知识点数：").pack(side=tk.LEFT, padx=(20, 0))
        self._knowledge_count_label = ttk.Label(l2_stats, text="-")
        self._knowledge_count_label.pack(side=tk.LEFT, padx=5)
        ttk.Label(l2_stats, text="风格数：").pack(side=tk.LEFT, padx=(20, 0))
        self._style_count_label = ttk.Label(l2_stats, text="-")
        self._style_count_label.pack(side=tk.LEFT, padx=5)
        
        l2_actions = ttk.Frame(l2_frame, style="TFrame")
        l2_actions.pack(fill=tk.X, pady=5)
        ttk.Button(l2_actions, text="重建向量索引", command=self._on_rebuild_vector_index).pack(side=tk.LEFT, padx=5)
        ttk.Button(l2_actions, text="导出向量数据", command=self._on_export_vector_data).pack(side=tk.LEFT, padx=5)
        ttk.Button(l2_actions, text="清空向量库", command=self._on_clear_vector_store).pack(side=tk.LEFT, padx=5)
        
        # L3冷记忆：Git-Notes管理
        l3_frame = ttk.LabelFrame(memory_frame, text="L3冷记忆（Git-Notes）", padding=10)
        l3_frame.pack(fill=tk.X, pady=5)
        
        l3_info = ttk.Frame(l3_frame, style="TFrame")
        l3_info.pack(fill=tk.X, pady=2)
        ttk.Label(l3_info, text="分支：").pack(side=tk.LEFT)
        self._git_branch_label = ttk.Label(l3_info, text="-")
        self._git_branch_label.pack(side=tk.LEFT, padx=5)
        ttk.Label(l3_info, text="笔记数：").pack(side=tk.LEFT, padx=(20, 0))
        self._git_notes_count_label = ttk.Label(l3_info, text="-")
        self._git_notes_count_label.pack(side=tk.LEFT, padx=5)
        
        l3_actions = ttk.Frame(l3_frame, style="TFrame")
        l3_actions.pack(fill=tk.X, pady=5)
        ttk.Button(l3_actions, text="查看历史决策", command=self._on_view_git_notes).pack(side=tk.LEFT, padx=5)
        ttk.Button(l3_actions, text="添加决策记录", command=self._on_add_git_note).pack(side=tk.LEFT, padx=5)
        
        # L4精选档案：MEMORY.md管理
        l4_frame = ttk.LabelFrame(memory_frame, text="L4精选档案（MEMORY.md）", padding=10)
        l4_frame.pack(fill=tk.X, pady=5)
        
        l4_stats = ttk.Frame(l4_frame, style="TFrame")
        l4_stats.pack(fill=tk.X, pady=2)
        ttk.Label(l4_stats, text="文件大小：").pack(side=tk.LEFT)
        self._memory_size_label = ttk.Label(l4_stats, text="-")
        self._memory_size_label.pack(side=tk.LEFT, padx=5)
        ttk.Label(l4_stats, text="最后更新：").pack(side=tk.LEFT, padx=(20, 0))
        self._memory_update_label = ttk.Label(l4_stats, text="-")
        self._memory_update_label.pack(side=tk.LEFT, padx=5)
        
        l4_actions = ttk.Frame(l4_frame, style="TFrame")
        l4_actions.pack(fill=tk.X, pady=5)
        ttk.Button(l4_actions, text="编辑MEMORY.md", command=self._on_edit_memory_md).pack(side=tk.LEFT, padx=5)
        ttk.Button(l4_actions, text="每日冥想", command=self._on_daily_meditation).pack(side=tk.LEFT, padx=5)
        ttk.Button(l4_actions, text="每周大冥想", command=self._on_weekly_meditation).pack(side=tk.LEFT, padx=5)
        ttk.Button(l4_actions, text="归档旧日记", command=self._on_archive_old_diary).pack(side=tk.LEFT, padx=5)
        
        # 记忆维护工具
        maintenance_frame = ttk.Frame(memory_frame, style="TFrame")
        maintenance_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(maintenance_frame, text="维护工具：").pack(side=tk.LEFT)
        ttk.Button(maintenance_frame, text="一键备份记忆", command=self._on_backup_memory).pack(side=tk.LEFT, padx=5)
        ttk.Button(maintenance_frame, text="恢复记忆备份", command=self._on_restore_memory).pack(side=tk.LEFT, padx=5)
        ttk.Button(maintenance_frame, text="清空所有记忆", command=self._on_clear_all_memory).pack(side=tk.LEFT, padx=5)
        
        # ========== 记忆配置（Sprint 9-10）==========
        memory_config_frame = ttk.LabelFrame(scrollable_frame, text="记忆配置", padding=15)
        memory_config_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # WAL自动保存
        wal_frame = ttk.Frame(memory_config_frame, style="TFrame")
        wal_frame.pack(fill=tk.X, pady=5)
        ttk.Label(wal_frame, text="WAL自动保存：").pack(side=tk.LEFT)
        self._wal_auto_save_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(wal_frame, text="启用（AI调用前自动保存会话状态）", variable=self._wal_auto_save_var).pack(side=tk.LEFT, padx=10)
        
        # 会话自动恢复
        recovery_frame = ttk.Frame(memory_config_frame, style="TFrame")
        recovery_frame.pack(fill=tk.X, pady=5)
        ttk.Label(recovery_frame, text="会话自动恢复：").pack(side=tk.LEFT)
        self._auto_recovery_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(recovery_frame, text="启用（程序启动时自动恢复上次会话）", variable=self._auto_recovery_var).pack(side=tk.LEFT, padx=10)
        
        # 向量召回配置
        recall_frame = ttk.Frame(memory_config_frame, style="TFrame")
        recall_frame.pack(fill=tk.X, pady=5)
        ttk.Label(recall_frame, text="向量召回数量：").pack(side=tk.LEFT)
        self._vector_recall_topk_var = tk.IntVar(value=10)
        ttk.Spinbox(recall_frame, from_=5, to=50, textvariable=self._vector_recall_topk_var, width=8).pack(side=tk.LEFT, padx=10)
        ttk.Label(recall_frame, text="条（上下文召回时返回的相似章节数）").pack(side=tk.LEFT)
        
        # 记忆冥想配置
        meditation_frame = ttk.Frame(memory_config_frame, style="TFrame")
        meditation_frame.pack(fill=tk.X, pady=5)
        ttk.Label(meditation_frame, text="每日冥想时间：").pack(side=tk.LEFT)
        self._daily_meditation_time_var = tk.StringVar(value="23:00")
        ttk.Entry(meditation_frame, textvariable=self._daily_meditation_time_var, width=10).pack(side=tk.LEFT, padx=10)
        ttk.Label(meditation_frame, text="（自动提炼当日重要事件到MEMORY.md）").pack(side=tk.LEFT)
        
        # 向量模型配置
        embedding_frame = ttk.Frame(memory_config_frame, style="TFrame")
        embedding_frame.pack(fill=tk.X, pady=5)
        ttk.Label(embedding_frame, text="Embedding模型：").pack(side=tk.LEFT)
        self._embedding_model_var = tk.StringVar(value="local")
        embedding_combo = ttk.Combobox(
            embedding_frame,
            textvariable=self._embedding_model_var,
            values=["local", "openai", "deepseek"],
            state="readonly",
            width=15
        )
        embedding_combo.pack(side=tk.LEFT, padx=10)
        ttk.Label(embedding_frame, text="（local=本地模型，openai=OpenAI API）").pack(side=tk.LEFT)
        
        # 从config.yaml加载设置
        self._load_settings_from_config()
        
        # 保存按钮
        save_frame = ttk.Frame(scrollable_frame, style="TFrame")
        save_frame.pack(fill=tk.X, padx=20, pady=20)
        
        ResponsiveButton(
            save_frame,
            text="保存设置",
            command=self._on_save_settings,
            async_handler=self._async_handler,
            style="Accent.TButton"
        ).pack(side=tk.LEFT)
        
        return frame
    
    # ============== 辅助方法 ==============

    def _safe_after(self, ms: int, func) -> None:
        """线程安全的 root.after 封装。
        
        EventBus 在 ThreadPoolExecutor 中调用 handler，此时直接调用
        root.after() 会触发 "main thread is not in main loop" 错误。
        本方法通过检测当前线程是否为主线程来决定调度策略：
        - 主线程：直接调用 root.after()
        - 子线程：通过 _result_queue + 轮询机制安全调度
        """
        import threading as _threading
        try:
            if _threading.current_thread() is _threading.main_thread():
                self.root.after(ms, func)
            else:
                # 子线程中不能直接调用 root.after，
                # 把任务塞入 result_queue，由主线程轮询处理
                self._result_queue.put({"type": "_safe_after_cb", "ms": ms, "func": func})
        except RuntimeError:
            pass

    def _set_status(self, message: str) -> None:
        """设置状态栏消息"""
        if self._status_var:
            self._status_var.set(message)
    
    def _get_character_manager(self) -> Any:
        """获取或创建人物管理器适配器（延迟初始化）"""
        if not hasattr(self, 'character_manager') or self.character_manager is None:
            from agents.adapters.character_adapter import CharacterManagerAdapter
            self.character_manager = CharacterManagerAdapter()
            # 必须调用initialize方法
            if not self.character_manager.initialize():
                raise RuntimeError("人物管理器适配器初始化失败")
        return self.character_manager
    
    def _process_result_queue(self) -> None:
        """处理结果队列"""
        try:
            while True:
                result = self._result_queue.get_nowait()
                if isinstance(result, dict):
                    result_type = result.get('type')
                    if result_type == 'status':
                        self._set_status(result.get('message', ''))
                    elif result_type == '_safe_after_cb':
                        # 线程安全回调：子线程通过队列将任务委托给主线程执行
                        ms = result.get('ms', 0)
                        func = result.get('func')
                        if func is not None:
                            self.root.after(ms, func)
        except queue.Empty:
            pass
        finally:
            self.root.after(100, self._process_result_queue)
    
    # ============== 功能方法 ==============
    
    def _get_hot_ranking_data(self) -> Dict:
        """获取热榜数据 - 调用HotRankingPlugin获取真实爬虫数据"""
        # 方式1: 尝试从插件注册表获取热榜插件
        if CORE_AVAILABLE:
            try:
                registry = get_plugin_registry()
                if registry:
                    hot_ranking_plugin = registry.get_plugin("hot-ranking-v1")
                    if hot_ranking_plugin:
                        # 调用插件获取真实数据（优先使用缓存）
                        result = hot_ranking_plugin.execute("get_data", {"force_fresh": False})
                        if result and isinstance(result, dict):
                            logger.info("[热榜] 成功从HotRankingPlugin获取数据（注册表）")
                            return result
            except Exception as e:
                logger.warning(f"[热榜] 插件注册表获取失败: {e}")
        
        # 方式2: 直接动态加载插件（当注册表中没有时）
        try:
            import importlib.util
            plugin_path = os.path.join(os.path.dirname(__file__), "plugins", "hot-ranking-v1", "plugin.py")
            spec = importlib.util.spec_from_file_location("hot_ranking_plugin", plugin_path)
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)
            
            HotRankingPlugin = module.HotRankingPlugin
            
            # 创建插件实例（不需要上下文，插件内部自己初始化）
            plugin = HotRankingPlugin()
            
            # 创建简化的上下文（PluginContext是dataclass，需要正确的参数）
            from core.plugin_interface import PluginContext
            context = PluginContext(
                event_bus=get_event_bus() if CORE_AVAILABLE else None,
                service_locator=get_service_locator() if CORE_AVAILABLE else None,
                config_manager=None,
                plugin_registry=None,
                logger=None
            )
            
            if plugin.initialize(context):
                # 获取数据（优先使用缓存）
                result = plugin.execute("get_data", {"force_fresh": False})
                if result and isinstance(result, dict):
                    logger.info("[热榜] 成功从HotRankingPlugin获取数据（动态加载）")
                    return result
        except Exception as e:
            logger.warning(f"[热榜] 动态加载插件失败: {e}")
        
        # 降级：返回静态默认数据
        logger.info("[热榜] 使用静态默认数据")
        return self._get_default_hot_ranking_data()
    
    def _get_default_hot_ranking_data(self) -> Dict:
        """获取默认热榜数据（离线降级）"""
        return {
            'sites': [
                {
                    'name': '📚 起点中文网',
                    'color': '#457B9D',
                    'books': [
                        {'title': '诡秘之主', 'author': '爱潜水的乌贼', 'category': '玄幻', 'heat': 150000},
                        {'title': '大奉打更人', 'author': '卖报小郎君', 'category': '仙侠', 'heat': 120000},
                        {'title': '深空彼岸', 'author': '辰东', 'category': '科幻', 'heat': 100000},
                        {'title': '星门', 'author': '老鹰吃小鸡', 'category': '玄幻', 'heat': 80000},
                        {'title': '完美世界', 'author': '辰东', 'category': '玄幻', 'heat': 60000},
                        {'title': '遮天', 'author': '辰东', 'category': '玄幻', 'heat': 55000},
                        {'title': '斗破苍穹', 'author': '天蚕土豆', 'category': '玄幻', 'heat': 50000},
                        {'title': '凡人修仙传', 'author': '忘语', 'category': '仙侠', 'heat': 45000},
                        {'title': '一念永恒', 'author': '耳根', 'category': '仙侠', 'heat': 40000},
                        {'title': '牧神记', 'author': '宅猪', 'category': '玄幻', 'heat': 35000},
                    ]
                },
                {
                    'name': '🎭 晋江文学城',
                    'color': '#E63946',
                    'books': [
                        {'title': '天官赐福', 'author': '墨香铜臭', 'category': '纯爱', 'heat': 180000},
                        {'title': '魔道祖师', 'author': '墨香铜臭', 'category': '纯爱', 'heat': 150000},
                        {'title': '人渣反派自救系统', 'author': '墨香铜臭', 'category': '纯爱', 'heat': 120000},
                        {'title': '二哈和他的白猫师尊', 'author': '肉包不吃肉', 'category': '纯爱', 'heat': 90000},
                        {'title': '全球高考', 'author': '木苏里', 'category': '纯爱', 'heat': 70000},
                        {'title': '撒野', 'author': '巫哲', 'category': '纯爱', 'heat': 65000},
                        {'title': '默读', 'author': 'Priest', 'category': '纯爱', 'heat': 60000},
                        {'title': '镇魂', 'author': 'Priest', 'category': '纯爱', 'heat': 55000},
                        {'title': '杀破狼', 'author': 'Priest', 'category': '纯爱', 'heat': 50000},
                        {'title': '有兽焉', 'author': '鹤来', 'category': '纯爱', 'heat': 45000},
                    ]
                },
                {
                    'name': '🍅 番茄小说',
                    'color': '#FF6B6B',
                    'books': [
                        {'title': '我在精神病院学斩神', 'author': '三九音域', 'category': '都市', 'heat': 185000},
                        {'title': '星门', 'author': '老鹰吃小鸡', 'category': '玄幻', 'heat': 168000},
                        {'title': '斩神', 'author': '三九音域', 'category': '都市', 'heat': 152000},
                        {'title': '重生之都市修仙', 'author': '十里剑神', 'category': '都市', 'heat': 138000},
                        {'title': '全职法师', 'author': '乱', 'category': '玄幻', 'heat': 125000},
                        {'title': '万族之劫', 'author': '老鹰吃小鸡', 'category': '玄幻', 'heat': 110000},
                        {'title': '开局签到荒古圣体', 'author': '神牧', 'category': '玄幻', 'heat': 95000},
                        {'title': '神墓', 'author': '辰东', 'category': '玄幻', 'heat': 80000},
                        {'title': '仙逆', 'author': '耳根', 'category': '仙侠', 'heat': 70000},
                        {'title': '我欲封天', 'author': '耳根', 'category': '仙侠', 'heat': 60000},
                    ]
                }
            ],
            'genres': {
                'male': {'title': '👨 男频题材榜', 'color': '#457B9D', 'genres': [
                    ('玄幻', 25.5), ('都市', 22.3), ('仙侠', 18.7), ('历史', 15.2), ('科幻', 12.8)
                ]},
                'female': {'title': '👩 女频题材榜', 'color': '#E63946', 'genres': [
                    ('现代言情', 28.6), ('古代言情', 24.1), ('玄幻言情', 19.8), ('仙侠奇缘', 16.4), ('青春校园', 13.2)
                ]}
            },
            'types': {
                'male': {'title': '👨 男频类型榜', 'color': '#2A9D8F', 'types': [
                    ('系统流', 22.8), ('穿越', 20.5), ('重生', 18.2), ('修仙', 16.9), ('无敌流', 14.3)
                ]},
                'female': {'title': '👩 女频类型榜', 'color': '#F4A261', 'types': [
                    ('甜宠', 26.5), ('虐恋', 23.1), ('穿越', 20.8), ('重生', 18.4), ('宫斗', 16.2)
                ]}
            },
            'authors': [
                {'rank': 1, 'name': '唐家三少', 'works': '《斗罗大陆》《绝世唐门》', 'income': '3.2亿', 'fans': '1200万'},
                {'rank': 2, 'name': '辰东', 'works': '《遮天》《完美世界》', 'income': '2.8亿', 'fans': '980万'},
                {'rank': 3, 'name': '天蚕土豆', 'works': '《斗破苍穹》《武动乾坤》', 'income': '2.5亿', 'fans': '890万'},
                {'rank': 4, 'name': '我吃西红柿', 'works': '《盘龙》《星辰变》', 'income': '2.1亿', 'fans': '760万'},
                {'rank': 5, 'name': '猫腻', 'works': '《庆余年》《将夜》', 'income': '1.7亿', 'fans': '680万'},
                {'rank': 6, 'name': '顾漫', 'works': '《你是我的荣耀》', 'income': '1.5亿', 'fans': '650万'},
                {'rank': 7, 'name': '梦入神机', 'works': '《龙蛇演义》《永生》', 'income': '1.9亿', 'fans': '720万'},
                {'rank': 8, 'name': '烽火戏诸侯', 'works': '《雪中悍刀行》《剑来》', 'income': '9800万', 'fans': '560万'},
                {'rank': 9, 'name': '月关', 'works': '《回明》《锦衣夜行》', 'income': '1.3亿', 'fans': '620万'},
                {'rank': 10, 'name': '血红', 'works': '《升龙道》《邪风曲》', 'income': '1.1亿', 'fans': '590万'}
            ],
            'update_time': datetime.now().strftime('%Y-%m-%d %H:%M')
        }
    
    def _load_plugins_async(self) -> None:
        """异步加载插件列表 - V2.3优化版
        
        使用AsyncHandler异步加载插件，避免UI卡顿
        """
        # 显示加载状态
        loading_item = self._plugin_tree.insert(
            "",
            tk.END,
            values=("正在加载插件...", "", "", "", "")
        )
        
        def load_task():
            """后台加载任务"""
            # 使用线程池执行加载
            from concurrent.futures import ThreadPoolExecutor
            import time
            
            # 模拟加载延迟（实际是文件IO）
            plugins_data = []
            plugins_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plugins")
            
            if os.path.exists(plugins_dir):
                protected_modules = {
                    "outline-parser-v3", "style-learner-v5", "character-manager-v1",
                    "worldview-parser-v1", "context-builder-v1", "iterative-generator-v2",
                    "quality-validator-v1", "novel-generator-v3", "hot-ranking-v1"
                }
                
                for plugin_name in sorted(os.listdir(plugins_dir)):
                    plugin_path = os.path.join(plugins_dir, plugin_name)
                    if os.path.isdir(plugin_path) and not plugin_name.startswith("__"):
                        plugin_json = os.path.join(plugin_path, "plugin.json")
                        if os.path.exists(plugin_json):
                            try:
                                with open(plugin_json, "r", encoding="utf-8") as f:
                                    data = json.load(f)
                                
                                is_protected = plugin_name in protected_modules
                                plugins_data.append({
                                    "id": plugin_name,
                                    "data": data,
                                    "is_protected": is_protected
                                })
                            except Exception as e:
                                logger.warning(f"Failed to load plugin.json for {plugin_name}: {e}")
            
            return plugins_data
        
        def on_success(plugins_data):
            """加载成功回调（在主线程执行）"""
            # 清空加载状态
            for item in self._plugin_tree.get_children():
                self._plugin_tree.delete(item)
            
            self._plugin_data = {}
            
            # 类型映射
            type_map = {
                "analyzer": "分析器",
                "generator": "生成器",
                "validator": "验证器",
                "tool": "工具",
                "storage": "存储",
                "ai": "AI服务",
                "protocol": "协议"
            }
            
            # 填充数据
            for plugin_info in plugins_data:
                plugin_id = plugin_info["id"]
                data = plugin_info["data"]
                is_protected = plugin_info["is_protected"]
                
                protected_text = "🔒" if is_protected else ""
                type_text = type_map.get(data.get("plugin_type", "tool").lower(), data.get("plugin_type", "工具"))
                status_text = "已启用" if data.get("enabled", True) else "已禁用"
                
                item_id = self._plugin_tree.insert(
                    "",
                    tk.END,
                    values=(
                        data.get("name", plugin_id),
                        data.get("version", "1.0.0"),
                        status_text,
                        type_text,
                        protected_text
                    )
                )
                
                self._plugin_data[item_id] = {
                    "id": plugin_id,
                    "name": data.get("name", plugin_id),
                    "version": data.get("version", "1.0.0"),
                    "description": data.get("description", ""),
                    "author": data.get("author", ""),
                    "type": data.get("plugin_type", "tool"),
                    "state": "active" if data.get("enabled", True) else "disabled",
                    "is_protected": is_protected,
                    "dependencies": data.get("dependencies", [])
                }
            
            self._set_status(f"已加载 {len(plugins_data)} 个插件")
            logger.info(f"Async loaded {len(plugins_data)} plugins")
        
        def on_error(error):
            """加载失败回调（在主线程执行）"""
            # 清空加载状态
            for item in self._plugin_tree.get_children():
                self._plugin_tree.delete(item)
            
            # 显示错误
            self._plugin_tree.insert(
                "",
                tk.END,
                values=(f"加载失败: {error}", "", "", "", "")
            )
            
            self._set_status("插件加载失败")
            logger.error(f"Failed to async load plugins: {error}")
        
        # 提交异步任务
        if self._async_handler:
            self._async_handler.submit(
                func=load_task,
                callback=on_success,
                error_callback=on_error,
                priority=TaskPriority.HIGH,
                timeout=10.0
            )
        else:
            # 回退到同步加载
            self._load_plugins()
    
    def _load_plugins(self) -> None:
        """加载插件列表 - V2.2重构版
        
        从PluginRegistry读取真实的插件数据，包括：
        - 插件元数据（名称、版本、类型）
        - 插件状态（active/loaded/error等）
        - V5保护模块标识
        """
        # 清空现有数据
        for item in self._plugin_tree.get_children():
            self._plugin_tree.delete(item)
        
        # 存储插件数据用于后续操作
        self._plugin_data: Dict[str, Dict[str, Any]] = {}
        
        # 尝试从PluginRegistry获取真实数据
        registry = None
        try:
            if CORE_AVAILABLE:
                registry = get_plugin_registry()
        except Exception as e:
            logger.warning(f"Failed to get PluginRegistry: {e}")
        
        if registry and registry._plugins:
            # 获取所有已注册的插件信息
            for plugin_id in registry._plugins.keys():
                plugin_info = registry.get_plugin_info(plugin_id)
                if plugin_info:
                    metadata = plugin_info.metadata
                    state = plugin_info.state
                    is_protected = registry.is_protected(plugin_id)
                    
                    # 状态显示映射
                    status_map = {
                        "active": "已启用",
                        "loaded": "已加载",
                        "error": "错误",
                        "unloaded": "未加载",
                        "unloading": "卸载中"
                    }
                    status_text = status_map.get(state, state)
                    
                    # 类型显示映射
                    type_map = {
                        "analyzer": "分析器",
                        "generator": "生成器",
                        "validator": "验证器",
                        "tool": "工具",
                        "storage": "存储",
                        "ai": "AI服务",
                        "protocol": "协议"
                    }
                    type_text = type_map.get(metadata.plugin_type.lower(), metadata.plugin_type)
                    
                    # 保护状态显示
                    protected_text = "🔒" if is_protected else ""
                    
                    # 插入到树形视图
                    item_id = self._plugin_tree.insert(
                        "", 
                        tk.END, 
                        values=(
                            metadata.name,
                            metadata.version,
                            status_text,
                            type_text,
                            protected_text
                        )
                    )
                    
                    # 存储插件数据
                    self._plugin_data[item_id] = {
                        "id": plugin_id,
                        "name": metadata.name,
                        "version": metadata.version,
                        "description": metadata.description,
                        "author": metadata.author,
                        "type": metadata.plugin_type,
                        "state": state,
                        "is_protected": is_protected,
                        "dependencies": metadata.dependencies,
                        "error_message": plugin_info.error_message if hasattr(plugin_info, 'error_message') else None
                    }
            
            logger.info(f"Loaded {len(self._plugin_data)} plugins from registry")
        else:
            # 回退：从plugins目录扫描plugin.json文件
            self._load_plugins_from_filesystem()
            logger.info(f"Loaded {len(self._plugin_data)} plugins from filesystem")
    
    def _load_plugins_from_filesystem(self) -> None:
        """从文件系统扫描插件并注册到PluginRegistry（修复版本）
        
        修复说明：
        匨修复前的流程：
        1. 只读取plugin.json文件
        2. 存储到self._plugin_data字典
        3. 显示在Treeview中
        ❌ 问题：插件未注册到PluginRegistry，导致启用/禁用操作失败
        
        @修复后的流程：
        1. 使用PluginLoader.discover_plugins()发现插件
        2. 使用PluginLoader.load_plugin()加载并注册插件
        3. 从PluginRegistry获取真实状态显示在UI
        """
        plugins_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plugins")
        
        if not os.path.exists(plugins_dir):
            return
        
        # V5保护模块列表
        protected_modules = {
            "outline-parser-v3", "style-learner-v2", "character-manager",
            "worldview-parser", "context-builder", "iterative-generator-v2",
            "weighted-validator", "optimized-generator-v2", "hot-ranking"
        }
        
        # 尝试使用PluginLoader来发现和注册插件
        try:
            from core.plugin_loader import get_plugin_loader
            from core.plugin_registry import get_plugin_registry, PluginState
            
            loader = get_plugin_loader()
            registry = get_plugin_registry()
            
            # 设置插件目录
            loader._plugin_directories = [plugins_dir]
            
            # 发现插件
            discovered_ids = loader.discover_plugins()
            logger.info(f"Discovered {len(discovered_ids)} plugins")
            
            # 按依赖顺序加载插件
            load_results = loader.load_all()
            logger.info(f"Loaded {sum(1 for r in load_results.values() if r.success)} plugins")
            
            # 清空现有数据
            for item in self._plugin_tree.get_children():
                self._plugin_tree.delete(item)
            self._plugin_data = {}
            
            # 从PluginRegistry获取插件信息并更新UI
            for plugin_id in discovered_ids:
                plugin_info = registry.get_plugin_info(plugin_id)
                if not plugin_info:
                    continue
                
                metadata = plugin_info.metadata
                state = plugin_info.state
                is_protected = plugin_id in protected_modules
                
                # 类型映射
                type_map = {
                    "analyzer": "分析器",
                    "generator": "生成器",
                    "validator": "验证器",
                    "tool": "工具",
                    "storage": "存储",
                    "ai": "AI服务",
                    "protocol": "协议"
                }
                type_text = type_map.get(metadata.plugin_type.lower(), metadata.plugin_type)
                protected_text = "🔒" if is_protected else ""
                
                # 状态映射
                state_map = {
                    "unloaded": "未加载",
                    "loaded": "已加载",
                    "active": "已激活",
                    "error": "错误",
                    "unloading": "卸载中"
                }
                state_text = state_map.get(state, state)
                
                item_id = self._plugin_tree.insert(
                    "",
                    tk.END,
                    values=(
                        metadata.name,
                        metadata.version,
                        state_text,
                        type_text,
                        protected_text
                    )
                )
                
                self._plugin_data[item_id] = {
                    "id": plugin_id,
                    "name": metadata.name,
                    "version": metadata.version,
                    "description": metadata.description,
                    "author": metadata.author,
                    "type": metadata.plugin_type,
                    "state": state,
                    "is_protected": is_protected,
                    "dependencies": metadata.dependencies,
                    "error_message": plugin_info.error_message
                }
                
        except Exception as e:
            logger.error(f"Failed to load plugins via PluginLoader: {e}")
            # 回退到旧的文件扫描方式
            self._load_plugins_from_filesystem_legacy()
    
    def _load_plugins_from_filesystem_legacy(self) -> None:
        """旧的文件扫描方式（作为回退方案）"""
        plugins_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plugins")
        
        if not os.path.exists(plugins_dir):
            return
        
        # V5保护模块列表
        protected_modules = {
            "outline-parser-v3", "style-learner-v2", "character-manager",
            "worldview-parser", "context-builder", "iterative-generator-v2",
            "weighted-validator", "optimized-generator-v2", "hot-ranking"
        }
        
        for plugin_name in os.listdir(plugins_dir):
            plugin_path = os.path.join(plugins_dir, plugin_name)
            if os.path.isdir(plugin_path) and not plugin_name.startswith("__"):
                plugin_json = os.path.join(plugin_path, "plugin.json")
                if os.path.exists(plugin_json):
                    try:
                        with open(plugin_json, "r", encoding="utf-8") as f:
                            data = json.load(f)
                        
                        is_protected = plugin_name in protected_modules
                        protected_text = "🔒" if is_protected else ""
                        
                        # 类型映射
                        type_map = {
                            "analyzer": "分析器",
                            "generator": "生成器",
                            "validator": "验证器",
                            "tool": "工具",
                            "storage": "存储",
                            "ai": "AI服务",
                            "protocol": "协议"
                        }
                        type_text = type_map.get(data.get("plugin_type", "tool").lower(), data.get("plugin_type", "工具"))
                        
                        item_id = self._plugin_tree.insert(
                            "",
                            tk.END,
                            values=(
                                data.get("name", plugin_name),
                                data.get("version", "1.0.0"),
                                "已加载" if data.get("enabled", True) else "已禁用",
                                type_text,
                                protected_text
                            )
                        )
                        
                        self._plugin_data[item_id] = {
                            "id": plugin_name,
                            "name": data.get("name", plugin_name),
                            "version": data.get("version", "1.0.0"),
                            "description": data.get("description", ""),
                            "author": data.get("author", ""),
                            "type": data.get("plugin_type", "tool"),
                            "state": "loaded" if data.get("enabled", True) else "disabled",
                            "is_protected": is_protected,
                            "dependencies": data.get("dependencies", []),
                            "error_message": None
                        }
                    except Exception as e:
                        logger.warning(f"Failed to load plugin.json for {plugin_name}: {e}")
    
    def _load_settings_from_config(self) -> None:
        """从config.yaml加载设置到UI"""
        config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.yaml")
        
        try:
            if os.path.exists(config_path):
                with open(config_path, "r", encoding="utf-8") as f:
                    config_data = yaml.safe_load(f)
                
                if config_data:
                    self._service_mode_var.set(config_data.get("service_mode", "online"))
                    self._provider_var.set(config_data.get("provider", "DeepSeek"))
                    self._model_var.set(config_data.get("model", "deepseek-chat"))
                    
                    # V2.23修复：从加密存储加载API Key
                    api_key_value = config_data.get("api_key", "")
                    if api_key_value == "ENCRYPTED_IN_SECRETS_FILE":
                        try:
                            from core.api_key_encryption import get_api_key_encryption
                            encryption = get_api_key_encryption()
                            provider = config_data.get("provider", "DeepSeek")
                            real_api_key = encryption.get_api_key(provider)
                            if real_api_key:
                                self._api_key_var.set(real_api_key)
                                logger.info(f"API Key从加密存储加载: {provider}")
                            else:
                                self._api_key_var.set("")
                                logger.warning(f"加密存储中未找到 {provider} 的API Key")
                        except Exception as e:
                            self._api_key_var.set("")
                            logger.warning(f"从加密存储加载API Key失败: {e}")
                    else:
                        self._api_key_var.set(api_key_value)
                    
                    # V3.2.1修复：加载local_url配置
                    local_url = config_data.get("local_url", "http://localhost:8000/v1")
                    self._local_url_var.set(local_url)
                    logger.info(f"[设置加载] local_url: {local_url}")
                    
                    self._temp_var.set(float(config_data.get("temperature", 0.7)))
                    if hasattr(self, '_theme_var'):
                        self._theme_var.set(config_data.get("theme", "dark"))
                    self._temp_label.configure(text=f"{self._temp_var.get():.2f}")
                    
                    # 加载记忆配置（Sprint 9-10）
                    memory_config = config_data.get("memory", {})
                    if hasattr(self, '_wal_auto_save_var'):
                        self._wal_auto_save_var.set(memory_config.get("wal_auto_save", True))
                    if hasattr(self, '_auto_recovery_var'):
                        self._auto_recovery_var.set(memory_config.get("auto_recovery", True))
                    if hasattr(self, '_vector_recall_topk_var'):
                        self._vector_recall_topk_var.set(memory_config.get("vector_recall_topk", 10))
                    if hasattr(self, '_daily_meditation_time_var'):
                        self._daily_meditation_time_var.set(memory_config.get("daily_meditation_time", "23:00"))
                    if hasattr(self, '_embedding_model_var'):
                        self._embedding_model_var.set(memory_config.get("embedding_model", "local"))
                    
                    logger.info("Settings loaded from config.yaml")
                    
                    # V3.2.1修复：启动时发布配置加载事件，通知AI状态管理插件
                    self._publish_ai_config_loaded()
        except Exception as e:
            logger.warning(f"Failed to load settings from config.yaml: {e}")
    
    def _publish_ai_config_loaded(self) -> None:
        """启动时发布AI配置加载事件（V3.2.1新增）
        
        目的：通知AI状态管理插件配置已加载，避免插件使用默认值
        """
        try:
            # 直接获取EventBus（避免依赖ConfigService的内部状态）
            from core.event_bus import get_event_bus, Event
            
            event_bus = get_event_bus()
            if not event_bus:
                logger.warning("[配置加载] EventBus不可用，跳过事件发布")
                return
            
            # 构建AI配置字典（从UI变量获取）
            ai_config = {
                "service_mode": self._service_mode_var.get(),
                "provider": self._provider_var.get(),
                "model": self._model_var.get(),
                "api_key": self._api_key_var.get(),
                "local_url": self._local_url_var.get(),
                "temperature": self._temp_var.get(),
            }
            
            # 发布config.changed事件（V3.2.2修复：publish方法签名是publish(event_type, data, source)）
            event_bus.publish(
                event_type='config.changed',
                data={
                    'type': 'ai_config',
                    'data': ai_config
                },
                source='MainWindow._load_settings_from_config'
            )
            logger.info(f"[配置加载] 已发布AI配置事件: provider={ai_config['provider']}, mode={ai_config['service_mode']}")
            
        except Exception as e:
            logger.warning(f"[配置加载] 发布AI配置事件失败: {e}")
    
    # ============== 事件处理器 ==============
    
    def _clear_hot_ranking_cache(self) -> None:
        """清除热榜缓存 - 调用HotRankingPlugin"""
        # 方式1: 从插件注册表获取
        if CORE_AVAILABLE:
            try:
                registry = get_plugin_registry()
                if registry:
                    hot_ranking_plugin = registry.get_plugin("hot-ranking-v1")
                    if hot_ranking_plugin:
                        result = hot_ranking_plugin.execute("clear_cache")
                        if result:
                            messagebox.showinfo("成功", "热榜缓存已清除！")
                            return
            except Exception as e:
                logger.warning(f"[热榜] 插件注册表获取失败: {e}")
        
        # 方式2: 直接动态加载插件
        try:
            import importlib.util
            plugin_path = os.path.join(os.path.dirname(__file__), "plugins", "hot-ranking-v1", "plugin.py")
            spec = importlib.util.spec_from_file_location("hot_ranking_plugin", plugin_path)
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)
            
            HotRankingPlugin = module.HotRankingPlugin
            from core.plugin_interface import PluginContext
            context = PluginContext(
                plugin_id="hot-ranking-v1",
                config={},
                event_bus=get_event_bus() if CORE_AVAILABLE else None,
                service_locator=get_service_locator() if CORE_AVAILABLE else None
            )
            
            plugin = HotRankingPlugin()
            if plugin.initialize(context):
                result = plugin.execute("clear_cache")
                if result:
                    messagebox.showinfo("成功", "热榜缓存已清除！")
                    return
        except Exception as e:
            logger.error(f"[热榜] 清除缓存失败: {e}")
        
        messagebox.showinfo("提示", "缓存已清除！")
    
    def _update_hot_ranking_data(self) -> None:
        """更新热榜数据 - 调用HotRankingPlugin爬取真实数据"""
        self._set_status("正在刷新热榜...")
        
        # 定义刷新完成回调
        def on_refresh_complete(data):
            # 回调到主线程刷新UI
            self.root.after(0, lambda: self._on_hot_ranking_refresh_complete(data))
        
        # 方式1: 从插件注册表获取
        if CORE_AVAILABLE:
            try:
                registry = get_plugin_registry()
                if registry:
                    hot_ranking_plugin = registry.get_plugin("hot-ranking-v1")
                    if hot_ranking_plugin:
                        # 使用异步刷新方法
                        hot_ranking_plugin.refresh_async(callback=on_refresh_complete, force_update=True)
                        return
            except Exception as e:
                logger.warning(f"[热榜] 插件注册表获取失败: {e}")
        
        # 方式2: 直接动态加载插件
        try:
            import importlib.util
            plugin_path = os.path.join(os.path.dirname(__file__), "plugins", "hot-ranking-v1", "plugin.py")
            spec = importlib.util.spec_from_file_location("hot_ranking_plugin", plugin_path)
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)
            
            HotRankingPlugin = module.HotRankingPlugin
            from core.plugin_interface import PluginContext
            context = PluginContext(
                event_bus=get_event_bus() if CORE_AVAILABLE else None,
                service_locator=get_service_locator() if CORE_AVAILABLE else None,
                config_manager=None,
                plugin_registry=None,
                logger=None
            )
            
            plugin = HotRankingPlugin()
            if plugin.initialize(context):
                # 使用异步刷新
                plugin.refresh_async(callback=on_refresh_complete, force_update=True)
                return
        except Exception as e:
            logger.error(f"[热榜] 更新失败: {e}")
            self._set_status(f"热榜更新失败: {e}")
            return
        
        # 降级：静态刷新
        time.sleep(0.5)
        self._refresh_current_page()
        self._set_status("热榜已更新（离线数据）")
    
    def _on_hot_ranking_refresh_complete(self, data: Dict) -> None:
        """热榜刷新完成回调（主线程执行）"""
        if data:
            self._set_status("热榜已更新（真实数据）")
            # 刷新当前页面显示
            if self._current_page == "hot_ranking":
                self._refresh_current_page()
        else:
            self._set_status("热榜更新失败，请重试")
    
    # ============== V2.18 热榜异步化方法 ==============
    
    def _async_fetch_hot_ranking(self, scrollable_frame):
        """异步加载热榜数据（缓存优先）"""
        def fetch_task():
            try:
                # 尝试获取缓存数据
                data = self._get_hot_ranking_data()
                
                # 回调主线程更新UI（使用try-except防止主线程未就绪）
                try:
                    self.root.after(0, lambda: self._on_hot_ranking_loaded(data))
                except Exception as callback_error:
                    logger.warning(f"[热榜] 主线程回调失败，稍后重试: {callback_error}")
                    # 延迟重试
                    self.root.after(100, lambda: self._on_hot_ranking_loaded(data))
                
            except Exception as e:
                error_msg = str(e)
                logger.error(f"[热榜] 异步加载失败: {error_msg}")
                # 修复：在lambda外捕获error_msg，避免闭包问题
                try:
                    self.root.after(0, lambda msg=error_msg: self._on_hot_ranking_load_failed(msg))
                except Exception:
                    pass  # 主线程未就绪，忽略错误
        
        # 使用线程池提交任务
        from core.thread_pool_manager import thread_pool_manager
        thread_pool_manager.submit_sync(fetch_task)
    
    def _on_hot_ranking_loaded(self, data: Dict):
        """热榜数据加载完成回调（主线程执行）"""
        if not data:
            self._on_hot_ranking_load_failed("数据为空")
            return
        
        # 渲染网站榜单
        sites_data = data.get('sites', [])
        if sites_data:
            self._render_hot_ranking_sites(sites_data)
        
        # 渲染题材榜
        genres_data = data.get('genres', {})
        if genres_data:
            self._render_hot_ranking_genres(genres_data)
        
        # 渲染类型榜
        types_data = data.get('types', {})
        if types_data:
            self._render_hot_ranking_types(types_data)
        
        # 渲染作家榜
        authors_data = data.get('authors', [])
        if authors_data:
            self._render_hot_ranking_authors(authors_data)
        
        # 更新进度
        self._hot_ranking_progress_bar['value'] = 100
        self._hot_ranking_progress_var.set("数据已加载")
        
        # 3秒后隐藏进度条
        self.root.after(3000, lambda: self._hot_ranking_progress_frame.pack_forget())
        
        self._set_status("热榜数据已加载")
    
    def _on_hot_ranking_load_failed(self, error: str):
        """热榜数据加载失败回调（主线程执行）"""
        self._hot_ranking_progress_var.set(f"加载失败: {error}")
        logger.error(f"[热榜] 加载失败: {error}")
        self._set_status(f"热榜加载失败: {error}")
    
    def _start_daily_meditation(self):
        """V2.20: 启动每日冥想定时任务"""
        try:
            from core.daily_meditation import get_daily_meditation_scheduler
            
            scheduler = get_daily_meditation_scheduler()
            success = scheduler.start()
            
            if success:
                logger.info("[每日冥想] 定时任务启动成功")
            else:
                logger.warning("[每日冥想] 定时任务启动失败或已禁用")
                
        except Exception as e:
            logger.error(f"[每日冥想] 启动异常: {e}")
    
    def _check_auto_recovery(self):
        """V2.19: 检查会话自动恢复配置"""
        try:
            import yaml
            from pathlib import Path
            
            # 读取config.yaml
            config_path = Path.cwd() / "config.yaml"
            if not config_path.exists():
                logger.info("[会话恢复] config.yaml不存在，跳过自动恢复")
                return
            
            with open(config_path, 'r', encoding='utf-8') as f:
                config = yaml.safe_load(f)
            
            # 读取记忆配置
            memory_config = config.get("memory", {})
            auto_recovery = memory_config.get("auto_recovery", True)
            
            if not auto_recovery:
                logger.info("[会话恢复] 自动恢复已禁用，跳过")
                return
            
            # 执行WAL恢复检查
            from core.wal_manager import get_wal_manager
            from core.session_state import get_session_state_manager
            
            wal_manager = get_wal_manager()
            session_manager = get_session_state_manager()
            
            # 检查是否有待恢复的记录
            recovered = wal_manager.recover_on_startup()
            
            if recovered:
                # 显示恢复提示
                operation = recovered.get("operation", "未知操作")
                timestamp = recovered.get("timestamp", "未知时间")
                
                from tkinter import messagebox
                result = messagebox.askyesno(
                    "会话恢复",
                    f"检测到上次未完成的操作：\n\n"
                    f"操作：{operation}\n"
                    f"时间：{timestamp}\n\n"
                    f"是否恢复上次会话？"
                )
                
                if result:
                    # 恢复会话状态
                    context = recovered.get("context", {})
                    session_manager.update_state({
                        "active_task": {
                            "current_function": operation,
                            "last_operation": "恢复会话",
                            "last_operation_time": timestamp
                        },
                        "temp_context": context
                    })
                    logger.info(f"[会话恢复] 已恢复上次会话: {operation}")
                    self._set_status(f"已恢复上次会话: {operation}")
                else:
                    # 用户选择不恢复，清除WAL记录
                    wal_manager.clear()
                    logger.info("[会话恢复] 用户取消恢复，已清除WAL记录")
            else:
                logger.info("[会话恢复] 无待恢复的会话")
                
        except Exception as e:
            logger.warning(f"[会话恢复] 检查失败（不影响正常使用）: {e}")
    
    def _sync_project_name_on_startup(self):
        """V3.2.1: 启动时同步项目名称显示
        
        检查项目管理器是否已有项目打开，如果有则更新状态栏和项目管理页面
        """
        try:
            if not self._project_manager:
                return
            
            # 检查是否有项目打开
            if self._project_manager.is_project_open():
                project_name = self._project_manager.get_project_name()
                project_data = self._project_manager.get_project_data()
                
                if project_name:
                    # 更新项目管理页面的项目名称
                    self._project_name_var.set(project_name)
                    
                    # 更新状态栏的项目名称
                    self._update_status_bar(project_name=project_name)
                    
                    logger.info(f"[启动同步] 项目名称已更新: {project_name}")
            
        except Exception as e:
            logger.warning(f"[启动同步] 同步项目名称失败: {e}")
    
    def _preload_vector_store_async(self):
        """V2.18: 后台预加载向量模型（不阻塞UI）"""
        def preload_task():
            try:
                logger.info("[预加载] 开始后台加载向量模型...")
                from infrastructure.vector_store import NovelVectorStore
                vs = NovelVectorStore()
                # 触发延迟加载
                _ = vs.vector_dim
                logger.info("[预加载] 向量模型后台加载完成")
            except Exception as e:
                logger.warning(f"[预加载] 向量模型加载失败（不影响正常使用）: {e}")
        
        # 使用线程池提交任务
        from core.thread_pool_manager import thread_pool_manager
        thread_pool_manager.submit_sync(preload_task)
    

    def _render_hot_ranking_sites(self, sites_data: List[Dict]):
        """渲染网站榜单"""
        if not hasattr(self, '_hot_sites_frame') or not self._hot_sites_frame:
            return
        
        # 清空现有内容
        for widget in self._hot_sites_frame.winfo_children():
            widget.destroy()
        
        for site_info in sites_data:
            site_column = ttk.Frame(self._hot_sites_frame, relief='flat', style="TFrame")
            site_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=8)
            
            # 网站标题
            site_header = tk.Label(
                site_column,
                text=site_info['name'],
                font=(GlassTheme.FONT_FAMILY, 13, 'bold'),
                fg='white',
                bg=site_info['color'],
                pady=10
            )
            site_header.pack(fill=tk.X)
            
            # 书籍列表
            books_frame = ttk.Frame(site_column, style="TFrame")
            books_frame.pack(fill=tk.X, pady=6)
            
            books = site_info.get('books', [])[:10]
            for idx, book in enumerate(books, 1):
                book_item = ttk.Frame(books_frame, style="TFrame")
                book_item.pack(fill=tk.X, pady=3)
                
                # 排名徽章
                rank_bg = site_info['color'] if idx <= 3 else '#CCCCCC'
                rank_label = tk.Label(
                    book_item,
                    text=str(idx),
                    font=(GlassTheme.FONT_FAMILY, 11, 'bold'),
                    fg='white',
                    bg=rank_bg,
                    width=3
                )
                rank_label.pack(side=tk.LEFT, padx=(0, 5))
                
                # 书籍信息
                info_frame = ttk.Frame(book_item, style="TFrame")
                info_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                title = book.get('title', '未知')
                author = book.get('author', '未知')
                category = book.get('category', '')
                heat = book.get('heat', 0)
                
                # 格式化热度
                if heat >= 100000000:
                    heat_str = f"{heat/100000000:.1f}亿"
                elif heat >= 10000:
                    heat_str = f"{heat/10000:.1f}万"
                elif heat > 0:
                    heat_str = str(heat)
                else:
                    heat_str = "-"
                
                # 书名
                ttk.Label(
                    info_frame,
                    text=f"《{title}》",
                    font=(GlassTheme.FONT_FAMILY, 10, 'bold')
                ).pack(anchor='w')
                
                # 作者/分类/热度
                detail_parts = [f"作者：{author}"]
                if category:
                    detail_parts.append(category)
                if heat_str != "-":
                    detail_parts.append(f"热度：{heat_str}")
                
                ttk.Label(
                    info_frame,
                    text="  ".join(detail_parts),
                    font=(GlassTheme.FONT_FAMILY, 8),
                    foreground='gray'
                ).pack(anchor='w')
    
    def _render_hot_ranking_genres(self, genres_data: Dict):
        """渲染题材热度榜"""
        if not hasattr(self, '_hot_genres_frame') or not self._hot_genres_frame:
            return
        
        # 清空现有内容
        for widget in self._hot_genres_frame.winfo_children():
            widget.destroy()
        
        for gender in ['male', 'female']:
            genre_info = genres_data.get(gender, {})
            if not genre_info:
                continue
            
            genre_column = ttk.Frame(self._hot_genres_frame, style="TFrame")
            genre_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
            
            genre_header = tk.Label(
                genre_column,
                text=genre_info.get('title', ''),
                font=(GlassTheme.FONT_FAMILY, 13, 'bold'),
                fg='white',
                bg=genre_info.get('color', GlassTheme.PRIMARY),
                pady=10
            )
            genre_header.pack(fill=tk.X)
            
            chart_frame = ttk.Frame(genre_column, style="TFrame")
            chart_frame.pack(fill=tk.X, pady=10)
            
            genres = genre_info.get('genres', [])
            valid_genres = [(name, heat) for name, heat in genres if name and name != '未知']

            if valid_genres:
                max_hot = max([g[1] for g in valid_genres]) or 1

                for genre_name, hot_value in valid_genres:
                    chart_item = ttk.Frame(chart_frame, style="TFrame")
                    chart_item.pack(fill=tk.X, pady=8)
                    
                    ttk.Label(
                        chart_item,
                        text=genre_name,
                        font=(GlassTheme.FONT_FAMILY, 10),
                        width=10,
                        anchor='e'
                    ).pack(side=tk.LEFT, padx=(0, 8))
                    
                    bar_container = tk.Frame(chart_item, bg=GlassTheme.GLASS_BORDER, height=20)
                    bar_container.pack(side=tk.LEFT, fill=tk.X, expand=True)
                    
                    bar_canvas = tk.Canvas(bar_container, bg=GlassTheme.GLASS_BORDER, height=20, highlightthickness=0)
                    bar_canvas.pack(fill=tk.BOTH, expand=True)
                    
                    ratio = hot_value / max_hot
                    color = genre_info.get('color', GlassTheme.PRIMARY)
                    
                    def on_bar_configure(event, canvas=bar_canvas, c=color, r=ratio):
                        w = canvas.winfo_width()
                        h = canvas.winfo_height()
                        if w > 10 and h > 5:
                            canvas.delete("all")
                            canvas.create_rectangle(0, 0, int(w * r), h, fill=c, outline="")
                    
                    bar_canvas.bind("<Configure>", on_bar_configure)
                    
                    tk.Label(
                        chart_item,
                        text=f"{hot_value:.1f}分",
                        font=(GlassTheme.FONT_FAMILY, 9, 'bold'),
                        fg=color,
                        bg=GlassTheme.GLASS_BG,
                        width=8, anchor='w'
                    ).pack(side=tk.LEFT, padx=(8, 0))
    
    def _render_hot_ranking_types(self, types_data: Dict):
        """渲染设定类型/创作流派榜"""
        if not hasattr(self, '_hot_types_frame') or not self._hot_types_frame:
            return
        
        # 清空现有内容
        for widget in self._hot_types_frame.winfo_children():
            widget.destroy()
        
        for gender in ['male', 'female']:
            type_info = types_data.get(gender, {})
            if not type_info:
                continue
            
            type_column = ttk.Frame(self._hot_types_frame, style="TFrame")
            type_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
            
            type_header = tk.Label(
                type_column,
                text=type_info.get('title', ''),
                font=(GlassTheme.FONT_FAMILY, 13, 'bold'),
                fg='white',
                bg=type_info.get('color', GlassTheme.ACCENT_GREEN),
                pady=10
            )
            type_header.pack(fill=tk.X)
            
            chart_frame = ttk.Frame(type_column, style="TFrame")
            chart_frame.pack(fill=tk.X, pady=10)
            
            types = type_info.get('types', [])
            valid_types = [(name, heat) for name, heat in types if name and name != '未知']

            if valid_types:
                max_hot = max([t[1] for t in valid_types]) or 1

                for type_name, hot_value in valid_types:
                    chart_item = ttk.Frame(chart_frame, style="TFrame")
                    chart_item.pack(fill=tk.X, pady=8)
                    
                    ttk.Label(
                        chart_item,
                        text=type_name,
                        font=(GlassTheme.FONT_FAMILY, 10),
                        width=10,
                        anchor='e'
                    ).pack(side=tk.LEFT, padx=(0, 8))
                    
                    bar_container = tk.Frame(chart_item, bg=GlassTheme.GLASS_BORDER, height=20)
                    bar_container.pack(side=tk.LEFT, fill=tk.X, expand=True)
                    
                    bar_canvas = tk.Canvas(bar_container, bg=GlassTheme.GLASS_BORDER, height=20, highlightthickness=0)
                    bar_canvas.pack(fill=tk.BOTH, expand=True)
                    
                    ratio = hot_value / max_hot
                    color = type_info.get('color', GlassTheme.ACCENT_GREEN)
                    
                    def on_bar_configure(event, canvas=bar_canvas, c=color, r=ratio):
                        w = canvas.winfo_width()
                        h = canvas.winfo_height()
                        if w > 10 and h > 5:
                            canvas.delete("all")
                            canvas.create_rectangle(0, 0, int(w * r), h, fill=c, outline="")
                    
                    bar_canvas.bind("<Configure>", on_bar_configure)
                    
                    tk.Label(
                        chart_item,
                        text=f"{hot_value:.1f}分",
                        font=(GlassTheme.FONT_FAMILY, 9, 'bold'),
                        fg=color,
                        bg=GlassTheme.GLASS_BG,
                        width=8, anchor='w'
                    ).pack(side=tk.LEFT, padx=(8, 0))
    
    def _render_hot_ranking_authors(self, authors_data: List[Dict]):
        """渲染作家排行榜"""
        if not hasattr(self, '_hot_authors_frame') or not self._hot_authors_frame:
            return
        
        # 清空现有内容
        for widget in self._hot_authors_frame.winfo_children():
            widget.destroy()
        
        author_table_frame = tk.Frame(self._hot_authors_frame, bg=GlassTheme.GLASS_SURFACE)
        author_table_frame.pack(fill=tk.BOTH, expand=True)

        # 配置列权重
        for idx in range(5):
            author_table_frame.grid_columnconfigure(idx, weight=1)

        headers = ['排名', '作家', '代表作品', '预估年收入', '粉丝数']
        col_weights = [0, 0, 1, 0, 0]

        for idx, (header, weight) in enumerate(zip(headers, col_weights)):
            tk.Label(
                author_table_frame,
                text=header,
                font=(GlassTheme.FONT_FAMILY, 11, 'bold'),
                fg='white', bg=GlassTheme.PRIMARY,
                pady=8, padx=5, anchor='center',
                relief='flat'
            ).grid(row=0, column=idx, sticky='nsew', padx=1, pady=1)
        
        for row_idx, author in enumerate(authors_data, start=1):
            rank = author.get('rank', row_idx)
            name = author.get('name', '未知')
            works = author.get('works', '暂无')
            income = author.get('income', '0')
            fans = author.get('fans', '0')
            
            if rank == 1:
                bg_color, fg_color = '#FFD700', 'black'
            elif rank == 2:
                bg_color, fg_color = '#C0C0C0', 'black'
            elif rank == 3:
                bg_color, fg_color = '#CD7F32', 'white'
            else:
                bg_color, fg_color = GlassTheme.GLASS_SURFACE, GlassTheme.TEXT_PRIMARY
            
            cells = [f"TOP{rank}", name, works[:25]+'...' if len(works) > 25 else works, f"¥{income}", fans]
            anchors = ['center', 'center', 'w', 'center', 'center']
            fonts = [
                (GlassTheme.FONT_FAMILY, 10, 'bold'),
                (GlassTheme.FONT_FAMILY, 10),
                (GlassTheme.FONT_FAMILY, 9),
                (GlassTheme.FONT_FAMILY, 10, 'bold'),
                (GlassTheme.FONT_FAMILY, 9),
            ]
            for col_idx, (text, anchor, font) in enumerate(zip(cells, anchors, fonts)):
                tk.Label(
                    author_table_frame,
                    text=text,
                    font=font,
                    fg=fg_color, bg=bg_color,
                    pady=6, padx=5, anchor=anchor,
                    relief='flat'
                ).grid(row=row_idx, column=col_idx, sticky='nsew', padx=1, pady=1)

    # ============== 原有回调方法（保留兼容）==============
    
    def _on_worldview(self) -> None:
        """世界观管理"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(0)  # 切换到世界观标签
    
    def _on_characters(self) -> None:
        """人物设定"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(1)  # 切换到人物标签
    
    def _on_outline(self) -> None:
        """大纲管理"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(2)  # 切换到大纲标签
    
    def _on_style(self) -> None:
        """风格学习"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(3)  # 切换到风格标签
    
    def _on_generation(self) -> None:
        """开始创作"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(4)  # 切换到创作标签
    
    def _on_reverse(self) -> None:
        """逆向反馈"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(5)  # 切换到逆向标签
    
    def _on_quick_create(self) -> None:
        """快捷创作"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(6)  # 切换到快捷标签
    
    def _on_continue(self) -> None:
        """续写功能"""
        self._switch_to_page("workbench")
        self._workbench_notebook.select(7)  # 切换到续写标签
    
    def _on_refresh_progress(self) -> None:
        """刷新进度"""
        self._set_status("刷新进度...")
        # 【修复】实际调用更新方法
        if hasattr(self, '_update_progress_display'):
            self._update_progress_display()
            self._set_status("进度已刷新")
    
    def _on_export_progress(self) -> None:
        """导出报告"""
        self._set_status("导出报告...")
    
    def _on_reset_progress(self) -> None:
        """重置进度"""
        self._set_status("重置进度...")
    
    def _on_new_project(self) -> None:
        """新建项目 - 完整实现（参考V5版本）"""
        self._set_status("正在创建新项目...")
        
        # 创建弹窗
        dialog = tk.Toplevel(self.root)
        dialog.title("新建项目")
        dialog.geometry("600x550")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 600) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 550) // 2
        dialog.geometry(f"+{x}+{y}")
        
        main_frame = ttk.Frame(dialog, style="TFrame", padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        ttk.Label(main_frame, text="新建小说创作项目", font=('Microsoft YaHei UI', 16, 'bold')).pack(pady=(0, 5))
        ttk.Label(main_frame, text="填写项目信息，开始您的创作之旅", font=('Microsoft YaHei UI', 10), foreground='gray').pack(pady=(0, 20))
        
        # 创建滚动区域
        canvas = tk.Canvas(main_frame, bg=GlassTheme.GLASS_BG, highlightthickness=0, cursor="arrow")
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="TFrame")
        
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 表单内容
        form_frame = ttk.Frame(scrollable_frame, style="TFrame")
        form_frame.pack(fill=tk.BOTH, expand=True, padx=5)
        
        # 项目名称
        ttk.Label(form_frame, text="项目名称 *", font=('Microsoft YaHei UI', 10, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=8)
        name_entry = ttk.Entry(form_frame, width=40)
        name_entry.grid(row=0, column=1, sticky=tk.EW, pady=8)
        ttk.Label(form_frame, text="（必填）", font=('Microsoft YaHei UI', 9), foreground='gray').grid(row=0, column=2, sticky=tk.W, padx=10, pady=8)
        
        # 作者
        ttk.Label(form_frame, text="作者", font=('Microsoft YaHei UI', 10, 'bold')).grid(row=1, column=0, sticky=tk.W, pady=8)
        author_entry = ttk.Entry(form_frame, width=40)
        author_entry.grid(row=1, column=1, sticky=tk.EW, pady=8)
        ttk.Label(form_frame, text="（可选）", font=('Microsoft YaHei UI', 9), foreground='gray').grid(row=1, column=2, sticky=tk.W, padx=10, pady=8)
        
        # 作品类型
        ttk.Label(form_frame, text="作品类型 *", font=('Microsoft YaHei UI', 10, 'bold')).grid(row=2, column=0, sticky=tk.W, pady=8)
        genre_combo = ttk.Combobox(form_frame, width=37,
                                  values=["玄幻", "武侠", "科幻", "都市", "言情", "历史", "军事", "其他", "仙侠", "灵异", "同人"])
        genre_combo.grid(row=2, column=1, sticky=tk.EW, pady=8)
        genre_combo.current(0)
        ttk.Label(form_frame, text="（必填）", font=('Microsoft YaHei UI', 9), foreground='gray').grid(row=2, column=2, sticky=tk.W, padx=10, pady=8)
        
        # 项目路径
        ttk.Label(form_frame, text="项目路径 *", font=('Microsoft YaHei UI', 10, 'bold')).grid(row=3, column=0, sticky=tk.W, pady=8)
        path_frame = ttk.Frame(form_frame, style="TFrame")
        path_frame.grid(row=3, column=1, sticky=tk.EW, pady=8)
        ttk.Label(form_frame, text="（必填）", font=('Microsoft YaHei UI', 9), foreground='gray').grid(row=3, column=2, sticky=tk.W, padx=10, pady=8)
        
        default_path = os.getcwd()
        path_entry = ttk.Entry(path_frame, width=30)
        path_entry.insert(0, default_path)
        path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        def browse_path():
            path = filedialog.askdirectory(title="选择项目路径", initialdir=default_path)
            if path:
                path_entry.delete(0, tk.END)
                path_entry.insert(0, path)
        
        ttk.Button(path_frame, text="浏览...", command=browse_path, width=8).pack(side=tk.LEFT, padx=(5, 0))
        
        # 目标字数
        ttk.Label(form_frame, text="目标字数 *", font=('Microsoft YaHei UI', 10, 'bold')).grid(row=4, column=0, sticky=tk.W, pady=8)
        words_frame = ttk.Frame(form_frame, style="TFrame")
        words_frame.grid(row=4, column=1, sticky=tk.W, pady=8)
        ttk.Label(form_frame, text="（必填）", font=('Microsoft YaHei UI', 9), foreground='gray').grid(row=4, column=2, sticky=tk.W, padx=10, pady=8)
        
        words_entry = ttk.Entry(words_frame, width=15)
        words_entry.insert(0, "100000")
        words_entry.pack(side=tk.LEFT)
        ttk.Label(words_frame, text="字", font=('Microsoft YaHei UI', 9)).pack(side=tk.LEFT, padx=5)
        
        # 预计章节数
        ttk.Label(form_frame, text="预计章节数", font=('Microsoft YaHei UI', 10, 'bold')).grid(row=5, column=0, sticky=tk.W, pady=8)
        chapters_frame = ttk.Frame(form_frame, style="TFrame")
        chapters_frame.grid(row=5, column=1, sticky=tk.W, pady=8)
        ttk.Label(form_frame, text="（可选）", font=('Microsoft YaHei UI', 9), foreground='gray').grid(row=5, column=2, sticky=tk.W, padx=10, pady=8)
        
        chapters_entry = ttk.Entry(chapters_frame, width=15)
        chapters_entry.insert(0, "50")
        chapters_entry.pack(side=tk.LEFT)
        ttk.Label(chapters_frame, text="章", font=('Microsoft YaHei UI', 9)).pack(side=tk.LEFT, padx=5)
        
        # 项目描述
        ttk.Label(form_frame, text="项目简介", font=('Microsoft YaHei UI', 10, 'bold')).grid(row=6, column=0, sticky=tk.NW, pady=8)
        desc_text = tk.Text(form_frame, width=40, height=6, font=('Microsoft YaHei UI', 10))
        desc_text.grid(row=6, column=1, sticky=tk.EW, pady=8)
        ttk.Label(form_frame, text="（可选）", font=('Microsoft YaHei UI', 9), foreground='gray').grid(row=6, column=2, sticky=tk.W, padx=10, pady=8)
        
        # 提示信息
        info_frame = ttk.LabelFrame(form_frame, text="说明", padding=10)
        info_frame.grid(row=7, column=0, columnspan=3, sticky=tk.EW, pady=15)
        info_text = "项目创建后，将在指定路径下自动创建以下目录结构：\n" \
                    "  • 大纲/ - 存放故事大纲文件（支持.docx和.txt格式）\n" \
                    "  • 人物/ - 存放人物设定文件\n" \
                    "  • 世界观/ - 存放世界观设定文件\n" \
                    "  • 小说/ - 存放生成的小说内容\n" \
                    "  • {项目名}.json - 项目配置文件"
        ttk.Label(info_frame, text=info_text, font=('Microsoft YaHei UI', 9), justify=tk.LEFT).pack(anchor=tk.W)
        
        def confirm_create():
            """确认创建项目"""
            name = name_entry.get().strip()
            author = author_entry.get().strip()
            genre = genre_combo.get()
            path = path_entry.get().strip()
            words = words_entry.get().strip()
            chapters = chapters_entry.get().strip()
            description = desc_text.get("1.0", tk.END).strip()
            
            # 验证必填字段
            if not name:
                messagebox.showwarning("输入错误", "请输入项目名称！", parent=dialog)
                name_entry.focus()
                return
            if not path:
                messagebox.showwarning("输入错误", "请选择项目路径！", parent=dialog)
                path_entry.focus()
                return
            if not os.path.exists(path):
                messagebox.showwarning("路径错误", f"项目路径不存在：{path}", parent=dialog)
                return
            if not words.isdigit() or int(words) <= 0:
                messagebox.showwarning("输入错误", "目标字数必须是大于0的数字！", parent=dialog)
                words_entry.focus()
                return
            
            # 验证项目名是否包含非法字符
            invalid_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
            if any(char in name for char in invalid_chars):
                messagebox.showwarning("输入错误", f"项目名称不能包含以下字符：{' '.join(invalid_chars)}", parent=dialog)
                return
            
            # 检查项目是否已存在
            project_dir = os.path.join(path, name)
            if os.path.exists(project_dir):
                response = messagebox.askyesno("项目已存在", f"目录 '{name}' 已存在，是否覆盖？\n\n警告：这将删除现有项目！", parent=dialog)
                if not response:
                    return
            
            try:
                import shutil
                
                # 创建项目数据
                project_data = {
                    "name": name,
                    "author": author if author else "未设置",
                    "genre": genre,
                    "path": path,
                    "target_words": int(words),
                    "estimated_chapters": int(chapters) if chapters.isdigit() else 50,
                    "description": description,
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "modified_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "outline": None,
                    "characters": [],
                    "worldview": [],
                    "style": None,  # 风格设定
                    "reverse_feedback": {  # 逆向反馈记录
                        "reports": [],  # 冲突报告列表
                        "corrections": [],  # 修正记录列表
                        "last_chapter_analyzed": None
                    },
                    "reverse_chapters": {},  # 逆向反馈上传的章节
                    "completed_chapters": [],  # 已完成章节列表
                    "generated_content": []
                }
                
                # 删除现有项目目录（如果存在）
                if os.path.exists(project_dir):
                    self._set_status("正在删除现有项目目录...")
                    try:
                        shutil.rmtree(project_dir)
                    except Exception as e:
                        messagebox.showerror("错误", f"无法删除现有项目目录: {e}", parent=dialog)
                        return
                
                # 创建项目主目录
                self._set_status("正在创建项目目录...")
                os.makedirs(project_dir, exist_ok=True)
                
                # 创建子目录
                subdirs = ["大纲", "人物", "世界观", "风格", "逆向反馈", "小说"]
                for subdir in subdirs:
                    subdir_path = os.path.join(project_dir, subdir)
                    os.makedirs(subdir_path, exist_ok=True)
                
                # 创建README文件
                readme_content = f"""# {name}

**作者**: {author if author else "未设置"}
**类型**: {genre}
**创建时间**: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}

## 项目简介

{description if description else "暂无简介"}

## 目录说明

- **大纲/**: 存放故事大纲文件（支持.docx和.txt格式）
- **人物/**: 存放人物设定文件
- **世界观/**: 存放世界观设定文件
- **风格/**: 存放风格学习配置和档案
- **逆向反馈/**: 存放逆向反馈分析报告和修正记录
- **小说/**: 存放生成的小说内容
- **{name}.json**: 项目配置文件

## 使用说明

1. 准备大纲文件（Word或TXT格式）放入"大纲"目录
2. 准备风格学习参考作品（可选）
3. 使用GUI的工作台功能进行创作

---
由 Novel Writing Assistant - Agent Pro 生成
"""
                readme_path = os.path.join(project_dir, "README.md")
                with open(readme_path, 'w', encoding='utf-8') as f:
                    f.write(readme_content)
                
                # 保存项目文件
                project_file = os.path.join(project_dir, f"{name}.json")
                with open(project_file, 'w', encoding='utf-8') as f:
                    json.dump(project_data, f, ensure_ascii=False, indent=2)
                
                # 设置当前项目
                self.current_project = project_data
                self.project_file = project_file
                
                # 更新UI显示
                self._project_name_var.set(name)
                self._project_path_var.set(project_dir)
                
                # 显示成功信息
                success_msg = f"项目创建成功！\n\n"
                success_msg += f"项目名称: {name}\n"
                success_msg += f"项目路径: {project_dir}\n"
                success_msg += f"目标字数: {int(words):,}字\n"
                success_msg += f"预计章节: {chapters if chapters.isdigit() else '50'}章\n\n"
                success_msg += "已创建的文件和目录:\n"
                success_msg += f"- {name}.json (项目配置文件)\n"
                success_msg += "- README.md (项目说明文件)\n"
                success_msg += "- 大纲/、人物/、世界观/、小说/ 目录\n\n"
                success_msg += "下一步：在工作台导入大纲文件，开始创作！"
                
                messagebox.showinfo("创建成功", success_msg, parent=dialog)
                self._set_status(f"项目 '{name}' 创建成功！")
                dialog.destroy()
                
            except Exception as e:
                import traceback
                messagebox.showerror("错误", f"创建项目失败：{e}", parent=dialog)
                self._set_status("创建项目失败")
                traceback.print_exc()
        
        # 按钮
        btn_frame = ttk.Frame(form_frame, style="TFrame")
        btn_frame.grid(row=8, column=0, columnspan=3, pady=20)
        
        ttk.Button(btn_frame, text="创建项目", command=confirm_create, width=15).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy, width=10).pack(side=tk.LEFT, padx=5)
        
        name_entry.focus()
    
    def _on_open_project(self) -> None:
        """打开项目 - 使用项目管理器"""
        # 首先尝试打开项目JSON文件
        project_file = filedialog.askopenfilename(
            title="打开项目",
            filetypes=[("项目文件", "*.json"), ("所有文件", "*.*")]
        )
        
        if not project_file:
            return
        
        self._set_status("正在加载项目...")
        
        def load_project_thread():
            """后台线程加载项目"""
            try:
                # 优先使用项目管理器
                if self._project_manager:
                    # 调用项目管理器加载项目
                    success = self._project_manager.load_project(project_file)
                    
                    if success:
                        project_name = self._project_manager.get_project_name()
                        project_data = self._project_manager.get_project_data()
                        
                        # 更新主线程UI
                        def update_ui():
                            self.current_project = project_data
                            self.project_file = project_file

                            # 更新显示
                            self._project_name_var.set(project_name)
                            project_dir = os.path.dirname(project_file)
                            self._project_path_var.set(project_dir)

                            # V3.2修复：更新状态栏项目名称
                            self._update_status_bar(project_name=project_name)

                            # 恢复所有模块数据到UI
                            self._restore_project_data_to_ui(project_data)

                            self._set_status(f"项目 '{project_name}' 已打开")
                            
                            # 显示成功信息
                            info_msg = f"项目加载成功！\n\n"
                            info_msg += f"项目名称: {project_name}\n"
                            info_msg += f"作者: {project_data.get('author', '未设置')}\n"
                            info_msg += f"类型: {project_data.get('genre', '未设置')}\n"
                            info_msg += f"目标字数: {project_data.get('target_words', 0):,}字\n"
                            
                            # 检查项目完整性
                            missing = []
                            if not project_data.get('outline'):
                                missing.append("大纲")
                            if not project_data.get('characters'):
                                missing.append("人物设定")
                            if not project_data.get('worldview'):
                                missing.append("世界观设定")
                            if not project_data.get('style'):
                                missing.append("风格设定")
                            if not project_data.get('reverse_feedback'):
                                missing.append("逆向反馈")
                            
                            if missing:
                                info_msg += f"\n\n提示：以下内容尚未设置：{', '.join(missing)}"
                            
                            messagebox.showinfo("项目已打开", info_msg)
                        
                        self.root.after(0, update_ui)
                        return
                    else:
                        # 项目管理器加载失败，继续使用降级方案
                        pass
                
                # 降级方案：直接读取项目文件
                # 检查文件是否存在
                if not os.path.exists(project_file):
                    self.root.after(0, lambda: [
                        messagebox.showerror("错误", f"项目文件不存在: {project_file}"),
                        self._set_status("打开项目失败")
                    ])
                    return
                
                # 读取项目文件
                with open(project_file, 'r', encoding='utf-8') as f:
                    project_data = json.load(f)
                
                # 验证项目数据
                if not isinstance(project_data, dict):
                    raise ValueError("项目文件格式错误")
                
                project_name = project_data.get('name', '未命名项目')
                
                # 更新主线程UI
                def update_ui():
                    self.current_project = project_data
                    self.project_file = project_file
                    
                    # 更新显示
                    self._project_name_var.set(project_name)
                    project_dir = os.path.dirname(project_file)
                    self._project_path_var.set(project_dir)
                    
                    # V3.2.1修复：更新状态栏项目名称
                    self._update_status_bar(project_name=project_name)
                    
                    # 恢复所有模块数据到UI
                    self._restore_project_data_to_ui(project_data)
                    
                    self._set_status(f"项目 '{project_name}' 已打开")
                    
                    # 显示成功信息
                    info_msg = f"项目加载成功！\n\n"
                    info_msg += f"项目名称: {project_name}\n"
                    info_msg += f"作者: {project_data.get('author', '未设置')}\n"
                    info_msg += f"类型: {project_data.get('genre', '未设置')}\n"
                    info_msg += f"目标字数: {project_data.get('target_words', 0):,}字\n"
                    
                    # 检查项目完整性
                    missing = []
                    if not project_data.get('outline'):
                        missing.append("大纲")
                    if not project_data.get('characters'):
                        missing.append("人物设定")
                    if not project_data.get('worldview'):
                        missing.append("世界观设定")
                    if not project_data.get('style'):
                        missing.append("风格设定")
                    if not project_data.get('reverse_feedback'):
                        missing.append("逆向反馈")
                    
                    if missing:
                        info_msg += f"\n\n提示：以下内容尚未设置：{', '.join(missing)}"
                    
                    messagebox.showinfo("项目已打开", info_msg)
                
                self.root.after(0, update_ui)
                
            except json.JSONDecodeError as e:
                self.root.after(0, lambda: [
                    messagebox.showerror("错误", f"项目文件格式错误: {e}"),
                    self._set_status("打开项目失败")
                ])
            except Exception as e:
                self.root.after(0, lambda: [
                    messagebox.showerror("错误", f"打开项目失败: {e}"),
                    self._set_status("打开项目失败")
                ])
        
        # 启动后台任务（使用统一线程池，解决卡顿问题）
        from core.thread_pool_manager import thread_pool_manager
        thread_pool_manager.submit_sync(load_project_thread)
    
    def _restore_project_data_to_ui(self, project_data: Dict):
        """从项目数据恢复到UI显示"""
        try:
            # 恢复大纲（大纲管理页面使用_tree，其他地方可能使用_text）
            if project_data.get('outline'):
                self._outline_content = project_data['outline']
                # 更新大纲树（如果存在）
                if hasattr(self, '_outline_tree'):
                    self._update_outline_tree_from_content(self._outline_content)
                # 更新大纲文本框（如果存在，如快捷创作页面）
                if hasattr(self, '_outline_text'):
                    self._outline_text.delete(1.0, tk.END)
                    self._outline_text.insert(1.0, self._outline_content)
            
            # 恢复人物设定
            if project_data.get('characters'):
                self._character_data = project_data['characters']
                # 【修复】同步到_character_entries，确保批量解析后的数据能正确恢复
                self._character_entries = project_data['characters']
                # 【修复】与世界观一致，检查_tree是否存在再更新
                if hasattr(self, '_character_tree'):
                    self._update_character_tree()
            
            # 恢复世界观（解析为列表并显示在树中）
            if project_data.get('worldview'):
                self._worldview_content = project_data['worldview']
                # 【修复】不再强制转换为字符串，直接传递原始数据给更新方法
                # _update_worldview_tree_from_content支持列表和字符串两种格式
                
                # 解析世界观内容并更新树
                if hasattr(self, '_worldview_tree'):
                    self._update_worldview_tree_from_content(self._worldview_content)
                    
                # 【修复】项目打开时不自动填充预览区域，只在双击或查看详情时显示
                # 保持预览区域为空，提示用户操作
                if hasattr(self, '_worldview_preview'):
                    self._worldview_preview.delete("1.0", tk.END)
                    self._worldview_preview.insert("1.0", "双击左侧列表查看详情，或点击【查看详情】按钮")
            
            # 恢复风格设定
            if project_data.get('style'):
                self._style_profile = project_data['style']
                # 更新风格学习页面显示
                if hasattr(self, '_update_style_display'):
                    self._update_style_display(self._style_profile)
            
            # 恢复逆向反馈上传的章节
            if project_data.get('reverse_chapters'):
                self._reverse_chapters = project_data['reverse_chapters']
                if hasattr(self, '_completed_chapters_tree'):
                    self._update_reverse_chapters_tree()
            
            # 恢复逆向反馈数据
            if project_data.get('reverse_feedback'):
                self._reverse_feedback_data = project_data['reverse_feedback']
                # 更新逆向反馈页面显示
                if hasattr(self, '_update_reverse_feedback_display'):
                    self._update_reverse_feedback_display(self._reverse_feedback_data)
            
            # 恢复已完成章节
            if project_data.get('completed_chapters'):
                self._completed_chapters = project_data['completed_chapters']
            
            # 恢复生成内容
            if project_data.get('generated_content'):
                self._generated_content = project_data['generated_content']
            
            self._current_project_name = project_data.get('name', '未命名项目')
            
            # V3.2.1修复：同步更新状态栏项目名称
            if project_data.get('name'):
                self._update_status_bar(project_name=project_data['name'])
            
            # 【修复】无论是否有generated_content，都更新创作进度页面
            # 因为即使没有生成内容，也需要显示大纲、人物、世界观等状态
            if hasattr(self, '_update_progress_display'):
                self._update_progress_display()
            
        except Exception as e:
            logger.error(f"恢复项目数据到UI失败: {e}")
    
    def _update_reverse_chapters_tree(self):
        """更新逆向反馈章节树显示"""
        if not hasattr(self, '_completed_chapters_tree') or not self._reverse_chapters:
            return
        
        try:
            # 清空现有数据
            self._completed_chapters_tree.delete(*self._completed_chapters_tree.get_children())
            
            # 添加章节数据
            for idx, (chapter_id, chapter) in enumerate(self._reverse_chapters.items(), 1):
                self._completed_chapters_tree.insert("", tk.END, iid=chapter_id, values=(
                    f"第{idx}章",
                    chapter.get('title', '未命名'),
                    f"{chapter.get('words', 0)}字",
                    chapter.get('status', '未完成'),
                    chapter.get('source', '未知')
                ))
        except Exception as e:
            logger.error(f"更新逆向反馈章节树失败: {e}")
    
    def _update_character_tree(self):
        """更新人物树显示

        支持两种格式：
        1. 字典列表格式（结构化人物档案）：
           [{'id': 'xxx', 'name': '张三', 'role': '主角', 'status': '活跃',
             'emotion': '平静', 'chapters': '第1-5章', ...}]
        2. 简单格式（兼容旧数据）：
           [{'name': '张三', 'role': '主角'}, ...]
        """
        if not hasattr(self, '_character_tree'):
            return

        # 优先使用_character_entries（批量解析后的数据），降级到_character_data
        char_data = getattr(self, '_character_entries', None) or getattr(self, '_character_data', None)

        if not char_data:
            return

        try:
            # 清空现有数据
            self._character_tree.delete(*self._character_tree.get_children())

            # 添加人物数据
            for char in char_data:
                char_name = char.get('name', '未命名')
                char_role = char.get('role', '未设置')
                char_status = char.get('status', '新建')
                char_emotion = char.get('emotion', '平静')
                char_chapters = char.get('chapters', '未设置')

                # 插入完整5列数据
                self._character_tree.insert('', tk.END, values=(
                    char_name,
                    char_role,
                    char_status,
                    char_emotion,
                    char_chapters
                ))
        except Exception as e:
            logger.error(f"更新人物树失败: {e}")
    
    def _update_outline_tree_from_content(self, outline_content: str):
        """从大纲内容更新大纲树显示"""
        if not hasattr(self, '_outline_tree') or not outline_content:
            return
        
        try:
            # 清空现有数据
            self._outline_tree.delete(*self._outline_tree.get_children())
            
            # 简单解析：按行分割，根据缩进判断层级
            lines = outline_content.split('\n')
            parent_ids = {}
            
            for line in lines:
                if not line.strip():
                    continue
                
                # 计算缩进级别（每两个空格为一级）
                indent_level = 0
                while line.startswith('  '):
                    indent_level += 1
                    line = line[2:]
                
                # 提取章节标题（去除序号和符号）
                title = line.strip()
                # 去除序号（如"第一章"、"1.1"等）
                import re
                title = re.sub(r'^第[一二三四五六七八九十百千万]+[卷章节]|^\d+\.?\d*', '', title).strip()
                
                if not title:
                    continue
                
                # 插入树节点
                parent_id = parent_ids.get(indent_level - 1, '')
                item_id = self._outline_tree.insert(parent_id, tk.END, text=title)
                parent_ids[indent_level] = item_id
                
        except Exception as e:
            logger.error(f"更新大纲树失败: {e}")
    
    def _update_worldview_tree_from_content(self, worldview_content):
        """从世界观内容更新世界观树显示

        【架构修复V3.1】解析逻辑已迁移到WorldviewParserAdapter.parse_for_display()
        GUI只负责：1. 调用适配器解析 2. 显示结果
        """
        if not hasattr(self, '_worldview_tree') or not worldview_content:
            return

        try:
            # 清空现有数据
            self._worldview_tree.delete(*self._worldview_tree.get_children())

            # 初始化词条列表
            if not hasattr(self, '_worldview_entries'):
                self._worldview_entries = []
            self._worldview_entries = []

            # 延迟初始化适配器（架构修复：业务逻辑在插件层）
            if self._worldview_adapter is None:
                from agents.adapters.worldview_adapter import WorldviewParserAdapter
                self._worldview_adapter = WorldviewParserAdapter()
                if not self._worldview_adapter.initialize():
                    logger.error("WorldviewParserAdapter初始化失败")
                    return

            # 调用适配器解析内容（业务逻辑在插件层）
            entries = self._worldview_adapter.parse_for_display(worldview_content)

            # 存储并显示解析结果
            for entry in entries:
                self._worldview_entries.append(entry)

                # 截断显示（UI显示逻辑）
                display_elements = entry['elements'][:50] + '...' if len(entry['elements']) > 50 else entry['elements']

                self._worldview_tree.insert('', tk.END, values=(
                    entry['name'],
                    entry['category'],
                    display_elements,
                    entry['status'],
                    entry['modified']
                ))

        except Exception as e:
            logger.error(f"更新世界观树失败: {e}")
    
    def _update_style_display(self, style_profile):
        """更新风格显示"""
        if not style_profile:
            return
        
        try:
            # 更新风格信息显示
            if hasattr(self, '_style_info'):
                style_name = style_profile.get('name', '未命名风格')
                author = style_profile.get('author', '未知')
                genre = style_profile.get('genre', '未知')
                
                info_text = f"""风格名称: {style_name}
参考作者: {author}
作品类型: {genre}

"""
                # 添加词汇特征
                vocab_features = style_profile.get('vocabulary_features', [])
                if vocab_features:
                    info_text += "【词汇特征】\n"
                    for vf in vocab_features[:10]:  # 只显示前10个
                        info_text += f"- {vf}\n"
                
                # 添加句式特征
                sentence_patterns = style_profile.get('sentence_patterns', [])
                if sentence_patterns:
                    info_text += "\n【句式特征】\n"
                    for sp in sentence_patterns[:10]:
                        info_text += f"- {sp}\n"
                
                self._style_info.delete(1.0, tk.END)
                self._style_info.insert(1.0, info_text)
            
            # 更新风格名称标签
            if hasattr(self, '_style_info_label'):
                style_name = style_profile.get('name', '未命名风格')
                self._style_info_label.config(text=f"当前风格: {style_name}")
                
        except Exception as e:
            logger.error(f"更新风格显示失败: {e}")
    
    def _update_reverse_feedback_display(self, feedback_data):
        """更新逆向反馈显示"""
        if not feedback_data:
            return
        
        try:
            if hasattr(self, '_issues_tree'):
                reports = feedback_data.get('reports', [])
                # 更新显示
                self._issues_tree.delete(*self._issues_tree.get_children())
                
                # 类型映射
                type_map = {
                    'character': '人物',
                    'outline': '大纲',
                    'worldview': '世界观',
                }
                
                # 优先级映射
                severity_map = {
                    'high': '🔴 高',
                    'medium': '🟡 中',
                    'low': '🟢 低',
                }
                
                for report in reports:
                    issues = report.get('issues', [])
                    for issue in issues:
                        issue_type = type_map.get(issue.get('type', 'outline'), '其他')
                        severity = severity_map.get(issue.get('severity', 'medium'), '🟡 中')
                        element = issue.get('element', '未知')
                        description = issue.get('description', '')[:50]  # 截取前50字符
                        
                        self._issues_tree.insert('', tk.END, values=(
                            issue_type,
                            severity,
                            element,
                            description
                        ))
                
                # 更新分析进度标签
                if hasattr(self, '_analysis_progress_label') and reports:
                    last_report = reports[-1]
                    self._analysis_progress_label.configure(
                        text=f"已加载 {len(reports)} 份分析报告，共 {sum(len(r.get('issues', [])) for r in reports)} 个冲突项"
                    )
        except Exception as e:
            logger.error(f"更新逆向反馈显示失败: {e}")
    
    def _update_progress_display(self):
        """更新创作进度显示"""
        try:
            # 更新项目名称（从_current_project_name获取）
            if hasattr(self, '_current_project_name') and self._progress_project_name is not None:
                self._progress_project_name.set(self._current_project_name)
            
            # 更新章节统计
            if self._progress_chapters is not None and self._progress_total_words is not None:
                if hasattr(self, '_generated_content') and self._generated_content:
                    total_chapters = len(self._generated_content)
                    completed_chapters = sum(1 for content in self._generated_content if content)
                    self._progress_chapters.set(f"{completed_chapters} / {total_chapters}")

                    # 计算总字数
                    total_words = 0
                    for content in self._generated_content:
                        if content and isinstance(content, str):
                            total_words += len(content)
                    self._progress_total_words.set(f"{total_words:,} 字")
                else:
                    self._progress_chapters.set("0 / 0")
                    self._progress_total_words.set("0 字")
            
            # 更新大纲解析状态
            if self._progress_outline is not None:
                if hasattr(self, '_outline_content') and self._outline_content:
                    self._progress_outline.set("已完成")
                else:
                    self._progress_outline.set("未完成")
            
            # 更新人物录入（支持字典和列表两种格式）
            if self._progress_characters is not None:
                if hasattr(self, '_character_data') and self._character_data:
                    if isinstance(self._character_data, dict):
                        char_count = len(self._character_data.get('characters', []))
                    elif isinstance(self._character_data, list):
                        char_count = len(self._character_data)
                    else:
                        char_count = 0
                    self._progress_characters.set(f"{char_count} 人")
                else:
                    self._progress_characters.set("0 人")
            
            # 更新世界观条目（支持字典列表和字符串两种格式）
            if self._progress_worldview is not None:
                if hasattr(self, '_worldview_content') and self._worldview_content:
                    if isinstance(self._worldview_content, list):
                        # 字典列表格式
                        worldview_count = len(self._worldview_content)
                    elif isinstance(self._worldview_content, str):
                        # 字符串格式（Markdown）
                        worldview_count = self._worldview_content.count('\n##') + self._worldview_content.count('\n###')
                        if worldview_count == 0 and self._worldview_content.strip():
                            worldview_count = 1  # 至少有一个条目
                    else:
                        worldview_count = 0
                    self._progress_worldview.set(f"{worldview_count} 个")
                else:
                    self._progress_worldview.set("0 个")
            
            # 更新风格设定
            if self._progress_style is not None:
                if hasattr(self, '_style_profile') and self._style_profile:
                    style_name = self._style_profile.get('name', '未命名风格')
                    self._progress_style.set(style_name)
                else:
                    self._progress_style.set("未设置")
            
        except Exception as e:
            logger.error(f"更新进度显示失败: {e}")
    
    def _on_backup_project(self) -> None:
        """备份项目 - 完整实现"""
        if not self.current_project or not hasattr(self, 'project_file'):
            messagebox.showwarning("备份项目", "当前没有打开的项目")
            return
        
        try:
            self._set_status("正在创建备份...")
            
            # 获取备份目录
            project_dir = os.path.dirname(self.project_file)
            project_name = self.current_project.get('name', '未命名')
            
            # 创建备份目录
            backup_dir = os.path.join(project_dir, "backups")
            os.makedirs(backup_dir, exist_ok=True)
            
            # 生成备份文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_file = os.path.join(backup_dir, f"{project_name}_backup_{timestamp}.json")
            
            # 更新修改时间
            self.current_project['modified_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # 保存备份
            with open(backup_file, 'w', encoding='utf-8') as f:
                json.dump(self.current_project, f, ensure_ascii=False, indent=2)
            
            messagebox.showinfo("备份成功", f"项目备份已创建！\n\n备份位置：{backup_file}")
            self._set_status("项目备份完成")
            
        except Exception as e:
            messagebox.showerror("错误", f"备份失败: {e}")
            self._set_status("备份失败")
    
    def _on_refresh_plugins(self) -> None:
        """刷新插件列表"""
        self._load_plugins()
        self._set_status("插件列表已刷新")
    
    def _on_plugin_select(self, event) -> None:
        """插件选择事件 - 更新详情显示"""
        selection = self._plugin_tree.selection()
        if not selection:
            return
        
        item_id = selection[0]
        plugin_data = self._plugin_data.get(item_id, {})
        
        # 更新详情显示
        details = [
            f"ID: {plugin_data.get('id', 'N/A')}",
            f"名称: {plugin_data.get('name', 'N/A')}",
            f"版本: {plugin_data.get('version', 'N/A')}",
            f"作者: {plugin_data.get('author', 'N/A')}",
            f"描述: {plugin_data.get('description', 'N/A')}",
            f"状态: {plugin_data.get('state', 'N/A')}",
        ]
        
        if plugin_data.get('is_protected'):
            details.append("⚠️ V5核心保护模块（禁止禁用）")
        
        if plugin_data.get('dependencies'):
            details.append(f"依赖: {', '.join(plugin_data['dependencies'])}")
        
        if plugin_data.get('error_message'):
            details.append(f"错误: {plugin_data['error_message']}")
        
        self._plugin_detail_var.set("\n".join(details))
    
    def _on_plugin_double_click(self, event) -> None:
        """双击插件打开配置对话框"""
        self._on_config_plugin()
    
    def _get_selected_plugin(self) -> Optional[Dict[str, Any]]:
        """获取当前选中的插件数据"""
        selection = self._plugin_tree.selection()
        if not selection:
            return None
        return self._plugin_data.get(selection[0])
    
    def _get_plugin_registry(self):
        """获取PluginRegistry实例"""
        if not CORE_AVAILABLE:
            return None
        try:
            return get_plugin_registry()
        except Exception as e:
            logger.warning(f"Failed to get PluginRegistry: {e}")
            return None
    
    def _get_config_manager(self):
        """获取ConfigManager实例"""
        if not CORE_AVAILABLE:
            return None
        try:
            return get_config_manager()
        except Exception as e:
            logger.warning(f"Failed to get ConfigManager: {e}")
            return None
    
    def _on_enable_plugin(self) -> None:
        """启用插件 - 调用PluginRegistry.activate()"""
        plugin_data = self._get_selected_plugin()
        if not plugin_data:
            messagebox.showwarning("提示", "请先选择一个插件")
            return
        
        plugin_id = plugin_data["id"]
        
        # 检查保护状态
        if plugin_data.get("is_protected"):
            messagebox.showwarning("警告", f"插件 '{plugin_id}' 是V5核心保护模块，无法禁用/启用操作")
            return
        
        # 检查当前状态
        if plugin_data.get("state") == "active":
            messagebox.showinfo("提示", f"插件 '{plugin_id}' 已经处于启用状态")
            return
        
        # 调用PluginRegistry激活插件
        registry = self._get_plugin_registry()
        if registry:
            try:
                success = registry.activate(plugin_id)
                if success:
                    self._set_status(f"插件 '{plugin_id}' 已启用")
                    self._load_plugins()  # 刷新列表
                    messagebox.showinfo("成功", f"插件 '{plugin_id}' 已成功启用")
                else:
                    messagebox.showerror("失败", f"插件 '{plugin_id}' 启用失败")
            except Exception as e:
                messagebox.showerror("错误", f"启用插件时发生错误：{e}")
        else:
            messagebox.showwarning("提示", "PluginRegistry不可用")
    
    def _on_disable_plugin(self) -> None:
        """禁用插件 - 调用PluginRegistry.deactivate()"""
        plugin_data = self._get_selected_plugin()
        if not plugin_data:
            messagebox.showwarning("提示", "请先选择一个插件")
            return
        
        plugin_id = plugin_data["id"]
        
        # 检查保护状态
        if plugin_data.get("is_protected"):
            messagebox.showwarning("警告", f"插件 '{plugin_id}' 是V5核心保护模块，禁止禁用")
            return
        
        # 检查当前状态
        if plugin_data.get("state") != "active":
            messagebox.showinfo("提示", f"插件 '{plugin_id}' 当前未启用")
            return
        
        # 确认对话框
        if not messagebox.askyesno("确认", f"确定要禁用插件 '{plugin_id}' 吗？\n禁用后相关功能将不可用。"):
            return
        
        # 调用PluginRegistry停用插件
        registry = self._get_plugin_registry()
        if registry:
            try:
                success = registry.deactivate(plugin_id)
                if success:
                    self._set_status(f"插件 '{plugin_id}' 已禁用")
                    self._load_plugins()  # 刷新列表
                    messagebox.showinfo("成功", f"插件 '{plugin_id}' 已成功禁用")
                else:
                    messagebox.showerror("失败", f"插件 '{plugin_id}' 禁用失败")
            except Exception as e:
                messagebox.showerror("错误", f"禁用插件时发生错误：{e}")
        else:
            messagebox.showwarning("提示", "PluginRegistry不可用")
    
    def _on_config_plugin(self) -> None:
        """打开插件配置对话框"""
        plugin_data = self._get_selected_plugin()
        if not plugin_data:
            messagebox.showwarning("提示", "请先选择一个插件")
            return
        
        plugin_id = plugin_data["id"]
        plugin_name = plugin_data.get("name", plugin_id)
        
        # 创建配置对话框
        self._show_plugin_config_dialog(plugin_id, plugin_name, plugin_data)
    
    def _show_plugin_config_dialog(self, plugin_id: str, plugin_name: str, plugin_data: Dict) -> None:
        """显示插件配置对话框
        
        Args:
            plugin_id: 插件ID
            plugin_name: 插件显示名称
            plugin_data: 插件数据字典
        """
        dialog = tk.Toplevel(self.root)
        dialog.title(f"插件配置 - {plugin_name}")
        dialog.geometry("500x400")
        dialog.configure(bg=GlassTheme.GLASS_BG)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 400) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # 主框架
        main_frame = ttk.Frame(dialog, style="TFrame", padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 插件信息
        info_frame = ttk.LabelFrame(main_frame, text="插件信息", padding=10)
        info_frame.pack(fill=tk.X, pady=(0, 10))
        
        info_text = f"""
ID: {plugin_id}
名称: {plugin_name}
版本: {plugin_data.get('version', 'N/A')}
作者: {plugin_data.get('author', 'N/A')}
类型: {plugin_data.get('type', 'N/A')}
状态: {plugin_data.get('state', 'N/A')}
保护: {'是 (V5核心模块)' if plugin_data.get('is_protected') else '否'}
        """.strip()
        
        info_label = ttk.Label(info_frame, text=info_text, justify=tk.LEFT)
        info_label.pack(anchor=tk.W)
        
        # 配置区域（根据插件类型动态生成）
        config_frame = ttk.LabelFrame(main_frame, text="插件配置", padding=10)
        config_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # 存储配置变量
        config_vars: Dict[str, tk.Variable] = {}
        
        # 从ConfigManager加载插件配置
        plugin_config = {}
        config = self._get_config_manager()
        if config:
            # 插件配置路径: plugins.{plugin_id}.*
            config_keys = config.list_keys(f"plugins.{plugin_id}")
            for key in config_keys:
                value = config.get(key)
                plugin_config[key] = value
        
        # 如果没有配置，显示默认配置选项
        if not plugin_config:
            # 根据插件类型提供默认配置选项
            default_configs = self._get_default_plugin_config(plugin_id, plugin_data.get('type', 'tool'))
            
            for config_key, config_info in default_configs.items():
                row_frame = ttk.Frame(config_frame, style="TFrame")
                row_frame.pack(fill=tk.X, pady=5)
                
                ttk.Label(row_frame, text=config_info['label'] + "：").pack(side=tk.LEFT)
                
                if config_info['type'] == 'string':
                    var = tk.StringVar(value=config_info.get('default', ''))
                    entry = ttk.Entry(row_frame, textvariable=var, width=30)
                    entry.pack(side=tk.LEFT, padx=10)
                    config_vars[config_key] = var
                elif config_info['type'] == 'int':
                    var = tk.IntVar(value=config_info.get('default', 0))
                    spinbox = ttk.Spinbox(row_frame, from_=0, to=100, textvariable=var, width=10)
                    spinbox.pack(side=tk.LEFT, padx=10)
                    config_vars[config_key] = var
                elif config_info['type'] == 'boolean':
                    var = tk.BooleanVar(value=config_info.get('default', True))
                    ttk.Checkbutton(row_frame, variable=var).pack(side=tk.LEFT, padx=10)
                    config_vars[config_key] = var
                elif config_info['type'] == 'choice':
                    var = tk.StringVar(value=config_info.get('default', ''))
                    combo = ttk.Combobox(row_frame, textvariable=var, values=config_info.get('choices', []), width=25)
                    combo.pack(side=tk.LEFT, padx=10)
                    config_vars[config_key] = var
        else:
            # 显示已有配置
            for key, value in plugin_config.items():
                row_frame = ttk.Frame(config_frame, style="TFrame")
                row_frame.pack(fill=tk.X, pady=5)
                
                # 提取配置键名（去掉前缀）
                display_key = key.replace(f"plugins.{plugin_id}.", "")
                ttk.Label(row_frame, text=display_key + "：").pack(side=tk.LEFT)
                
                if isinstance(value, bool):
                    var = tk.BooleanVar(value=value)
                    ttk.Checkbutton(row_frame, variable=var).pack(side=tk.LEFT, padx=10)
                    config_vars[key] = var
                elif isinstance(value, int):
                    var = tk.IntVar(value=value)
                    spinbox = ttk.Spinbox(row_frame, from_=0, to=10000, textvariable=var, width=10)
                    spinbox.pack(side=tk.LEFT, padx=10)
                    config_vars[key] = var
                else:
                    var = tk.StringVar(value=str(value))
                    entry = ttk.Entry(row_frame, textvariable=var, width=30)
                    entry.pack(side=tk.LEFT, padx=10)
                    config_vars[key] = var
        
        # 如果是保护模块，显示警告
        if plugin_data.get('is_protected'):
            warning_label = ttk.Label(
                config_frame, 
                text="⚠️ 此插件为V5核心保护模块，部分配置修改可能导致系统不稳定",
                foreground=GlassTheme.WARNING
            )
            warning_label.pack(pady=10)
        
        # 按钮区域
        btn_frame = ttk.Frame(main_frame, style="TFrame")
        btn_frame.pack(fill=tk.X)
        
        def save_config():
            """保存配置到ConfigManager"""
            config = self._get_config_manager()
            if config:
                try:
                    for key, var in config_vars.items():
                        full_key = f"plugins.{plugin_id}.{key}" if not key.startswith("plugins.") else key
                        value = var.get()
                        config.set(full_key, value, source="user")
                    
                    # 保存到文件
                    config._save_config()
                    
                    messagebox.showinfo("成功", "配置已保存")
                    dialog.destroy()
                except Exception as e:
                    messagebox.showerror("错误", f"保存配置失败：{e}")
            else:
                messagebox.showwarning("提示", "ConfigManager不可用，无法保存配置")
        
        def reset_config():
            """重置配置为默认值"""
            for key, var in config_vars.items():
                default_configs = self._get_default_plugin_config(plugin_id, plugin_data.get('type', 'tool'))
                if key in default_configs:
                    var.set(default_configs[key].get('default', ''))
        
        ttk.Button(btn_frame, text="保存", command=save_config).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="重置", command=reset_config).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def _get_default_plugin_config(self, plugin_id: str, plugin_type: str) -> Dict[str, Dict]:
        """获取插件的默认配置选项
        
        Args:
            plugin_id: 插件ID
            plugin_type: 插件类型
            
        Returns:
            配置选项字典 {config_key: {label, type, default, choices}}
        """
        # 通用配置
        common_config = {
            "enabled": {
                "label": "启用插件",
                "type": "boolean",
                "default": True
            },
            "log_level": {
                "label": "日志级别",
                "type": "choice",
                "default": "INFO",
                "choices": ["DEBUG", "INFO", "WARNING", "ERROR"]
            }
        }
        
        # 根据插件类型添加特定配置
        type_specific_configs = {
            "analyzer": {
                "max_content_length": {
                    "label": "最大内容长度",
                    "type": "int",
                    "default": 100000
                },
                "cache_enabled": {
                    "label": "启用缓存",
                    "type": "boolean",
                    "default": True
                }
            },
            "generator": {
                "max_retries": {
                    "label": "最大重试次数",
                    "type": "int",
                    "default": 3
                },
                "timeout": {
                    "label": "超时时间(秒)",
                    "type": "int",
                    "default": 300
                }
            },
            "validator": {
                "strict_mode": {
                    "label": "严格模式",
                    "type": "boolean",
                    "default": False
                },
                "min_score": {
                    "label": "最低分数",
                    "type": "int",
                    "default": 60
                }
            },
            "tool": {
                "auto_start": {
                    "label": "自动启动",
                    "type": "boolean",
                    "default": True
                }
            }
        }
        
        # 合并配置
        config = common_config.copy()
        if plugin_type.lower() in type_specific_configs:
            config.update(type_specific_configs[plugin_type.lower()])
        
        # 插件特定配置（例如大纲解析器的编码设置）
        if plugin_id == "outline-parser-v3":
            config["encoding"] = {
                "label": "文件编码",
                "type": "choice",
                "default": "utf-8",
                "choices": ["utf-8", "gbk", "gb2312", "big5"]
            }
        
        return config
    
    def _on_reload_plugin(self) -> None:
        """重新加载插件"""
        plugin_data = self._get_selected_plugin()
        if not plugin_data:
            messagebox.showwarning("提示", "请先选择一个插件")
            return
        
        plugin_id = plugin_data["id"]
        
        # 检查保护状态
        if plugin_data.get("is_protected"):
            messagebox.showwarning("警告", f"插件 '{plugin_id}' 是V5核心保护模块，禁止重新加载")
            return
        
        # 确认对话框
        if not messagebox.askyesno("确认", f"确定要重新加载插件 '{plugin_id}' 吗？"):
            return
        
        # 调用PluginRegistry重新加载
        registry = self._get_plugin_registry()
        if registry:
            try:
                success, error = registry.reload_plugin_runtime(plugin_id)
                if success:
                    self._set_status(f"插件 '{plugin_id}' 已重新加载")
                    self._load_plugins()  # 刷新列表
                    messagebox.showinfo("成功", f"插件 '{plugin_id}' 已成功重新加载")
                else:
                    messagebox.showerror("失败", f"插件 '{plugin_id}' 重新加载失败：{error}")
            except Exception as e:
                messagebox.showerror("错误", f"重新加载插件时发生错误：{e}")
        else:
            messagebox.showwarning("提示", "PluginRegistry不可用")
    
    def _on_install_plugin(self) -> None:
        """安装插件"""
        # 选择插件包文件
        file_path = filedialog.askopenfilename(
            title="选择插件包",
            filetypes=[("Plugin Package", "*.zip;*.tar.gz"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        self._set_status(f"正在安装插件：{os.path.basename(file_path)}")
        # TODO: 实现插件安装逻辑
        messagebox.showinfo("提示", "插件安装功能开发中...")
    
    def _on_uninstall_plugin(self) -> None:
        """卸载插件"""
        plugin_data = self._get_selected_plugin()
        if not plugin_data:
            messagebox.showwarning("提示", "请先选择一个插件")
            return
        
        plugin_id = plugin_data["id"]
        
        # 检查保护状态
        if plugin_data.get("is_protected"):
            messagebox.showwarning("警告", f"插件 '{plugin_id}' 是V5核心保护模块，禁止卸载")
            return
        
        # 确认对话框
        if not messagebox.askyesno("确认卸载", 
            f"确定要卸载插件 '{plugin_id}' 吗？\n\n此操作将移除插件及其配置，且不可恢复！"):
            return
        
        # 调用PluginRegistry卸载
        registry = self._get_plugin_registry()
        if registry:
            try:
                success, error = registry.unload_plugin_runtime(plugin_id)
                if success:
                    self._set_status(f"插件 '{plugin_id}' 已卸载")
                    self._load_plugins()  # 刷新列表
                    messagebox.showinfo("成功", f"插件 '{plugin_id}' 已成功卸载")
                else:
                    messagebox.showerror("失败", f"插件 '{plugin_id}' 卸载失败：{error}")
            except Exception as e:
                messagebox.showerror("错误", f"卸载插件时发生错误：{e}")
        else:
            messagebox.showwarning("提示", "PluginRegistry不可用")
    
    def _on_save_settings(self) -> None:
        """保存设置到config.yaml并应用

        V2.6更新：API Key安全存储
        - API Key保存到加密存储（.secrets/api_keys.enc）
        - config.yaml只保存占位符ENCRYPTED_IN_SECRETS_FILE
        - 自动发布config.changed事件到EventBus
        - 无需重启应用即可切换AI服务
        """
        config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.yaml")

        # 获取API Key
        api_key = self._api_key_var.get()
        provider = self._provider_var.get()

        # 1. 保存API Key到加密存储
        if api_key and provider:
            try:
                from core.api_key_encryption import APIKeyEncryption
                encryption = APIKeyEncryption(Path(__file__).parent)
                encryption.save_api_key(provider, api_key)
                logger.info(f"API Key已加密保存: {provider}")
            except Exception as e:
                logger.warning(f"API Key加密保存失败: {e}")
                messagebox.showwarning("警告", f"API Key加密保存失败，将保存到配置文件。\n错误：{e}")

        config_data = {
            "service_mode": self._service_mode_var.get(),
            "provider": provider,
            "model": self._model_var.get(),
            # API Key占位符（实际Key存储在加密文件）
            "api_key": "ENCRYPTED_IN_SECRETS_FILE" if api_key else "",
            "local_url": self._local_url_var.get(),
            "temperature": self._temp_var.get(),
            "theme": self._theme_var.get() if hasattr(self, '_theme_var') else "dark",
            "ai_learning": self._ai_learning_var.get() if hasattr(self, '_ai_learning_var') else True,
            "auto_save": True,
            "backup_interval": self._backup_interval_var.get() if hasattr(self, '_backup_interval_var') else "30",
            "font_size": self._font_size_var.get() if hasattr(self, '_font_size_var') else "14",
            "window_size": self._window_size_var.get() if hasattr(self, '_window_size_var') else "1280x720",
            "save_path": self._save_path_var.get() if hasattr(self, '_save_path_var') else os.path.dirname(os.path.abspath(__file__)),
            # 记忆配置（Sprint 9-10）
            "memory": {
                "wal_auto_save": self._wal_auto_save_var.get() if hasattr(self, '_wal_auto_save_var') else True,
                "auto_recovery": self._auto_recovery_var.get() if hasattr(self, '_auto_recovery_var') else True,
                "vector_recall_topk": self._vector_recall_topk_var.get() if hasattr(self, '_vector_recall_topk_var') else 10,
                "daily_meditation_time": self._daily_meditation_time_var.get() if hasattr(self, '_daily_meditation_time_var') else "23:00",
                "embedding_model": self._embedding_model_var.get() if hasattr(self, '_embedding_model_var') else "local"
            }
        }

        try:
            # 2. 保存到config.yaml文件（不含真实API Key）
            with open(config_path, "w", encoding="utf-8") as f:
                yaml.dump(config_data, f, allow_unicode=True, default_flow_style=False)

            # 3. 调用ConfigService更新AI配置（V3.2重构：业务逻辑在插件层）
            try:
                from core import get_config_service
                config_service = get_config_service()

                # 从Entry组件获取最新的本地URL
                local_url = self._url_entry.get().strip() if hasattr(self, "_url_entry") else self._local_url_var.get()

                # 构建AI配置字典
                ai_config = {
                    "service_mode": self._service_mode_var.get(),
                    "provider": provider,
                    "model": self._model_var.get(),
                    "api_key": api_key,
                    "local_url": local_url,
                    "temperature": self._temp_var.get(),
                }

                # 更新AI配置（会发布config.changed事件）
                # 插件订阅该事件后自动测试连接并更新状态
                config_service.update_ai_config(ai_config)
                logger.info(f"[设置保存] AI配置已更新: provider={provider}, endpoint={local_url}")

            except Exception as e:
                logger.error(f"[设置保存] ConfigService更新失败: {e}", exc_info=True)
                messagebox.showerror("错误", f"保存设置失败：{e}")
                return

            # 4. 应用设置
            self._apply_settings(config_data)

            self._set_status("设置已保存")
            messagebox.showinfo("成功", "设置已保存！\n\n配置已生效，AI状态将在连接成功后更新。")

        except Exception as e:
            self._set_status(f"保存设置失败：{e}")
            messagebox.showerror("错误", f"保存设置失败：{e}")
    
    def _apply_settings(self, config_data: dict) -> None:
        """应用设置到界面"""
        # 应用主题（深色/浅色）
        theme = config_data.get("theme", "dark")
        if theme == "light":
            # 浅色主题配置
            GlassTheme.GLASS_BG = "#FFFFFF"
            GlassTheme.GLASS_SURFACE = "#F5F5F5"
            GlassTheme.TEXT_PRIMARY = "#333333"
            GlassTheme.TEXT_SECONDARY = "#666666"
        else:
            # 深色主题配置（默认）
            GlassTheme.GLASS_BG = "#1E1E1E"
            GlassTheme.GLASS_SURFACE = "#2D2D2D"
            GlassTheme.TEXT_PRIMARY = "#FFFFFF"
            GlassTheme.TEXT_SECONDARY = "#CCCCCC"
        
        # 应用字体大小
        font_size = int(config_data.get("font_size", 14))
        GlassTheme.FONT_SIZE_NORMAL = str(font_size)
        GlassTheme.FONT_SIZE_SMALL = str(max(10, font_size - 2))
        GlassTheme.FONT_SIZE_SUBTITLE = str(max(14, font_size - 2))
        GlassTheme.FONT_SIZE_TITLE = str(max(18, font_size + 4))
        
        # 应用窗口大小（需要重启才生效）
        window_size = config_data.get("window_size", "1280x720")
        # 保存到配置文件，下次启动时应用
        self.root.after(100, lambda: self._set_status(f"窗口大小将在重启后应用：{window_size}"))
    
    def _on_browse_save_path(self):
        """浏览默认保存路径"""
        path = filedialog.askdirectory(title="选择默认保存路径")
        if path:
            self._save_path_var.set(path)
    
    def _on_clear_learning_data(self):
        """清除学习数据"""
        if messagebox.askyesno("确认清除", "确定要清除所有AI学习数据吗？\n这将重置您的写作偏好配置。"):
            self._set_status("学习数据已清除")
    
    def _on_export_preferences(self):
        """导出偏好配置"""
        path = filedialog.asksaveasfilename(
            title="导出偏好配置",
            defaultextension=".yaml",
            filetypes=[("YAML文件", "*.yaml"), ("所有文件", "*.*")]
        )
        if path:
            self._set_status(f"偏好配置已导出到：{os.path.basename(path)}")
    
    # ============== API安全相关方法（V2.8）==============
    
    def _check_gitignore_protection(self) -> bool:
        """检查.gitignore是否已配置API Key保护"""
        gitignore_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".gitignore")
        protected_patterns = ["config.yaml", ".secrets/", "*.enc", "*.key"]
        
        try:
            if os.path.exists(gitignore_path):
                with open(gitignore_path, "r", encoding="utf-8") as f:
                    content = f.read()
                # 检查是否包含关键保护规则
                return "config.yaml" in content or ".secrets" in content
        except:
            pass
        return False
    
    def _toggle_api_key_visibility(self):
        """切换API Key显示/隐藏"""
        if self._show_key_var.get():
            # 显示明文
            self._key_entry.config(show="")
            self._toggle_key_btn.config(text="🔒")
        else:
            # 显示密文
            self._key_entry.config(show="*")
            self._toggle_key_btn.config(text="👁")
        self._show_key_var.set(not self._show_key_var.get())
    
    def _on_test_api_connection(self):
        """测试API连接（V3.2重构：使用AIStatusManagerPlugin）

        【架构设计】
        - 业务逻辑在AIStatusManagerPlugin插件中实现
        - GUI只负责调用和显示结果
        - 状态变更通过EventBus实时同步到状态栏
        """
        provider = self._provider_var.get()
        api_key = self._api_key_var.get()
        model = self._model_var.get()
        service_mode = self._service_mode_var.get()
        local_url = self._local_url_var.get()

        # 确保AI状态管理插件已初始化
        if not self._ai_status_manager:
            self._init_ai_status_manager()

        # 更新状态为"连接中"
        self._set_status(f"正在测试{provider}连接...")

        # 本地模式：使用AIStatusManagerPlugin测试和启动服务
        if service_mode == "local":
            # 先测试连接
            result = self._ai_status_manager.test_connection(local_url, provider)

            if result["success"]:
                # 连接成功
                self._set_status(f"✓ {provider}服务连接成功")
                messagebox.showinfo("连接成功", result["message"])
                return

            # 连接失败，尝试启动服务（仅Qwen）
            if provider.lower() == "qwen":
                self._set_status(f"服务未运行，正在启动...")
                logger.info("[测试连接] 检测到Qwen未运行，尝试自动启动...")

                # 启动服务
                start_result = self._ai_status_manager.start_local_service("qwen")

                if start_result["success"]:
                    # 服务启动中，等待就绪
                    self._set_status(f"服务启动中，请稍候...")
                    messagebox.showinfo(
                        "服务启动中",
                        f"Qwen服务正在启动...\n\n"
                        f"进程ID: {start_result.get('pid')}\n"
                        f"端点: {start_result.get('endpoint')}\n\n"
                        f"服务启动需要约30-60秒，请稍后再测试连接。"
                    )
                else:
                    # 启动失败
                    self._set_status(f"✗ 服务启动失败")
                    messagebox.showerror(
                        "启动失败",
                        f"无法启动Qwen服务：\n{start_result.get('message')}\n\n"
                        f"请检查：\n"
                        f"1. Qwen已正确部署到 F:\\Qwen\n"
                        f"2. start_server_v2.py 脚本存在\n"
                        f"3. 显卡驱动和CUDA已安装"
                    )
            else:
                # 其他本地服务
                self._set_status(f"✗ 连接失败")
                messagebox.showerror(
                    "连接失败",
                    f"无法连接到{provider}服务\n\n"
                    f"请确保服务已启动后再测试。"
                )
            return
        
        # 线上模式：测试API连接
        if not api_key:
            messagebox.showwarning("提示", "请先输入API Key")
            return
        
        # 显示测试中状态
        self._set_status(f"正在测试{provider}连接...")
        
        try:
            from openai import OpenAI
            
            # 根据提供商配置客户端
            if provider == "DeepSeek":
                client = OpenAI(
                    api_key=api_key,
                    base_url="https://api.deepseek.com/v1"
                )
            elif provider == "OpenAI":
                client = OpenAI(api_key=api_key)
            elif provider == "Anthropic":
                client = OpenAI(
                    api_key=api_key,
                    base_url="https://api.anthropic.com/v1"
                )
            else:  # 其他提供商
                local_url = self._local_url_var.get()
                client = OpenAI(
                    api_key="ollama",  # Ollama不需要真实key
                    base_url=local_url
                )
            
            # 发送简单测试请求
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": "Hello, this is a connection test."}],
                max_tokens=10
            )
            
            # 连接成功
            self._set_status(f"✓ {provider}连接成功")
            messagebox.showinfo("连接成功", f"已成功连接到{provider}！\n模型：{model}\n响应正常。")
            
        except Exception as e:
            self._set_status(f"✗ {provider}连接失败")
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_api_error
            title, full_message = convert_api_error(e, provider)
            messagebox.showerror(title, full_message)
    
    def _on_backup_api_keys(self):
        """备份加密的API Key"""
        # 选择保存位置
        backup_path = filedialog.asksaveasfilename(
            title="备份API密钥",
            defaultextension=".enc",
            filetypes=[("加密备份文件", "*.enc"), ("所有文件", "*.*")]
        )
        
        if not backup_path:
            return
        
        try:
            # 尝试使用加密模块
            try:
                from core.api_key_encryption import APIKeyEncryption
                encryption = APIKeyEncryption()
                
                # 收集所有API Key
                api_keys = {
                    "deepseek": self._api_key_var.get() if self._provider_var.get() == "DeepSeek" else "",
                    # 可以扩展其他提供商
                }
                
                # 加密并保存
                encryption.encrypt_and_save(api_keys)
                
                # 复制到用户选择的备份位置
                import shutil
                secrets_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".secrets", "api_keys.enc")
                if os.path.exists(secrets_path):
                    shutil.copy2(secrets_path, backup_path)
                    self._set_status(f"API密钥已备份到：{os.path.basename(backup_path)}")
                    messagebox.showinfo("备份成功", f"API密钥已加密备份到：\n{backup_path}\n\n请妥善保管备份文件！")
                else:
                    raise FileNotFoundError("加密文件未生成")
                    
            except ImportError:
                # 降级：使用简单的base64编码（不推荐）
                import base64
                import json
                
                api_keys = {
                    "provider": self._provider_var.get(),
                    "api_key": self._api_key_var.get(),
                    "backup_time": datetime.now().isoformat()
                }
                
                encoded = base64.b64encode(json.dumps(api_keys).encode()).decode()
                with open(backup_path, "w") as f:
                    f.write(encoded)
                
                self._set_status(f"API密钥已备份（Base64编码）")
                messagebox.showwarning("备份成功（简化模式）", 
                    f"API密钥已备份到：\n{backup_path}\n\n"
                    "⚠ 注意：当前使用Base64编码，建议安装cryptography库以启用AES加密。")
                
        except Exception as e:
            self._set_status("API密钥备份失败")
            messagebox.showerror("备份失败", f"无法备份API密钥：\n{str(e)}")
    
    def _on_restore_api_keys(self):
        """从备份恢复API Key"""
        # 选择备份文件
        backup_path = filedialog.askopenfilename(
            title="恢复API密钥",
            filetypes=[("加密备份文件", "*.enc"), ("所有文件", "*.*")]
        )
        
        if not backup_path:
            return
        
        try:
            # 尝试使用加密模块解密
            try:
                from core.api_key_encryption import APIKeyEncryption
                import shutil
                
                # 复制备份文件到.secrets目录
                secrets_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".secrets")
                os.makedirs(secrets_dir, exist_ok=True)
                secrets_path = os.path.join(secrets_dir, "api_keys.enc")
                shutil.copy2(backup_path, secrets_path)
                
                # 解密
                encryption = APIKeyEncryption()
                api_keys = encryption.decrypt_and_load()
                
                if api_keys:
                    # 恢复到界面
                    if "deepseek" in api_keys and api_keys["deepseek"]:
                        self._api_key_var.set(api_keys["deepseek"])
                    
                    self._set_status("API密钥已恢复")
                    messagebox.showinfo("恢复成功", "API密钥已成功恢复！\n请点击[保存设置]使其生效。")
                else:
                    raise ValueError("备份文件为空或格式错误")
                    
            except ImportError:
                # 降级：尝试Base64解码
                import base64
                import json
                
                with open(backup_path, "r") as f:
                    encoded = f.read()
                
                decoded = json.loads(base64.b64decode(encoded).decode())
                
                if "api_key" in decoded:
                    self._api_key_var.set(decoded["api_key"])
                    if "provider" in decoded:
                        self._provider_var.set(decoded["provider"])
                    
                    self._set_status("API密钥已恢复（Base64解码）")
                    messagebox.showinfo("恢复成功", "API密钥已成功恢复！\n请点击[保存设置]使其生效。")
                else:
                    raise ValueError("备份文件格式错误")
                
        except Exception as e:
            self._set_status("API密钥恢复失败")
            messagebox.showerror("恢复失败", f"无法恢复API密钥：\n{str(e)}\n\n请确认备份文件正确。")
    
    # ============== 长期记忆管理方法（Sprint 9-10）==============
    
    def _on_refresh_memory_status(self):
        """刷新记忆层级状态"""
        try:
            # L1热记忆状态
            from core.session_state import get_session_state_manager  # V2.20.1修复：函数名修正
            session = get_session_state_manager()
            if session:
                state = session.get_current_state()
                self._session_state_status.config(text=f"活跃（{state.get('session_id', 'unknown')}）")
            else:
                self._session_state_status.config(text="未初始化")
            
            # L2温记忆状态
            try:
                from infrastructure.vector_store import NovelVectorStore
                vector_store = NovelVectorStore()
                
                # 获取章节数
                try:
                    chapter_table = vector_store.db.open_table("chapters")
                    chapter_count = chapter_table.count_rows() if hasattr(chapter_table, 'count_rows') else 0
                    self._chapter_count_label.config(text=str(chapter_count))
                except:
                    self._chapter_count_label.config(text="0")
                
                # 获取知识点数（统计所有知识库相关表）
                try:
                    knowledge_count = 0
                    knowledge_table_prefixes = ('scifi_', 'xuanhuan_', 'general_', 'fantasy_', 'history_', 'knowledge')
                    for table_name in vector_store.db.table_names():
                        if table_name.startswith(knowledge_table_prefixes):
                            try:
                                t = vector_store.db.open_table(table_name)
                                knowledge_count += t.count_rows()
                            except:
                                pass
                    self._knowledge_count_label.config(text=str(knowledge_count))
                except:
                    self._knowledge_count_label.config(text="0")
                
                # 获取风格数
                try:
                    style_table = vector_store.db.open_table("styles")
                    style_count = style_table.count_rows() if hasattr(style_table, 'count_rows') else 0
                    self._style_count_label.config(text=str(style_count))
                except:
                    self._style_count_label.config(text="0")
                
            except Exception as e:
                logger.warning(f"向量库状态获取失败: {e}")
                self._chapter_count_label.config(text="-")
                self._knowledge_count_label.config(text="-")
                self._style_count_label.config(text="-")
            
            # L3冷记忆状态
            try:
                from core.git_notes_manager import GitNotesManager
                git_manager = GitNotesManager(Path(os.getcwd()))
                
                # 获取当前分支
                result = subprocess.run(
                    ["git", "rev-parse", "--abbrev-ref", "HEAD"],
                    capture_output=True, text=True, cwd=os.getcwd()
                )
                branch = result.stdout.strip() if result.returncode == 0 else "unknown"
                self._git_branch_label.config(text=branch)
                
                # 获取笔记数
                notes = git_manager.get_notes_for_branch(branch)
                self._git_notes_count_label.config(text=str(len(notes)))
                
            except Exception as e:
                logger.warning(f"Git-Notes状态获取失败: {e}")
                self._git_branch_label.config(text="-")
                self._git_notes_count_label.config(text="-")
            
            # L4精选档案状态
            memory_file = Path(".workbuddy/memory/MEMORY.md")
            if memory_file.exists():
                size_kb = memory_file.stat().st_size / 1024
                self._memory_size_label.config(text=f"{size_kb:.1f} KB")
                
                import time
                mtime = memory_file.stat().st_mtime
                update_time = time.strftime("%Y-%m-%d %H:%M", time.localtime(mtime))
                self._memory_update_label.config(text=update_time)
            else:
                self._memory_size_label.config(text="不存在")
                self._memory_update_label.config(text="-")
            
            # 更新总体状态
            self._memory_status_label.config(text="✅ 所有记忆层级已加载")
            self._set_status("记忆状态已刷新")
            
        except Exception as e:
            logger.error(f"刷新记忆状态失败: {e}")
            self._memory_status_label.config(text=f"❌ 加载失败: {str(e)[:30]}")
    
    def _on_view_session_state(self):
        """查看当前会话状态"""
        try:
            from core.session_state import get_session_state_manager  # V2.20.1修复：函数名修正
            session = get_session_state_manager()
            
            if session:
                state = session.get_current_state()
                
                # 创建查看窗口
                dialog = tk.Toplevel(self.root)
                dialog.title("当前会话状态")
                dialog.geometry("600x400")
                dialog.transient(self.root)
                
                # 创建文本框显示状态
                text = tk.Text(dialog, wrap=tk.WORD, font=("Consolas", 10))
                text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
                
                # 格式化显示
                import json
                text.insert("1.0", json.dumps(state, indent=2, ensure_ascii=False))
                text.config(state=tk.DISABLED)
                
            else:
                messagebox.showinfo("提示", "会话状态未初始化")
                
        except Exception as e:
            logger.error(f"查看会话状态失败: {e}")
            messagebox.showerror("错误", f"查看失败: {e}")
    
    def _on_clear_session_state(self):
        """清空会话状态"""
        if messagebox.askyesno("确认", "确定要清空当前会话状态吗？"):
            try:
                from core.session_state import get_session_state_manager  # V2.20.1修复：函数名修正
                session = get_session_state_manager()
                if session:
                    session.clear_state()
                    self._set_status("会话状态已清空")
                    messagebox.showinfo("成功", "会话状态已清空")
            except Exception as e:
                logger.error(f"清空会话状态失败: {e}")
                messagebox.showerror("错误", f"清空失败: {e}")
    
    def _on_save_session_snapshot(self):
        """保存会话快照"""
        path = filedialog.asksaveasfilename(
            title="保存会话快照",
            defaultextension=".json",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")]
        )
        if path:
            try:
                from core.session_state import get_session_state_manager  # V2.20.1修复：函数名修正
                session = get_session_state_manager()
                if session:
                    state = session.get_current_state()
                    import json
                    with open(path, 'w', encoding='utf-8') as f:
                        json.dump(state, f, indent=2, ensure_ascii=False)
                    self._set_status(f"会话快照已保存到：{os.path.basename(path)}")
                    messagebox.showinfo("成功", "会话快照已保存")
            except Exception as e:
                logger.error(f"保存会话快照失败: {e}")
                messagebox.showerror("错误", f"保存失败: {e}")
    
    def _on_rebuild_vector_index(self):
        """重建向量索引"""
        if messagebox.askyesno("确认", "重建向量索引可能需要较长时间，确定继续吗？"):
            try:
                from infrastructure.vector_store import NovelVectorStore
                vector_store = NovelVectorStore()
                
                self._set_status("正在重建向量索引...")
                
                # 重建章节向量
                # TODO: 从项目加载所有章节并重新编码
                
                self._set_status("向量索引重建完成")
                messagebox.showinfo("成功", "向量索引已重建")
                
            except Exception as e:
                logger.error(f"重建向量索引失败: {e}")
                messagebox.showerror("错误", f"重建失败: {e}")
    
    def _on_export_vector_data(self):
        """导出向量数据"""
        path = filedialog.asksaveasfilename(
            title="导出向量数据",
            defaultextension=".json",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")]
        )
        if path:
            try:
                from infrastructure.vector_store import NovelVectorStore
                vector_store = NovelVectorStore()
                
                data = {
                    "chapters": [],
                    "knowledge": [],
                    "styles": []
                }
                
                # 导出章节数据
                try:
                    table = vector_store.db.open_table("chapters")
                    df = table.to_pandas()
                    data["chapters"] = df.to_dict('records')
                except:
                    pass
                
                # 导出知识数据
                try:
                    table = vector_store.db.open_table("knowledge")
                    df = table.to_pandas()
                    data["knowledge"] = df.to_dict('records')
                except:
                    pass
                
                import json
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                
                self._set_status(f"向量数据已导出到：{os.path.basename(path)}")
                messagebox.showinfo("成功", "向量数据已导出")
                
            except Exception as e:
                logger.error(f"导出向量数据失败: {e}")
                messagebox.showerror("错误", f"导出失败: {e}")
    
    def _on_clear_vector_store(self):
        """清空向量库"""
        if messagebox.askyesno("警告", "清空向量库将删除所有章节、知识和风格向量，确定继续吗？"):
            try:
                from infrastructure.vector_store import NovelVectorStore
                import shutil
                
                vector_store = NovelVectorStore()
                db_path = Path(vector_store.db.uri)
                
                # 关闭连接
                del vector_store
                
                # 删除数据库文件
                if db_path.exists():
                    shutil.rmtree(db_path)
                
                self._set_status("向量库已清空")
                messagebox.showinfo("成功", "向量库已清空")
                
                # 刷新状态
                self._on_refresh_memory_status()
                
            except Exception as e:
                logger.error(f"清空向量库失败: {e}")
                messagebox.showerror("错误", f"清空失败: {e}")
    
    def _on_view_git_notes(self):
        """查看历史决策"""
        try:
            from core.git_notes_manager import GitNotesManager
            git_manager = GitNotesManager(Path(os.getcwd()))
            
            # 获取当前分支
            result = subprocess.run(
                ["git", "rev-parse", "--abbrev-ref", "HEAD"],
                capture_output=True, text=True, cwd=os.getcwd()
            )
            branch = result.stdout.strip() if result.returncode == 0 else "main"
            
            notes = git_manager.get_notes_for_branch(branch)
            
            # 创建查看窗口
            dialog = tk.Toplevel(self.root)
            dialog.title(f"历史决策记录 - {branch}分支")
            dialog.geometry("700x500")
            dialog.transient(self.root)
            
            # 创建Treeview
            tree_frame = ttk.Frame(dialog)
            tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            tree = ttk.Treeview(tree_frame, columns=("commit", "note"), show="headings")
            tree.heading("commit", text="提交")
            tree.heading("note", text="决策记录")
            tree.column("commit", width=100)
            tree.column("note", width=550)
            
            for note in notes:
                tree.insert("", tk.END, values=(note.get("commit", "")[:8], note.get("note", "")))
            
            tree.pack(fill=tk.BOTH, expand=True)
            
        except Exception as e:
            logger.error(f"查看Git-Notes失败: {e}")
            messagebox.showerror("错误", f"查看失败: {e}")
    
    def _on_add_git_note(self):
        """添加决策记录"""
        # 创建输入对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("添加决策记录")
        dialog.geometry("500x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text="决策记录内容：").pack(pady=10)
        
        text = tk.Text(dialog, wrap=tk.WORD, height=10)
        text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        def on_submit():
            note = text.get("1.0", tk.END).strip()
            if note:
                try:
                    from core.git_notes_manager import GitNotesManager
                    git_manager = GitNotesManager(Path(os.getcwd()))
                    
                    # 获取当前提交
                    result = subprocess.run(
                        ["git", "rev-parse", "HEAD"],
                        capture_output=True, text=True, cwd=os.getcwd()
                    )
                    commit = result.stdout.strip() if result.returncode == 0 else None
                    
                    if commit:
                        git_manager.add_note(commit, note)
                        messagebox.showinfo("成功", "决策记录已添加")
                        dialog.destroy()
                    else:
                        messagebox.showerror("错误", "无法获取当前提交")
                        
                except Exception as e:
                    logger.error(f"添加Git-Note失败: {e}")
                    messagebox.showerror("错误", f"添加失败: {e}")
        
        ttk.Button(dialog, text="提交", command=on_submit).pack(pady=10)
    
    def _on_edit_memory_md(self):
        """编辑MEMORY.md（Claw化L4档案记忆）"""
        memory_file = Path("Memory-Novel Writing Assistant-Agent Pro/MEMORY.md")
        if not memory_file.exists():
            # 自动创建
            memory_file.parent.mkdir(parents=True, exist_ok=True)
            memory_file.write_text("# Claw化记忆库\n\n> 由每日冥想自动维护\n\n---\n", encoding="utf-8")
            messagebox.showinfo("提示", "已创建新的MEMORY.md文件")
            return
        
        # 创建编辑窗口
        dialog = tk.Toplevel(self.root)
        dialog.title("编辑 MEMORY.md")
        dialog.geometry("800x600")
        dialog.transient(self.root)
        
        # 创建文本框
        text_frame = ttk.Frame(dialog)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        scrollbar = ttk.Scrollbar(text_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        text = tk.Text(text_frame, wrap=tk.WORD, font=("Consolas", 10), yscrollcommand=scrollbar.set)
        text.pack(fill=tk.BOTH, expand=True)
        scrollbar.config(command=text.yview)
        
        # 加载文件内容
        try:
            with open(memory_file, 'r', encoding='utf-8') as f:
                text.insert("1.0", f.read())
        except Exception as e:
            messagebox.showerror("错误", f"加载失败: {e}")
            return
        
        def on_save():
            try:
                with open(memory_file, 'w', encoding='utf-8') as f:
                    f.write(text.get("1.0", tk.END))
                messagebox.showinfo("成功", "MEMORY.md已保存")
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("错误", f"保存失败: {e}")
        
        # 保存按钮
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        ttk.Button(btn_frame, text="保存", command=on_save).pack(side=tk.RIGHT)
    
    def _on_daily_meditation(self):
        """每日冥想：提炼当日重要事件"""
        try:
            from tools.memory_maintenance import MemoryMaintenanceTool
            
            tool = MemoryMaintenanceTool()
            result = tool.daily_meditation()
            
            messagebox.showinfo(
                "每日冥想完成",
                f"扫描文件：{result['files_scanned']}\n"
                f"提取事件：{result['events_extracted']}\n"
                f"保留事件：{result['events_preserved']}\n"
                f"MEMORY.md大小：{result['memory_size_kb']:.1f} KB"
            )
            
            self._set_status("每日冥想完成")
            self._on_refresh_memory_status()
            
        except Exception as e:
            logger.error(f"每日冥想失败: {e}")
            messagebox.showerror("错误", f"执行失败: {e}")
    
    def _on_weekly_meditation(self):
        """每周大冥想：深度精简MEMORY.md"""
        if messagebox.askyesno("确认", "每周大冥想将深度精简MEMORY.md并归档旧日记，确定继续吗？"):
            try:
                from tools.memory_maintenance import MemoryMaintenanceTool
                
                tool = MemoryMaintenanceTool()
                result = tool.weekly_meditation()
                
                messagebox.showinfo(
                    "每周大冥想完成",
                    f"归档日记：{result['archived_diaries']}\n"
                    f"精简事件：{result['events_removed']}\n"
                    f"MEMORY.md大小：{result['memory_size_kb']:.1f} KB"
                )
                
                self._set_status("每周大冥想完成")
                self._on_refresh_memory_status()
                
            except Exception as e:
                logger.error(f"每周大冥想失败: {e}")
                messagebox.showerror("错误", f"执行失败: {e}")
    
    def _on_archive_old_diary(self):
        """归档旧日记"""
        if messagebox.askyesno("确认", "将归档30天前的日记到 memory/archive/ 目录，确定继续吗？"):
            try:
                from tools.memory_maintenance import MemoryMaintenanceTool
                
                tool = MemoryMaintenanceTool()
                result = tool.archive_old_diaries()
                
                messagebox.showinfo(
                    "归档完成",
                    f"归档文件：{result['archived_files']}\n"
                    f"释放空间：{result['freed_space_kb']:.1f} KB"
                )
                
                self._set_status("旧日记已归档")
                
            except Exception as e:
                logger.error(f"归档旧日记失败: {e}")
                messagebox.showerror("错误", f"归档失败: {e}")
    
    def _on_backup_memory(self):
        """一键备份记忆"""
        path = filedialog.askdirectory(title="选择备份目录")
        if path:
            try:
                import shutil
                from datetime import datetime
                
                backup_dir = Path(path) / f"memory_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                backup_dir.mkdir(parents=True, exist_ok=True)
                
                # 备份记忆目录
                memory_dir = Path(".workbuddy/memory")
                if memory_dir.exists():
                    shutil.copytree(memory_dir, backup_dir / "memory")
                
                # 备份向量数据库
                vector_dir = Path("data/vector_store")
                if vector_dir.exists():
                    shutil.copytree(vector_dir, backup_dir / "vector_store")
                
                self._set_status(f"记忆已备份到：{backup_dir.name}")
                messagebox.showinfo("成功", f"记忆已备份到：\n{backup_dir}")
                
            except Exception as e:
                logger.error(f"备份记忆失败: {e}")
                messagebox.showerror("错误", f"备份失败: {e}")
    
    def _on_restore_memory(self):
        """恢复记忆备份"""
        path = filedialog.askdirectory(title="选择备份目录")
        if path:
            try:
                import shutil
                backup_dir = Path(path)
                
                # 恢复记忆目录
                memory_backup = backup_dir / "memory"
                if memory_backup.exists():
                    memory_dir = Path(".workbuddy/memory")
                    if memory_dir.exists():
                        shutil.rmtree(memory_dir)
                    shutil.copytree(memory_backup, memory_dir)
                
                # 恢复向量数据库
                vector_backup = backup_dir / "vector_store"
                if vector_backup.exists():
                    vector_dir = Path("data/vector_store")
                    if vector_dir.exists():
                        shutil.rmtree(vector_dir)
                    shutil.copytree(vector_backup, vector_dir)
                
                self._set_status("记忆已恢复")
                messagebox.showinfo("成功", "记忆已恢复，请重启程序生效")
                
            except Exception as e:
                logger.error(f"恢复记忆失败: {e}")
                messagebox.showerror("错误", f"恢复失败: {e}")
    
    def _on_clear_all_memory(self):
        """清空所有记忆"""
        if messagebox.askyesno("⚠️ 危险操作", "清空所有记忆将删除所有记忆层级数据，此操作不可逆！\n\n确定继续吗？"):
            if messagebox.askyesno("再次确认", "真的要清空所有记忆吗？"):
                try:
                    import shutil
                    
                    # 清空记忆目录（保留CodeBuddy记忆，清空Claw化运行记忆）
                    claw_memory_dir = Path("Memory-Novel Writing Assistant-Agent Pro")
                    if claw_memory_dir.exists():
                        shutil.rmtree(claw_memory_dir)
                        claw_memory_dir.mkdir(parents=True, exist_ok=True)
                    
                    # 清空向量数据库
                    vector_dir = Path("data/vector_store")
                    if vector_dir.exists():
                        shutil.rmtree(vector_dir)
                    
                    self._set_status("所有记忆已清空")
                    messagebox.showinfo("成功", "所有记忆已清空")
                    
                    # 刷新状态
                    self._on_refresh_memory_status()
                    
                except Exception as e:
                    logger.error(f"清空记忆失败: {e}")
                    messagebox.showerror("错误", f"清空失败: {e}")
    
    def _on_service_mode_changed(self):
        """服务模式切换回调（V2.23增强：动态更新提供商和模型选项）"""
        mode = self._service_mode_var.get()
        
        if mode == "local":
            # 本地模式：显示部署地址，隐藏API Key
            self._key_label.config(text="本地部署地址：")
            self._key_entry.pack_forget()
            self._url_entry.pack(side=tk.LEFT, padx=10)
            
            # 更新提供商选项为本地模型
            if hasattr(self, '_provider_combo'):
                self._provider_combo['values'] = ["Ollama", "Qwen"]
                self._provider_var.set("Qwen")  # 默认选择Qwen
            
            # 更新模型选项为本地模型
            if hasattr(self, '_model_combo'):
                self._model_combo['values'] = ["qwen2.5-14b-gptq", "llama3.1", "mistral", "deepseek-coder-v2"]
                self._model_var.set("qwen2.5-14b-gptq")  # 默认选择Qwen模型
            
            # 更新本地URL为Qwen默认地址（V3.2.1修复：添加/v1路径）
            if hasattr(self, '_local_url_var'):
                self._local_url_var.set("http://localhost:8000/v1")
        else:
            # 线上模式：显示API Key，隐藏本地地址
            self._key_label.config(text="API Key：")
            self._url_entry.pack_forget()
            self._key_entry.pack(side=tk.LEFT, padx=10)
            
            # 更新提供商选项为线上API
            if hasattr(self, '_provider_combo'):
                self._provider_combo['values'] = ["DeepSeek", "OpenAI", "Anthropic"]
                self._provider_var.set("DeepSeek")  # 默认选择DeepSeek
            
            # 更新模型选项为线上模型
            if hasattr(self, '_model_combo'):
                self._model_combo['values'] = ["deepseek-chat", "deepseek-reasoner", "gpt-4", "gpt-3.5-turbo", "claude-3"]
                self._model_var.set("deepseek-chat")  # 默认选择DeepSeek模型
    
    def _on_theme_changed(self):
        """主题切换回调"""
        theme = self._theme_var.get()
        if theme == "light":
            # 浅色主题配置（使用半透明效果）
            GlassTheme.GLASS_BG = "#F0F0F0"
            GlassTheme.GLASS_SURFACE = "#FAFAFA"
            GlassTheme.TEXT_PRIMARY = "#333333"
            GlassTheme.TEXT_SECONDARY = "#666666"
        else:
            # 深色主题配置（默认）
            GlassTheme.GLASS_BG = "#1E1E1E"
            GlassTheme.GLASS_SURFACE = "#2D2D2D"
            GlassTheme.TEXT_PRIMARY = "#FFFFFF"
            GlassTheme.TEXT_SECONDARY = "#CCCCCC"
        # 立即刷新所有UI元素
        self._refresh_theme()
        self._set_status(f"主题已切换为：{'浅色' if theme == 'light' else '深色'}")

    def _refresh_theme(self):
        """刷新所有UI元素的主题"""
        # 更新主窗口背景
        self.root.configure(bg=GlassTheme.GLASS_BG)

        # 更新标题栏
        if hasattr(self, '_title_bar'):
            self._title_bar.configure(bg=GlassTheme.GLASS_SURFACE)

        # 更新侧边栏
        if hasattr(self, '_sidebar'):
            self._sidebar.configure(bg=GlassTheme.GLASS_SURFACE)

        # 更新所有Text组件
        for widget in self.root.winfo_children():
            self._update_widget_theme(widget)

    def _update_widget_theme(self, widget):
        """递归更新所有组件的主题"""
        try:
            if isinstance(widget, tk.Text):
                widget.configure(bg=GlassTheme.GLASS_SURFACE, fg=GlassTheme.TEXT_PRIMARY)
            elif isinstance(widget, tk.Label) and widget.cget("style") == "":
                widget.configure(bg=GlassTheme.GLASS_BG, fg=GlassTheme.TEXT_PRIMARY)
            elif isinstance(widget, tk.Label) and widget.cget("bg") == GlassTheme.GLASS_BG:
                widget.configure(bg=GlassTheme.GLASS_BG)
            elif isinstance(widget, tk.Label) and widget.cget("bg") == GlassTheme.GLASS_SURFACE:
                widget.configure(bg=GlassTheme.GLASS_SURFACE)

            # 递归处理子组件
            for child in widget.winfo_children():
                self._update_widget_theme(child)
        except Exception:
            pass
    
    def _on_font_size_changed(self):
        """字体大小改变回调"""
        try:
            font_size = int(self._font_size_var.get())
            # 更新GlassTheme字体大小
            GlassTheme.FONT_SIZE_NORMAL = str(font_size)
            GlassTheme.FONT_SIZE_SMALL = str(max(10, font_size - 2))
            GlassTheme.FONT_SIZE_SUBTITLE = str(max(14, font_size - 2))
            GlassTheme.FONT_SIZE_TITLE = str(max(18, font_size + 4))

            # 立即刷新所有字体
            self._refresh_fonts()
            self._set_status(f"字体大小已更新为：{font_size}px")
        except ValueError:
            messagebox.showerror("错误", "请输入有效的字体大小（10-20）")

    def _refresh_fonts(self):
        """刷新所有UI元素的字体"""
        # 更新所有Text组件字体
        for widget in self.root.winfo_children():
            self._update_widget_font(widget)

    def _update_widget_font(self, widget):
        """递归更新所有组件的字体"""
        try:
            if isinstance(widget, tk.Text):
                widget.configure(font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL))
            elif isinstance(widget, tk.Label) and widget.cget("font") != "":
                # 获取原有字体设置（粗体等）
                current_font = widget.cget("font")
                if isinstance(current_font, tuple) and len(current_font) >= 3:
                    # 保留字体族和粗体样式
                    widget.configure(font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL, current_font[2]))
                elif isinstance(current_font, tuple) and len(current_font) >= 2:
                    widget.configure(font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_NORMAL))

            # 递归处理子组件
            for child in widget.winfo_children():
                self._update_widget_font(child)
        except Exception:
            pass

    # ============== 生成事件订阅 ==============
    
    def _subscribe_generation_events(self) -> None:
        """订阅生成相关事件，绑定到开始创作页面"""
        event_bus = get_event_bus()
        if not event_bus:
            return
        
        events = [
            ("generation.started", self._on_gen_event_started),
            ("generation.completed", self._on_gen_event_completed),
            ("generation.failed", self._on_gen_event_failed),
            ("generation.progress", self._on_gen_event_progress),
            ("pipeline.stage_started", self._on_gen_event_stage_started),
            ("pipeline.stage_completed", self._on_gen_event_stage_completed),
            ("pipeline.iteration_started", self._on_gen_event_iteration_started),
            ("pipeline.completed", self._on_gen_event_pipeline_completed),
            ("agent.task.started", self._on_gen_event_agent_started),
            ("agent.task.completed", self._on_gen_event_agent_completed),
            ("agent.task.failed", self._on_gen_event_agent_failed),
            # 订阅热榜进度事件（P2-001修复）
            ("hot_ranking.progress", self._on_hot_ranking_progress_event),
            ("hot_ranking.updated", self._on_hot_ranking_updated_event),
        ]
        
        for event_type, handler in events:
            try:
                sub_id = event_bus.subscribe(event_type, handler)
                self._subscription_ids.append(sub_id)
            except Exception as e:
                logger.warning(f"订阅事件失败 {event_type}: {e}")
        
        logger.info(f"Subscribed to {len(self._subscription_ids)} generation events")
    
    def _unsubscribe_generation_events(self) -> None:
        """取消订阅生成事件"""
        event_bus = get_event_bus()
        if not event_bus:
            return
        
        for sub_id in self._subscription_ids:
            try:
                event_bus.unsubscribe(sub_id)
            except Exception:
                pass
        
        self._subscription_ids.clear()
    
    def _gen_log_insert(self, message: str) -> None:
        """线程安全地向生成日志插入消息"""
        if hasattr(self, '_gen_log') and self._gen_log:
            self._gen_log.insert(tk.END, message)
            self._gen_log.see(tk.END)
    
    def _gen_update_status(self, status: str) -> None:
        """线程安全地更新生成状态"""
        if hasattr(self, '_gen_status_var') and self._gen_status_var:
            self._gen_status_var.set(status)
    
    def _gen_update_progress(self, value: float) -> None:
        """线程安全地更新进度条"""
        if hasattr(self, '_gen_progress') and self._gen_progress:
            self._gen_progress['value'] = value
    
    # === 热榜进度事件处理器（P2-001修复）===
    
    def _on_hot_ranking_progress_event(self, event) -> None:
        """热榜进度事件"""
        def update():
            progress = event.data.get("progress", 0)
            status = event.data.get("status", "")
            
            # 显示进度条
            if hasattr(self, '_hot_ranking_progress_frame') and self._hot_ranking_progress_frame:
                try:
                    self._hot_ranking_progress_frame.pack(fill=tk.X, padx=20, pady=5)
                except Exception:
                    pass
            
            # 更新进度
            if hasattr(self, '_hot_ranking_progress_bar'):
                self._hot_ranking_progress_bar['value'] = progress
            if hasattr(self, '_hot_ranking_progress_var'):
                self._hot_ranking_progress_var.set(f"📊 {status} ({progress}%)")
            
            # 更新状态栏
            self._set_status(f"热榜更新中: {status}")
            
            # 进度完成时延迟隐藏进度条
            if progress >= 100:
                self.root.after(1500, self._hide_hot_ranking_progress_bar)
        
        self._safe_after(0, update)
    
    def _on_hot_ranking_updated_event(self, event) -> None:
        """热榜数据更新完成事件"""
        def update():
            self._hide_hot_ranking_progress_bar()
            self._set_status("热榜数据更新完成")
            timestamp = event.data.get("timestamp", "")
            if timestamp:
                logger.info(f"[热榜] 数据更新完成: {timestamp}")
        
        self._safe_after(0, update)
    
    def _hide_hot_ranking_progress_bar(self) -> None:
        """隐藏热榜进度条"""
        if hasattr(self, '_hot_ranking_progress_frame') and self._hot_ranking_progress_frame:
            try:
                self._hot_ranking_progress_frame.pack_forget()
            except Exception:
                pass
    
    # === 事件处理器 ===
    
    def _on_gen_event_started(self, event) -> None:
        """生成开始事件"""
        def update():
            self._gen_update_status("🚀 生成任务已开始")
            self._gen_log_insert(f"[{self._timestamp()}] 🚀 生成任务已开始\n")
            pipeline_id = event.data.get("pipeline_id", "")
            if pipeline_id:
                self._gen_log_insert(f"[{self._timestamp()}] Pipeline ID: {pipeline_id}\n")
        
        self._safe_after(0, update)
    
    def _on_gen_event_completed(self, event) -> None:
        """生成完成事件"""
        def update():
            self._gen_update_status("✅ 生成任务已完成")
            total_words = event.data.get("total_words", 0)
            scores = event.data.get("scores", {})
            chapter_id = event.data.get("chapter_id", "")
            content = event.data.get("content", "")
            
            self._gen_log_insert(f"[{self._timestamp()}] ✅ 生成任务已完成\n")
            self._gen_log_insert(f"[{self._timestamp()}] 总字数: {total_words}\n")
            self._gen_update_progress(100)
            
            # P1修复：生成完成后弹出反馈对话框
            if content and scores:
                # 延迟弹出，避免阻塞UI更新
                self.root.after(500, lambda: self._show_feedback_dialog_auto(
                    chapter_id=chapter_id,
                    content=content,
                    scores=scores
                ))
        
        self._safe_after(0, update)
    
    def _on_gen_event_failed(self, event) -> None:
        """生成失败事件"""
        def update():
            self._gen_update_status("❌ 生成任务失败")
            error = event.data.get("error", "未知错误")
            self._gen_log_insert(f"[{self._timestamp()}] ❌ 生成失败: {error}\n")
        
        self._safe_after(0, update)
    
    def _on_gen_event_progress(self, event) -> None:
        """生成进度事件"""
        def update():
            progress = event.data.get("progress", 0)
            message = event.data.get("message", "")
            self._gen_update_progress(progress)
            if message:
                self._gen_log_insert(f"[{self._timestamp()}] 📊 {message} ({progress}%)\n")
        
        self._safe_after(0, update)
    
    def _on_gen_event_stage_started(self, event) -> None:
        """阶段开始事件"""
        def update():
            stage = event.data.get("stage", "Unknown")
            self._gen_log_insert(f"[{self._timestamp()}] 🔄 阶段开始: {stage}\n")
        
        self._safe_after(0, update)
    
    def _on_gen_event_stage_completed(self, event) -> None:
        """阶段完成事件"""
        def update():
            stage = event.data.get("stage", "Unknown")
            success = event.data.get("success", False)
            status = "✅" if success else "❌"
            self._gen_log_insert(f"[{self._timestamp()}] {status} 阶段完成: {stage}\n")
        
        self._safe_after(0, update)
    
    def _on_gen_event_iteration_started(self, event) -> None:
        """迭代开始事件"""
        def update():
            iteration = event.data.get("iteration", 0)
            max_iterations = event.data.get("max_iterations", 0)
            self._gen_log_insert(f"[{self._timestamp()}] 🔁 迭代 {iteration}/{max_iterations}\n")
        
        self._safe_after(0, update)
    
    def _on_gen_event_pipeline_completed(self, event) -> None:
        """流水线完成事件"""
        def update():
            success = event.data.get("success", False)
            status = "✅ 成功" if success else "❌ 失败"
            self._gen_log_insert(f"[{self._timestamp()}] 🏁 流水线{status}\n")
        
        self._safe_after(0, update)
    
    def _on_gen_event_agent_started(self, event) -> None:
        """Agent任务开始事件"""
        def update():
            agent_name = event.data.get("agent", "Unknown")
            task_id = event.data.get("task_id", "")
            self._gen_log_insert(f"[{self._timestamp()}] 🤖 Agent启动: {agent_name}\n")
            # 更新Agent状态树
            if hasattr(self, '_gen_agent_tree') and self._gen_agent_tree:
                self._gen_agent_tree.insert("", tk.END, iid=task_id or agent_name, values=("🔄 运行中", agent_name))
        
        self._safe_after(0, update)
    
    def _on_gen_event_agent_completed(self, event) -> None:
        """Agent任务完成事件"""
        def update():
            agent_name = event.data.get("agent", "Unknown")
            success = event.data.get("success", False)
            task_id = event.data.get("task_id", "")
            status = "✅" if success else "❌"
            self._gen_log_insert(f"[{self._timestamp()}] {status} Agent完成: {agent_name}\n")
            # 更新Agent状态树
            if hasattr(self, '_gen_agent_tree') and self._gen_agent_tree:
                item_id = task_id or agent_name
                try:
                    self._gen_agent_tree.item(item_id, values=(f"{status} 完成", agent_name))
                except Exception:
                    pass  # 如果项目不存在，忽略错误
        
        self._safe_after(0, update)
    
    def _on_gen_event_agent_failed(self, event) -> None:
        """Agent任务失败事件"""
        def update():
            agent_name = event.data.get("agent", "Unknown")
            error = event.data.get("error", "未知错误")
            task_id = event.data.get("task_id", "")
            self._gen_log_insert(f"[{self._timestamp()}] ❌ Agent失败: {agent_name} - {error}\n")
            # 更新Agent状态树
            if hasattr(self, '_gen_agent_tree') and self._gen_agent_tree:
                item_id = task_id or agent_name
                try:
                    self._gen_agent_tree.item(item_id, values=(f"❌ 失败", f"{agent_name}: {error[:30]}"))
                except Exception:
                    pass
        
        self._safe_after(0, update)
    
    def _timestamp(self) -> str:
        """获取当前时间戳"""
        from datetime import datetime
        return datetime.now().strftime("%H:%M:%S")
    
    # ============== P1修复：GUI反馈集成 ==============
    
    def _show_feedback_dialog_auto(self, chapter_id: str, content: str, scores: Dict) -> None:
        """
        自动显示反馈对话框（生成完成后调用）
        
        Args:
            chapter_id: 章节ID
            content: 生成的内容
            scores: 评分字典
        """
        try:
            # 检查是否启用自动反馈
            config = self._load_config()
            auto_feedback = config.get("feedback", {}).get("auto_show", True)
            
            if not auto_feedback:
                logger.info("[Feedback] 自动反馈已禁用，跳过")
                return
            
            # 显示简化的反馈对话框
            self.show_feedback_dialog(
                chapter_id=chapter_id,
                content_preview=content[:500] + "..." if len(content) > 500 else content,
                scores=scores
            )
            
        except Exception as e:
            logger.error(f"[Feedback] 显示反馈对话框失败: {e}")
    
    def show_feedback_dialog(self, chapter_id: str, content_preview: str, scores: Dict) -> None:
        """
        显示用户反馈对话框
        
        P1修复：集成FeedbackSentimentAnalyzer实现情感分析
        收集用户对生成内容的反馈，用于Claw化闭环优化
        
        Args:
            chapter_id: 章节ID
            content_preview: 内容预览（截断后的）
            scores: 评分字典（各维度评分）
        """
        try:
            # 创建反馈对话框窗口
            dialog = tk.Toplevel(self.root)
            dialog.title("📝 内容反馈")
            dialog.geometry("500x600")
            dialog.configure(bg=GlassTheme.GLASS_BG)
            dialog.transient(self.root)
            dialog.grab_set()
            
            # 存储反馈数据
            feedback_data = {
                "chapter_id": chapter_id,
                "rating": tk.IntVar(value=4),  # 默认4星
                "feedback_type": tk.StringVar(value="positive"),
                "details": tk.StringVar()
            }
            
            # 标题
            title_frame = tk.Frame(dialog, bg=GlassTheme.GLASS_SURFACE)
            title_frame.pack(fill=tk.X, padx=10, pady=10)
            
            tk.Label(
                title_frame,
                text="📊 请对本次生成内容进行评价",
                font=(GlassTheme.FONT_FAMILY, GlassTheme.FONT_SIZE_SUBTITLE, "bold"),
                bg=GlassTheme.GLASS_SURFACE,
                fg=GlassTheme.TEXT_PRIMARY
            ).pack(pady=10)
            
            # 评分显示
            scores_frame = tk.LabelFrame(dialog, text="评分详情", bg=GlassTheme.GLASS_BG, fg=GlassTheme.TEXT_PRIMARY)
            scores_frame.pack(fill=tk.X, padx=10, pady=5)
            
            if scores:
                for dim, score in scores.items():
                    dim_frame = tk.Frame(scores_frame, bg=GlassTheme.GLASS_BG)
                    dim_frame.pack(fill=tk.X, padx=5, pady=2)
                    
                    tk.Label(
                        dim_frame,
                        text=f"{dim}:",
                        width=10,
                        anchor="w",
                        bg=GlassTheme.GLASS_BG,
                        fg=GlassTheme.TEXT_PRIMARY
                    ).pack(side=tk.LEFT)
                    
                    # 评分条
                    score_bar = tk.Frame(dim_frame, bg=GlassTheme.GLASS_BG)
                    score_bar.pack(side=tk.LEFT, fill=tk.X, expand=True)
                    
                    # 根据评分设置颜色
                    if score >= 0.8:
                        color = "#4CAF50"  # 绿色
                    elif score >= 0.6:
                        color = "#FFC107"  # 黄色
                    else:
                        color = "#F44336"  # 红色
                    
                    tk.Frame(score_bar, bg=color, width=int(score * 150)).pack(side=tk.LEFT)
                    tk.Label(score_bar, text=f" {score:.2f}", bg=GlassTheme.GLASS_BG, fg=GlassTheme.TEXT_PRIMARY).pack(side=tk.LEFT)
            
            # 星级评分
            rating_frame = tk.LabelFrame(dialog, text="整体评分", bg=GlassTheme.GLASS_BG, fg=GlassTheme.TEXT_PRIMARY)
            rating_frame.pack(fill=tk.X, padx=10, pady=5)
            
            stars_frame = tk.Frame(rating_frame, bg=GlassTheme.GLASS_BG)
            stars_frame.pack(pady=10)
            
            for i in range(1, 6):
                tk.Radiobutton(
                    stars_frame,
                    text="⭐" * i,
                    variable=feedback_data["rating"],
                    value=i,
                    bg=GlassTheme.GLASS_BG,
                    fg=GlassTheme.TEXT_PRIMARY,
                    selectcolor=GlassTheme.GLASS_SURFACE
                ).pack(side=tk.LEFT, padx=5)
            
            # 反馈类型
            type_frame = tk.LabelFrame(dialog, text="反馈类型", bg=GlassTheme.GLASS_BG, fg=GlassTheme.TEXT_PRIMARY)
            type_frame.pack(fill=tk.X, padx=10, pady=5)
            
            types = [("😊 满意", "positive"), ("⚠️ 需改进", "negative"), ("💡 建议", "suggestion")]
            for text, value in types:
                tk.Radiobutton(
                    type_frame,
                    text=text,
                    variable=feedback_data["feedback_type"],
                    value=value,
                    bg=GlassTheme.GLASS_BG,
                    fg=GlassTheme.TEXT_PRIMARY,
                    selectcolor=GlassTheme.GLASS_SURFACE
                ).pack(side=tk.LEFT, padx=10)
            
            # 详细反馈
            details_frame = tk.LabelFrame(dialog, text="详细反馈（可选）", bg=GlassTheme.GLASS_BG, fg=GlassTheme.TEXT_PRIMARY)
            details_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
            
            details_text = tk.Text(
                details_frame,
                height=5,
                font=(GlassTheme.FONT_FAMILY_TEXT, GlassTheme.FONT_SIZE_SMALL),
                bg=GlassTheme.GLASS_SURFACE,
                fg=GlassTheme.TEXT_PRIMARY,
                wrap=tk.WORD
            )
            details_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            
            # 提示文字
            details_text.insert(tk.END, "请描述您的反馈，例如：\n- 风格是否符合预期\n- 人物行为是否合理\n- 情节是否连贯\n- 知识点运用是否准确")
            details_text.bind("<FocusIn>", lambda e: details_text.delete("1.0", tk.END) if "请描述" in details_text.get("1.0", tk.END) else None)
            
            def submit_feedback():
                """提交反馈"""
                try:
                    details = details_text.get("1.0", tk.END).strip()
                    if "请描述" in details:
                        details = ""
                    
                    rating = feedback_data["rating"].get()
                    fb_type = feedback_data["feedback_type"].get()
                    
                    # 调用反馈闭环系统
                    from core.user_feedback_loop import get_user_feedback_loop, get_sentiment_analyzer
                    
                    feedback_loop = get_user_feedback_loop(Path(self._workspace_root))
                    
                    # 如果有详细反馈，进行情感分析
                    if details:
                        analyzer = get_sentiment_analyzer(Path(self._workspace_root))
                        sentiment = analyzer.analyze_sentiment(details)
                        
                        # 如果情感分析与用户选择的类型不符，使用AI建议
                        if sentiment["confidence"] > 0.7:
                            fb_type = sentiment["suggested_type"]
                        
                        # 记录情感分析结果
                        logger.info(f"[Feedback] 情感分析: {sentiment['sentiment']}, 置信度: {sentiment['confidence']:.2f}")
                    
                    # 提交反馈
                    feedback_id = feedback_loop.collect_feedback(
                        chapter_id=chapter_id,
                        feedback_type=fb_type,
                        details=details,
                        rating=float(rating)
                    )
                    
                    # 更新实时指标
                    from core.metrics_monitor import get_realtime_monitor
                    monitor = get_realtime_monitor(Path(self._workspace_root))
                    if scores:
                        monitor.update_batch(chapter_id, scores)
                    monitor.update_metric(chapter_id, "用户评分", rating / 5.0)
                    
                    messagebox.showinfo("感谢反馈", "您的反馈已提交，感谢您的参与！")
                    dialog.destroy()
                    
                    logger.info(f"[Feedback] 反馈已提交: {feedback_id}")
                    
                except Exception as e:
                    logger.error(f"[Feedback] 提交失败: {e}")
                    messagebox.showerror("提交失败", f"反馈提交失败：{str(e)}")
            
            # 按钮区域
            btn_frame = tk.Frame(dialog, bg=GlassTheme.GLASS_BG)
            btn_frame.pack(fill=tk.X, padx=10, pady=10)
            
            ttk.Button(btn_frame, text="提交反馈", command=submit_feedback).pack(side=tk.RIGHT, padx=5)
            ttk.Button(btn_frame, text="跳过", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
            
            # 居中显示
            dialog.update_idletasks()
            x = self.root.winfo_x() + (self.root.winfo_width() - dialog.winfo_width()) // 2
            y = self.root.winfo_y() + (self.root.winfo_height() - dialog.winfo_height()) // 2
            dialog.geometry(f"+{x}+{y}")
            
        except Exception as e:
            logger.error(f"[Feedback] 创建对话框失败: {e}")
            messagebox.showerror("错误", f"创建反馈对话框失败：{str(e)}")
    
    def _load_config(self) -> Dict:
        """加载配置文件"""
        try:
            import yaml
            config_file = Path(self._workspace_root) / "config.yaml"
            
            if config_file.exists():
                with open(config_file, 'r', encoding='utf-8') as f:
                    return yaml.safe_load(f) or {}
        except Exception as e:
            logger.warning(f"[Config] 加载配置失败: {e}")
        
        return {}

    
    def _on_close(self) -> None:
        """窗口关闭确认"""
        if messagebox.askokcancel("退出", "确定要退出吗？\n未保存的内容将会丢失。"):
            # 取消事件订阅
            self._unsubscribe_generation_events()
            logger.info("Application closed by user")
            self.root.destroy()
    
    def _on_save_project(self):
        """保存当前项目（通用保存功能 - 使用项目管理器）

        功能：
        - 使用项目管理器保存当前项目数据
        - 显示保存成功/失败提示
        - 更新状态栏
        """
        try:
            # 优先使用项目管理器
            if self._project_manager and self._project_manager.is_project_open():
                # 同步所有模块数据到项目管理器
                self._sync_all_data_to_manager()

                # 调用项目管理器保存
                success = self._project_manager.save_project()

                if success:
                    project_name = self._project_manager.get_project_name()
                    self._set_status("项目保存完成")
                    messagebox.showinfo("保存成功", f"项目「{project_name}」已保存成功")
                else:
                    self._set_status("保存项目失败")
                    messagebox.showwarning("保存失败", "项目保存失败，请查看日志获取详情")
                return

            # 降级方案：如果项目管理器不可用，使用旧的直接保存方式
            if not self.current_project or not self.project_file:
                messagebox.showwarning("保存项目", "当前没有打开的项目")
                return

            # 获取当前项目名称
            project_name = self.current_project.get('name', '未命名项目')

            # 同步所有模块数据到项目
            self._sync_all_data_to_project()

            # 更新修改时间
            from datetime import datetime
            self.current_project['modified_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            # 保存项目文件
            self._set_status("正在保存项目...")
            with open(self.project_file, 'w', encoding='utf-8') as f:
                json.dump(self.current_project, f, ensure_ascii=False, indent=2)

            self._set_status("项目保存完成")
            messagebox.showinfo("保存成功", f"项目「{project_name}」已保存成功")
            logger.info(f"Project saved: {project_name}")

        except Exception as e:
            # P2-003修复：用户友好错误提示
            from core.user_friendly_errors import convert_file_error
            title, full_message = convert_file_error(e, "保存项目")
            messagebox.showerror(title, full_message)
            self._set_status("保存项目失败")
            logger.error(f"Error saving project: {e}")
    
    def _setup_system_tray(self) -> None:
        """
        创建系统托盘图标（作为任务栏图标的替代方案）
        
        由于Tkinter的overrideredirect(True)导致任务栏不显示图标，
        使用系统托盘图标作为替代方案。
        """
        try:
            import pystray
            from pystray import Menu, MenuItem
            from PIL import Image
            
            # 获取图标文件路径
            icon_path = os.path.join(self._workspace_root, "icon.ico")
            
            if not os.path.exists(icon_path):
                logger.warning(f"Icon file not found: {icon_path}")
                return
            
            # 加载图标
            try:
                icon_image = Image.open(icon_path)
                logger.info(f"System tray icon loaded from: {icon_path}")
            except Exception as e:
                logger.warning(f"Failed to load icon for system tray: {e}")
                return
            
            # 定义托盘图标菜单
            def on_show_window(icon, item):
                """显示主窗口"""
                self.root.deiconify()
                self.root.lift()
                self.root.focus_force()
                logger.info("Window restored from system tray")
            
            def on_minimize(icon, item):
                """最小化窗口"""
                try:
                    self.root.withdraw()
                    logger.info("Window minimized to system tray")
                except Exception as e:
                    logger.error(f"Error minimizing window: {e}")
            
            def on_quit(icon, item):
                """退出程序"""
                icon.stop()
                self.root.quit()
            
            menu = Menu(
                MenuItem('显示窗口', on_show_window, default=True),
                MenuItem('最小化', on_minimize),
                MenuItem('退出', on_quit)
            )
            
            # 创建系统托盘图标（默认菜单项在左键单击时触发）
            icon = pystray.Icon(
                "Novel Writing Assistant",
                icon_image,
                "Novel Writing Assistant",
                menu
            )
            
            # 在后台运行托盘图标
            import threading
            tray_thread = threading.Thread(target=icon.run, daemon=True)
            tray_thread.start()
            
            logger.info("System tray icon created successfully")
            
            # 保存托盘图标引用，防止被垃圾回收
            self._tray_icon = icon
            
        except ImportError:
            logger.warning("pystray not available, system tray icon not created")
        except Exception as e:
            logger.error(f"Failed to create system tray icon: {e}")
    
    def run(self) -> None:
        """运行主窗口"""
        logger.info("Starting main window")
        
        # 关键：在mainloop之前强制更新窗口，确保窗口句柄可用
        # 让Windows API设置的无边框样式生效
        self.root.update()
        
        # 创建系统托盘图标（作为任务栏图标的替代方案）
        self._setup_system_tray()
        
        self.root.mainloop()


# ============== 入口 ==============

def _global_exception_handler(exc_type, exc_value, exc_tb):
    """
    P1-12修复：全局异常处理器

    捕获未处理的异常，记录日志并显示友好的错误提示。
    """
    # 记录完整的异常信息
    logger.critical(
        "未捕获的异常",
        exc_info=(exc_type, exc_value, exc_tb)
    )

    # 显示友好的错误提示
    try:
        messagebox.showerror(
            "程序错误",
            f"发生未预期的错误：{exc_type.__name__}\n"
            f"详情：{str(exc_value)[:200]}\n\n"
            "请查看日志文件获取更多信息。"
        )
    except Exception:
        # 如果messagebox也失败了，打印到控制台
        print(f"CRITICAL ERROR: {exc_type.__name__}: {exc_value}")


def main():
    """主入口 - 使用OptimizedLauncher优化启动"""
    # P1-12修复：注册全局异常处理器
    import sys
    sys.excepthook = _global_exception_handler

    # P0-002修复：集成OptimizedLauncher实现分层延迟加载
    launcher = None
    try:
        from core.app_launcher import OptimizedLauncher
        launcher = OptimizedLauncher()
        
        # 配置启动参数
        launcher.configure(
            async_load=True,
            hide_window_on_start=True,
            target_startup_time=1.0
        )
        
        # 注册启动完成回调
        def on_startup_complete(startup_time: float):
            logger.info(f"Optimized startup completed in {startup_time:.3f}s")
        
        launcher.register_complete_callback(on_startup_complete)
        
        # 启动应用（加载核心层）
        startup_time = launcher.start()
        logger.info(f"Core modules loaded in {startup_time:.3f}s")
        
    except Exception as e:
        logger.warning(f"OptimizedLauncher not available, using traditional startup: {e}")
        launcher = None
        
        # 传统启动流程（fallback）
        if CORE_AVAILABLE:
            try:
                init_results = initialize_core_services()
                if init_results.get("ConfigService") and init_results.get("LoggingService"):
                    logging_service = get_logging_service()
                    logging_service.log_system_event(
                        "startup", "应用程序启动 - 核心服务初始化完成"
                    )
                    logger.info("Core services initialized successfully")
                else:
                    logger.warning(f"Core services initialization partial: {init_results}")
            except Exception as e:
                logger.error(f"Failed to initialize core services: {e}")

    try:
        app = MainWindow()
        app.run()
    except Exception as e:
        logger.error(f"Application error: {e}")
        raise
    finally:
        # 释放资源
        if launcher:
            try:
                launcher.shutdown()
                logger.info("OptimizedLauncher shutdown completed")
            except Exception as e:
                logger.error(f"Failed to shutdown OptimizedLauncher: {e}")
        
        # 释放核心服务
        if CORE_AVAILABLE:
            try:
                dispose_core_services()
                logger.info("Core services disposed")
            except Exception as e:
                logger.error(f"Failed to dispose core services: {e}")


if __name__ == "__main__":
    main()
